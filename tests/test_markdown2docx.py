from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

from mdparse.markdown2docx import (
    BULLETED_LIST_INDENT,
    add_hyperlink,
    combine_documents,
    get_or_create_hyperlink_style,
    insert_hr,
    markdown_to_docx,
)


def test_get_or_create_hyperlink_style():
    doc = Document()
    # Style does not exist yet
    style1 = get_or_create_hyperlink_style(doc)
    assert style1.name == "Hyperlink"
    assert style1.font.color.rgb == RGBColor(0, 0, 255)
    assert style1.font.underline is True
    # Calling again returns same style
    style2 = get_or_create_hyperlink_style(doc)
    assert style1 == style2


def test_add_hyperlink():
    doc = Document()
    p = doc.add_paragraph()
    hyperlink = add_hyperlink(p, "Example", "http://example.com")
    assert hyperlink.tag.endswith("hyperlink")
    r_id = hyperlink.get(qn("r:id"))
    assert r_id is not None
    # The paragraph text should include the link text
    assert "Example" in p.text


def test_insert_hr():
    doc = Document()
    p = doc.add_paragraph()
    insert_hr(p)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    assert pBdr is not None
    bottom = pBdr.find(qn("w:bottom"))
    assert bottom is not None
    assert bottom.get(qn("w:val")) == "single"
    assert bottom.get(qn("w:sz")) == "6"
    assert bottom.get(qn("w:space")) == "1"
    assert bottom.get(qn("w:color")) == "auto"


def test_markdown_to_docx_code_block():
    md = "```\nline1\nline2\n```"
    doc = markdown_to_docx(md)
    code_paras = [p for p in doc.paragraphs if p.style.name == "Code"]
    texts = [p.text for p in code_paras]
    assert texts == ["line1", "line2"]


def test_markdown_to_docx_inline_and_headings_and_blockquote_and_hr_and_image():
    md = "\n".join(
        [
            "# Heading1",
            "",
            "**bold**",
            "",
            "*italic*",
            "",
            "`code`",
            "",
            "[link](http://example.com)",
            "",
            "> quote",
            "",
            "---",
            "",
            "![AltText](no_image.png)",
        ]
    )
    doc = markdown_to_docx(md)
    # Heading
    heading_paras = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any(p.text == "Heading1" for p in heading_paras)
    # Bold
    bold_runs = [run for p in doc.paragraphs for run in p.runs if run.text == "bold"]
    assert any(run.bold for run in bold_runs)
    # Italic
    italic_runs = [run for p in doc.paragraphs for run in p.runs if run.text == "italic"]
    assert any(run.italic for run in italic_runs)
    # Code font
    code_runs = [run for p in doc.paragraphs for run in p.runs if run.text == "code"]
    assert any(run.font.name == "Courier New" for run in code_runs)
    # Blockquote
    quote_para = next(p for p in doc.paragraphs if p.text == "quote")
    assert quote_para.style.name == "Intense Quote"
    # Horizontal rule
    hr_paras = [p for p in doc.paragraphs if p._p.pPr is not None and p._p.pPr.find(qn("w:pBdr")) is not None]
    assert hr_paras, "Expected a paragraph with a horizontal rule"
    # Image not found branch
    img_para = next(p for p in doc.paragraphs if "Image:" in p.text)
    assert "AltText" in img_para.text and "no_image.png" in img_para.text


def test_markdown_to_docx_lists():
    md = "\n".join(["- item1", "  - nested", "1. num1", "2. num2"])
    doc = markdown_to_docx(md)
    # Bullet items
    bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert any(p.text == "item1" for p in bullets)
    # Nested bullet indentation
    nested = next(p for p in doc.paragraphs if p.text == "nested")
    assert nested.style.name == "List Bullet"
    expected_indent = (1 + 1) * BULLETED_LIST_INDENT
    assert nested.paragraph_format.left_indent.pt == expected_indent
    # Numbered items
    nums = [p for p in doc.paragraphs if p.style.name == "List Number"]
    texts = [p.text for p in nums]
    assert "num1" in texts and "num2" in texts


def test_markdown_to_docx_table():
    md = "\n".join(["| Col1 | Col2 |", "|------|------|", "| A **bold** | B *italic* |"])
    doc = markdown_to_docx(md)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    # Header bold
    header = table.rows[0]
    for cell in header.cells:
        for run in cell.paragraphs[0].runs:
            assert run.bold
    # Data row formatting
    data = table.rows[1]
    a_cell_runs = data.cells[0].paragraphs[0].runs
    b_cell_runs = data.cells[1].paragraphs[0].runs
    assert any(run.bold for run in a_cell_runs if run.text == "bold")
    assert any(run.italic for run in b_cell_runs if run.text == "italic")


def test_combine_documents_without_page_breaks():
    doc1 = Document()
    doc1.add_paragraph("Doc1")
    doc2 = Document()
    doc2.add_paragraph("Doc2")
    merged = combine_documents([doc1, doc2], add_page_breaks=False)
    texts = [p.text for p in merged.paragraphs]
    assert "Doc1" in texts and "Doc2" in texts
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    brs = merged.element.body.xpath('.//w:br[@w:type="page"]')
    assert len(brs) == 0


def test_combine_documents_with_page_breaks():
    doc1 = Document()
    doc1.add_paragraph("Doc1")
    doc2 = Document()
    doc2.add_paragraph("Doc2")
    doc3 = Document()
    doc3.add_paragraph("Doc3")
    merged = combine_documents([doc1, doc2, doc3], add_page_breaks=True)
    texts = [p.text for p in merged.paragraphs]
    assert texts.index("Doc1") < texts.index("Doc2") < texts.index("Doc3")
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    brs = merged.element.body.xpath('.//w:br[@w:type="page"]')
    assert len(brs) == 1
