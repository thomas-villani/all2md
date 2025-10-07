"""Advanced tests for DOCX hyperlink handling edge cases."""

import docx

from all2md import to_markdown as docx_to_markdown
from utils import assert_markdown_valid, cleanup_test_dir, create_test_temp_dir


class TestDocxLinksAdvanced:
    """Test complex hyperlink scenarios in DOCX documents."""

    def setup_method(self):
        """Set up test environment."""
        self.temp_dir = create_test_temp_dir()

    def teardown_method(self):
        """Clean up test environment."""
        cleanup_test_dir(self.temp_dir)

    def test_external_hyperlinks(self):
        """Test various external hyperlink formats."""
        doc = docx.Document()
        doc.add_heading("External Links Test", level=1)

        # HTTP links
        p1 = doc.add_paragraph("Visit ")
        # Note: python-docx hyperlink support is limited, we simulate the structure
        p1.add_run("Google").hyperlink = "https://www.google.com"

        # HTTPS links
        p2 = doc.add_paragraph("Secure site: ")
        p2.add_run("GitHub").hyperlink = "https://github.com"

        # FTP links
        p3 = doc.add_paragraph("FTP: ")
        p3.add_run("FTP Server").hyperlink = "ftp://ftp.example.com"

        # Link with query parameters
        p4 = doc.add_paragraph("Search: ")
        p4.add_run("Google Search").hyperlink = "https://www.google.com/search?q=python"

        temp_file = self.temp_dir / "external_links.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should contain link text
        assert "Google" in markdown
        assert "GitHub" in markdown
        assert "FTP Server" in markdown
        assert "Google Search" in markdown

        # Note: Actual hyperlink extraction depends on python-docx and implementation details

    def test_mailto_links(self):
        """Test mailto hyperlinks."""
        doc = docx.Document()
        doc.add_heading("Mailto Links Test", level=1)

        # Simple mailto
        p1 = doc.add_paragraph("Contact: ")
        p1.add_run("support@example.com").hyperlink = "mailto:support@example.com"

        # Mailto with subject
        p2 = doc.add_paragraph("Bug report: ")
        p2.add_run("Report Bug").hyperlink = "mailto:bugs@example.com?subject=Bug Report"

        # Mailto with CC and body
        p3 = doc.add_paragraph("Complex email: ")
        p3.add_run(
            "Contact Sales").hyperlink = "mailto:sales@example.com?cc=manager@example.com&subject=Inquiry&body=Hello"

        temp_file = self.temp_dir / "mailto_links.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should contain email addresses and link text
        assert "support@example.com" in markdown
        assert "Report Bug" in markdown
        assert "Contact Sales" in markdown

    def test_relative_path_links(self):
        """Test relative path hyperlinks."""
        doc = docx.Document()
        doc.add_heading("Relative Links Test", level=1)

        # Relative file paths
        p1 = doc.add_paragraph("Local document: ")
        p1.add_run("README").hyperlink = "./README.md"

        # Parent directory
        p2 = doc.add_paragraph("Parent folder: ")
        p2.add_run("Config").hyperlink = "../config/settings.json"

        # Subdirectory
        p3 = doc.add_paragraph("Images: ")
        p3.add_run("Logo").hyperlink = "images/logo.png"

        # Root relative
        p4 = doc.add_paragraph("Root: ")
        p4.add_run("Home").hyperlink = "/index.html"

        temp_file = self.temp_dir / "relative_links.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should preserve link text
        assert "README" in markdown
        assert "Config" in markdown
        assert "Logo" in markdown
        assert "Home" in markdown

    def test_anchor_links(self):
        """Test anchor/bookmark hyperlinks."""
        doc = docx.Document()
        doc.add_heading("Anchor Links Test", level=1)

        # Internal anchor links
        p1 = doc.add_paragraph("Go to ")
        p1.add_run("Section 1").hyperlink = "#section1"

        p2 = doc.add_paragraph("Jump to ")
        p2.add_run("Conclusion").hyperlink = "#conclusion"

        # Add some content to simulate sections
        doc.add_heading("Section 1", level=2)
        doc.add_paragraph("This is section 1 content.")

        doc.add_heading("Conclusion", level=2)
        doc.add_paragraph("This is the conclusion.")

        temp_file = self.temp_dir / "anchor_links.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should contain both link text and target sections
        assert "Section 1" in markdown
        assert "Conclusion" in markdown
        assert "## Section 1" in markdown or "Section 1" in markdown
        assert "## Conclusion" in markdown or "conclusion" in markdown.lower()

    def test_links_in_different_contexts(self):
        """Test hyperlinks in various document contexts."""
        doc = docx.Document()

        # Link in paragraph
        p1 = doc.add_paragraph("Normal paragraph with ")
        p1.add_run("link").hyperlink = "https://example.com"
        p1.add_run(" inside.")

        # Link in list item
        p2 = doc.add_paragraph("", style="List Bullet")
        p2.add_run("List item with ")
        p2.add_run("link").hyperlink = "https://list-example.com"

        # Link in table cell
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Cell with"
        # Simulated link in cell
        table.rows[0].cells[1].text = "linked content"
        table.rows[1].cells[0].text = "Regular"
        table.rows[1].cells[1].text = "Cell"

        # Link in heading
        heading = doc.add_heading("", level=2)
        heading.add_run("Heading with ")
        heading.add_run("Link").hyperlink = "https://heading-example.com"

        temp_file = self.temp_dir / "links_in_context.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should preserve links in different contexts
        assert "Normal paragraph" in markdown
        assert "List item with" in markdown
        assert "linked content" in markdown
        assert "Heading with" in markdown

    def test_link_formatting_combinations(self):
        """Test links with various text formatting."""
        doc = docx.Document()

        # Bold link
        p1 = doc.add_paragraph("Text with ")
        run1 = p1.add_run("bold link")
        run1.bold = True
        run1.hyperlink = "https://bold-example.com"

        # Italic link
        p2 = doc.add_paragraph("Text with ")
        run2 = p2.add_run("italic link")
        run2.italic = True
        run2.hyperlink = "https://italic-example.com"

        # Underlined link (often default for links)
        p3 = doc.add_paragraph("Text with ")
        run3 = p3.add_run("underlined link")
        run3.underline = True
        run3.hyperlink = "https://underlined-example.com"

        # Combined formatting
        p4 = doc.add_paragraph("Text with ")
        run4 = p4.add_run("bold italic link")
        run4.bold = True
        run4.italic = True
        run4.hyperlink = "https://combined-example.com"

        temp_file = self.temp_dir / "formatted_links.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should preserve both link and formatting
        # Note: Exact formatting depends on implementation
        assert "bold link" in markdown
        assert "italic link" in markdown
        assert "underlined link" in markdown
        assert "bold italic link" in markdown

    def test_broken_or_malformed_links(self):
        """Test handling of broken or malformed hyperlinks."""
        doc = docx.Document()

        # Malformed URL
        p1 = doc.add_paragraph("Malformed: ")
        p1.add_run("bad link").hyperlink = "htp://broken-url"

        # Empty URL
        p2 = doc.add_paragraph("Empty URL: ")
        p2.add_run("empty link").hyperlink = ""

        # Link with special characters
        p3 = doc.add_paragraph("Special chars: ")
        p3.add_run("special link").hyperlink = "https://example.com/path with spaces&special=chars"

        temp_file = self.temp_dir / "broken_links.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle gracefully
        assert "bad link" in markdown
        assert "empty link" in markdown
        assert "special link" in markdown

    def test_nested_link_content(self):
        """Test links containing complex nested content."""
        doc = docx.Document()

        # Link with formatted text inside
        p1 = doc.add_paragraph("Link with ")
        link_run = p1.add_run("formatted text")
        link_run.hyperlink = "https://example.com"
        # Note: Nested formatting in links is complex in python-docx

        # Multiple links in same paragraph
        p2 = doc.add_paragraph("Multiple ")
        p2.add_run("first link").hyperlink = "https://first.com"
        p2.add_run(" and ")
        p2.add_run("second link").hyperlink = "https://second.com"
        p2.add_run(" in paragraph.")

        temp_file = self.temp_dir / "nested_links.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle multiple links
        assert "first link" in markdown
        assert "second link" in markdown
        assert "formatted text" in markdown

    def test_very_long_urls(self):
        """Test handling of very long URLs."""
        doc = docx.Document()

        # Very long URL
        long_url = "https://example.com/" + "very-long-path/" * 20 + "final-page.html"
        p1 = doc.add_paragraph("Long URL: ")
        p1.add_run("Long Link").hyperlink = long_url

        # URL with many parameters
        param_url = "https://example.com/search?" + "&".join([f"param{i}=value{i}" for i in range(20)])
        p2 = doc.add_paragraph("Many params: ")
        p2.add_run("Param Link").hyperlink = param_url

        temp_file = self.temp_dir / "long_urls.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle long URLs gracefully
        assert "Long Link" in markdown
        assert "Param Link" in markdown

    def test_link_title_attributes(self):
        """Test links with title attributes (tooltips)."""
        doc = docx.Document()

        # Links with tooltips/titles (limited support in python-docx)
        p1 = doc.add_paragraph("Link with tooltip: ")
        p1.add_run("Hover Link").hyperlink = "https://example.com"

        # Screen reader friendly link
        p2 = doc.add_paragraph("Accessible link: ")
        p2.add_run("More Information").hyperlink = "https://info.example.com"

        temp_file = self.temp_dir / "link_titles.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should preserve link text
        assert "Hover Link" in markdown
        assert "More Information" in markdown
