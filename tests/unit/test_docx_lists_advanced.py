"""Advanced tests for DOCX list handling edge cases."""

import docx
from docx.shared import Inches

from all2md import to_markdown as docx_to_markdown
from all2md.options import DocxOptions, MarkdownOptions
from tests.utils import assert_markdown_valid, cleanup_test_dir, create_test_temp_dir


class TestDocxListsAdvanced:
    """Test complex list scenarios in DOCX documents."""

    def setup_method(self):
        """Set up test environment."""
        self.temp_dir = create_test_temp_dir()

    def teardown_method(self):
        """Clean up test environment."""
        cleanup_test_dir(self.temp_dir)

    def test_mixed_numbering_styles(self):
        """Test lists with mixed numbering styles (roman, alpha, numeric)."""
        doc = docx.Document()

        # Different numbering styles - python-docx has limitations but we can simulate
        doc.add_paragraph("First item", style="List Number")
        doc.add_paragraph("Second item", style="List Number")
        doc.add_paragraph("Third item", style="List Number")

        # Nested with different style (simulated)
        p4 = doc.add_paragraph("Nested alpha item", style="List Number")
        p4.paragraph_format.left_indent = Inches(0.5)

        temp_file = self.temp_dir / "mixed_numbering.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should contain numbered items
        assert "1. First item" in markdown
        assert "2. Second item" in markdown
        assert "3. Third item" in markdown
        # Nested item should be properly indented (correct behavior)
        assert "   1. Nested alpha item" in markdown

    def test_restart_numbering(self):
        """Test lists with restart numbering."""
        doc = docx.Document()

        doc.add_heading("First List", level=2)
        doc.add_paragraph("Item 1", style="List Number")
        doc.add_paragraph("Item 2", style="List Number")

        doc.add_paragraph("Interrupting paragraph")

        doc.add_heading("Restarted List", level=2)
        # Simulate restarted numbering
        doc.add_paragraph("New item 1", style="List Number")
        doc.add_paragraph("New item 2", style="List Number")

        temp_file = self.temp_dir / "restart_numbering.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should have restarted numbering
        lines = markdown.split('\n')
        first_list_items = [line for line in lines if "Item" in line and line.strip().startswith('1.')]
        new_list_items = [line for line in lines if "New item" in line and line.strip().startswith('1.')]

        assert len(first_list_items) >= 1
        assert len(new_list_items) >= 1

    def test_mixed_bullet_symbols(self):
        """Test lists with mixed bullet symbols."""
        doc = docx.Document()

        # Different bullet levels
        doc.add_paragraph("First level bullet", style="List Bullet")

        p2 = doc.add_paragraph("Second level bullet", style="List Bullet")
        p2.paragraph_format.left_indent = Inches(0.5)

        p3 = doc.add_paragraph("Third level bullet", style="List Bullet")
        p3.paragraph_format.left_indent = Inches(1.0)

        temp_file = self.temp_dir / "mixed_bullets.docx"
        doc.save(str(temp_file))

        options = DocxOptions(markdown_options=MarkdownOptions(bullet_symbols="*+-"))
        markdown = docx_to_markdown(str(temp_file), options=options)
        assert_markdown_valid(markdown)

        # Should use different symbols for different levels
        assert "* First level bullet" in markdown
        assert "  + Second level bullet" in markdown or "  - Second level bullet" in markdown
        assert "Third level bullet" in markdown

    def test_large_indent_levels(self):
        """Test lists with large indent levels (5+ levels deep)."""
        doc = docx.Document()

        levels = ["First", "Second", "Third", "Fourth", "Fifth", "Sixth"]

        for i, level in enumerate(levels):
            p = doc.add_paragraph(f"{level} level item", style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.5 * i)

        temp_file = self.temp_dir / "large_indent.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        lines = markdown.split('\n')
        list_lines = [line for line in lines if "level item" in line]

        # Should handle deep nesting appropriately
        assert len(list_lines) == 6

        # Check indentation increases
        for i, line in enumerate(list_lines[:4]):  # Check first 4 levels
            leading_spaces = len(line) - len(line.lstrip())
            # Should have increasing indentation
            if i > 0:
                assert leading_spaces > 0

    def test_paragraphs_with_mixed_content(self):
        """Test paragraphs with both text and inline images/hyperlinks."""
        doc = docx.Document()

        # Create paragraph with mixed content
        p1 = doc.add_paragraph("Text before ")

        # Add hyperlink
        p1.add_run("hyperlink")
        # Note: python-docx hyperlink support is limited, but we test the structure

        p1.add_run(" and text after with more content")

        # List item with mixed content
        p2 = doc.add_paragraph("", style="List Bullet")
        p2.add_run("List item with ")
        p2.add_run("embedded link")
        p2.add_run(" and more text")

        temp_file = self.temp_dir / "mixed_content.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should preserve text content
        assert "Text before" in markdown
        assert "hyperlink" in markdown
        assert "text after" in markdown
        assert "List item with" in markdown
        assert "embedded link" in markdown

    def test_list_formatting_preservation(self):
        """Test that formatting within list items is preserved."""
        doc = docx.Document()

        # List with bold text
        p1 = doc.add_paragraph("", style="List Bullet")
        bold_run = p1.add_run("Bold list item")
        bold_run.bold = True

        # List with italic text
        p2 = doc.add_paragraph("", style="List Bullet")
        italic_run = p2.add_run("Italic list item")
        italic_run.italic = True

        # List with mixed formatting
        p3 = doc.add_paragraph("", style="List Bullet")
        p3.add_run("Mixed ")
        bold_run = p3.add_run("bold")
        bold_run.bold = True
        p3.add_run(" and ")
        italic_run = p3.add_run("italic")
        italic_run.italic = True
        p3.add_run(" text")

        temp_file = self.temp_dir / "formatted_lists.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should preserve formatting
        assert "**Bold list item**" in markdown
        assert "*Italic list item*" in markdown
        assert "Mixed **bold** and *italic* text" in markdown

    def test_list_with_multiple_paragraphs(self):
        """Test list items that span multiple paragraphs."""
        doc = docx.Document()

        # First list item
        doc.add_paragraph("First paragraph of list item", style="List Bullet")

        # Continuation paragraph (no list style but indented)
        p2 = doc.add_paragraph("Second paragraph of same list item")
        p2.paragraph_format.left_indent = Inches(0.5)

        # Second list item
        doc.add_paragraph("New list item", style="List Bullet")

        temp_file = self.temp_dir / "multi_paragraph_list.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle multiple paragraphs appropriately
        assert "First paragraph of list item" in markdown
        assert "Second paragraph" in markdown
        assert "New list item" in markdown

    def test_empty_list_items(self):
        """Test handling of empty or whitespace-only list items."""
        doc = docx.Document()

        doc.add_paragraph("Valid list item", style="List Bullet")
        doc.add_paragraph("", style="List Bullet")  # Empty
        doc.add_paragraph("   ", style="List Bullet")  # Whitespace only
        doc.add_paragraph("Another valid item", style="List Bullet")

        temp_file = self.temp_dir / "empty_list_items.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle empty items gracefully
        assert "Valid list item" in markdown
        assert "Another valid item" in markdown

        # Empty items might be skipped or rendered as empty bullets
        lines = [line.strip() for line in markdown.split('\n') if line.strip()]
        bullet_lines = [line for line in lines if line.startswith('*')]
        assert len(bullet_lines) >= 2  # At least the non-empty items

    def test_list_style_detection_edge_cases(self):
        """Test edge cases in list style detection."""
        doc = docx.Document()

        # Paragraph that looks like a list but isn't styled as one
        doc.add_paragraph("1. This looks like a list but isn't styled")
        doc.add_paragraph("2. Another fake list item")

        # Actual styled list
        doc.add_paragraph("Real list item", style="List Number")
        doc.add_paragraph("Another real item", style="List Number")

        # Mixed case
        doc.add_paragraph("* Bullet-like text")
        doc.add_paragraph("Actual bullet", style="List Bullet")

        temp_file = self.temp_dir / "style_detection.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should distinguish between styled and unstyled content
        lines = markdown.split('\n')

        # Check that actual list items are formatted as lists
        list_lines = [line for line in lines if line.strip().startswith(('*', '1.', '2.'))]
        assert len(list_lines) >= 2  # Should have at least the styled items
