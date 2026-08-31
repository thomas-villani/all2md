"""OMML to LaTeX conversion, and the bare ``m:oMath`` shapes Word actually writes.

There were no tests here at all, which is how a radical handler that looked for a
child element the OMML schema does not define survived: every radical resolved to the
empty string and vanished, leaving the surrounding expression intact and plausible.
`\\frac{-b±}{2a}` is wrong in a way that still looks like an equation.
"""

from __future__ import annotations

import pytest
from lxml import etree

from all2md.parsers.docx import MATH_NS, _omml_to_latex

pytestmark = [pytest.mark.unit, pytest.mark.docx]

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSMAP = {"m": MATH_NS, "w": W_NS}


def omml(fragment: str) -> etree._Element:
    """Parse an OMML fragment, with the maths and wordprocessing namespaces bound."""
    xml = f'<m:oMath xmlns:m="{MATH_NS}" xmlns:w="{W_NS}">{fragment}</m:oMath>'
    return etree.fromstring(xml.encode("utf-8"))


def run(text: str) -> str:
    return f"<m:r><m:t>{text}</m:t></m:r>"


# ------------------------------------------------------------------- radicals (#481)
def test_a_square_root_keeps_its_radicand():
    """The regression: the radicand is `m:e`, not `m:base`."""
    latex = _omml_to_latex(
        omml(f"<m:rad><m:radPr><m:degHide m:val=\"1\"/></m:radPr><m:deg/><m:e>{run('x')}</m:e></m:rad>")
    )
    assert latex == "\\sqrt{x}"


def test_a_radical_with_a_degree_keeps_both_parts():
    latex = _omml_to_latex(omml(f"<m:rad><m:deg>{run('3')}</m:deg><m:e>{run('x')}</m:e></m:rad>"))
    assert latex == "\\sqrt[3]{x}"


def test_a_radical_inside_a_fraction_survives_whole():
    """The shape that was failing in the wild: `\\frac{-b±}{2a}` looked fine."""
    numerator = (
        f"{run('-b±')}<m:rad><m:radPr><m:degHide m:val=\"1\"/></m:radPr><m:deg/><m:e>{run('b^2-4ac')}</m:e></m:rad>"
    )
    latex = _omml_to_latex(omml(f"<m:f><m:num>{numerator}</m:num><m:den>{run('2a')}</m:den></m:f>"))
    assert "\\sqrt{b^2-4ac}" in latex
    assert latex == "\\frac{-b±\\sqrt{b^2-4ac}}{2a}"


# ---------------------------------------------------------- the constructs that worked
def test_a_fraction_converts():
    assert _omml_to_latex(omml(f"<m:f><m:num>{run('a')}</m:num><m:den>{run('b')}</m:den></m:f>")) == "\\frac{a}{b}"


def test_a_superscript_converts():
    latex = _omml_to_latex(omml(f"<m:sSup><m:e>{run('x')}</m:e><m:sup>{run('2')}</m:sup></m:sSup>"))
    assert latex == "x^{2}"


def test_a_subscript_converts():
    latex = _omml_to_latex(omml(f"<m:sSub><m:e>{run('x')}</m:e><m:sub>{run('i')}</m:sub></m:sSub>"))
    assert latex == "x_{i}"


# ------------------------------------------------- bare m:oMath placement (#481 proper)
def paragraph_xml(body: str) -> bytes:
    return (f'<w:p xmlns:w="{W_NS}" xmlns:m="{MATH_NS}">{body}</w:p>').encode("utf-8")


def w_run(text: str) -> str:
    return f'<w:r><w:t xml:space="preserve">{text}</w:t></w:r>'


def bare_math(inner: str) -> str:
    return f"<m:oMath>{inner}</m:oMath>"


@pytest.fixture()
def parser():
    from all2md.parsers.docx import DocxToAstConverter

    return DocxToAstConverter.__new__(DocxToAstConverter)


class FakeParagraph:
    """Just enough of `python-docx`'s Paragraph for the two helpers under test."""

    def __init__(self, xml: bytes) -> None:
        self._element = etree.fromstring(xml)

    @property
    def text(self) -> str:
        return "".join(node.text or "" for node in self._element.iter(f"{{{W_NS}}}t"))


def test_an_equation_alone_in_its_paragraph_is_display_math(parser):
    """Word writes a bare m:oMath for a standalone display equation."""
    para = FakeParagraph(
        paragraph_xml(bare_math(f"<m:f><m:num>{run('a')}</m:num><m:den>{run('b')}</m:den></m:f>") + w_run(" "))
    )
    blocks = parser._extract_math_blocks_from_paragraph(para)
    assert len(blocks) == 1
    assert blocks[0].content == "\\frac{a}{b}"


def test_an_equation_sharing_a_paragraph_with_text_is_not_a_block(parser):
    """Otherwise a formula mid-sentence gets torn out and re-emitted after it."""
    para = FakeParagraph(paragraph_xml(w_run("see ") + bare_math(run("x")) + w_run(" here")))
    assert parser._extract_math_blocks_from_paragraph(para) == []


def test_a_bare_omath_is_found_among_the_paragraph_children(parser):
    para = FakeParagraph(paragraph_xml(w_run("a") + bare_math(run("x")) + w_run("b")))
    assert len(parser._bare_omath_children(para)) == 1


def test_a_paragraph_with_no_math_reports_none(parser):
    para = FakeParagraph(paragraph_xml(w_run("just text")))
    assert parser._bare_omath_children(para) == []
    assert parser._extract_math_blocks_from_paragraph(para) == []


# ------------------------------------------------------------------ end-to-end, corpus
def test_the_corpus_display_equation_converts_completely():
    """The whole point: the committed document, through the real parser."""
    from pathlib import Path

    import all2md

    document = (
        Path(__file__).resolve().parents[2] / "benchmarks" / "docx" / "corpus" / "baseline" / "display-equation.docx"
    )
    if not document.exists():  # pragma: no cover - the corpus is committed
        pytest.skip("corpus not present")
    out = all2md.to_markdown(str(document))
    assert "$$" in out, "a standalone equation is displayed maths"
    assert "\\frac{-b±\\sqrt{b^{2}-4ac}}{2a}" in out


def test_inline_math_lands_between_the_words_it_sits_between():
    """A bare m:oMath among runs is inline maths, and position matters.

    Appending it to the end of the paragraph would be a different sentence.
    """
    import io

    import docx

    import all2md

    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("see ")
    paragraph._p.append(
        etree.fromstring(
            f'<m:oMath xmlns:m="{MATH_NS}"><m:f><m:num>{run("a")}</m:num>'
            f'<m:den>{run("b")}</m:den></m:f></m:oMath>'.encode("utf-8")
        )
    )
    paragraph.add_run(" here")

    buffer = io.BytesIO()
    document.save(buffer)
    assert all2md.to_markdown(buffer.getvalue()) == r"see $\frac{a}{b}$ here"
