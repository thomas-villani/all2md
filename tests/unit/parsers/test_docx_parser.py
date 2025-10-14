#  Copyright (c) 2025 Tom Villani, Ph.D.
#
# tests/unit/test_docx_ast.py
"""Unit tests for DOCX to AST converter.

Tests cover:
- DOCX paragraph to AST node conversion
- Heading detection and conversion
- Text formatting (bold, italic, underline, etc.)
- Hyperlink processing
- List detection and conversion
- Table structure conversion
- Image markdown parsing
- Run grouping and formatting preservation

"""


import docx
import pytest
from docx.oxml import parse_xml
from fixtures import FIXTURES_PATH

from all2md.ast import (
    BlockQuote,
    Document,
    Emphasis,
    FootnoteDefinition,
    FootnoteReference,
    Heading,
    List,
    MathBlock,
    MathInline,
    Paragraph,
    Strikethrough,
    Strong,
    Subscript,
    Superscript,
    Table,
    Text,
    Underline,
)
from all2md.ast.transforms import extract_nodes
from all2md.options import DocxOptions
from all2md.parsers.docx import DocxToAstConverter

FIXTURE_FOOTNOTES_DOC = (
    FIXTURES_PATH / "documents" / "footnotes-endnotes-comments.docx"
)
FIXTURE_MATH_DOC = (
    FIXTURES_PATH / "documents" / "math-basic.docx"
)


def _inline_text(nodes: list) -> str:
    parts: list[str] = []
    for node in nodes:
        if isinstance(node, Text):
            parts.append(node.content)
        elif hasattr(node, "content"):
            child = node.content
            if isinstance(child, list):
                parts.append(_inline_text(child))
    return "".join(parts)


@pytest.mark.unit
class TestBasicElements:
    """Tests for basic DOCX element conversion."""

    def test_simple_paragraph(self) -> None:
        """Test converting a simple paragraph."""
        doc = docx.Document()
        doc.add_paragraph("Hello world")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        assert isinstance(ast_doc, Document)
        assert len(ast_doc.children) == 1
        assert isinstance(ast_doc.children[0], Paragraph)
        para = ast_doc.children[0]
        assert len(para.content) == 1
        assert isinstance(para.content[0], Text)
        assert para.content[0].content == "Hello world"

    def test_multiple_paragraphs(self) -> None:
        """Test converting multiple paragraphs."""
        doc = docx.Document()
        doc.add_paragraph("First")
        doc.add_paragraph("Second")
        doc.add_paragraph("Third")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        assert len(ast_doc.children) == 3
        assert all(isinstance(child, Paragraph) for child in ast_doc.children)

    def test_headings_1_to_6(self) -> None:
        """Test converting all heading levels."""
        doc = docx.Document()
        doc.add_heading("Heading 1", level=1)
        doc.add_heading("Heading 2", level=2)
        doc.add_heading("Heading 3", level=3)
        doc.add_heading("Heading 4", level=4)
        doc.add_heading("Heading 5", level=5)
        doc.add_heading("Heading 6", level=6)

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        assert len(ast_doc.children) == 6
        for i, child in enumerate(ast_doc.children):
            assert isinstance(child, Heading)
            assert child.level == i + 1
            assert isinstance(child.content[0], Text)
            assert child.content[0].content == f"Heading {i + 1}"

    def test_empty_paragraphs_skipped(self) -> None:
        """Test that empty paragraphs are skipped."""
        doc = docx.Document()
        doc.add_paragraph("First")
        doc.add_paragraph("")  # Empty
        doc.add_paragraph("   ")  # Whitespace only
        doc.add_paragraph("Second")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        # Only 2 paragraphs should be in AST (empty ones skipped)
        assert len(ast_doc.children) == 2
        assert ast_doc.children[0].content[0].content == "First"
        assert ast_doc.children[1].content[0].content == "Second"

    def test_inline_math_extraction(self) -> None:
        """Inline OMML equations should become MathInline nodes."""
        doc = docx.Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Inline ")
        math_run = paragraph.add_run()
        math_element = parse_xml(
            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
            "<m:r><m:t>x</m:t></m:r>"
            "</m:oMath>"
        )
        math_run._element.append(math_element)
        paragraph.add_run(" equals three")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        assert isinstance(ast_doc.children[0], Paragraph)
        para = ast_doc.children[0]
        assert any(isinstance(node, MathInline) for node in para.content)

        texts = [node for node in para.content if isinstance(node, Text)]
        assert texts[0].content == "Inline "
        math_nodes = [node for node in para.content if isinstance(node, MathInline)]
        assert math_nodes[0].content == "x"
        assert math_nodes[0].notation == "latex"
        assert math_nodes[0].representations["latex"] == "x"


@pytest.mark.unit
class TestTextFormatting:
    """Tests for text formatting conversion."""

    def test_bold_text(self) -> None:
        """Test bold text conversion."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold text")
        run.bold = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        assert isinstance(para_node.content[0], Strong)
        assert isinstance(para_node.content[0].content[0], Text)
        assert para_node.content[0].content[0].content == "Bold text"

    def test_italic_text(self) -> None:
        """Test italic text conversion."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Italic text")
        run.italic = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        assert isinstance(para_node.content[0], Emphasis)
        assert isinstance(para_node.content[0].content[0], Text)
        assert para_node.content[0].content[0].content == "Italic text"

    def test_underline_text(self) -> None:
        """Test underline text conversion."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Underlined text")
        run.underline = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        assert isinstance(para_node.content[0], Underline)
        assert isinstance(para_node.content[0].content[0], Text)
        assert para_node.content[0].content[0].content == "Underlined text"

    def test_strikethrough_text(self) -> None:
        """Test strikethrough text conversion."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Strikethrough text")
        run.font.strike = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        assert isinstance(para_node.content[0], Strikethrough)
        assert isinstance(para_node.content[0].content[0], Text)
        assert para_node.content[0].content[0].content == "Strikethrough text"

    def test_subscript_text(self) -> None:
        """Test subscript text conversion."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("H2O")
        run.font.subscript = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        assert isinstance(para_node.content[0], Subscript)

    def test_superscript_text(self) -> None:
        """Test superscript text conversion."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("E=mc^2")
        run.font.superscript = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        assert isinstance(para_node.content[0], Superscript)

    def test_multiple_formatting(self) -> None:
        """Test text with multiple formatting applied."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold and italic")
        run.bold = True
        run.italic = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        # Should be nested: Emphasis -> Strong -> Text (italic applied after bold)
        assert isinstance(para_node.content[0], Emphasis)
        assert isinstance(para_node.content[0].content[0], Strong)
        assert isinstance(para_node.content[0].content[0].content[0], Text)
        assert para_node.content[0].content[0].content[0].content == "Bold and italic"

    def test_mixed_formatting_runs(self) -> None:
        """Test paragraph with multiple runs having different formatting."""
        doc = docx.Document()
        para = doc.add_paragraph()
        para.add_run("Normal ")
        bold_run = para.add_run("bold ")
        bold_run.bold = True
        para.add_run("normal ")
        italic_run = para.add_run("italic")
        italic_run.italic = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        # Should have 4 inline nodes
        assert len(para_node.content) == 4
        assert isinstance(para_node.content[0], Text)
        assert isinstance(para_node.content[1], Strong)
        assert isinstance(para_node.content[2], Text)
        assert isinstance(para_node.content[3], Emphasis)


@pytest.mark.unit
class TestLists:
    """Tests for list conversion."""

    def test_bullet_list(self) -> None:
        """Test bullet list conversion."""
        doc = docx.Document()
        doc.add_paragraph("Item 1", style="List Bullet")
        doc.add_paragraph("Item 2", style="List Bullet")
        doc.add_paragraph("Item 3", style="List Bullet")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        # Should have one List node
        assert len(ast_doc.children) == 1
        assert isinstance(ast_doc.children[0], List)
        list_node = ast_doc.children[0]
        assert not list_node.ordered
        assert len(list_node.items) == 3

    def test_numbered_list(self) -> None:
        """Test numbered list conversion."""
        doc = docx.Document()
        doc.add_paragraph("First", style="List Number")
        doc.add_paragraph("Second", style="List Number")
        doc.add_paragraph("Third", style="List Number")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        # Should have one List node
        assert len(ast_doc.children) == 1
        assert isinstance(ast_doc.children[0], List)
        list_node = ast_doc.children[0]
        assert list_node.ordered
        assert len(list_node.items) == 3

    def test_list_followed_by_paragraph(self) -> None:
        """Test list properly closed when followed by regular paragraph."""
        doc = docx.Document()
        doc.add_paragraph("Item 1", style="List Bullet")
        doc.add_paragraph("Item 2", style="List Bullet")
        doc.add_paragraph("Regular paragraph")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        # Should have 2 nodes: List and Paragraph
        assert len(ast_doc.children) == 2
        assert isinstance(ast_doc.children[0], List)
        assert isinstance(ast_doc.children[1], Paragraph)

    def test_multiple_separate_lists(self) -> None:
        """Test multiple lists separated by paragraphs."""
        doc = docx.Document()
        doc.add_paragraph("List 1 Item 1", style="List Bullet")
        doc.add_paragraph("List 1 Item 2", style="List Bullet")
        doc.add_paragraph("Separator")
        doc.add_paragraph("List 2 Item 1", style="List Number")
        doc.add_paragraph("List 2 Item 2", style="List Number")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        # Should have 3 nodes: List, Paragraph, List
        assert len(ast_doc.children) == 3
        assert isinstance(ast_doc.children[0], List)
        assert not ast_doc.children[0].ordered  # Bullet
        assert isinstance(ast_doc.children[1], Paragraph)
        assert isinstance(ast_doc.children[2], List)
        assert ast_doc.children[2].ordered  # Numbered


@pytest.mark.unit
class TestTables:
    """Tests for table conversion."""

    def test_simple_table(self) -> None:
        """Test simple table conversion."""
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "H1"
        table.rows[0].cells[1].text = "H2"
        table.rows[1].cells[0].text = "R1C1"
        table.rows[1].cells[1].text = "R1C2"

        converter = DocxToAstConverter(options=DocxOptions(preserve_tables=True))
        ast_doc = converter.convert_to_ast(doc)

        # Should have one Table node
        assert len(ast_doc.children) == 1
        assert isinstance(ast_doc.children[0], Table)
        table_node = ast_doc.children[0]

        # Check header
        assert table_node.header is not None
        assert len(table_node.header.cells) == 2

        # Check data rows
        assert len(table_node.rows) == 1
        assert len(table_node.rows[0].cells) == 2

    def test_table_with_formatted_cells(self) -> None:
        """Test table with formatted cell content."""
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)

        # Add bold text to header
        para = table.rows[0].cells[0].paragraphs[0]
        run = para.add_run("Bold Header")
        run.bold = True

        # Add normal text to data cell
        table.rows[1].cells[0].text = "Normal text"

        converter = DocxToAstConverter(options=DocxOptions(preserve_tables=True))
        ast_doc = converter.convert_to_ast(doc)

        table_node = ast_doc.children[0]
        # Check that formatting is preserved
        header_cell = table_node.header.cells[0]
        assert isinstance(header_cell.content[0], Strong)

    def test_table_flattening(self) -> None:
        """Test table flattening when preserve_tables=False."""
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "H1"
        table.rows[0].cells[1].text = "H2"
        table.rows[1].cells[0].text = "R1C1"
        table.rows[1].cells[1].text = "R1C2"

        converter = DocxToAstConverter(options=DocxOptions(preserve_tables=False))
        ast_doc = converter.convert_to_ast(doc)

        # Should have multiple Paragraph nodes (flattened)
        assert all(isinstance(child, Paragraph) for child in ast_doc.children)
        assert len(ast_doc.children) >= 4  # At least 4 cells


@pytest.mark.unit
class TestHyperlinks:
    """Tests for hyperlink conversion."""

    def test_hyperlink_parsing(self) -> None:
        """Test hyperlink detection and URL extraction."""
        # Note: Creating actual hyperlinks in python-docx is complex,
        # so we test the conversion logic with mocked data
        converter = DocxToAstConverter()

        # Test URL extraction
        result = converter._process_hyperlink(None)
        assert result == (None, None)


@pytest.mark.unit
class TestFormattingKey:
    """Tests for run formatting key generation."""

    def test_formatting_key_plain_text(self) -> None:
        """Test formatting key for plain text."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Plain text")

        converter = DocxToAstConverter()
        key = converter._get_run_formatting_key(run, False)

        # All False except is_hyperlink
        assert key == (False, False, False, False, False, False, False)

    def test_formatting_key_bold(self) -> None:
        """Test formatting key for bold text."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold")
        run.bold = True

        converter = DocxToAstConverter()
        key = converter._get_run_formatting_key(run, False)

        assert key[0] is True  # bold
        assert key[1] is False  # italic

    def test_formatting_key_all_formatting(self) -> None:
        """Test formatting key with all formatting applied."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Formatted")
        run.bold = True
        run.italic = True
        run.underline = True
        run.font.strike = True
        run.font.subscript = True

        converter = DocxToAstConverter()
        key = converter._get_run_formatting_key(run, True)

        assert key == (True, True, True, True, True, False, True)


@pytest.mark.unit
class TestListFinalization:
    """Tests for list finalization logic."""

    def test_list_at_end_of_document(self) -> None:
        """Test that list at end of document is properly finalized."""
        doc = docx.Document()
        doc.add_paragraph("Item 1", style="List Bullet")
        doc.add_paragraph("Item 2", style="List Bullet")
        # No paragraph after list

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        # List should still be in the document
        assert len(ast_doc.children) == 1
        assert isinstance(ast_doc.children[0], List)
        assert len(ast_doc.children[0].items) == 2

    def test_empty_list_stack(self) -> None:
        """Test finalization with empty list stack."""
        converter = DocxToAstConverter()
        result = converter._finalize_current_list()
        assert result is None


@pytest.mark.unit
class TestRunGrouping:
    """Tests for run grouping optimization."""

    def test_runs_with_same_formatting_grouped(self) -> None:
        """Test that consecutive runs with same formatting are grouped."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run1 = para.add_run("Bold ")
        run1.bold = True
        run2 = para.add_run("text")
        run2.bold = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        # Should be grouped into single Strong node
        assert len(para_node.content) == 1
        assert isinstance(para_node.content[0], Strong)
        # Check text is combined
        text_content = para_node.content[0].content[0].content
        assert text_content == "Bold text"

    def test_runs_with_different_formatting_separate(self) -> None:
        """Test that runs with different formatting are kept separate."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run1 = para.add_run("Bold")
        run1.bold = True
        run2 = para.add_run("Italic")
        run2.italic = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        # Should have 2 separate inline nodes
        assert len(para_node.content) == 2
        assert isinstance(para_node.content[0], Strong)
        assert isinstance(para_node.content[1], Emphasis)


@pytest.mark.unit
class TestComplexStructures:
    """Tests for complex document structures."""

    def test_mixed_content_document(self) -> None:
        """Test document with mixed content types."""
        doc = docx.Document()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Introduction paragraph")
        doc.add_paragraph("Item 1", style="List Bullet")
        doc.add_paragraph("Item 2", style="List Bullet")
        doc.add_paragraph("Conclusion paragraph")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        # Should have: Heading, Paragraph, List, Paragraph
        assert len(ast_doc.children) == 4
        assert isinstance(ast_doc.children[0], Heading)
        assert isinstance(ast_doc.children[1], Paragraph)
        assert isinstance(ast_doc.children[2], List)
        assert isinstance(ast_doc.children[3], Paragraph)

    def test_heading_with_formatting(self) -> None:
        """Test heading with formatted text."""
        doc = docx.Document()
        heading = doc.add_heading(level=1)
        run = heading.add_run("Bold Title")
        run.bold = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        heading_node = ast_doc.children[0]
        assert isinstance(heading_node, Heading)
        assert heading_node.level == 1
        # Check formatting is preserved
        assert isinstance(heading_node.content[0], Strong)


@pytest.mark.unit
class TestFootnotes:
    """Tests for DOCX footnote and endnote extraction."""

    def test_footnote_reference_and_definition(self) -> None:
        """Parser should emit references and collect matching definitions."""
        ast_doc = self._convert_with_fixture()
        references, definitions = self._extract_notes(ast_doc)

        footnote_refs = [ref for ref in references if not ref.identifier.startswith("end")]
        footnote_defs = [definition for definition in definitions if definition.metadata.get("note_type") == "footnote"]

        assert footnote_refs, "Expected at least one footnote reference"
        assert footnote_defs, "Expected at least one captured footnote definition"

        footnote_ids = {definition.identifier for definition in footnote_defs}
        assert {ref.identifier for ref in footnote_refs} <= footnote_ids

        for definition in footnote_defs:
            assert definition.content, "Footnote definition should contain block content"
            assert all(isinstance(block, Paragraph) for block in definition.content)
            assert all(block.content for block in definition.content)

    def test_endnote_reference_and_definition(self) -> None:
        """Endnotes should use prefixed identifiers and be collected."""
        ast_doc = self._convert_with_fixture()
        references, definitions = self._extract_notes(ast_doc)

        endnote_refs = [ref for ref in references if ref.identifier.startswith("end")]
        endnote_defs = [definition for definition in definitions if definition.metadata.get("note_type") == "endnote"]

        assert endnote_refs, "Expected at least one endnote reference"
        assert endnote_defs, "Expected at least one captured endnote definition"

        endnote_ids = {definition.identifier for definition in endnote_defs}
        assert {ref.identifier for ref in endnote_refs} <= endnote_ids

        for definition in endnote_defs:
            assert definition.content, "Endnote definition should contain block content"
            assert all(isinstance(block, Paragraph) for block in definition.content)
            assert all(block.content for block in definition.content)

    def test_footnotes_can_be_disabled(self) -> None:
        """Footnotes should be omitted entirely when option disabled."""
        ast_doc = self._convert_with_fixture(include_footnotes=False)
        references, definitions = self._extract_notes(ast_doc)

        footnote_refs = [ref for ref in references if not ref.identifier.startswith("end")]
        assert not footnote_refs
        assert all(definition.metadata.get("note_type") != "footnote" for definition in definitions)
        assert any(ref.identifier.startswith("end") for ref in references), "Endnote references should remain"

    def test_endnotes_can_be_disabled(self) -> None:
        """Endnotes should be omitted when include_endnotes is False."""
        ast_doc = self._convert_with_fixture(include_endnotes=False)
        references, definitions = self._extract_notes(ast_doc)

        endnote_refs = [ref for ref in references if ref.identifier.startswith("end")]
        assert not endnote_refs
        assert all(definition.metadata.get("note_type") != "endnote" for definition in definitions)
        assert any(not ref.identifier.startswith("end") for ref in references), "Footnote references should remain"

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _convert_with_fixture(
        include_footnotes: bool = True,
        include_endnotes: bool = True,
    ) -> Document:
        options = DocxOptions(
            include_footnotes=include_footnotes,
            include_endnotes=include_endnotes,
            include_comments=False,
        )
        converter = DocxToAstConverter(options=options)
        fixture_doc = docx.Document(FIXTURE_FOOTNOTES_DOC)
        return converter.convert_to_ast(fixture_doc)

    @staticmethod
    def _extract_notes(document: Document) -> tuple[list[FootnoteReference], list[FootnoteDefinition]]:
        references = extract_nodes(document, FootnoteReference)
        definitions = extract_nodes(document, FootnoteDefinition)
        return references, definitions


@pytest.mark.unit
class TestComments:
    """Tests for DOCX comment rendering options."""

    def test_comments_append_when_position_is_footnotes(self) -> None:
        doc = docx.Document(str(FIXTURE_FOOTNOTES_DOC))
        options = DocxOptions(
            include_comments=True,
            comments_position="footnotes",
            comment_mode="blockquote",
        )
        converter = DocxToAstConverter(options=options)
        ast_doc = converter.convert_to_ast(doc)

        blockquotes = [node for node in ast_doc.children if isinstance(node, BlockQuote)]
        assert blockquotes, "Expected comments to be appended as blockquotes"

        quoted_text = _inline_text(blockquotes[0].children[0].content)
        assert "I decided not to think of something funny." in quoted_text

        paragraph = next(
            child
            for child in ast_doc.children
            if isinstance(child, Paragraph)
            and "However, it will have a comment" in _inline_text(child.content)
        )
        inline_text = _inline_text(paragraph.content)
        assert "comment1" not in inline_text

    def test_comments_inline_when_requested(self) -> None:
        doc = docx.Document(str(FIXTURE_FOOTNOTES_DOC))
        options = DocxOptions(
            include_comments=True,
            comments_position="inline",
            comment_mode="blockquote",
        )
        converter = DocxToAstConverter(options=options)
        ast_doc = converter.convert_to_ast(doc)

        paragraph = next(
            child
            for child in ast_doc.children
            if isinstance(child, Paragraph)
            and "However, it will have a comment" in _inline_text(child.content)
        )
        inline_text = _inline_text(paragraph.content)
        assert "I decided not to think of something funny." in inline_text
        assert "comment1" in inline_text

        blockquotes = [node for node in ast_doc.children if isinstance(node, BlockQuote)]
        assert not blockquotes, "Inline comments should not append trailing blockquotes"

@pytest.mark.unit
class TestMathExtraction:
    """Tests for DOCX math conversion."""

    def test_math_block_from_fixture(self) -> None:
        """Math blocks should convert to MathBlock nodes with LaTeX content."""
        converter = DocxToAstConverter()
        ast_doc = converter.parse(FIXTURE_MATH_DOC)

        math_nodes = extract_nodes(ast_doc, MathBlock)
        assert math_nodes, "Expected at least one MathBlock from math fixture"

        content = math_nodes[0].content.replace(" ", "")
        assert content.startswith("e^{πi}")
        assert "=-1" in content
        assert math_nodes[0].notation == "latex"
        assert math_nodes[0].representations["latex"].replace(" ", "") == content
