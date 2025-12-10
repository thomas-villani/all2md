#  Copyright (c) 2025 Tom Villani, Ph.D.
#
# tests/unit/test_docx_renderer.py
"""Unit tests for DocxRenderer.

Tests cover:
- Rendering all node types to DOCX
- Document structure and styles
- Table formatting
- Image handling
- Edge cases and nested structures

Note: These tests require python-docx to be installed.

"""

from io import BytesIO

import pytest

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from all2md.ast import (
    BlockQuote,
    Code,
    CodeBlock,
    Comment,
    CommentInline,
    DefinitionDescription,
    DefinitionList,
    DefinitionTerm,
    Document,
    Emphasis,
    FootnoteDefinition,
    FootnoteReference,
    Heading,
    LineBreak,
    Link,
    List,
    ListItem,
    MathBlock,
    MathInline,
    Paragraph,
    Strikethrough,
    Strong,
    Table,
    TableCell,
    TableRow,
    Text,
    ThematicBreak,
    Underline,
)
from all2md.options import DocxRendererOptions

if DOCX_AVAILABLE:
    from all2md.renderers.docx import DocxRenderer

pytestmark = pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")


@pytest.mark.unit
@pytest.mark.docx
class TestBasicRendering:
    """Tests for basic DOCX rendering."""

    def test_render_empty_document(self, tmp_path):
        """Test rendering an empty document."""
        doc = Document()
        renderer = DocxRenderer()
        output_file = tmp_path / "empty.docx"
        renderer.render(doc, output_file)

        assert output_file.exists()
        docx_doc = DocxDocument(str(output_file))
        # Empty document still has default styles
        assert docx_doc is not None

    def test_render_text_only(self, tmp_path):
        """Test rendering plain text."""
        doc = Document(children=[Paragraph(content=[Text(content="Hello world")])])
        renderer = DocxRenderer()
        output_file = tmp_path / "text.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert len(docx_doc.paragraphs) >= 1
        assert "Hello world" in docx_doc.paragraphs[0].text

    def test_render_multiple_paragraphs(self, tmp_path):
        """Test rendering multiple paragraphs."""
        doc = Document(
            children=[
                Paragraph(content=[Text(content="First paragraph")]),
                Paragraph(content=[Text(content="Second paragraph")]),
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "paras.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        paragraphs = [p for p in docx_doc.paragraphs if p.text.strip()]
        assert len(paragraphs) >= 2
        assert "First paragraph" in paragraphs[0].text
        assert "Second paragraph" in paragraphs[1].text

    def test_render_to_bytes_io(self):
        """Test rendering to BytesIO."""
        doc = Document(children=[Paragraph(content=[Text(content="Test")])])
        renderer = DocxRenderer()
        output = BytesIO()
        renderer.render(doc, output)

        output.seek(0)
        docx_doc = DocxDocument(output)
        assert len(docx_doc.paragraphs) >= 1


@pytest.mark.unit
@pytest.mark.docx
class TestHeadingRendering:
    """Tests for heading rendering."""

    def test_heading_level_1(self, tmp_path):
        """Test rendering h1."""
        doc = Document(children=[Heading(level=1, content=[Text(content="Title")])])
        renderer = DocxRenderer()
        output_file = tmp_path / "h1.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        # Find heading paragraph
        heading_para = None
        for para in docx_doc.paragraphs:
            if para.text == "Title":
                heading_para = para
                break
        assert heading_para is not None

    def test_heading_level_2(self, tmp_path):
        """Test rendering h2."""
        doc = Document(children=[Heading(level=2, content=[Text(content="Subtitle")])])
        renderer = DocxRenderer()
        output_file = tmp_path / "h2.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        found = False
        for para in docx_doc.paragraphs:
            if "Subtitle" in para.text:
                found = True
                break
        assert found

    def test_heading_with_formatting(self, tmp_path):
        """Test heading with inline formatting."""
        doc = Document(
            children=[Heading(level=1, content=[Text(content="Bold "), Strong(content=[Text(content="title")])])]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "h_format.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        found = False
        for para in docx_doc.paragraphs:
            if "Bold" in para.text and "title" in para.text:
                found = True
                break
        assert found


@pytest.mark.unit
@pytest.mark.docx
class TestInlineFormatting:
    """Tests for inline formatting."""

    def test_strong(self, tmp_path):
        """Test bold text rendering."""
        doc = Document(children=[Paragraph(content=[Strong(content=[Text(content="bold")])])])
        renderer = DocxRenderer()
        output_file = tmp_path / "bold.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "bold" in para.text
        # Check if any run is bold
        has_bold = any(run.bold for run in para.runs)
        assert has_bold

    def test_emphasis(self, tmp_path):
        """Test italic text rendering."""
        doc = Document(children=[Paragraph(content=[Emphasis(content=[Text(content="italic")])])])
        renderer = DocxRenderer()
        output_file = tmp_path / "italic.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "italic" in para.text
        # Check if any run is italic
        has_italic = any(run.italic for run in para.runs)
        assert has_italic

    def test_code(self, tmp_path):
        """Test inline code rendering."""
        doc = Document(children=[Paragraph(content=[Code(content="code")])])
        renderer = DocxRenderer()
        output_file = tmp_path / "code.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "code" in para.text

    def test_strikethrough(self, tmp_path):
        """Test strikethrough rendering."""
        doc = Document(children=[Paragraph(content=[Strikethrough(content=[Text(content="deleted")])])])
        renderer = DocxRenderer()
        output_file = tmp_path / "strike.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "deleted" in para.text

    def test_underline(self, tmp_path):
        """Test underline rendering."""
        doc = Document(children=[Paragraph(content=[Underline(content=[Text(content="underlined")])])])
        renderer = DocxRenderer()
        output_file = tmp_path / "underline.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "underlined" in para.text

    def test_nested_formatting(self, tmp_path):
        """Test nested inline formatting."""
        doc = Document(
            children=[Paragraph(content=[Strong(content=[Emphasis(content=[Text(content="bold italic")])])])]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "nested.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "bold italic" in para.text


@pytest.mark.unit
@pytest.mark.docx
class TestLinkRendering:
    """Tests for hyperlink rendering."""

    def test_simple_link(self, tmp_path):
        """Test basic hyperlink rendering."""
        doc = Document(
            children=[
                Paragraph(
                    content=[Text(content="Visit "), Link(url="https://example.com", content=[Text(content="Example")])]
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "link.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "Visit" in para.text
        assert "Example" in para.text

        # Check that hyperlink exists in XML structure
        from docx.oxml.ns import qn

        hyperlinks = para._element.findall(qn("w:hyperlink"))
        assert len(hyperlinks) > 0, "No hyperlinks found in paragraph XML"

        # Verify the hyperlink has proper w:r and w:t structure
        hyperlink = hyperlinks[0]
        runs = hyperlink.findall(qn("w:r"))
        assert len(runs) > 0, "Hyperlink has no w:r elements"

        # Check that text is in w:t element (not directly in w:r)
        text_elements = runs[0].findall(qn("w:t"))
        assert len(text_elements) > 0, "Run has no w:t element - text not properly structured"
        assert text_elements[0].text == "Example", f"Expected 'Example' but got '{text_elements[0].text}'"

    def test_link_with_formatting(self, tmp_path):
        """Test hyperlink with nested formatting."""
        doc = Document(
            children=[
                Paragraph(
                    content=[Link(url="https://example.com", content=[Strong(content=[Text(content="Bold Link")])])]
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "link_formatted.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "Bold Link" in para.text


@pytest.mark.unit
@pytest.mark.docx
class TestListRendering:
    """Tests for list rendering."""

    def test_unordered_list(self, tmp_path):
        """Test unordered list rendering."""
        doc = Document(
            children=[
                List(
                    ordered=False,
                    items=[
                        ListItem(children=[Paragraph(content=[Text(content="Item 1")])]),
                        ListItem(children=[Paragraph(content=[Text(content="Item 2")])]),
                    ],
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "ul.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        paragraphs = docx_doc.paragraphs
        text_content = [p.text for p in paragraphs]
        assert "Item 1" in text_content
        assert "Item 2" in text_content

    def test_ordered_list(self, tmp_path):
        """Test ordered list rendering."""
        doc = Document(
            children=[
                List(
                    ordered=True,
                    items=[
                        ListItem(children=[Paragraph(content=[Text(content="First")])]),
                        ListItem(children=[Paragraph(content=[Text(content="Second")])]),
                    ],
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "ol.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        paragraphs = docx_doc.paragraphs
        text_content = [p.text for p in paragraphs]
        assert "First" in text_content
        assert "Second" in text_content


@pytest.mark.unit
@pytest.mark.docx
class TestTableRendering:
    """Tests for table rendering."""

    def test_simple_table(self, tmp_path):
        """Test basic table rendering."""
        doc = Document(
            children=[
                Table(
                    header=TableRow(
                        cells=[TableCell(content=[Text(content="Name")]), TableCell(content=[Text(content="Age")])]
                    ),
                    rows=[
                        TableRow(
                            cells=[TableCell(content=[Text(content="Alice")]), TableCell(content=[Text(content="30")])]
                        )
                    ],
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "table.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert len(docx_doc.tables) >= 1
        table = docx_doc.tables[0]

        # Check table has correct dimensions
        assert len(table.rows) == 2
        assert len(table.rows[0].cells) == 2

        # Check header content
        assert "Name" in table.rows[0].cells[0].text
        assert "Age" in table.rows[0].cells[1].text

        # Check body content
        assert "Alice" in table.rows[1].cells[0].text
        assert "30" in table.rows[1].cells[1].text

    def test_table_without_header(self, tmp_path):
        """Test table without header row."""
        doc = Document(
            children=[
                Table(
                    rows=[
                        TableRow(cells=[TableCell(content=[Text(content="A")]), TableCell(content=[Text(content="B")])])
                    ]
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "table_no_header.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert len(docx_doc.tables) >= 1


@pytest.mark.unit
@pytest.mark.docx
class TestBlockElements:
    """Tests for block-level elements."""

    def test_code_block(self, tmp_path):
        """Test code block rendering."""
        doc = Document(children=[CodeBlock(content="def hello():\n    print('world')", language="python")])
        renderer = DocxRenderer()
        output_file = tmp_path / "codeblock.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        found = False
        for para in docx_doc.paragraphs:
            if "def hello()" in para.text:
                found = True
                break
        assert found

    def test_blockquote(self, tmp_path):
        """Test blockquote rendering with indentation."""
        doc = Document(
            children=[
                Paragraph(content=[Text(content="Normal text")]),
                BlockQuote(children=[Paragraph(content=[Text(content="Quoted text")])]),
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "blockquote.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))

        # Find the quoted paragraph
        quoted_para = None
        normal_para = None
        for para in docx_doc.paragraphs:
            if "Quoted text" in para.text:
                quoted_para = para
            elif "Normal text" in para.text:
                normal_para = para

        assert quoted_para is not None, "Quoted text not found"
        assert normal_para is not None, "Normal text not found"

        # Verify blockquote has indentation
        quoted_indent = quoted_para.paragraph_format.left_indent
        normal_indent = normal_para.paragraph_format.left_indent or 0

        assert quoted_indent is not None, "Blockquote paragraph has no indentation"
        assert (
            quoted_indent > normal_indent
        ), f"Blockquote indent ({quoted_indent}) should be greater than normal ({normal_indent})"

    def test_nested_blockquote(self, tmp_path):
        """Test nested blockquote rendering with increased indentation."""
        doc = Document(
            children=[
                BlockQuote(
                    children=[
                        Paragraph(content=[Text(content="Level 1")]),
                        BlockQuote(children=[Paragraph(content=[Text(content="Level 2")])]),
                    ]
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "nested_blockquote.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))

        # Find both paragraphs
        level1_para = None
        level2_para = None
        for para in docx_doc.paragraphs:
            if "Level 1" in para.text:
                level1_para = para
            elif "Level 2" in para.text:
                level2_para = para

        assert level1_para is not None, "Level 1 text not found"
        assert level2_para is not None, "Level 2 text not found"

        # Verify nested blockquote has more indentation
        level1_indent = level1_para.paragraph_format.left_indent
        level2_indent = level2_para.paragraph_format.left_indent

        assert level1_indent is not None, "Level 1 has no indentation"
        assert level2_indent is not None, "Level 2 has no indentation"
        assert (
            level2_indent > level1_indent
        ), f"Level 2 indent ({level2_indent}) should be greater than Level 1 ({level1_indent})"

    def test_thematic_break(self, tmp_path):
        """Test horizontal rule rendering."""
        doc = Document(children=[ThematicBreak()])
        renderer = DocxRenderer()
        output_file = tmp_path / "hr.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        # Thematic break creates a paragraph with border
        assert len(docx_doc.paragraphs) >= 1


@pytest.mark.unit
@pytest.mark.docx
class TestDocumentMetadata:
    """Tests for document metadata."""

    def test_metadata_properties(self, tmp_path):
        """Test document properties from metadata."""
        doc = Document(
            metadata={"title": "Test Document", "author": "Test Author", "subject": "Testing"},
            children=[Paragraph(content=[Text(content="Content")])],
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "metadata.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert docx_doc.core_properties.title == "Test Document"
        assert docx_doc.core_properties.author == "Test Author"
        assert docx_doc.core_properties.subject == "Testing"


@pytest.mark.unit
@pytest.mark.docx
class TestDefinitionLists:
    """Tests for definition list rendering."""

    def test_definition_list(self, tmp_path):
        """Test definition list rendering."""
        doc = Document(
            children=[
                DefinitionList(
                    items=[
                        (
                            DefinitionTerm(content=[Text(content="Term")]),
                            [DefinitionDescription(content=[Text(content="Description")])],
                        )
                    ]
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "deflist.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        text_content = " ".join(p.text for p in docx_doc.paragraphs)
        assert "Term" in text_content
        assert "Description" in text_content


@pytest.mark.unit
@pytest.mark.docx
class TestRendererOptions:
    """Tests for renderer options."""

    def test_custom_fonts(self, tmp_path):
        """Test custom font settings."""
        doc = Document(children=[Paragraph(content=[Text(content="Custom font")])])
        options = DocxRendererOptions(default_font="Arial", default_font_size=14)
        renderer = DocxRenderer(options)
        output_file = tmp_path / "custom_font.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert docx_doc is not None

    def test_code_font_option(self, tmp_path):
        """Test code font option."""
        doc = Document(children=[CodeBlock(content="code here")])
        options = DocxRendererOptions(code_font="Monaco")
        renderer = DocxRenderer(options)
        output_file = tmp_path / "code_font.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert docx_doc is not None


@pytest.mark.unit
@pytest.mark.docx
class TestMathRendering:
    """Tests for math rendering."""

    def test_inline_math(self, tmp_path):
        """Test inline math rendering."""
        doc = Document(children=[Paragraph(content=[MathInline(content="x^2", notation="latex")])])
        renderer = DocxRenderer()
        output_file = tmp_path / "math_inline.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        # Math is rendered as plain text in simple implementation
        found = False
        for para in docx_doc.paragraphs:
            if "x^2" in para.text:
                found = True
                break
        assert found

    def test_block_math(self, tmp_path):
        """Test block math rendering."""
        doc = Document(children=[MathBlock(content="E = mc^2", notation="latex")])
        renderer = DocxRenderer()
        output_file = tmp_path / "math_block.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        # Math is rendered as code block
        found = False
        for para in docx_doc.paragraphs:
            if "E = mc^2" in para.text:
                found = True
                break
        assert found


@pytest.mark.unit
@pytest.mark.docx
class TestFootnotes:
    """Tests for footnote rendering."""

    def test_footnote_reference(self, tmp_path):
        """Test footnote reference rendering."""
        doc = Document(children=[Paragraph(content=[Text(content="Text"), FootnoteReference(identifier="1")])])
        renderer = DocxRenderer()
        output_file = tmp_path / "footnote_ref.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        # Footnote reference is rendered as superscript
        text_content = "".join(p.text for p in docx_doc.paragraphs)
        assert "Text" in text_content

    def test_footnote_definition(self, tmp_path):
        """Test footnote definition rendering."""
        doc = Document(children=[FootnoteDefinition(identifier="1", content=[Text(content="Footnote text")])])
        renderer = DocxRenderer()
        output_file = tmp_path / "footnote_def.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        text_content = "".join(p.text for p in docx_doc.paragraphs)
        assert "Footnote text" in text_content


@pytest.mark.unit
@pytest.mark.docx
class TestLineBreaks:
    """Tests for line break rendering."""

    def test_hard_line_break(self, tmp_path):
        """Test hard line break rendering."""
        doc = Document(
            children=[Paragraph(content=[Text(content="Line 1"), LineBreak(soft=False), Text(content="Line 2")])]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "linebreak.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "Line 1" in para.text
        assert "Line 2" in para.text


@pytest.mark.unit
@pytest.mark.docx
class TestHeadingFormattingWithoutStyles:
    """Tests for heading rendering with use_styles=False."""

    def test_heading_bold_without_styles(self, tmp_path):
        """Test that headings are bold when use_styles=False."""
        doc = Document(children=[Heading(level=1, content=[Text(content="Test Heading")])])
        options = DocxRendererOptions(use_styles=False)
        renderer = DocxRenderer(options)
        output_file = tmp_path / "heading_no_styles.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        # Find the heading paragraph
        heading_para = None
        for para in docx_doc.paragraphs:
            if "Test Heading" in para.text:
                heading_para = para
                break

        assert heading_para is not None, "Heading not found"
        # At least one run should be bold
        has_bold = any(run.bold for run in heading_para.runs)
        assert has_bold, "Heading should have bold formatting when use_styles=False"

    def test_heading_font_size_without_styles(self, tmp_path):
        """Test that custom heading font sizes are applied when use_styles=False."""
        doc = Document(
            children=[
                Heading(level=1, content=[Text(content="Large Heading")]),
                Heading(level=2, content=[Text(content="Medium Heading")]),
            ]
        )

        # Custom font sizes for headings
        options = DocxRendererOptions(use_styles=False, heading_font_sizes={1: 24, 2: 18})
        renderer = DocxRenderer(options)
        output_file = tmp_path / "heading_custom_sizes.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))

        # Find h1
        h1_para = None
        h2_para = None
        for para in docx_doc.paragraphs:
            if "Large Heading" in para.text:
                h1_para = para
            elif "Medium Heading" in para.text:
                h2_para = para

        assert h1_para is not None, "H1 not found"
        assert h2_para is not None, "H2 not found"

        # Check font sizes are applied
        h1_sizes = [run.font.size for run in h1_para.runs if run.font.size]
        h2_sizes = [run.font.size for run in h2_para.runs if run.font.size]

        # Import Pt for comparison
        from docx.shared import Pt

        assert any(size == Pt(24) for size in h1_sizes), "H1 should have 24pt font size"
        assert any(size == Pt(18) for size in h2_sizes), "H2 should have 18pt font size"

    def test_heading_with_formatting_without_styles(self, tmp_path):
        """Test heading with inline formatting when use_styles=False."""
        doc = Document(
            children=[
                Heading(level=1, content=[Text(content="Normal "), Emphasis(content=[Text(content="emphasized")])])
            ]
        )
        options = DocxRendererOptions(use_styles=False)
        renderer = DocxRenderer(options)
        output_file = tmp_path / "heading_formatted_no_styles.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        heading_para = None
        for para in docx_doc.paragraphs:
            if "Normal" in para.text and "emphasized" in para.text:
                heading_para = para
                break

        assert heading_para is not None, "Heading not found"
        # Should have both regular and italic runs, all bold
        has_italic = any(run.italic for run in heading_para.runs)
        all_bold = all(run.bold for run in heading_para.runs if run.text.strip())
        assert has_italic, "Should have italic formatting"
        assert all_bold, "All runs should be bold (heading-level formatting)"


@pytest.mark.unit
@pytest.mark.docx
class TestHyperlinkWhitespace:
    """Tests for hyperlink whitespace preservation."""

    def test_hyperlink_with_leading_space(self, tmp_path):
        """Test hyperlink text with leading space is preserved."""
        doc = Document(
            children=[Paragraph(content=[Link(url="https://example.com", content=[Text(content=" Link Text")])])]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "link_leading_space.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]

        # Check hyperlink element in XML
        from docx.oxml.ns import qn

        hyperlinks = para._element.findall(qn("w:hyperlink"))
        assert len(hyperlinks) > 0, "No hyperlinks found"

        hyperlink = hyperlinks[0]
        runs = hyperlink.findall(qn("w:r"))
        text_elements = runs[0].findall(qn("w:t"))

        # Verify xml:space="preserve" is present
        space_attr = text_elements[0].get(qn("xml:space"))
        assert space_attr == "preserve", "xml:space='preserve' attribute should be present"

        # Verify text has leading space
        assert text_elements[0].text == " Link Text", "Leading space should be preserved"

    def test_hyperlink_with_trailing_space(self, tmp_path):
        """Test hyperlink text with trailing space is preserved."""
        doc = Document(
            children=[Paragraph(content=[Link(url="https://example.com", content=[Text(content="Link Text ")])])]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "link_trailing_space.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]

        # Check hyperlink element in XML
        from docx.oxml.ns import qn

        hyperlinks = para._element.findall(qn("w:hyperlink"))
        text_elements = hyperlinks[0].findall(qn("w:r"))[0].findall(qn("w:t"))

        # Verify xml:space="preserve" is present
        space_attr = text_elements[0].get(qn("xml:space"))
        assert space_attr == "preserve", "xml:space='preserve' attribute should be present"

        # Verify text has trailing space
        assert text_elements[0].text == "Link Text ", "Trailing space should be preserved"

    def test_hyperlink_with_multiple_spaces(self, tmp_path):
        """Test hyperlink text with multiple consecutive spaces is preserved."""
        doc = Document(
            children=[Paragraph(content=[Link(url="https://example.com", content=[Text(content="Link  Text")])])]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "link_multiple_spaces.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]

        # Check hyperlink element in XML
        from docx.oxml.ns import qn

        hyperlinks = para._element.findall(qn("w:hyperlink"))
        text_elements = hyperlinks[0].findall(qn("w:r"))[0].findall(qn("w:t"))

        # Verify xml:space="preserve" is present
        space_attr = text_elements[0].get(qn("xml:space"))
        assert space_attr == "preserve", "xml:space='preserve' attribute should be present"

        # Verify double space is preserved
        assert text_elements[0].text == "Link  Text", "Multiple spaces should be preserved"


@pytest.mark.unit
@pytest.mark.docx
class TestComments:
    """Tests for comment rendering."""

    def test_block_comment_with_author(self, tmp_path):
        """Test block comment with author metadata."""
        doc = Document(
            children=[
                Comment(
                    content="This is a review comment",
                    metadata={"author": "John Doe", "date": "2025-01-20", "comment_type": "docx_review"},
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "block_comment.docx"
        renderer.render(doc, output_file)

        assert output_file.exists()
        docx_doc = DocxDocument(str(output_file))
        # Comment should be rendered (either as native comment or fallback)
        assert docx_doc is not None

    def test_block_comment_without_author(self, tmp_path):
        """Test block comment without author metadata."""
        doc = Document(children=[Comment(content="Anonymous comment", metadata={"comment_type": "generic"})])
        renderer = DocxRenderer()
        output_file = tmp_path / "block_comment_no_author.docx"
        renderer.render(doc, output_file)

        assert output_file.exists()
        docx_doc = DocxDocument(str(output_file))
        assert docx_doc is not None

    def test_inline_comment_with_author(self, tmp_path):
        """Test inline comment with author metadata."""
        doc = Document(
            children=[
                Paragraph(
                    content=[
                        Text(content="This is text "),
                        CommentInline(
                            content="Needs clarification",
                            metadata={"author": "Jane Smith", "comment_type": "docx_review"},
                        ),
                        Text(content=" more text."),
                    ]
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "inline_comment.docx"
        renderer.render(doc, output_file)

        assert output_file.exists()
        docx_doc = DocxDocument(str(output_file))
        # Comment should be rendered
        para = docx_doc.paragraphs[0]
        assert "This is text" in para.text
        assert "more text" in para.text

    def test_inline_comment_without_author(self, tmp_path):
        """Test inline comment without author metadata."""
        doc = Document(
            children=[
                Paragraph(
                    content=[
                        Text(content="Text with "),
                        CommentInline(content="anonymous comment", metadata={}),
                        Text(content=" here."),
                    ]
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "inline_comment_no_author.docx"
        renderer.render(doc, output_file)

        assert output_file.exists()
        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "Text with" in para.text

    def test_comment_with_label(self, tmp_path):
        """Test comment with label metadata."""
        doc = Document(
            children=[
                Comment(
                    content="Important review note",
                    metadata={"author": "Reviewer", "label": "1", "date": "2025-01-20"},
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "comment_with_label.docx"
        renderer.render(doc, output_file)

        assert output_file.exists()
        docx_doc = DocxDocument(str(output_file))
        assert docx_doc is not None

    def test_multiple_comments(self, tmp_path):
        """Test multiple comments in document."""
        doc = Document(
            children=[
                Comment(content="First comment", metadata={"author": "User 1"}),
                Paragraph(content=[Text(content="Text here")]),
                Comment(content="Second comment", metadata={"author": "User 2"}),
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "multiple_comments.docx"
        renderer.render(doc, output_file)

        assert output_file.exists()
        docx_doc = DocxDocument(str(output_file))
        assert len(docx_doc.paragraphs) >= 1

    def test_comment_initials_generation(self, tmp_path):
        """Test automatic generation of initials from author name."""
        doc = Document(
            children=[
                Comment(content="Test comment", metadata={"author": "John David Smith", "comment_type": "docx_review"})
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "comment_initials.docx"
        renderer.render(doc, output_file)

        assert output_file.exists()
        # Initials should be generated as "JDS" from "John David Smith"
        docx_doc = DocxDocument(str(output_file))
        assert docx_doc is not None

    def test_comment_with_explicit_initials(self, tmp_path):
        """Test comment with explicitly provided initials."""
        doc = Document(
            children=[
                Comment(
                    content="Test comment",
                    metadata={"author": "John Smith", "initials": "JS", "comment_type": "docx_review"},
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "comment_explicit_initials.docx"
        renderer.render(doc, output_file)

        assert output_file.exists()
        docx_doc = DocxDocument(str(output_file))
        assert docx_doc is not None


@pytest.mark.unit
@pytest.mark.docx
class TestCommentModes:
    """Tests for comment rendering modes."""

    def test_comment_mode_ignore(self, tmp_path):
        """Test that comments are ignored when comment_mode is 'ignore'."""
        doc = Document(
            children=[
                Paragraph(content=[Text(content="Visible text")]),
                Comment(content="This should be ignored", metadata={"author": "Test"}),
            ]
        )
        options = DocxRendererOptions(comment_mode="ignore")
        renderer = DocxRenderer(options)
        output_file = tmp_path / "comment_ignore.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        text_content = " ".join(p.text for p in docx_doc.paragraphs)
        assert "Visible text" in text_content
        # Comment should not appear in text
        assert "This should be ignored" not in text_content

    def test_comment_mode_visible(self, tmp_path):
        """Test that comments are visible when comment_mode is 'visible'."""
        doc = Document(
            children=[
                Comment(
                    content="Visible comment",
                    metadata={"author": "Author", "date": "2025-01-20", "comment_type": "review"},
                )
            ]
        )
        options = DocxRendererOptions(comment_mode="visible")
        renderer = DocxRenderer(options)
        output_file = tmp_path / "comment_visible.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        text_content = " ".join(p.text for p in docx_doc.paragraphs)
        assert "Visible comment" in text_content
        assert "Author" in text_content

    def test_inline_comment_mode_ignore(self, tmp_path):
        """Test inline comment is ignored when comment_mode is 'ignore'."""
        doc = Document(
            children=[
                Paragraph(
                    content=[
                        Text(content="Before "),
                        CommentInline(content="ignored inline", metadata={}),
                        Text(content=" after"),
                    ]
                )
            ]
        )
        options = DocxRendererOptions(comment_mode="ignore")
        renderer = DocxRenderer(options)
        output_file = tmp_path / "inline_ignore.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "Before" in para.text
        assert "after" in para.text
        assert "ignored inline" not in para.text

    def test_inline_comment_mode_visible(self, tmp_path):
        """Test inline comment is visible when comment_mode is 'visible'."""
        doc = Document(
            children=[
                Paragraph(
                    content=[
                        Text(content="Text "),
                        CommentInline(
                            content="visible inline",
                            metadata={"author": "Test Author", "date": "2025-01-20", "label": "1"},
                        ),
                    ]
                )
            ]
        )
        options = DocxRendererOptions(comment_mode="visible")
        renderer = DocxRenderer(options)
        output_file = tmp_path / "inline_visible.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "visible inline" in para.text
        assert "Test Author" in para.text


@pytest.mark.unit
@pytest.mark.docx
class TestRenderToBytes:
    """Tests for render_to_bytes method."""

    def test_render_to_bytes_simple(self):
        """Test render_to_bytes returns valid DOCX bytes."""
        doc = Document(children=[Paragraph(content=[Text(content="Byte content")])])
        renderer = DocxRenderer()
        result = renderer.render_to_bytes(doc)

        assert isinstance(result, bytes)
        assert len(result) > 0

        # Verify it's a valid DOCX by loading it
        from io import BytesIO

        docx_doc = DocxDocument(BytesIO(result))
        assert "Byte content" in docx_doc.paragraphs[0].text


@pytest.mark.unit
@pytest.mark.docx
class TestMetadataKeywords:
    """Tests for document metadata with keywords."""

    def test_keywords_as_list(self, tmp_path):
        """Test keywords metadata when provided as a list."""
        doc = Document(
            metadata={"title": "Test", "keywords": ["keyword1", "keyword2", "keyword3"]},
            children=[Paragraph(content=[Text(content="Content")])],
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "keywords_list.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert "keyword1" in docx_doc.core_properties.keywords
        assert "keyword2" in docx_doc.core_properties.keywords
        assert "keyword3" in docx_doc.core_properties.keywords

    def test_keywords_as_string(self, tmp_path):
        """Test keywords metadata when provided as a string."""
        doc = Document(
            metadata={"keywords": "single keyword string"},
            children=[Paragraph(content=[Text(content="Content")])],
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "keywords_string.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert docx_doc.core_properties.keywords == "single keyword string"


@pytest.mark.unit
@pytest.mark.docx
class TestTableSpanning:
    """Tests for table colspan and rowspan."""

    def test_table_with_colspan(self, tmp_path):
        """Test table with column spanning."""
        doc = Document(
            children=[
                Table(
                    header=TableRow(
                        cells=[
                            TableCell(content=[Text(content="Merged Header")], colspan=2),
                            TableCell(content=[Text(content="Single")]),
                        ]
                    ),
                    rows=[
                        TableRow(
                            cells=[
                                TableCell(content=[Text(content="A")]),
                                TableCell(content=[Text(content="B")]),
                                TableCell(content=[Text(content="C")]),
                            ]
                        )
                    ],
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "table_colspan.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert len(docx_doc.tables) == 1
        table = docx_doc.tables[0]
        assert len(table.rows) == 2

    def test_table_with_rowspan(self, tmp_path):
        """Test table with row spanning."""
        doc = Document(
            children=[
                Table(
                    rows=[
                        TableRow(
                            cells=[
                                TableCell(content=[Text(content="Merged")], rowspan=2),
                                TableCell(content=[Text(content="Row 1")]),
                            ]
                        ),
                        TableRow(cells=[TableCell(content=[Text(content="Row 2")])]),
                    ]
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "table_rowspan.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert len(docx_doc.tables) == 1


@pytest.mark.unit
@pytest.mark.docx
class TestSubscriptSuperscript:
    """Tests for subscript and superscript rendering."""

    def test_superscript(self, tmp_path):
        """Test superscript rendering."""
        from all2md.ast import Superscript

        doc = Document(
            children=[
                Paragraph(
                    content=[
                        Text(content="E = mc"),
                        Superscript(content=[Text(content="2")]),
                    ]
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "superscript.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "E = mc" in para.text
        assert "2" in para.text
        # Check that some run has superscript
        has_superscript = any(run.font.superscript for run in para.runs if run.text == "2")
        assert has_superscript

    def test_subscript(self, tmp_path):
        """Test subscript rendering."""
        from all2md.ast import Subscript

        doc = Document(
            children=[
                Paragraph(
                    content=[
                        Text(content="H"),
                        Subscript(content=[Text(content="2")]),
                        Text(content="O"),
                    ]
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "subscript.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        text = para.text
        assert "H" in text
        assert "2" in text
        assert "O" in text


@pytest.mark.unit
@pytest.mark.docx
class TestNestedLists:
    """Tests for nested list rendering."""

    def test_nested_unordered_list(self, tmp_path):
        """Test nested unordered list."""
        doc = Document(
            children=[
                List(
                    ordered=False,
                    items=[
                        ListItem(
                            children=[
                                Paragraph(content=[Text(content="Item 1")]),
                                List(
                                    ordered=False,
                                    items=[
                                        ListItem(children=[Paragraph(content=[Text(content="Nested 1.1")])]),
                                        ListItem(children=[Paragraph(content=[Text(content="Nested 1.2")])]),
                                    ],
                                ),
                            ]
                        ),
                        ListItem(children=[Paragraph(content=[Text(content="Item 2")])]),
                    ],
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "nested_list.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        paragraphs = docx_doc.paragraphs
        text_content = [p.text for p in paragraphs]
        assert "Item 1" in text_content
        assert "Nested 1.1" in text_content
        assert "Item 2" in text_content

    def test_multi_paragraph_list_item(self, tmp_path):
        """Test list item with multiple paragraphs."""
        doc = Document(
            children=[
                List(
                    ordered=False,
                    items=[
                        ListItem(
                            children=[
                                Paragraph(content=[Text(content="First paragraph")]),
                                Paragraph(content=[Text(content="Second paragraph")]),
                            ]
                        )
                    ],
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "multi_para_list.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        text_content = " ".join(p.text for p in docx_doc.paragraphs)
        assert "First paragraph" in text_content
        assert "Second paragraph" in text_content


@pytest.mark.unit
@pytest.mark.docx
class TestBlockquoteWithElements:
    """Tests for blockquote with different element types."""

    def test_blockquote_with_heading(self, tmp_path):
        """Test heading inside blockquote has indentation."""
        doc = Document(children=[BlockQuote(children=[Heading(level=2, content=[Text(content="Quoted Heading")])])])
        renderer = DocxRenderer()
        output_file = tmp_path / "bq_heading.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        heading_para = None
        for para in docx_doc.paragraphs:
            if "Quoted Heading" in para.text:
                heading_para = para
                break

        assert heading_para is not None
        assert heading_para.paragraph_format.left_indent is not None

    def test_blockquote_with_code_block(self, tmp_path):
        """Test code block inside blockquote has indentation."""
        doc = Document(children=[BlockQuote(children=[CodeBlock(content="code inside quote", language="python")])])
        renderer = DocxRenderer()
        output_file = tmp_path / "bq_code.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        code_para = None
        for para in docx_doc.paragraphs:
            if "code inside quote" in para.text:
                code_para = para
                break

        assert code_para is not None
        assert code_para.paragraph_format.left_indent is not None


@pytest.mark.unit
@pytest.mark.docx
class TestTableStyle:
    """Tests for table style option."""

    def test_table_with_custom_style(self, tmp_path):
        """Test table with custom table style."""
        doc = Document(
            children=[
                Table(
                    header=TableRow(cells=[TableCell(content=[Text(content="Header")])]),
                    rows=[TableRow(cells=[TableCell(content=[Text(content="Data")])])],
                )
            ]
        )
        options = DocxRendererOptions(table_style="Table Grid")
        renderer = DocxRenderer(options)
        output_file = tmp_path / "table_style.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert len(docx_doc.tables) == 1
        # Table should exist and have content
        assert "Header" in docx_doc.tables[0].rows[0].cells[0].text


@pytest.mark.unit
@pytest.mark.docx
class TestEmptyTable:
    """Tests for empty table handling."""

    def test_empty_table_no_rows(self, tmp_path):
        """Test table with no rows is handled gracefully."""
        doc = Document(children=[Table(rows=[])])
        renderer = DocxRenderer()
        output_file = tmp_path / "empty_table.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        # Empty table should not be created
        assert len(docx_doc.tables) == 0


@pytest.mark.unit
@pytest.mark.docx
class TestDefinitionListAdvanced:
    """Advanced tests for definition list rendering."""

    def test_definition_list_with_paragraph_content(self, tmp_path):
        """Test definition list where content is wrapped in Paragraph."""
        doc = Document(
            children=[
                DefinitionList(
                    items=[
                        (
                            DefinitionTerm(content=[Text(content="API")]),
                            [
                                DefinitionDescription(
                                    content=[Paragraph(content=[Text(content="Application Programming Interface")])]
                                )
                            ],
                        )
                    ]
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "deflist_para.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        text_content = " ".join(p.text for p in docx_doc.paragraphs)
        assert "API" in text_content
        assert "Application Programming Interface" in text_content

    def test_definition_list_multiple_descriptions(self, tmp_path):
        """Test term with multiple descriptions."""
        doc = Document(
            children=[
                DefinitionList(
                    items=[
                        (
                            DefinitionTerm(content=[Text(content="Term")]),
                            [
                                DefinitionDescription(content=[Text(content="First description")]),
                                DefinitionDescription(content=[Text(content="Second description")]),
                            ],
                        )
                    ]
                )
            ]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "deflist_multi_desc.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        text_content = " ".join(p.text for p in docx_doc.paragraphs)
        assert "Term" in text_content
        assert "First description" in text_content
        assert "Second description" in text_content
