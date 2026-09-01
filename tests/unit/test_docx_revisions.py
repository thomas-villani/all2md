"""Tracked changes: the three resolutions, and the markup shapes Word actually writes.

A document left with Track Changes on used to lose every insertion silently (#480),
because ``python-docx`` cannot see a run nested inside ``w:ins``. The output stayed
plausible -- "The quick  fox jumps." -- which is why these tests assert the *text*
each policy should produce rather than that parsing merely succeeded.

The XML fragments here are copied from the shapes Word wrote into the
``benchmarks/docx`` corpus, not invented: the ``w:delText`` rename, the paragraph
mark carrying its own ``w:ins``, and a whole paragraph wrapped in one ``w:ins``.
"""

from __future__ import annotations

import io

import docx
import pytest
from lxml import etree

from all2md.ast import Strikethrough
from all2md.ast.transforms import extract_nodes
from all2md.options.docx import DocxOptions
from all2md.parsers.docx import DocxToAstConverter
from all2md.parsers.docx_revisions import (
    W,
    document_has_revisions,
    resolve_revisions,
    run_revision,
)

pytestmark = [pytest.mark.unit, pytest.mark.docx]

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REV = 'w:id="1" w:author="A Reviewer" w:date="2026-08-31T18:31:00Z"'


def body(fragment: str) -> etree._Element:
    """Parse a ``w:body`` fragment with the wordprocessing namespace bound."""
    return etree.fromstring(f'<w:body xmlns:w="{W_NS}">{fragment}</w:body>'.encode())


def text_of(root: etree._Element) -> list[str]:
    """The visible text of each paragraph, as any reader of ``w:t`` would see it."""
    return ["".join(t.text or "" for t in p.iter(f"{W}t")) for p in root.iter(f"{W}p")]


def run_xml(text: str) -> str:
    return f'<w:r><w:t xml:space="preserve">{text}</w:t></w:r>'


def deleted_run(text: str) -> str:
    return f'<w:r><w:delText xml:space="preserve">{text}</w:delText></w:r>'


# --------------------------------------------------------------- the cheap fast path
def test_a_document_with_no_revisions_is_recognised_as_such():
    """Nothing is copied or rewritten for the documents that are the vast majority."""
    assert not document_has_revisions(body(f"<w:p>{run_xml('Plain prose.')}</w:p>"))


def test_an_insertion_is_recognised():
    assert document_has_revisions(body(f"<w:p><w:ins {REV}>{run_xml('new')}</w:ins></w:p>"))


# ------------------------------------------------------------------ inline revisions
SUBSTITUTION = (
    "<w:p>"
    + run_xml("The quick ")
    + f"<w:ins {REV}>{run_xml('crimson')}</w:ins>"
    + f"<w:del {REV}>{deleted_run('brown')}</w:del>"
    + run_xml(" fox jumps.")
    + "</w:p>"
)


def test_accept_takes_the_insertion_and_drops_the_deletion():
    root = body(SUBSTITUTION)
    resolve_revisions(root, "accept")
    assert text_of(root) == ["The quick crimson fox jumps."]


def test_reject_takes_the_deletion_and_drops_the_insertion():
    """The deleted text lives in ``w:delText``, so keeping the run is not enough."""
    root = body(SUBSTITUTION)
    resolve_revisions(root, "reject")
    assert text_of(root) == ["The quick brown fox jumps."]


def test_mark_keeps_both_halves_in_document_order():
    root = body(SUBSTITUTION)
    resolve_revisions(root, "mark")
    assert text_of(root) == ["The quick crimsonbrown fox jumps."]


def test_mark_stamps_each_run_with_its_revision():
    root = body(SUBSTITUTION)
    resolve_revisions(root, "mark")
    stamped = [run_revision(r) for r in root.iter(f"{W}r")]
    assert stamped[0] is None
    assert stamped[1] == {"type": "insert", "author": "A Reviewer", "date": "2026-08-31T18:31:00Z", "id": "1"}
    assert stamped[2] == {"type": "delete", "author": "A Reviewer", "date": "2026-08-31T18:31:00Z", "id": "1"}


def test_an_unknown_policy_leaves_the_tree_alone():
    root = body(SUBSTITUTION)
    resolve_revisions(root, "nonsense")
    assert document_has_revisions(root)


# ------------------------------------------------------- a paragraph that is all one
ALL_INSERTED = f"<w:p>{run_xml('Kept.')}</w:p><w:p><w:ins {REV}>{run_xml('Every word inserted.')}</w:ins></w:p>"


def test_a_wholly_inserted_paragraph_survives_accept():
    """The shape that vanished entirely: zero direct ``w:r`` children, empty text."""
    root = body(ALL_INSERTED)
    resolve_revisions(root, "accept")
    assert text_of(root) == ["Kept.", "Every word inserted."]


def test_a_wholly_inserted_paragraph_is_empty_under_reject():
    root = body(ALL_INSERTED)
    resolve_revisions(root, "reject")
    assert text_of(root) == ["Kept.", ""]


# ----------------------------------------------------------------- paragraph marks
def mark_revision(tag: str) -> str:
    return f"<w:pPr><w:rPr><w:{tag} {REV}/></w:rPr></w:pPr>"


def test_accepting_a_deleted_paragraph_mark_joins_the_paragraphs():
    """Deleting a pilcrow joins two paragraphs; resolving must actually join them."""
    root = body(
        f"<w:p>{mark_revision('del')}{run_xml('First half ')}</w:p><w:p>{run_xml('second half.')}</w:p>",
    )
    resolve_revisions(root, "accept")
    assert text_of(root) == ["First half second half."]


def test_rejecting_an_inserted_paragraph_mark_joins_the_paragraphs():
    """Splitting a paragraph inserts a mark; rejecting the split un-splits it."""
    root = body(
        f"<w:p>{mark_revision('ins')}{run_xml('First half ')}</w:p><w:p>{run_xml('second half.')}</w:p>",
    )
    resolve_revisions(root, "reject")
    assert text_of(root) == ["First half second half."]


def test_a_paragraph_mark_the_policy_keeps_leaves_the_split_alone():
    root = body(
        f"<w:p>{mark_revision('ins')}{run_xml('First half ')}</w:p><w:p>{run_xml('second half.')}</w:p>",
    )
    resolve_revisions(root, "accept")
    assert text_of(root) == ["First half ", "second half."]


def test_a_vanishing_mark_on_the_last_paragraph_keeps_its_text():
    """There is nothing to merge into, and dropping the paragraph would lose text."""
    root = body(f"<w:p>{run_xml('Only. ')}</w:p><w:p>{mark_revision('del')}{run_xml('Last.')}</w:p>")
    resolve_revisions(root, "accept")
    assert text_of(root) == ["Only. ", "Last."]


def test_the_merged_paragraph_keeps_the_surviving_marks_properties():
    """Word's rule: the mark that survives owns the merged paragraph's properties."""
    root = body(
        f"<w:p>{mark_revision('del')}{run_xml('Heading text ')}</w:p>"
        f'<w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr>{run_xml("and more.")}</w:p>'
    )
    resolve_revisions(root, "accept")
    paragraphs = list(root.iter(f"{W}p"))
    assert len(paragraphs) == 1
    assert paragraphs[0].find(f"{W}pPr/{W}pStyle").get(f"{W}val") == "Heading1"


# ------------------------------------------------------------------------ table rows
def test_accepting_a_row_deletion_removes_the_row():
    """The cells are marked individually, so an unremoved row is a phantom empty one."""
    root = body(
        "<w:tbl>"
        f"<w:tr><w:tc><w:p>{run_xml('kept')}</w:p></w:tc></w:tr>"
        f"<w:tr><w:trPr><w:del {REV}/></w:trPr><w:tc><w:p>"
        f"<w:del {REV}>{deleted_run('gone')}</w:del></w:p></w:tc></w:tr>"
        "</w:tbl>"
    )
    resolve_revisions(root, "accept")
    assert len(list(root.iter(f"{W}tr"))) == 1
    assert text_of(root) == ["kept"]


def test_rejecting_a_row_insertion_removes_the_row():
    root = body(
        "<w:tbl>"
        f"<w:tr><w:tc><w:p>{run_xml('kept')}</w:p></w:tc></w:tr>"
        f"<w:tr><w:trPr><w:ins {REV}/></w:trPr><w:tc><w:p>"
        f"<w:ins {REV}>{run_xml('added')}</w:ins></w:p></w:tc></w:tr>"
        "</w:tbl>"
    )
    resolve_revisions(root, "reject")
    assert text_of(root) == ["kept"]


# --------------------------------------------------------------- formatting changes
def test_rejecting_a_style_change_puts_the_old_style_back():
    """A paragraph restyled under review must stop being a heading when rejected."""
    root = body(
        '<w:p><w:pPr><w:pStyle w:val="Heading1"/>'
        f'<w:pPrChange {REV}><w:pPr><w:pStyle w:val="Normal"/></w:pPr></w:pPrChange>'
        f"</w:pPr>{run_xml('Was body text.')}</w:p>"
    )
    resolve_revisions(root, "reject")
    assert root.find(f".//{W}pStyle").get(f"{W}val") == "Normal"


def test_accepting_a_style_change_keeps_the_new_style_and_drops_the_record():
    root = body(
        '<w:p><w:pPr><w:pStyle w:val="Heading1"/>'
        f'<w:pPrChange {REV}><w:pPr><w:pStyle w:val="Normal"/></w:pPr></w:pPrChange>'
        f"</w:pPr>{run_xml('Now a heading.')}</w:p>"
    )
    resolve_revisions(root, "accept")
    assert root.find(f".//{W}pStyle").get(f"{W}val") == "Heading1"
    assert not document_has_revisions(root)


# ------------------------------------------------------------------------ end to end
def revised_document() -> bytes:
    """A real .docx whose second paragraph carries a substitution, as Word writes it."""
    document = docx.Document()
    document.add_paragraph("Untouched.")
    paragraph = document.add_paragraph()
    paragraph._p.append(etree.fromstring(f'<w:r xmlns:w="{W_NS}"><w:t xml:space="preserve">The quick </w:t></w:r>'))
    paragraph._p.append(etree.fromstring(f'<w:ins xmlns:w="{W_NS}" {REV}>{run_xml("crimson")}</w:ins>'))
    paragraph._p.append(etree.fromstring(f'<w:del xmlns:w="{W_NS}" {REV}>{deleted_run("brown")}</w:del>'))
    paragraph._p.append(etree.fromstring(f'<w:r xmlns:w="{W_NS}"><w:t xml:space="preserve"> fox.</w:t></w:r>'))
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.mark.parametrize(
    ("policy", "expected"),
    [
        ("accept", "The quick crimson fox."),
        ("reject", "The quick brown fox."),
        ("mark", "The quick crimson~~brown~~ fox."),
    ],
)
def test_each_policy_end_to_end(policy: str, expected: str):
    from all2md import to_markdown

    out = to_markdown(
        io.BytesIO(revised_document()), parser_options=DocxOptions(revisions=policy), source_format="docx"
    )
    assert expected in out


def test_accept_is_the_default():
    from all2md import to_markdown

    assert "The quick crimson fox." in to_markdown(io.BytesIO(revised_document()), source_format="docx")


def test_mark_carries_the_revision_facts_on_the_node():
    parser = DocxToAstConverter(DocxOptions(revisions="mark"))
    document = parser.parse(io.BytesIO(revised_document()))
    struck = extract_nodes(document, Strikethrough)
    assert len(struck) == 1
    assert struck[0].metadata["revision"] == {
        "type": "delete",
        "author": "A Reviewer",
        "date": "2026-08-31T18:31:00Z",
        "id": "1",
    }


def test_a_callers_document_is_not_edited_underneath_them():
    """Resolution rewrites the tree, so a borrowed ``Document`` is copied first."""
    live = docx.Document(io.BytesIO(revised_document()))
    before = len(live.element.findall(f".//{W}ins"))
    DocxToAstConverter(DocxOptions()).parse(live)
    assert len(live.element.findall(f".//{W}ins")) == before == 1
