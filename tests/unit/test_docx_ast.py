#  Copyright (c) 2025 Tom Villani, Ph.D.
#
# tests/unit/test_docx_ast.py
"""Unit tests for DOCX to AST converter.

Tests cover:
- DOCX paragraph to AST node conversion
- Heading detection and conversion
- Text formatting (bold, italic, underline, etc.)
- Hyperlink processing
- List detection and conversion
- Table structure conversion
- Image markdown parsing
- Run grouping and formatting preservation

"""

import docx
import pytest

from all2md.ast import (
    Document,
    Emphasis,
    Heading,
    Image,
    Link,
    List,
    ListItem,
    Paragraph,
    Strong,
    Strikethrough,
    Subscript,
    Superscript,
    Table,
    Text,
    Underline,
)
from all2md.converters.docx2ast import DocxToAstConverter
from all2md.options import DocxOptions


@pytest.mark.unit
class TestBasicElements:
    """Tests for basic DOCX element conversion."""

    def test_simple_paragraph(self) -> None:
        """Test converting a simple paragraph."""
        doc = docx.Document()
        doc.add_paragraph("Hello world")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        assert isinstance(ast_doc, Document)
        assert len(ast_doc.children) == 1
        assert isinstance(ast_doc.children[0], Paragraph)
        para = ast_doc.children[0]
        assert len(para.content) == 1
        assert isinstance(para.content[0], Text)
        assert para.content[0].content == "Hello world"

    def test_multiple_paragraphs(self) -> None:
        """Test converting multiple paragraphs."""
        doc = docx.Document()
        doc.add_paragraph("First")
        doc.add_paragraph("Second")
        doc.add_paragraph("Third")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        assert len(ast_doc.children) == 3
        assert all(isinstance(child, Paragraph) for child in ast_doc.children)

    def test_headings_1_to_6(self) -> None:
        """Test converting all heading levels."""
        doc = docx.Document()
        doc.add_heading("Heading 1", level=1)
        doc.add_heading("Heading 2", level=2)
        doc.add_heading("Heading 3", level=3)
        doc.add_heading("Heading 4", level=4)
        doc.add_heading("Heading 5", level=5)
        doc.add_heading("Heading 6", level=6)

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        assert len(ast_doc.children) == 6
        for i, child in enumerate(ast_doc.children):
            assert isinstance(child, Heading)
            assert child.level == i + 1
            assert isinstance(child.content[0], Text)
            assert child.content[0].content == f"Heading {i + 1}"

    def test_empty_paragraphs_skipped(self) -> None:
        """Test that empty paragraphs are skipped."""
        doc = docx.Document()
        doc.add_paragraph("First")
        doc.add_paragraph("")  # Empty
        doc.add_paragraph("   ")  # Whitespace only
        doc.add_paragraph("Second")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        # Only 2 paragraphs should be in AST (empty ones skipped)
        assert len(ast_doc.children) == 2
        assert ast_doc.children[0].content[0].content == "First"
        assert ast_doc.children[1].content[0].content == "Second"


@pytest.mark.unit
class TestTextFormatting:
    """Tests for text formatting conversion."""

    def test_bold_text(self) -> None:
        """Test bold text conversion."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold text")
        run.bold = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        assert isinstance(para_node.content[0], Strong)
        assert isinstance(para_node.content[0].content[0], Text)
        assert para_node.content[0].content[0].content == "Bold text"

    def test_italic_text(self) -> None:
        """Test italic text conversion."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Italic text")
        run.italic = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        assert isinstance(para_node.content[0], Emphasis)
        assert isinstance(para_node.content[0].content[0], Text)
        assert para_node.content[0].content[0].content == "Italic text"

    def test_underline_text(self) -> None:
        """Test underline text conversion."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Underlined text")
        run.underline = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        assert isinstance(para_node.content[0], Underline)
        assert isinstance(para_node.content[0].content[0], Text)
        assert para_node.content[0].content[0].content == "Underlined text"

    def test_strikethrough_text(self) -> None:
        """Test strikethrough text conversion."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Strikethrough text")
        run.font.strike = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        assert isinstance(para_node.content[0], Strikethrough)
        assert isinstance(para_node.content[0].content[0], Text)
        assert para_node.content[0].content[0].content == "Strikethrough text"

    def test_subscript_text(self) -> None:
        """Test subscript text conversion."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("H2O")
        run.font.subscript = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        assert isinstance(para_node.content[0], Subscript)

    def test_superscript_text(self) -> None:
        """Test superscript text conversion."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("E=mc^2")
        run.font.superscript = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        assert isinstance(para_node.content[0], Superscript)

    def test_multiple_formatting(self) -> None:
        """Test text with multiple formatting applied."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold and italic")
        run.bold = True
        run.italic = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        # Should be nested: Emphasis -> Strong -> Text (italic applied after bold)
        assert isinstance(para_node.content[0], Emphasis)
        assert isinstance(para_node.content[0].content[0], Strong)
        assert isinstance(para_node.content[0].content[0].content[0], Text)
        assert para_node.content[0].content[0].content[0].content == "Bold and italic"

    def test_mixed_formatting_runs(self) -> None:
        """Test paragraph with multiple runs having different formatting."""
        doc = docx.Document()
        para = doc.add_paragraph()
        para.add_run("Normal ")
        bold_run = para.add_run("bold ")
        bold_run.bold = True
        para.add_run("normal ")
        italic_run = para.add_run("italic")
        italic_run.italic = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        # Should have 4 inline nodes
        assert len(para_node.content) == 4
        assert isinstance(para_node.content[0], Text)
        assert isinstance(para_node.content[1], Strong)
        assert isinstance(para_node.content[2], Text)
        assert isinstance(para_node.content[3], Emphasis)


@pytest.mark.unit
class TestLists:
    """Tests for list conversion."""

    def test_bullet_list(self) -> None:
        """Test bullet list conversion."""
        doc = docx.Document()
        doc.add_paragraph("Item 1", style="List Bullet")
        doc.add_paragraph("Item 2", style="List Bullet")
        doc.add_paragraph("Item 3", style="List Bullet")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        # Should have one List node
        assert len(ast_doc.children) == 1
        assert isinstance(ast_doc.children[0], List)
        list_node = ast_doc.children[0]
        assert not list_node.ordered
        assert len(list_node.items) == 3

    def test_numbered_list(self) -> None:
        """Test numbered list conversion."""
        doc = docx.Document()
        doc.add_paragraph("First", style="List Number")
        doc.add_paragraph("Second", style="List Number")
        doc.add_paragraph("Third", style="List Number")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        # Should have one List node
        assert len(ast_doc.children) == 1
        assert isinstance(ast_doc.children[0], List)
        list_node = ast_doc.children[0]
        assert list_node.ordered
        assert len(list_node.items) == 3

    def test_list_followed_by_paragraph(self) -> None:
        """Test list properly closed when followed by regular paragraph."""
        doc = docx.Document()
        doc.add_paragraph("Item 1", style="List Bullet")
        doc.add_paragraph("Item 2", style="List Bullet")
        doc.add_paragraph("Regular paragraph")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        # Should have 2 nodes: List and Paragraph
        assert len(ast_doc.children) == 2
        assert isinstance(ast_doc.children[0], List)
        assert isinstance(ast_doc.children[1], Paragraph)

    def test_multiple_separate_lists(self) -> None:
        """Test multiple lists separated by paragraphs."""
        doc = docx.Document()
        doc.add_paragraph("List 1 Item 1", style="List Bullet")
        doc.add_paragraph("List 1 Item 2", style="List Bullet")
        doc.add_paragraph("Separator")
        doc.add_paragraph("List 2 Item 1", style="List Number")
        doc.add_paragraph("List 2 Item 2", style="List Number")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        # Should have 3 nodes: List, Paragraph, List
        assert len(ast_doc.children) == 3
        assert isinstance(ast_doc.children[0], List)
        assert not ast_doc.children[0].ordered  # Bullet
        assert isinstance(ast_doc.children[1], Paragraph)
        assert isinstance(ast_doc.children[2], List)
        assert ast_doc.children[2].ordered  # Numbered


@pytest.mark.unit
class TestTables:
    """Tests for table conversion."""

    def test_simple_table(self) -> None:
        """Test simple table conversion."""
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "H1"
        table.rows[0].cells[1].text = "H2"
        table.rows[1].cells[0].text = "R1C1"
        table.rows[1].cells[1].text = "R1C2"

        converter = DocxToAstConverter(options=DocxOptions(preserve_tables=True))
        ast_doc = converter.convert_to_ast(doc)

        # Should have one Table node
        assert len(ast_doc.children) == 1
        assert isinstance(ast_doc.children[0], Table)
        table_node = ast_doc.children[0]

        # Check header
        assert table_node.header is not None
        assert len(table_node.header.cells) == 2

        # Check data rows
        assert len(table_node.rows) == 1
        assert len(table_node.rows[0].cells) == 2

    def test_table_with_formatted_cells(self) -> None:
        """Test table with formatted cell content."""
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)

        # Add bold text to header
        para = table.rows[0].cells[0].paragraphs[0]
        run = para.add_run("Bold Header")
        run.bold = True

        # Add normal text to data cell
        table.rows[1].cells[0].text = "Normal text"

        converter = DocxToAstConverter(options=DocxOptions(preserve_tables=True))
        ast_doc = converter.convert_to_ast(doc)

        table_node = ast_doc.children[0]
        # Check that formatting is preserved
        header_cell = table_node.header.cells[0]
        assert isinstance(header_cell.content[0], Strong)

    def test_table_flattening(self) -> None:
        """Test table flattening when preserve_tables=False."""
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "H1"
        table.rows[0].cells[1].text = "H2"
        table.rows[1].cells[0].text = "R1C1"
        table.rows[1].cells[1].text = "R1C2"

        converter = DocxToAstConverter(options=DocxOptions(preserve_tables=False))
        ast_doc = converter.convert_to_ast(doc)

        # Should have multiple Paragraph nodes (flattened)
        assert all(isinstance(child, Paragraph) for child in ast_doc.children)
        assert len(ast_doc.children) >= 4  # At least 4 cells


@pytest.mark.unit
class TestHyperlinks:
    """Tests for hyperlink conversion."""

    def test_hyperlink_parsing(self) -> None:
        """Test hyperlink detection and URL extraction."""
        # Note: Creating actual hyperlinks in python-docx is complex,
        # so we test the conversion logic with mocked data
        converter = DocxToAstConverter()

        # Test URL extraction
        result = converter._process_hyperlink(None)
        assert result == (None, None)


@pytest.mark.unit
class TestImageParsing:
    """Tests for image markdown parsing."""

    def test_parse_full_markdown_image(self) -> None:
        """Test parsing complete markdown image with URL."""
        converter = DocxToAstConverter()

        img_node = converter._parse_markdown_image("![Alt text](https://example.com/image.png)")
        assert isinstance(img_node, Image)
        assert img_node.alt_text == "Alt text"
        assert img_node.url == "https://example.com/image.png"
        assert img_node.title is None

    def test_parse_image_with_title(self) -> None:
        """Test parsing image with title."""
        converter = DocxToAstConverter()

        img_node = converter._parse_markdown_image('![Alt](url.png "Title")')
        assert isinstance(img_node, Image)
        assert img_node.alt_text == "Alt"
        assert img_node.url == "url.png"
        assert img_node.title == "Title"

    def test_parse_alt_text_only_image(self) -> None:
        """Test parsing image with only alt text (no URL)."""
        converter = DocxToAstConverter()

        img_node = converter._parse_markdown_image("![Alt text only]")
        assert isinstance(img_node, Image)
        assert img_node.alt_text == "Alt text only"
        assert img_node.url == ""
        assert img_node.title is None

    def test_parse_base64_image(self) -> None:
        """Test parsing base64 data URI image."""
        converter = DocxToAstConverter()

        img_node = converter._parse_markdown_image("![image](data:image/png;base64,iVBORw0KGgo=)")
        assert isinstance(img_node, Image)
        assert img_node.alt_text == "image"
        assert img_node.url.startswith("data:image/png;base64,")

    def test_invalid_image_markdown(self) -> None:
        """Test that invalid markdown returns None."""
        converter = DocxToAstConverter()

        assert converter._parse_markdown_image("Not an image") is None
        assert converter._parse_markdown_image("[wrong](format)") is None
        assert converter._parse_markdown_image("![missing paren") is None


@pytest.mark.unit
class TestFormattingKey:
    """Tests for run formatting key generation."""

    def test_formatting_key_plain_text(self) -> None:
        """Test formatting key for plain text."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Plain text")

        converter = DocxToAstConverter()
        key = converter._get_run_formatting_key(run, False)

        # All False except is_hyperlink
        assert key == (False, False, False, False, False, False, False)

    def test_formatting_key_bold(self) -> None:
        """Test formatting key for bold text."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold")
        run.bold = True

        converter = DocxToAstConverter()
        key = converter._get_run_formatting_key(run, False)

        assert key[0] is True  # bold
        assert key[1] is False  # italic

    def test_formatting_key_all_formatting(self) -> None:
        """Test formatting key with all formatting applied."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Formatted")
        run.bold = True
        run.italic = True
        run.underline = True
        run.font.strike = True
        run.font.subscript = True

        converter = DocxToAstConverter()
        key = converter._get_run_formatting_key(run, True)

        assert key == (True, True, True, True, True, False, True)


@pytest.mark.unit
class TestListFinalization:
    """Tests for list finalization logic."""

    def test_list_at_end_of_document(self) -> None:
        """Test that list at end of document is properly finalized."""
        doc = docx.Document()
        doc.add_paragraph("Item 1", style="List Bullet")
        doc.add_paragraph("Item 2", style="List Bullet")
        # No paragraph after list

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        # List should still be in the document
        assert len(ast_doc.children) == 1
        assert isinstance(ast_doc.children[0], List)
        assert len(ast_doc.children[0].items) == 2

    def test_empty_list_stack(self) -> None:
        """Test finalization with empty list stack."""
        converter = DocxToAstConverter()
        result = converter._finalize_current_list()
        assert result is None


@pytest.mark.unit
class TestRunGrouping:
    """Tests for run grouping optimization."""

    def test_runs_with_same_formatting_grouped(self) -> None:
        """Test that consecutive runs with same formatting are grouped."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run1 = para.add_run("Bold ")
        run1.bold = True
        run2 = para.add_run("text")
        run2.bold = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        # Should be grouped into single Strong node
        assert len(para_node.content) == 1
        assert isinstance(para_node.content[0], Strong)
        # Check text is combined
        text_content = para_node.content[0].content[0].content
        assert text_content == "Bold text"

    def test_runs_with_different_formatting_separate(self) -> None:
        """Test that runs with different formatting are kept separate."""
        doc = docx.Document()
        para = doc.add_paragraph()
        run1 = para.add_run("Bold")
        run1.bold = True
        run2 = para.add_run("Italic")
        run2.italic = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        para_node = ast_doc.children[0]
        # Should have 2 separate inline nodes
        assert len(para_node.content) == 2
        assert isinstance(para_node.content[0], Strong)
        assert isinstance(para_node.content[1], Emphasis)


@pytest.mark.unit
class TestComplexStructures:
    """Tests for complex document structures."""

    def test_mixed_content_document(self) -> None:
        """Test document with mixed content types."""
        doc = docx.Document()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Introduction paragraph")
        doc.add_paragraph("Item 1", style="List Bullet")
        doc.add_paragraph("Item 2", style="List Bullet")
        doc.add_paragraph("Conclusion paragraph")

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        # Should have: Heading, Paragraph, List, Paragraph
        assert len(ast_doc.children) == 4
        assert isinstance(ast_doc.children[0], Heading)
        assert isinstance(ast_doc.children[1], Paragraph)
        assert isinstance(ast_doc.children[2], List)
        assert isinstance(ast_doc.children[3], Paragraph)

    def test_heading_with_formatting(self) -> None:
        """Test heading with formatted text."""
        doc = docx.Document()
        heading = doc.add_heading(level=1)
        run = heading.add_run("Bold Title")
        run.bold = True

        converter = DocxToAstConverter()
        ast_doc = converter.convert_to_ast(doc)

        heading_node = ast_doc.children[0]
        assert isinstance(heading_node, Heading)
        assert heading_node.level == 1
        # Check formatting is preserved
        assert isinstance(heading_node.content[0], Strong)
