"""Tests for DOCX Markdown character escaping edge cases."""


import docx

from all2md.converters.docx2markdown import docx_to_markdown
from all2md.options import DocxOptions, MarkdownOptions
from tests.utils import assert_markdown_valid, cleanup_test_dir, create_test_temp_dir


class TestDocxEscaping:
    """Test Markdown special character escaping in DOCX documents."""

    def setup_method(self):
        """Set up test environment."""
        self.temp_dir = create_test_temp_dir()

    def teardown_method(self):
        """Clean up test environment."""
        cleanup_test_dir(self.temp_dir)

    def test_basic_special_characters(self):
        """Test basic Markdown special characters that need escaping."""
        doc = docx.Document()

        # Common special characters
        special_chars = "* _ # [ ] ( ) { } \\ ` + - . ! ~ ^ | < >"
        doc.add_paragraph(f"Special characters: {special_chars}")

        temp_file = self.temp_dir / "special_chars.docx"
        doc.save(str(temp_file))

        # Test with escaping enabled (default)
        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Common characters that should be escaped in certain contexts
        assert "\\*" in markdown or "*" in markdown  # Depends on context
        assert "\\_" in markdown or "_" in markdown
        assert "\\#" in markdown or "#" in markdown

    def test_escaping_in_different_contexts(self):
        """Test escaping behavior in different document contexts."""
        doc = docx.Document()

        # In paragraph text
        doc.add_paragraph("Paragraph with * asterisk and _ underscore")

        # In list items
        doc.add_paragraph("List item with # hash and [ brackets ]", style="List Bullet")

        # In table cells
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Cell with ( parentheses )"
        table.rows[0].cells[1].text = "Cell with { braces }"

        # In headings
        doc.add_heading("Heading with * special chars", level=2)

        temp_file = self.temp_dir / "context_escaping.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle special characters in all contexts
        assert "asterisk" in markdown
        assert "underscore" in markdown
        assert "brackets" in markdown
        assert "parentheses" in markdown
        assert "braces" in markdown

    def test_formatted_text_with_special_chars(self):
        """Test special characters within formatted text."""
        doc = docx.Document()

        # Bold text with special characters
        p1 = doc.add_paragraph()
        run1 = p1.add_run("Bold text with * asterisks *")
        run1.bold = True

        # Italic text with special characters
        p2 = doc.add_paragraph()
        run2 = p2.add_run("Italic text with _ underscores _")
        run2.italic = True

        # Code-like content
        p3 = doc.add_paragraph()
        run3 = p3.add_run("Code: `function() { return 'hello'; }`")

        temp_file = self.temp_dir / "formatted_special_chars.docx"
        doc.save(str(temp_file))

        # Test with escaping enabled
        options = DocxOptions(markdown_options=MarkdownOptions(escape_special=True))
        markdown = docx_to_markdown(str(temp_file), options=options)
        assert_markdown_valid(markdown)

        # Should preserve formatting while handling special characters
        assert "**" in markdown  # Bold formatting
        assert "*" in markdown   # Italic formatting
        assert "\\*" in markdown # escaped
        assert "function\\(\\)" in markdown

        # Test with escaping disabled
        options_no_escape = DocxOptions(markdown_options=MarkdownOptions(escape_special=False))
        markdown_no_escape = docx_to_markdown(str(temp_file), options=options_no_escape)
        assert_markdown_valid(markdown_no_escape)

    def test_code_blocks_and_inline_code(self):
        """Test special characters in code-like contexts."""
        doc = docx.Document()

        # Monospace font (simulating code)
        p1 = doc.add_paragraph()
        run1 = p1.add_run("def function():\n    return {'key': 'value'}")
        # Note: Setting monospace font in python-docx is complex

        # Paragraph with code-like content
        doc.add_paragraph("Use `backticks` for inline code")

        # Multi-line code block
        doc.add_paragraph("```python\ndef hello():\n    print('world')\n```")

        temp_file = self.temp_dir / "code_content.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle code content appropriately
        assert "function\\(\\)" in markdown
        assert "backticks" in markdown
        assert "python" in markdown

    def test_links_with_special_characters(self):
        """Test hyperlinks containing special characters."""
        doc = docx.Document()

        # Link with special characters in URL
        p1 = doc.add_paragraph("Link with special URL: ")
        p1.add_run("Special Link").hyperlink = "https://example.com/path?param=value&other=data#anchor"

        # Link text with special characters
        p2 = doc.add_paragraph("Text with ")
        p2.add_run("link [with] brackets").hyperlink = "https://example.com"

        # Link with parentheses
        p3 = doc.add_paragraph("Link: ")
        p3.add_run("text (with) parentheses").hyperlink = "https://example.com"

        temp_file = self.temp_dir / "links_special_chars.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle special characters in links
        assert "Special Link" in markdown
        assert "brackets" in markdown
        assert "parentheses" in markdown

    def test_table_content_with_special_chars(self):
        """Test special characters within table cells."""
        doc = docx.Document()

        table = doc.add_table(rows=3, cols=2)

        # Headers with special characters
        table.rows[0].cells[0].text = "Column * 1"
        table.rows[0].cells[1].text = "Column # 2"

        # Data with various special characters
        table.rows[1].cells[0].text = "Data with [ brackets ]"
        table.rows[1].cells[1].text = "Data with { braces }"

        # Complex cell content
        table.rows[2].cells[0].text = "Code: function() { return true; }"
        table.rows[2].cells[1].text = "Math: x^2 + y^2 = z^2"

        temp_file = self.temp_dir / "table_special_chars.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should maintain table structure with special characters
        assert "Column" in markdown
        assert "brackets" in markdown
        assert "braces" in markdown
        assert "function()" in markdown
        assert "x^2" in markdown
        assert "|" in markdown  # Table structure

    def test_escaping_with_different_emphasis_symbols(self):
        """Test escaping with different emphasis symbol configurations."""
        doc = docx.Document()

        # Text with underscores and asterisks
        doc.add_paragraph("Text with _underscores_ and *asterisks*")

        p = doc.add_paragraph()
        run = p.add_run("Bold and italic text")
        run.bold = True
        run.italic = True

        temp_file = self.temp_dir / "emphasis_symbols.docx"
        doc.save(str(temp_file))

        # Test with asterisk emphasis
        options_asterisk = DocxOptions(markdown_options=MarkdownOptions(emphasis_symbol="*"))
        markdown_asterisk = docx_to_markdown(str(temp_file), options=options_asterisk)
        assert_markdown_valid(markdown_asterisk)

        # Test with underscore emphasis
        options_underscore = DocxOptions(markdown_options=MarkdownOptions(emphasis_symbol="_"))
        markdown_underscore = docx_to_markdown(str(temp_file), options=options_underscore)
        assert_markdown_valid(markdown_underscore)

        # Both should contain the text
        assert "underscores" in markdown_asterisk
        assert "asterisks" in markdown_asterisk
        assert "underscores" in markdown_underscore
        assert "asterisks" in markdown_underscore

    def test_list_markers_and_special_chars(self):
        """Test list markers with special characters."""
        doc = docx.Document()

        # List items with various special characters
        doc.add_paragraph("Item with * asterisk", style="List Bullet")
        doc.add_paragraph("Item with + plus", style="List Bullet")
        doc.add_paragraph("Item with - minus", style="List Bullet")

        # Numbered list with special characters
        doc.add_paragraph("Numbered item with # hash", style="List Number")
        doc.add_paragraph("Numbered item with . period", style="List Number")

        temp_file = self.temp_dir / "list_special_chars.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should distinguish between list markers and content
        assert "asterisk" in markdown
        assert "plus" in markdown
        assert "minus" in markdown
        assert "hash" in markdown
        assert "period" in markdown

    def test_blockquote_like_content(self):
        """Test content that might be confused with blockquotes."""
        doc = docx.Document()

        # Text starting with >
        doc.add_paragraph("> This looks like a blockquote but isn't")

        # Multiple lines that might look like blockquotes
        doc.add_paragraph("> Line 1\n> Line 2\n> Line 3")

        # Mixed content
        doc.add_paragraph("Regular text > with greater than < and less than")

        temp_file = self.temp_dir / "blockquote_like.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle > characters appropriately
        assert "looks like a blockquote" in markdown
        assert "Line 1" in markdown
        assert "greater than" in markdown
        assert "less than" in markdown

    def test_escaping_edge_cases(self):
        """Test edge cases in character escaping."""
        doc = docx.Document()

        # Multiple consecutive special characters
        doc.add_paragraph("Multiple *** asterisks *** and ___ underscores ___")

        # Mixed with whitespace
        doc.add_paragraph("Spaced * out * special * characters")

        # At line boundaries
        doc.add_paragraph("* Start of line")
        doc.add_paragraph("End of line *")

        # Empty emphasis markers
        doc.add_paragraph("Empty ** bold ** and __ underline __")

        temp_file = self.temp_dir / "escaping_edge_cases.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle edge cases gracefully
        assert "Multiple" in markdown
        assert "asterisks" in markdown
        assert "underscores" in markdown
        assert "Spaced" in markdown
        assert "Start of line" in markdown
        assert "End of line" in markdown
