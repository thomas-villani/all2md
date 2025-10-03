#  Copyright (c) 2025 Tom Villani, Ph.D.
#
# tests/unit/test_docx_renderer.py
"""Unit tests for DocxRenderer.

Tests cover:
- Rendering all node types to DOCX
- Document structure and styles
- Table formatting
- Image handling
- Edge cases and nested structures

Note: These tests require python-docx to be installed.

"""

import pytest
from io import BytesIO

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from all2md.options import DocxRendererOptions
from all2md.ast import (
    BlockQuote,
    Code,
    CodeBlock,
    DefinitionDescription,
    DefinitionList,
    DefinitionTerm,
    Document,
    Emphasis,
    FootnoteDefinition,
    FootnoteReference,
    Heading,
    Image,
    LineBreak,
    Link,
    List,
    ListItem,
    MathBlock,
    MathInline,
    Paragraph,
    Strikethrough,
    Strong,
    Subscript,
    Superscript,
    Table,
    TableCell,
    TableRow,
    Text,
    ThematicBreak,
    Underline,
)

if DOCX_AVAILABLE:
    from all2md.renderers.docx import DocxRenderer


pytestmark = pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")


@pytest.mark.unit
@pytest.mark.docx
class TestBasicRendering:
    """Tests for basic DOCX rendering."""

    def test_render_empty_document(self, tmp_path):
        """Test rendering an empty document."""
        doc = Document()
        renderer = DocxRenderer()
        output_file = tmp_path / "empty.docx"
        renderer.render(doc, output_file)

        assert output_file.exists()
        docx_doc = DocxDocument(str(output_file))
        # Empty document still has default styles
        assert docx_doc is not None

    def test_render_text_only(self, tmp_path):
        """Test rendering plain text."""
        doc = Document(children=[
            Paragraph(content=[Text(content="Hello world")])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "text.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert len(docx_doc.paragraphs) >= 1
        assert "Hello world" in docx_doc.paragraphs[0].text

    def test_render_multiple_paragraphs(self, tmp_path):
        """Test rendering multiple paragraphs."""
        doc = Document(children=[
            Paragraph(content=[Text(content="First paragraph")]),
            Paragraph(content=[Text(content="Second paragraph")])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "paras.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        paragraphs = [p for p in docx_doc.paragraphs if p.text.strip()]
        assert len(paragraphs) >= 2
        assert "First paragraph" in paragraphs[0].text
        assert "Second paragraph" in paragraphs[1].text

    def test_render_to_bytes_io(self):
        """Test rendering to BytesIO."""
        doc = Document(children=[
            Paragraph(content=[Text(content="Test")])
        ])
        renderer = DocxRenderer()
        output = BytesIO()
        renderer.render(doc, output)

        output.seek(0)
        docx_doc = DocxDocument(output)
        assert len(docx_doc.paragraphs) >= 1


@pytest.mark.unit
@pytest.mark.docx
class TestHeadingRendering:
    """Tests for heading rendering."""

    def test_heading_level_1(self, tmp_path):
        """Test rendering h1."""
        doc = Document(children=[
            Heading(level=1, content=[Text(content="Title")])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "h1.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        # Find heading paragraph
        heading_para = None
        for para in docx_doc.paragraphs:
            if para.text == "Title":
                heading_para = para
                break
        assert heading_para is not None

    def test_heading_level_2(self, tmp_path):
        """Test rendering h2."""
        doc = Document(children=[
            Heading(level=2, content=[Text(content="Subtitle")])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "h2.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        found = False
        for para in docx_doc.paragraphs:
            if "Subtitle" in para.text:
                found = True
                break
        assert found

    def test_heading_with_formatting(self, tmp_path):
        """Test heading with inline formatting."""
        doc = Document(children=[
            Heading(level=1, content=[
                Text(content="Bold "),
                Strong(content=[Text(content="title")])
            ])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "h_format.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        found = False
        for para in docx_doc.paragraphs:
            if "Bold" in para.text and "title" in para.text:
                found = True
                break
        assert found


@pytest.mark.unit
@pytest.mark.docx
class TestInlineFormatting:
    """Tests for inline formatting."""

    def test_strong(self, tmp_path):
        """Test bold text rendering."""
        doc = Document(children=[
            Paragraph(content=[Strong(content=[Text(content="bold")])])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "bold.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "bold" in para.text
        # Check if any run is bold
        has_bold = any(run.bold for run in para.runs)
        assert has_bold

    def test_emphasis(self, tmp_path):
        """Test italic text rendering."""
        doc = Document(children=[
            Paragraph(content=[Emphasis(content=[Text(content="italic")])])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "italic.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "italic" in para.text
        # Check if any run is italic
        has_italic = any(run.italic for run in para.runs)
        assert has_italic

    def test_code(self, tmp_path):
        """Test inline code rendering."""
        doc = Document(children=[
            Paragraph(content=[Code(content="code")])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "code.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "code" in para.text

    def test_strikethrough(self, tmp_path):
        """Test strikethrough rendering."""
        doc = Document(children=[
            Paragraph(content=[Strikethrough(content=[Text(content="deleted")])])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "strike.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "deleted" in para.text

    def test_underline(self, tmp_path):
        """Test underline rendering."""
        doc = Document(children=[
            Paragraph(content=[Underline(content=[Text(content="underlined")])])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "underline.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "underlined" in para.text

    def test_nested_formatting(self, tmp_path):
        """Test nested inline formatting."""
        doc = Document(children=[
            Paragraph(content=[
                Strong(content=[
                    Emphasis(content=[Text(content="bold italic")])
                ])
            ])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "nested.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "bold italic" in para.text


@pytest.mark.unit
@pytest.mark.docx
class TestListRendering:
    """Tests for list rendering."""

    def test_unordered_list(self, tmp_path):
        """Test unordered list rendering."""
        doc = Document(children=[
            List(ordered=False, items=[
                ListItem(children=[Paragraph(content=[Text(content="Item 1")])]),
                ListItem(children=[Paragraph(content=[Text(content="Item 2")])])
            ])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "ul.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        paragraphs = docx_doc.paragraphs
        text_content = [p.text for p in paragraphs]
        assert "Item 1" in text_content
        assert "Item 2" in text_content

    def test_ordered_list(self, tmp_path):
        """Test ordered list rendering."""
        doc = Document(children=[
            List(ordered=True, items=[
                ListItem(children=[Paragraph(content=[Text(content="First")])]),
                ListItem(children=[Paragraph(content=[Text(content="Second")])])
            ])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "ol.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        paragraphs = docx_doc.paragraphs
        text_content = [p.text for p in paragraphs]
        assert "First" in text_content
        assert "Second" in text_content


@pytest.mark.unit
@pytest.mark.docx
class TestTableRendering:
    """Tests for table rendering."""

    def test_simple_table(self, tmp_path):
        """Test basic table rendering."""
        doc = Document(children=[
            Table(
                header=TableRow(cells=[
                    TableCell(content=[Text(content="Name")]),
                    TableCell(content=[Text(content="Age")])
                ]),
                rows=[
                    TableRow(cells=[
                        TableCell(content=[Text(content="Alice")]),
                        TableCell(content=[Text(content="30")])
                    ])
                ]
            )
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "table.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert len(docx_doc.tables) >= 1
        table = docx_doc.tables[0]

        # Check table has correct dimensions
        assert len(table.rows) == 2
        assert len(table.rows[0].cells) == 2

        # Check header content
        assert "Name" in table.rows[0].cells[0].text
        assert "Age" in table.rows[0].cells[1].text

        # Check body content
        assert "Alice" in table.rows[1].cells[0].text
        assert "30" in table.rows[1].cells[1].text

    def test_table_without_header(self, tmp_path):
        """Test table without header row."""
        doc = Document(children=[
            Table(
                rows=[
                    TableRow(cells=[
                        TableCell(content=[Text(content="A")]),
                        TableCell(content=[Text(content="B")])
                    ])
                ]
            )
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "table_no_header.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert len(docx_doc.tables) >= 1


@pytest.mark.unit
@pytest.mark.docx
class TestBlockElements:
    """Tests for block-level elements."""

    def test_code_block(self, tmp_path):
        """Test code block rendering."""
        doc = Document(children=[
            CodeBlock(content="def hello():\n    print('world')", language="python")
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "codeblock.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        found = False
        for para in docx_doc.paragraphs:
            if "def hello()" in para.text:
                found = True
                break
        assert found

    def test_blockquote(self, tmp_path):
        """Test blockquote rendering."""
        doc = Document(children=[
            BlockQuote(children=[
                Paragraph(content=[Text(content="Quoted text")])
            ])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "blockquote.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        found = False
        for para in docx_doc.paragraphs:
            if "Quoted text" in para.text:
                found = True
                break
        assert found

    def test_thematic_break(self, tmp_path):
        """Test horizontal rule rendering."""
        doc = Document(children=[
            ThematicBreak()
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "hr.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        # Thematic break creates a paragraph with border
        assert len(docx_doc.paragraphs) >= 1


@pytest.mark.unit
@pytest.mark.docx
class TestDocumentMetadata:
    """Tests for document metadata."""

    def test_metadata_properties(self, tmp_path):
        """Test document properties from metadata."""
        doc = Document(
            metadata={
                "title": "Test Document",
                "author": "Test Author",
                "subject": "Testing"
            },
            children=[Paragraph(content=[Text(content="Content")])]
        )
        renderer = DocxRenderer()
        output_file = tmp_path / "metadata.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert docx_doc.core_properties.title == "Test Document"
        assert docx_doc.core_properties.author == "Test Author"
        assert docx_doc.core_properties.subject == "Testing"


@pytest.mark.unit
@pytest.mark.docx
class TestDefinitionLists:
    """Tests for definition list rendering."""

    def test_definition_list(self, tmp_path):
        """Test definition list rendering."""
        doc = Document(children=[
            DefinitionList(items=[
                (
                    DefinitionTerm(content=[Text(content="Term")]),
                    [DefinitionDescription(content=[Text(content="Description")])]
                )
            ])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "deflist.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        text_content = " ".join(p.text for p in docx_doc.paragraphs)
        assert "Term" in text_content
        assert "Description" in text_content


@pytest.mark.unit
@pytest.mark.docx
class TestRendererOptions:
    """Tests for renderer options."""

    def test_custom_fonts(self, tmp_path):
        """Test custom font settings."""
        doc = Document(children=[
            Paragraph(content=[Text(content="Custom font")])
        ])
        options = DocxRendererOptions(
            default_font="Arial",
            default_font_size=14
        )
        renderer = DocxRenderer(options)
        output_file = tmp_path / "custom_font.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert docx_doc is not None

    def test_code_font_option(self, tmp_path):
        """Test code font option."""
        doc = Document(children=[
            CodeBlock(content="code here")
        ])
        options = DocxRendererOptions(code_font="Monaco")
        renderer = DocxRenderer(options)
        output_file = tmp_path / "code_font.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert docx_doc is not None


@pytest.mark.unit
@pytest.mark.docx
class TestMathRendering:
    """Tests for math rendering."""

    def test_inline_math(self, tmp_path):
        """Test inline math rendering."""
        doc = Document(children=[
            Paragraph(content=[
                MathInline(content="x^2", notation="latex")
            ])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "math_inline.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        # Math is rendered as plain text in simple implementation
        found = False
        for para in docx_doc.paragraphs:
            if "x^2" in para.text:
                found = True
                break
        assert found

    def test_block_math(self, tmp_path):
        """Test block math rendering."""
        doc = Document(children=[
            MathBlock(content="E = mc^2", notation="latex")
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "math_block.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        # Math is rendered as code block
        found = False
        for para in docx_doc.paragraphs:
            if "E = mc^2" in para.text:
                found = True
                break
        assert found


@pytest.mark.unit
@pytest.mark.docx
class TestFootnotes:
    """Tests for footnote rendering."""

    def test_footnote_reference(self, tmp_path):
        """Test footnote reference rendering."""
        doc = Document(children=[
            Paragraph(content=[
                Text(content="Text"),
                FootnoteReference(identifier="1")
            ])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "footnote_ref.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        # Footnote reference is rendered as superscript
        text_content = "".join(p.text for p in docx_doc.paragraphs)
        assert "Text" in text_content

    def test_footnote_definition(self, tmp_path):
        """Test footnote definition rendering."""
        doc = Document(children=[
            FootnoteDefinition(identifier="1", content=[Text(content="Footnote text")])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "footnote_def.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        text_content = "".join(p.text for p in docx_doc.paragraphs)
        assert "Footnote text" in text_content


@pytest.mark.unit
@pytest.mark.docx
class TestLineBreaks:
    """Tests for line break rendering."""

    def test_hard_line_break(self, tmp_path):
        """Test hard line break rendering."""
        doc = Document(children=[
            Paragraph(content=[
                Text(content="Line 1"),
                LineBreak(soft=False),
                Text(content="Line 2")
            ])
        ])
        renderer = DocxRenderer()
        output_file = tmp_path / "linebreak.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        para = docx_doc.paragraphs[0]
        assert "Line 1" in para.text
        assert "Line 2" in para.text
