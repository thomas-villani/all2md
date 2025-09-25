"""DOCX test fixture generators using python-docx.

This module provides functions to programmatically create DOCX documents
for testing various aspects of DOCX-to-Markdown conversion.
"""

import tempfile
from io import BytesIO
from pathlib import Path
from typing import Optional

import docx
from docx.shared import Inches, Pt, RGBColor


def create_docx_with_formatting() -> docx.Document:
    """Create a DOCX with various text formatting for testing emphasis detection.

    Returns
    -------
    docx.Document
        Document with bold, italic, underline, and combined formatting.
    """
    doc = docx.Document()

    # Title
    title = doc.add_heading("Formatting Test Document", level=1)

    # Basic formatting examples
    p1 = doc.add_paragraph("This paragraph contains ")
    p1.add_run("bold text").bold = True
    p1.add_run(", ")
    p1.add_run("italic text").italic = True
    p1.add_run(", and ")
    p1.add_run("underlined text").underline = True
    p1.add_run(".")

    # Combined formatting
    p2 = doc.add_paragraph("This paragraph has ")
    run = p2.add_run("bold and italic text")
    run.bold = True
    run.italic = True
    p2.add_run(" together.")

    # Different heading levels
    doc.add_heading("Level 2 Heading", level=2)
    doc.add_paragraph("Content under level 2 heading.")

    doc.add_heading("Level 3 Heading", level=3)
    doc.add_paragraph("Content under level 3 heading.")

    # Font size variations
    p3 = doc.add_paragraph()
    p3.add_run("Large text ").font.size = Pt(18)
    p3.add_run("normal text ").font.size = Pt(12)
    p3.add_run("small text").font.size = Pt(8)

    # Colors (though these may not convert to Markdown)
    p4 = doc.add_paragraph("This text has ")
    red_run = p4.add_run("red color")
    red_run.font.color.rgb = RGBColor(255, 0, 0)
    p4.add_run(" and this is normal.")

    return doc


def create_docx_with_tables() -> docx.Document:
    """Create a DOCX with tables for testing table conversion.

    Returns
    -------
    docx.Document
        Document with various table structures and formatting.
    """
    doc = docx.Document()

    doc.add_heading("Table Test Document", level=1)

    # Simple 3x3 table
    doc.add_paragraph("Here is a simple table:")

    table1 = doc.add_table(rows=1, cols=3)
    table1.style = 'Table Grid'

    # Headers
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Age'
    hdr_cells[2].text = 'City'

    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    data = [
        ('Alice Johnson', '25', 'New York'),
        ('Bob Smith', '30', 'San Francisco'),
        ('Carol Davis', '28', 'Los Angeles')
    ]

    for name, age, city in data:
        row_cells = table1.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = age
        row_cells[2].text = city

    doc.add_paragraph("\nHere is a more complex table with merged cells:")

    # Complex table with merged cells
    table2 = doc.add_table(rows=4, cols=4)
    table2.style = 'Table Grid'

    # Headers
    headers = ['Product', 'Q1', 'Q2', 'Total']
    for i, header in enumerate(headers):
        table2.cell(0, i).text = header
        for paragraph in table2.cell(0, i).paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data
    table2.cell(1, 0).text = 'Product A'
    table2.cell(1, 1).text = '100'
    table2.cell(1, 2).text = '150'
    table2.cell(1, 3).text = '250'

    table2.cell(2, 0).text = 'Product B'
    table2.cell(2, 1).text = '200'
    table2.cell(2, 2).text = '180'
    table2.cell(2, 3).text = '380'

    table2.cell(3, 0).text = 'Total'
    table2.cell(3, 1).text = '300'
    table2.cell(3, 2).text = '330'
    table2.cell(3, 3).text = '630'

    # Make totals row bold
    for i in range(4):
        for paragraph in table2.cell(3, i).paragraphs:
            for run in paragraph.runs:
                run.bold = True

    return doc


def create_docx_with_lists() -> docx.Document:
    """Create a DOCX with various list types for testing list conversion.

    Returns
    -------
    docx.Document
        Document with bullet lists, numbered lists, and nested lists.
    """
    doc = docx.Document()

    doc.add_heading("List Test Document", level=1)

    # Bullet list
    doc.add_paragraph("Simple bullet list:")
    doc.add_paragraph("First bullet item", style='List Bullet')
    doc.add_paragraph("Second bullet item", style='List Bullet')
    doc.add_paragraph("Third bullet item", style='List Bullet')

    doc.add_paragraph("Numbered list:")
    doc.add_paragraph("First numbered item", style='List Number')
    doc.add_paragraph("Second numbered item", style='List Number')
    doc.add_paragraph("Third numbered item", style='List Number')

    # Nested lists
    doc.add_paragraph("\nNested list example:")
    doc.add_paragraph("Main item 1", style='List Bullet')

    # Create indented sub-items (manual approach since python-docx nested lists are complex)
    p1 = doc.add_paragraph("Sub-item 1.1", style='List Bullet')
    p1.paragraph_format.left_indent = Inches(0.5)

    p2 = doc.add_paragraph("Sub-item 1.2", style='List Bullet')
    p2.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph("Main item 2", style='List Bullet')

    p3 = doc.add_paragraph("Sub-item 2.1", style='List Bullet')
    p3.paragraph_format.left_indent = Inches(0.5)

    # Mixed list with formatting
    doc.add_paragraph("\nList with formatted text:")
    p4 = doc.add_paragraph("", style='List Bullet')
    p4.add_run("Item with ").bold = False
    p4.add_run("bold text").bold = True
    p4.add_run(" in it")

    p5 = doc.add_paragraph("", style='List Bullet')
    p5.add_run("Item with ").italic = False
    p5.add_run("italic text").italic = True
    p5.add_run(" in it")

    return doc


def create_docx_with_images() -> docx.Document:
    """Create a DOCX with placeholder images for testing image handling.

    Returns
    -------
    docx.Document
        Document with image placeholders and captions.
    """
    doc = docx.Document()

    doc.add_heading("Image Test Document", level=1)

    doc.add_paragraph("This document contains image placeholders.")

    # We can't easily create real images in this fixture, but we can
    # add text that represents where images would be
    doc.add_paragraph("[ Image 1: Sample chart showing data trends ]")
    doc.add_paragraph("Figure 1: This would be a caption for the first image.")

    doc.add_paragraph("\nSome text between images.")

    doc.add_paragraph("[ Image 2: Additional visualization ]")
    doc.add_paragraph("Figure 2: This would be a caption for the second image.")

    doc.add_paragraph("\nText after images to test proper spacing and formatting.")

    return doc


def create_docx_with_links() -> docx.Document:
    """Create a DOCX with hyperlinks for testing link conversion.

    Returns
    -------
    docx.Document
        Document with various types of hyperlinks.
    """
    doc = docx.Document()

    doc.add_heading("Link Test Document", level=1)

    # Simple paragraph with text describing links
    # Note: Creating actual hyperlinks in python-docx is complex,
    # so we'll create text that represents links for testing purposes
    doc.add_paragraph("This document contains various types of links:")

    p1 = doc.add_paragraph("Visit ")
    p1.add_run("https://www.example.com").underline = True
    p1.add_run(" for more information.")

    p2 = doc.add_paragraph("Email us at ")
    p2.add_run("contact@example.com").underline = True
    p2.add_run(" for support.")

    p3 = doc.add_paragraph("Check out ")
    run = p3.add_run("this link")
    run.underline = True
    run.bold = True
    p3.add_run(" which has both formatting and linking.")

    # Internal references
    doc.add_paragraph("\nInternal document references:")
    p4 = doc.add_paragraph("See ")
    p4.add_run("Section 2.1").italic = True
    p4.add_run(" for details.")

    return doc


def create_minimal_docx(title: str = "Test Document", content: str = "Test content") -> docx.Document:
    """Create a minimal DOCX document for basic testing.

    Parameters
    ----------
    title : str, optional
        Document title, by default "Test Document"
    content : str, optional
        Document content, by default "Test content"

    Returns
    -------
    docx.Document
        Simple document with title and content.
    """
    doc = docx.Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    return doc


def save_docx_to_bytes(doc: docx.Document) -> bytes:
    """Save a DOCX document to bytes for testing.

    Parameters
    ----------
    doc : docx.Document
        Document to convert to bytes

    Returns
    -------
    bytes
        Document as bytes
    """
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.read()


def create_docx_file(doc: docx.Document, file_path: Optional[Path] = None) -> Path:
    """Save a DOCX document to a temporary file.

    Parameters
    ----------
    doc : docx.Document
        Document to save
    file_path : Path, optional
        File path to save to, by default creates temporary file

    Returns
    -------
    Path
        Path to the saved file
    """
    if file_path is None:
        temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        file_path = Path(temp_file.name)
        temp_file.close()

    doc.save(str(file_path))
    return file_path
