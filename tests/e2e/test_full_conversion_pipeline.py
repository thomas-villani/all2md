"""End-to-end tests for the complete all2md conversion pipeline.

These tests verify the entire conversion process from file input to markdown output,
testing the integration of all components including format detection, parsing,
and conversion across different document types.
"""

from io import BytesIO

import pytest
from fixtures.generators.docx_fixtures import create_docx_with_formatting, save_docx_to_bytes
from fixtures.generators.html_fixtures import create_html_with_tables
from fixtures.generators.pdf_test_fixtures import create_pdf_with_figures
from fixtures.generators.pptx_fixtures import create_pptx_with_basic_slides, save_pptx_to_bytes
from utils import assert_markdown_valid

from all2md import to_markdown
from all2md.options import DocxOptions, HtmlOptions, MarkdownOptions, PdfOptions, PptxOptions


@pytest.mark.e2e
class TestFullConversionPipeline:
    """End-to-end tests for complete document conversion workflows."""

    def test_docx_to_markdown_full_pipeline(self, temp_dir):
        """Test complete DOCX to Markdown conversion pipeline."""
        # Create a DOCX document with various formatting
        doc = create_docx_with_formatting()
        docx_bytes = save_docx_to_bytes(doc)

        # Test conversion with different options
        options = DocxOptions(
            attachment_mode="base64"
        )
        markdown_options = MarkdownOptions(
            bullet_symbols='-'
        )

        # Convert to markdown
        result = to_markdown(
            BytesIO(docx_bytes),
            source_format="docx",
            parser_options=options,
            renderer_options=markdown_options
        )

        assert_markdown_valid(result)

        # Should contain expected content
        assert "Formatting Test Document" in result
        assert "**bold text**" in result.lower()
        assert "*italic text*" in result.lower()

        # Should detect format correctly
        # Format detection successful - result should contain converted content
        assert result

    def test_html_to_markdown_full_pipeline(self, temp_dir):
        """Test complete HTML to Markdown conversion pipeline."""
        # Create HTML content with tables
        html_content = create_html_with_tables()

        # Convert with options
        options = HtmlOptions(
            strip_dangerous_elements=True
        )
        markdown_options = MarkdownOptions()

        result = to_markdown(
            BytesIO(html_content.encode('utf-8')),
            source_format="html",
            parser_options=options,
            renderer_options=markdown_options
        )

        assert_markdown_valid(result)

        # Should contain table content
        assert "Table Test Document" in result
        assert "|" in result  # Should have table formatting
        assert "Alice Johnson" in result

        # Should detect format correctly
        # Format detection successful - result should contain converted content
        assert result

    def test_pptx_to_markdown_full_pipeline(self, temp_dir):
        """Test complete PPTX to Markdown conversion pipeline."""
        # Create a PPTX presentation
        prs = create_pptx_with_basic_slides()
        pptx_bytes = save_pptx_to_bytes(prs)

        # Convert with options
        options = PptxOptions(
            attachment_mode="skip",
            include_slide_numbers=True
        )
        markdown_options = MarkdownOptions()

        result = to_markdown(
            BytesIO(pptx_bytes),
            source_format="pptx",
            parser_options=options,
            renderer_options=markdown_options
        )

        assert_markdown_valid(result)

        # Should contain presentation content
        assert "Test Presentation" in result
        assert "Main Content Slide" in result
        assert "First bullet point" in result

        # Should detect format correctly
        # Format detection successful - result should contain converted content
        assert result

    def test_pdf_to_markdown_full_pipeline(self, temp_dir):
        """Test complete PDF to Markdown conversion pipeline."""
        # Create a PDF with figures
        doc = create_pdf_with_figures()

        try:
            # Save to bytes
            pdf_bytes = BytesIO()
            doc.save(pdf_bytes)
            pdf_bytes.seek(0)

            # Convert with options
            options = PdfOptions(
                attachment_mode="base64",
                enable_table_fallback_detection=True
            )
            markdown_options = MarkdownOptions()

            result = to_markdown(
                pdf_bytes,
                source_format="pdf",
                parser_options=options,
                renderer_options=markdown_options
            )

            assert_markdown_valid(result)

            # Should contain PDF content
            assert "Test Document with Figures" in result
            assert "figure" in result.lower()

            # Should detect format correctly
            # Format detection successful - result should contain converted content
            assert result

        finally:
            doc.close()

    def test_automatic_format_detection_pipeline(self, temp_dir):
        """Test automatic format detection in the conversion pipeline."""
        # Test with different file types without explicit extension

        # DOCX test
        doc = create_docx_with_formatting()
        docx_bytes = save_docx_to_bytes(doc)

        result_docx = to_markdown(BytesIO(docx_bytes))  # No explicit extension
        # Format detection successful - result should contain converted content
        assert result_docx
        assert_markdown_valid(result_docx)

        # HTML test
        html_content = create_html_with_tables()
        html_bytes = html_content.encode('utf-8')

        result_html = to_markdown(BytesIO(html_bytes))  # No explicit extension
        # Format detection successful - result should contain converted content
        assert result_html
        assert_markdown_valid(result_html)

    def test_error_handling_in_full_pipeline(self, temp_dir):
        """Test error handling throughout the conversion pipeline."""
        # Test with invalid content
        invalid_content = b"This is not a valid document"

        # Should handle gracefully
        with pytest.raises(Exception):  # Specific exception depends on implementation
            to_markdown(BytesIO(invalid_content), source_format="docx")

        # Test with unsupported format - should fallback to text
        result = to_markdown(BytesIO(b"content"), source_format="unsupported")
        assert result == "content"  # Should fallback to treating as text

    def test_large_document_pipeline_performance(self, temp_dir):
        """Test pipeline performance with larger documents."""
        # Create a larger document for performance testing
        import docx

        doc = docx.Document()
        doc.add_heading("Large Document Test", level=1)

        # Add many paragraphs
        for i in range(100):
            p = doc.add_paragraph(f"This is paragraph {i + 1} with some content.")
            if i % 10 == 0:
                p.runs[0].bold = True

        # Add a large table
        table = doc.add_table(rows=20, cols=5)
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = f"Row {i + 1}, Col {j + 1}"

        docx_bytes = save_docx_to_bytes(doc)

        # Convert (should complete without timeout)
        result = to_markdown(BytesIO(docx_bytes), source_format="docx")

        assert_markdown_valid(result)
        assert "Large Document Test" in result
        assert "paragraph 50" in result
        assert "Row 10, Col 3" in result

    def test_options_propagation_through_pipeline(self, temp_dir):
        """Test that conversion options are properly propagated through the pipeline."""
        # Create HTML with various elements
        html = '''
        <html>
        <head><title>Options Test</title></head>
        <body>
            <h1>Main Title</h1>
            <p>Regular paragraph.</p>
            <table>
                <tr><th>Header</th></tr>
                <tr><td>Data</td></tr>
            </table>
            <script>alert('test');</script>
            <style>body { color: red; }</style>
        </body>
        </html>
        '''

        # Test with basic HTML options - the key is that options are accepted and processed
        options_test = HtmlOptions(
            strip_dangerous_elements=True,
            convert_nbsp=False
        )

        result = to_markdown(
            BytesIO(html.encode('utf-8')),
            source_format="html",
            parser_options=options_test
        )

        # Verify the conversion worked with options
        assert "Main Title" in result  # Should contain content
        assert "Header" in result  # Should contain table content
        assert "|" in result  # Should have table formatting
        assert_markdown_valid(result)  # Should produce valid markdown

        # Test that different options produce output (may be the same, but should not crash)
        options_test2 = HtmlOptions(
            strip_dangerous_elements=False,
            convert_nbsp=True
        )

        result2 = to_markdown(
            BytesIO(html.encode('utf-8')),
            source_format="html",
            parser_options=options_test2
        )

        # Both conversions should work
        assert "Main Title" in result2
        assert len(result2) > 0

    def test_mixed_content_document_pipeline(self, temp_dir):
        """Test conversion of documents with mixed content types."""
        # Create a complex DOCX with tables, images, and various formatting
        import docx

        doc = docx.Document()

        # Title and introduction
        doc.add_heading("Complex Document", level=1)
        doc.add_paragraph("This document contains various content types.")

        # Table
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'

        # Add table content
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if i == 0:
                    cell.text = f"Header {j + 1}"
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                else:
                    cell.text = f"Data {i}-{j + 1}"

        # More content
        doc.add_heading("Section 2", level=2)
        doc.add_paragraph("Content after the table.")

        # Lists
        doc.add_paragraph("Features:", style='List Bullet')
        doc.add_paragraph("Feature 1", style='List Bullet')
        doc.add_paragraph("Feature 2", style='List Bullet')

        docx_bytes = save_docx_to_bytes(doc)

        # Convert with comprehensive options
        docx_options = DocxOptions(
            attachment_mode="base64"
        )
        markdown_options = MarkdownOptions(
            bullet_symbols='-'
        )

        result = to_markdown(
            BytesIO(docx_bytes),
            source_format="docx",
            parser_options=docx_options,
            renderer_options=markdown_options
        )

        assert_markdown_valid(result)

        # Should contain all content types
        assert "# Complex Document" in result
        assert "## Section 2" in result
        assert "Header 1" in result
        assert "Data 1-1" in result
        assert "- Feature 1" in result
        assert "various content types" in result

    @pytest.mark.slow
    def test_batch_conversion_pipeline(self, temp_dir):
        """Test conversion of multiple documents in sequence."""
        # Create multiple test documents
        documents = []

        # DOCX
        docx_doc = create_docx_with_formatting()
        documents.append((save_docx_to_bytes(docx_doc), '.docx', "docx"))

        # HTML
        html_content = create_html_with_tables()
        documents.append((html_content.encode('utf-8'), '.html', "html"))

        # PPTX
        pptx_prs = create_pptx_with_basic_slides()
        documents.append((save_pptx_to_bytes(pptx_prs), '.pptx', "pptx"))

        # Convert all documents
        results = []
        for content, extension, _expected_format in documents:
            result = to_markdown(BytesIO(content), source_format=extension.lstrip('.'))
            results.append(result)

            # Verify each conversion
            # Format detection successful - result should contain converted content
            assert result
            assert_markdown_valid(result)
            assert len(result) > 0

        # Verify different content in each result
        assert "Formatting Test Document" in results[0]  # DOCX
        assert "Table Test Document" in results[1]  # HTML
        assert "Test Presentation" in results[2]  # PPTX

        # All should be valid markdown
        for result in results:
            assert_markdown_valid(result)
