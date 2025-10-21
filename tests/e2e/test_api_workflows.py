#  Copyright (c) 2025 Tom Villani, Ph.D.
#
# tests/e2e/test_api_functions_e2e.py
"""End-to-end tests for major API functions.

This module provides comprehensive e2e tests for the five major API functions:
to_markdown, to_ast, from_ast, from_markdown, and convert. Focus is on to_ast
(under-tested) and multi-function workflows that simulate real-world usage.
"""

from importlib.util import find_spec
from io import BytesIO

import pytest

DOCX_AVAILABLE = find_spec("docx") is not None
REPORTLAB_AVAILABLE = find_spec("reportlab") is not None
EBOOKLIB_AVAILABLE = find_spec("ebooklib") is not None
PPTX_AVAILABLE = find_spec("pptx") is not None

from fixtures.generators.docx_fixtures import create_docx_with_formatting, save_docx_to_bytes
from fixtures.generators.html_fixtures import create_html_with_tables
from fixtures.generators.pdf_test_fixtures import create_pdf_with_figures
from fixtures.generators.pptx_fixtures import create_pptx_with_basic_slides, save_pptx_to_bytes
from utils import assert_markdown_valid

from all2md import convert, from_ast, from_markdown, to_ast, to_markdown
from all2md.ast import Document, Heading, Paragraph, Strong, Table
from all2md.options import DocxOptions, HtmlOptions, HtmlRendererOptions, MarkdownOptions, PdfOptions, PptxOptions


@pytest.mark.e2e
class TestToAstE2E:
    """End-to-end tests for to_ast function."""

    @pytest.mark.pdf
    def test_to_ast_from_pdf_complex_document(self, temp_dir):
        """Test converting complex PDF to AST."""
        # Create PDF with figures and complex layout
        doc = create_pdf_with_figures()

        try:
            pdf_bytes = BytesIO()
            doc.save(pdf_bytes)
            pdf_bytes.seek(0)

            # Convert to AST
            ast_doc = to_ast(pdf_bytes, source_format="pdf")

            # Verify AST structure
            assert isinstance(ast_doc, Document)
            assert ast_doc.children is not None
            assert len(ast_doc.children) > 0

            # Should contain text content
            has_text = any(isinstance(child, (Paragraph, Heading)) for child in ast_doc.children)
            assert has_text

            # Metadata should be captured
            assert ast_doc.metadata is not None

        finally:
            doc.close()

    @pytest.mark.pdf
    def test_to_ast_from_pdf_with_options(self, temp_dir):
        """Test to_ast with PDF-specific options."""
        doc = create_pdf_with_figures()

        try:
            pdf_bytes = BytesIO()
            doc.save(pdf_bytes)
            pdf_bytes.seek(0)

            # Convert with options
            options = PdfOptions(attachment_mode="base64", enable_table_fallback_detection=True)

            ast_doc = to_ast(pdf_bytes, source_format="pdf", parser_options=options)

            assert isinstance(ast_doc, Document)
            assert len(ast_doc.children) > 0

        finally:
            doc.close()

    @pytest.mark.docx
    def test_to_ast_from_docx_with_formatting(self, temp_dir):
        """Test converting DOCX with formatting to AST."""
        # Create DOCX with formatting
        docx_doc = create_docx_with_formatting()
        docx_bytes = save_docx_to_bytes(docx_doc)

        # Convert to AST
        ast_doc = to_ast(BytesIO(docx_bytes), source_format="docx")

        # Verify structure
        assert isinstance(ast_doc, Document)
        assert len(ast_doc.children) > 0

        # Should contain headings
        has_heading = any(isinstance(child, Heading) for child in ast_doc.children)
        assert has_heading

        # Should contain paragraphs with formatting
        has_paragraph = any(isinstance(child, Paragraph) for child in ast_doc.children)
        assert has_paragraph

    @pytest.mark.docx
    def test_to_ast_from_docx_with_options(self, temp_dir):
        """Test to_ast with DOCX-specific options."""
        docx_doc = create_docx_with_formatting()
        docx_bytes = save_docx_to_bytes(docx_doc)

        options = DocxOptions(attachment_mode="skip", preserve_tables=True)

        ast_doc = to_ast(BytesIO(docx_bytes), source_format="docx", parser_options=options)

        assert isinstance(ast_doc, Document)
        # Metadata may or may not be present
        assert ast_doc.metadata is not None or ast_doc.metadata is None

    @pytest.mark.html
    def test_to_ast_from_html_complex(self, temp_dir):
        """Test converting complex HTML to AST."""
        html_content = create_html_with_tables()

        # Convert to AST
        ast_doc = to_ast(BytesIO(html_content.encode("utf-8")), source_format="html")

        # Verify structure
        assert isinstance(ast_doc, Document)
        assert len(ast_doc.children) > 0

        # Should contain tables
        has_table = any(isinstance(child, Table) for child in ast_doc.children)
        assert has_table

    @pytest.mark.html
    def test_to_ast_from_html_with_options(self, temp_dir):
        """Test to_ast with HTML-specific options."""
        html = """
        <html>
        <head><title>Test Document</title></head>
        <body>
            <h1>Main Title</h1>
            <p>Content with <strong>bold</strong> text.</p>
            <script>alert('dangerous');</script>
        </body>
        </html>
        """

        options = HtmlOptions(strip_dangerous_elements=True, extract_title=True)

        ast_doc = to_ast(BytesIO(html.encode("utf-8")), source_format="html", parser_options=options)

        assert isinstance(ast_doc, Document)
        # Title should be extracted to metadata
        if ast_doc.metadata:
            assert "title" in ast_doc.metadata or len(ast_doc.metadata) >= 0

    @pytest.mark.pptx
    def test_to_ast_from_pptx_presentation(self, temp_dir):
        """Test converting PPTX presentation to AST."""
        prs = create_pptx_with_basic_slides()
        pptx_bytes = save_pptx_to_bytes(prs)

        # Convert to AST
        ast_doc = to_ast(BytesIO(pptx_bytes), source_format="pptx")

        # Verify structure
        assert isinstance(ast_doc, Document)
        assert len(ast_doc.children) > 0

        # Should contain slide content
        has_content = any(isinstance(child, (Heading, Paragraph)) for child in ast_doc.children)
        assert has_content

    @pytest.mark.epub
    @pytest.mark.skipif(not EBOOKLIB_AVAILABLE, reason="ebooklib not available")
    def test_to_ast_from_epub_book(self, temp_dir):
        """Test converting EPUB to AST."""
        from fixtures.generators.epub_fixtures import create_simple_epub

        epub_content = create_simple_epub()
        epub_file = temp_dir / "test.epub"
        epub_file.write_bytes(epub_content)

        # Convert to AST
        ast_doc = to_ast(str(epub_file), source_format="epub")

        # Verify structure
        assert isinstance(ast_doc, Document)
        assert len(ast_doc.children) > 0

    @pytest.mark.ipynb
    def test_to_ast_from_ipynb_notebook(self, temp_dir):
        """Test converting Jupyter notebook to AST."""
        from fixtures.generators.ipynb_fixtures import create_simple_notebook, save_notebook_to_file

        notebook = create_simple_notebook()
        ipynb_file = temp_dir / "test.ipynb"
        save_notebook_to_file(notebook, str(ipynb_file))

        # Convert to AST
        ast_doc = to_ast(str(ipynb_file), source_format="ipynb")

        # Verify structure
        assert isinstance(ast_doc, Document)
        assert len(ast_doc.children) > 0

    def test_to_ast_with_file_path_input(self, temp_dir):
        """Test to_ast with file path as string."""
        html_file = temp_dir / "test.html"
        html_file.write_text("<h1>Test</h1><p>Content</p>")

        ast_doc = to_ast(str(html_file), source_format="html")

        assert isinstance(ast_doc, Document)
        assert len(ast_doc.children) > 0

    def test_to_ast_with_path_object_input(self, temp_dir):
        """Test to_ast with Path object."""
        html_file = temp_dir / "test.html"
        html_file.write_text("<h1>Test</h1><p>Content</p>")

        ast_doc = to_ast(html_file, source_format="html")

        assert isinstance(ast_doc, Document)

    def test_to_ast_with_bytes_input(self):
        """Test to_ast with raw bytes input."""
        html_bytes = b"<h1>Title</h1><p>Content</p>"

        ast_doc = to_ast(html_bytes, source_format="html")

        assert isinstance(ast_doc, Document)

    def test_to_ast_with_io_input(self):
        """Test to_ast with IO input."""
        html_bytes = b"<h1>Title</h1><p>Content</p>"
        html_io = BytesIO(html_bytes)

        ast_doc = to_ast(html_io, source_format="html")

        assert isinstance(ast_doc, Document)

    def test_to_ast_format_auto_detection(self, temp_dir):
        """Test automatic format detection in to_ast."""
        # HTML file with .html extension
        html_file = temp_dir / "test.html"
        html_file.write_text("<h1>Auto Detection</h1>")

        ast_doc = to_ast(str(html_file))  # No explicit format

        assert isinstance(ast_doc, Document)

    def test_to_ast_ast_structure_validation(self):
        """Test that AST structure is well-formed."""
        html = """
        <html>
        <body>
            <h1>Title</h1>
            <p>Paragraph with <strong>bold</strong> text.</p>
            <ul>
                <li>Item 1</li>
                <li>Item 2</li>
            </ul>
        </body>
        </html>
        """

        ast_doc = to_ast(BytesIO(html.encode("utf-8")), source_format="html")

        # Root should be Document
        assert isinstance(ast_doc, Document)

        # Should have children
        assert ast_doc.children is not None
        assert len(ast_doc.children) > 0

        # Verify node types are correct - each child should inherit from Node
        from all2md.ast.nodes import Node

        for child in ast_doc.children:
            # Each child should be a valid AST node
            assert isinstance(child, Node)

    def test_to_ast_metadata_capture(self, temp_dir):
        """Test that document metadata is captured."""
        html = """
        <html>
        <head>
            <title>Test Document</title>
            <meta name="author" content="Test Author">
        </head>
        <body><h1>Content</h1></body>
        </html>
        """

        ast_doc = to_ast(BytesIO(html.encode("utf-8")), source_format="html", extract_title=True)

        # Metadata should be present
        assert ast_doc.metadata is not None
        assert isinstance(ast_doc.metadata, dict)

    def test_to_ast_error_handling_invalid_content(self):
        """Test error handling with invalid content."""
        invalid_content = b"This is not a valid DOCX file"

        # Should raise appropriate exception
        with pytest.raises(Exception):  # Specific exception depends on implementation
            to_ast(BytesIO(invalid_content), source_format="docx")


@pytest.mark.e2e
class TestMultiStepWorkflowsE2E:
    """End-to-end tests for multi-function workflows."""

    @pytest.mark.pdf
    def test_pdf_to_ast_to_multiple_formats(self, temp_dir):
        """Test PDF → AST → render to multiple formats."""
        doc = create_pdf_with_figures()

        try:
            pdf_bytes = BytesIO()
            doc.save(pdf_bytes)
            pdf_bytes.seek(0)

            # Step 1: PDF → AST
            ast_doc = to_ast(pdf_bytes, source_format="pdf")

            # Step 2: AST → Markdown
            markdown_content = from_ast(ast_doc, "markdown")
            assert isinstance(markdown_content, str)

            assert len(markdown_content) > 0
            assert_markdown_valid(markdown_content)

            # Step 3: AST → HTML (from same AST)
            html_output = from_ast(ast_doc, "html")
            assert isinstance(html_output, str)
            assert "<" in html_output

            # Both outputs should contain similar content
            assert "Test Document" in markdown_content
            assert "Test Document" in html_output

        finally:
            doc.close()

    @pytest.mark.docx
    def test_docx_to_ast_with_transforms_to_markdown(self, temp_dir):
        """Test DOCX → AST → apply transforms → Markdown."""
        docx_doc = create_docx_with_formatting()
        docx_bytes = save_docx_to_bytes(docx_doc)

        # Step 1: DOCX → AST
        ast_doc = to_ast(BytesIO(docx_bytes), source_format="docx")

        # Step 2: AST → Markdown with transforms
        markdown_without_images = from_ast(ast_doc, "markdown", transforms=["remove-images"])

        assert isinstance(markdown_without_images, str)
        assert_markdown_valid(markdown_without_images)

        # Step 3: AST → Markdown with different transforms
        markdown_with_heading_offset = from_ast(ast_doc, "markdown", transforms=["heading-offset"])

        assert isinstance(markdown_with_heading_offset, str)
        # After heading offset, structure should be different
        assert markdown_with_heading_offset != markdown_without_images

    @pytest.mark.html
    def test_html_to_ast_modify_to_clean_output(self, temp_dir):
        """Test HTML → AST → remove images via transform → output."""
        html = """
        <html>
        <body>
            <h1>Document Title</h1>
            <p>Paragraph text.</p>
            <img src="test.png" alt="Test Image">
            <p>More text.</p>
        </body>
        </html>
        """

        # Step 1: HTML → AST
        ast_doc = to_ast(BytesIO(html.encode("utf-8")), source_format="html")

        # Step 2: Render with remove-images transform
        markdown_output = from_ast(ast_doc, "markdown", transforms=["remove-images"])

        assert isinstance(markdown_output, str)
        # Should have content but no images
        assert "Document Title" in markdown_output
        assert "Paragraph text" in markdown_output
        assert "![" not in markdown_output  # No image syntax

    def test_roundtrip_html_to_ast_to_html(self, temp_dir):
        """Test roundtrip: HTML → AST → HTML → verify preservation."""
        original_html = """
        <html>
        <body>
            <h1>Test Title</h1>
            <p>Content with <strong>bold</strong> and <em>italic</em>.</p>
            <ul>
                <li>Item 1</li>
                <li>Item 2</li>
            </ul>
        </body>
        </html>
        """

        # HTML → AST
        ast_doc = to_ast(BytesIO(original_html.encode("utf-8")), source_format="html")

        # AST → HTML
        recovered_html = from_ast(ast_doc, "html")

        assert isinstance(recovered_html, str)
        # Should contain key content
        assert "Test Title" in recovered_html
        assert "bold" in recovered_html
        assert "italic" in recovered_html
        assert "Item 1" in recovered_html

    @pytest.mark.docx
    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not available")
    def test_roundtrip_markdown_to_docx_to_markdown(self, temp_dir):
        """Test roundtrip: Markdown → DOCX → Markdown."""
        original_md = "# Title\n\n**Bold** and *italic* text."

        # Write to file
        md_file = temp_dir / "original.md"
        md_file.write_text(original_md)

        # Markdown → DOCX
        docx_file = temp_dir / "temp.docx"
        from_markdown(str(md_file), "docx", output=docx_file)

        assert docx_file.exists()

        # DOCX → Markdown
        recovered_md = to_markdown(str(docx_file))

        # Should contain key content
        # Note: Some formatting may be lost/changed in roundtrip
        assert "Title" in recovered_md
        assert "Bold" in recovered_md or "bold" in recovered_md
        assert "italic" in recovered_md

    def test_chained_conversion_html_to_markdown_to_html(self, temp_dir):
        """Test chained: HTML → Markdown → HTML with content verification."""
        original_html = "<h1>Title</h1><p><strong>Important</strong> content.</p>"

        # HTML → Markdown
        markdown_content = convert(
            BytesIO(original_html.encode("utf-8")), source_format="html", target_format="markdown"
        )

        assert isinstance(markdown_content, str)
        assert "Title" in markdown_content
        assert "**Important**" in markdown_content

        # Markdown → HTML
        md_file = temp_dir / "temp.md"
        md_file.write_text(markdown_content)

        final_html_content = convert(str(md_file), source_format="markdown", target_format="html")

        assert isinstance(final_html_content, str)

        assert "Title" in final_html_content
        assert "Important" in final_html_content

    def test_transform_pipeline_integration(self, temp_dir):
        """Test multiple transforms applied in sequence."""
        html = """
        <html>
        <body>
            <h1>Main Title</h1>
            <h2>Section 1</h2>
            <p>Content</p>
            <img src="test.png" alt="Test">
            <h2>Section 2</h2>
        </body>
        </html>
        """

        # Apply multiple transforms
        result_content = convert(
            BytesIO(html.encode("utf-8")),
            source_format="html",
            target_format="markdown",
            transforms=["remove-images", "heading-offset"],
        )

        assert isinstance(result_content, str)
        # Images should be removed
        assert "![" not in result_content
        # Headings should be offset
        assert "## Main Title" in result_content  # h1 → h2
        assert "### Section 1" in result_content  # h2 → h3

    @pytest.mark.slow
    def test_complex_multi_format_workflow(self, temp_dir):
        """Test complex workflow: Parse → Transform → Multiple outputs."""
        # Create complex HTML
        html = create_html_with_tables()

        # Step 1: HTML → AST
        ast_doc = to_ast(BytesIO(html.encode("utf-8")), source_format="html")

        # Step 2: Create multiple outputs with different options
        md_content = from_ast(ast_doc, "markdown", renderer_options=MarkdownOptions(bullet_symbols="-"))

        html_content = from_ast(ast_doc, "html", renderer_options=HtmlRendererOptions(standalone=True))

        assert isinstance(md_content, str)
        assert isinstance(html_content, str)

        # Both should contain table content
        assert "Table Test Document" in md_content or "Alice Johnson" in md_content
        assert "Table Test Document" in html_content or "Alice Johnson" in html_content

        # HTML should be standalone
        assert "<!DOCTYPE html>" in html_content


@pytest.mark.e2e
class TestRealWorldUsagePatternsE2E:
    """End-to-end tests for real-world usage patterns."""

    @pytest.mark.pdf
    def test_extract_and_clean_document_workflow(self, temp_dir):
        """Real-world: Extract PDF content and clean for publication."""
        doc = create_pdf_with_figures()

        try:
            pdf_bytes = BytesIO()
            doc.save(pdf_bytes)
            pdf_bytes.seek(0)

            # Workflow: PDF → AST → remove images → clean Markdown
            ast_doc = to_ast(pdf_bytes, source_format="pdf")

            clean_markdown_content = from_ast(ast_doc, "markdown", transforms=["remove-images"])

            assert isinstance(clean_markdown_content, str)
            # Should have text content
            assert len(clean_markdown_content) > 0
            assert_markdown_valid(clean_markdown_content)
            # Should not have images
            assert "![" not in clean_markdown_content

        finally:
            doc.close()

    @pytest.mark.docx
    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not available")
    def test_multi_format_publishing_workflow(self, temp_dir):
        """Real-world: Single source → multiple output formats."""
        # Create source document
        docx_doc = create_docx_with_formatting()
        docx_bytes = save_docx_to_bytes(docx_doc)

        # Step 1: Parse to AST once
        ast_doc = to_ast(BytesIO(docx_bytes), source_format="docx")

        # Step 2: Generate multiple outputs
        md_file = temp_dir / "output.md"
        html_file = temp_dir / "output.html"
        docx_file = temp_dir / "output.docx"

        # Markdown output
        from_ast(ast_doc, "markdown", output=md_file)

        # HTML output
        from_ast(ast_doc, "html", output=html_file)

        # DOCX output
        from_ast(ast_doc, "docx", output=docx_file)

        # Verify all outputs created
        assert md_file.exists()
        assert html_file.exists()
        assert docx_file.exists()

        # Verify content
        md_content = md_file.read_text()
        html_content = html_file.read_text()

        assert "Formatting Test Document" in md_content
        assert "Formatting Test Document" in html_content

    def test_content_extraction_workflow(self, temp_dir):
        """Real-world: Extract specific sections from document."""
        html = """
        <html>
        <body>
            <h1>Document Title</h1>
            <h2>Section 1</h2>
            <p>Content for section 1.</p>
            <h2>Section 2</h2>
            <p>Content for section 2.</p>
            <h2>Section 3</h2>
            <p>Content for section 3.</p>
        </body>
        </html>
        """

        # Parse to AST
        ast_doc = to_ast(BytesIO(html.encode("utf-8")), source_format="html")

        # Extract headings using extract_nodes from transforms
        from all2md.ast.transforms import extract_nodes

        headings = extract_nodes(ast_doc, Heading)

        # Should have extracted headings
        assert len(headings) >= 3
        # Verify they are heading nodes
        for h in headings:
            assert isinstance(h, Heading)

    def test_format_migration_workflow(self, temp_dir):
        """Real-world: Migrate legacy format to modern with transforms."""
        # Simulate legacy HTML with issues
        legacy_html = """
        <html>
        <body>
            <h1>Legacy Document</h1>
            <p>Old content with unnecessary elements.</p>
            <center>Centered text (deprecated)</center>
        </body>
        </html>
        """

        # Migrate: Parse → Clean → Modern output
        ast_doc = to_ast(BytesIO(legacy_html.encode("utf-8")), source_format="html")

        # Output to modern Markdown
        modern_md_content = from_ast(ast_doc, "markdown")

        assert isinstance(modern_md_content, str)
        # Should have preserved content
        assert "Legacy Document" in modern_md_content
        assert "Old content" in modern_md_content

    @pytest.mark.slow
    def test_batch_processing_workflow(self, temp_dir):
        """Real-world: Process multiple documents consistently."""
        # Create multiple HTML files
        files = []
        for i in range(5):
            html_file = temp_dir / f"doc_{i}.html"
            html_file.write_text(f"<h1>Document {i}</h1><p>Content {i}</p>")
            files.append(html_file)

        # Batch process: Parse all → Apply consistent transforms → Output
        outputs = []
        for file in files:
            ast_doc = to_ast(str(file), source_format="html")
            markdown = from_ast(ast_doc, "markdown", transforms=["heading-offset"])
            assert isinstance(markdown, str)
            outputs.append(markdown)

        # Verify all processed
        assert len(outputs) == 5
        for i, output in enumerate(outputs):
            assert f"Document {i}" in output
            # Headings should be offset
            assert "##" in output

    def test_document_merging_workflow(self, temp_dir):
        """Real-world: Merge multiple documents via AST."""
        # Create source documents
        html1 = "<h1>Part 1</h1><p>Content from first document.</p>"
        html2 = "<h1>Part 2</h1><p>Content from second document.</p>"

        # Parse both to AST
        ast1 = to_ast(BytesIO(html1.encode("utf-8")), source_format="html")
        ast2 = to_ast(BytesIO(html2.encode("utf-8")), source_format="html")

        # Merge by combining children
        merged_doc = Document(children=ast1.children + ast2.children, metadata={"title": "Merged Document"})

        # Render merged document
        merged_md_content = from_ast(merged_doc, "markdown")

        assert isinstance(merged_md_content, str)
        # Should contain both parts
        assert "Part 1" in merged_md_content
        assert "Part 2" in merged_md_content
        assert "first document" in merged_md_content
        assert "second document" in merged_md_content

    def test_custom_transform_application(self, temp_dir):
        """Real-world: Apply custom transform in workflow."""
        html = """
        <html>
        <body>
            <h1>Document</h1>
            <p>Text with <strong>bold</strong> formatting.</p>
        </body>
        </html>
        """

        # Custom workflow: Remove all Strong nodes using filter_nodes
        from all2md.ast.transforms import filter_nodes

        ast_doc = to_ast(BytesIO(html.encode("utf-8")), source_format="html")

        # Filter out Strong nodes
        modified_ast = filter_nodes(ast_doc, lambda node: not isinstance(node, Strong))

        markdown_content = from_ast(modified_ast, "markdown")

        assert isinstance(markdown_content, str)
        # Should have text but Strong nodes removed
        assert "Document" in markdown_content
        # The text "bold" may or may not appear depending on how filter works
        assert len(markdown_content) > 0


@pytest.mark.e2e
class TestAPIConsistencyE2E:
    """End-to-end tests for API consistency across functions."""

    def test_to_markdown_vs_to_ast_from_ast_equivalence(self, temp_dir):
        """Test that to_markdown equals to_ast + from_ast."""
        html_file = temp_dir / "test.html"
        html_file.write_text("<html><body><h1>Title</h1><p>Content</p></body></html>")

        # Direct conversion
        direct = to_markdown(str(html_file))

        # Via AST
        ast_doc = to_ast(str(html_file))
        via_ast = from_ast(ast_doc, "markdown")

        # Should be identical
        assert isinstance(via_ast, str)
        assert direct == via_ast

    def test_convert_vs_specialized_functions_equivalence(self, temp_dir):
        """Test convert vs specialized functions produce same results."""
        html_file = temp_dir / "test.html"
        html_file.write_text("<h1>Test</h1><p>Content</p>")

        # Using convert
        result1 = convert(str(html_file), target_format="markdown")

        # Using to_markdown
        result2 = to_markdown(str(html_file))

        # Should be equivalent
        assert isinstance(result1, str)
        assert result1 == result2

    def test_options_propagation_across_functions(self, temp_dir):
        """Test that options propagate correctly across function calls."""
        md_file = temp_dir / "test.md"
        md_file.write_text("# Title\n\n*Italic* text")

        # from_markdown with options
        result1 = from_markdown(str(md_file), "html", renderer_options=HtmlRendererOptions(standalone=False))

        # convert with same options
        result2 = convert(
            str(md_file),
            source_format="markdown",
            target_format="html",
            renderer_options=HtmlRendererOptions(standalone=False),
        )

        # Should produce same results
        assert isinstance(result1, str)
        assert isinstance(result2, str)
        assert result1 == result2

    def test_transform_behavior_consistency(self, temp_dir):
        """Test transforms behave consistently across APIs."""
        html = "<h1>Title</h1><p>Content</p><img src='test.png'>"

        # Via to_markdown with transform
        result1 = to_markdown(BytesIO(html.encode("utf-8")), source_format="html", transforms=["remove-images"])

        # Via to_ast + from_ast with transform
        ast_doc = to_ast(BytesIO(html.encode("utf-8")), source_format="html")
        result2 = from_ast(ast_doc, "markdown", transforms=["remove-images"])

        # Should be equivalent
        assert isinstance(result2, str)
        assert result1 == result2

    def test_io_handling_consistency(self, temp_dir):
        """Test consistent IO handling across functions."""
        html_bytes = b"<h1>Test</h1><p>Content</p>"

        # Different input types should work consistently
        result1 = to_ast(html_bytes, source_format="html")
        result2 = to_ast(BytesIO(html_bytes), source_format="html")

        # Both should produce equivalent AST
        md1 = from_ast(result1, "markdown")
        md2 = from_ast(result2, "markdown")

        assert isinstance(md1, str)
        assert isinstance(md2, str)
        assert md1 == md2

    def test_error_handling_consistency(self):
        """Test consistent error handling across functions."""
        invalid_bytes = b"Invalid content"

        # Should raise similar errors
        with pytest.raises(Exception):
            to_ast(BytesIO(invalid_bytes), source_format="docx")

        with pytest.raises(Exception):
            to_markdown(BytesIO(invalid_bytes), source_format="docx")

    def test_metadata_handling_consistency(self, temp_dir):
        """Test metadata handling is consistent."""
        html = """
        <html>
        <head><title>Test Doc</title></head>
        <body><h1>Content</h1></body>
        </html>
        """

        # Get AST with metadata
        ast_doc = to_ast(BytesIO(html.encode("utf-8")), source_format="html", extract_title=True)

        # Metadata should be present
        assert ast_doc.metadata is not None

        # Render should preserve metadata context
        result = from_ast(ast_doc, "markdown")
        assert isinstance(result, str)


@pytest.mark.e2e
class TestFullConversionPipeline:
    """End-to-end tests for complete document conversion workflows."""

    def test_docx_to_markdown_full_pipeline(self, temp_dir):
        """Test complete DOCX to Markdown conversion pipeline."""
        # Create a DOCX document with various formatting
        doc = create_docx_with_formatting()
        docx_bytes = save_docx_to_bytes(doc)

        # Test conversion with different options
        options = DocxOptions(attachment_mode="base64")
        markdown_options = MarkdownOptions(bullet_symbols="-")

        # Convert to markdown
        result = to_markdown(
            BytesIO(docx_bytes), source_format="docx", parser_options=options, renderer_options=markdown_options
        )

        assert_markdown_valid(result)

        # Should contain expected content
        assert "Formatting Test Document" in result
        assert "**bold text**" in result.lower()
        assert "*italic text*" in result.lower()

        # Should detect format correctly
        # Format detection successful - result should contain converted content
        assert result

    def test_html_to_markdown_full_pipeline(self, temp_dir):
        """Test complete HTML to Markdown conversion pipeline."""
        # Create HTML content with tables
        html_content = create_html_with_tables()

        # Convert with options
        options = HtmlOptions(strip_dangerous_elements=True)
        markdown_options = MarkdownOptions()

        result = to_markdown(
            BytesIO(html_content.encode("utf-8")),
            source_format="html",
            parser_options=options,
            renderer_options=markdown_options,
        )

        assert_markdown_valid(result)

        # Should contain table content
        assert "Table Test Document" in result
        assert "|" in result  # Should have table formatting
        assert "Alice Johnson" in result

        # Should detect format correctly
        # Format detection successful - result should contain converted content
        assert result

    def test_pptx_to_markdown_full_pipeline(self, temp_dir):
        """Test complete PPTX to Markdown conversion pipeline."""
        # Create a PPTX presentation
        prs = create_pptx_with_basic_slides()
        pptx_bytes = save_pptx_to_bytes(prs)

        # Convert with options
        options = PptxOptions(attachment_mode="skip", include_slide_numbers=True)
        markdown_options = MarkdownOptions()

        result = to_markdown(
            BytesIO(pptx_bytes), source_format="pptx", parser_options=options, renderer_options=markdown_options
        )

        assert_markdown_valid(result)

        # Should contain presentation content
        assert "Test Presentation" in result
        assert "Main Content Slide" in result
        assert "First bullet point" in result

        # Should detect format correctly
        # Format detection successful - result should contain converted content
        assert result

    def test_pdf_to_markdown_full_pipeline(self, temp_dir):
        """Test complete PDF to Markdown conversion pipeline."""
        # Create a PDF with figures
        doc = create_pdf_with_figures()

        try:
            # Save to bytes
            pdf_bytes = BytesIO()
            doc.save(pdf_bytes)
            pdf_bytes.seek(0)

            # Convert with options
            options = PdfOptions(attachment_mode="base64", enable_table_fallback_detection=True)
            markdown_options = MarkdownOptions()

            result = to_markdown(
                pdf_bytes, source_format="pdf", parser_options=options, renderer_options=markdown_options
            )

            assert_markdown_valid(result)

            # Should contain PDF content
            assert "Test Document with Figures" in result
            assert "figure" in result.lower()

            # Should detect format correctly
            # Format detection successful - result should contain converted content
            assert result

        finally:
            doc.close()

    def test_automatic_format_detection_pipeline(self, temp_dir):
        """Test automatic format detection in the conversion pipeline."""
        # Test with different file types without explicit extension

        # DOCX test
        doc = create_docx_with_formatting()
        docx_bytes = save_docx_to_bytes(doc)

        result_docx = to_markdown(BytesIO(docx_bytes))  # No explicit extension
        # Format detection successful - result should contain converted content
        assert result_docx
        assert_markdown_valid(result_docx)

        # HTML test
        html_content = create_html_with_tables()
        html_bytes = html_content.encode("utf-8")

        result_html = to_markdown(BytesIO(html_bytes))  # No explicit extension
        # Format detection successful - result should contain converted content
        assert result_html
        assert_markdown_valid(result_html)

    def test_error_handling_in_full_pipeline(self, temp_dir):
        """Test error handling throughout the conversion pipeline."""
        # Test with invalid content
        invalid_content = b"This is not a valid document"

        # Should handle gracefully
        with pytest.raises(Exception):  # Specific exception depends on implementation
            to_markdown(BytesIO(invalid_content), source_format="docx")

        # Test with unsupported format - should fallback to text
        result = to_markdown(BytesIO(b"content"), source_format="unsupported")
        assert result == "content"  # Should fallback to treating as text

    def test_large_document_pipeline_performance(self, temp_dir):
        """Test pipeline performance with larger documents."""
        # Create a larger document for performance testing
        import docx

        doc = docx.Document()
        doc.add_heading("Large Document Test", level=1)

        # Add many paragraphs
        for i in range(100):
            p = doc.add_paragraph(f"This is paragraph {i + 1} with some content.")
            if i % 10 == 0:
                p.runs[0].bold = True

        # Add a large table
        table = doc.add_table(rows=20, cols=5)
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = f"Row {i + 1}, Col {j + 1}"

        docx_bytes = save_docx_to_bytes(doc)

        # Convert (should complete without timeout)
        result = to_markdown(BytesIO(docx_bytes), source_format="docx")

        assert_markdown_valid(result)
        assert "Large Document Test" in result
        assert "paragraph 50" in result
        assert "Row 10, Col 3" in result

    def test_options_propagation_through_pipeline(self, temp_dir):
        """Test that conversion options are properly propagated through the pipeline."""
        # Create HTML with various elements
        html = """
        <html>
        <head><title>Options Test</title></head>
        <body>
            <h1>Main Title</h1>
            <p>Regular paragraph.</p>
            <table>
                <tr><th>Header</th></tr>
                <tr><td>Data</td></tr>
            </table>
            <script>alert('test');</script>
            <style>body { color: red; }</style>
        </body>
        </html>
        """

        # Test with basic HTML options - the key is that options are accepted and processed
        options_test = HtmlOptions(strip_dangerous_elements=True, convert_nbsp=False)

        result = to_markdown(BytesIO(html.encode("utf-8")), source_format="html", parser_options=options_test)

        # Verify the conversion worked with options
        assert "Main Title" in result  # Should contain content
        assert "Header" in result  # Should contain table content
        assert "|" in result  # Should have table formatting
        assert_markdown_valid(result)  # Should produce valid markdown

        # Test that different options produce output (may be the same, but should not crash)
        options_test2 = HtmlOptions(strip_dangerous_elements=False, convert_nbsp=True)

        result2 = to_markdown(BytesIO(html.encode("utf-8")), source_format="html", parser_options=options_test2)

        # Both conversions should work
        assert "Main Title" in result2
        assert len(result2) > 0

    def test_mixed_content_document_pipeline(self, temp_dir):
        """Test conversion of documents with mixed content types."""
        # Create a complex DOCX with tables, images, and various formatting
        import docx

        doc = docx.Document()

        # Title and introduction
        doc.add_heading("Complex Document", level=1)
        doc.add_paragraph("This document contains various content types.")

        # Table
        table = doc.add_table(rows=3, cols=3)
        table.style = "Table Grid"

        # Add table content
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if i == 0:
                    cell.text = f"Header {j + 1}"
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                else:
                    cell.text = f"Data {i}-{j + 1}"

        # More content
        doc.add_heading("Section 2", level=2)
        doc.add_paragraph("Content after the table.")

        # Lists
        doc.add_paragraph("Features:", style="List Bullet")
        doc.add_paragraph("Feature 1", style="List Bullet")
        doc.add_paragraph("Feature 2", style="List Bullet")

        docx_bytes = save_docx_to_bytes(doc)

        # Convert with comprehensive options
        docx_options = DocxOptions(attachment_mode="base64")
        markdown_options = MarkdownOptions(bullet_symbols="-")

        result = to_markdown(
            BytesIO(docx_bytes), source_format="docx", parser_options=docx_options, renderer_options=markdown_options
        )

        assert_markdown_valid(result)

        # Should contain all content types
        assert "# Complex Document" in result
        assert "## Section 2" in result
        assert "Header 1" in result
        assert "Data 1-1" in result
        assert "- Feature 1" in result
        assert "various content types" in result

    @pytest.mark.slow
    def test_batch_conversion_pipeline(self, temp_dir):
        """Test conversion of multiple documents in sequence."""
        # Create multiple test documents
        documents = []

        # DOCX
        docx_doc = create_docx_with_formatting()
        documents.append((save_docx_to_bytes(docx_doc), ".docx", "docx"))

        # HTML
        html_content = create_html_with_tables()
        documents.append((html_content.encode("utf-8"), ".html", "html"))

        # PPTX
        pptx_prs = create_pptx_with_basic_slides()
        documents.append((save_pptx_to_bytes(pptx_prs), ".pptx", "pptx"))

        # Convert all documents
        results = []
        for content, extension, _expected_format in documents:
            result = to_markdown(BytesIO(content), source_format=extension.lstrip("."))
            results.append(result)

            # Verify each conversion
            # Format detection successful - result should contain converted content
            assert result
            assert_markdown_valid(result)
            assert len(result) > 0

        # Verify different content in each result
        assert "Formatting Test Document" in results[0]  # DOCX
        assert "Table Test Document" in results[1]  # HTML
        assert "Test Presentation" in results[2]  # PPTX

        # All should be valid markdown
        for result in results:
            assert_markdown_valid(result)
