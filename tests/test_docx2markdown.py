import base64

import docx

from all2md import docx2markdown as md


class FakeIndent:
    def __init__(self, pt):
        self.pt = pt


class FakeFormat:
    def __init__(self, left_indent=None):
        self.left_indent = left_indent


class FakeStyle:
    def __init__(self, name):
        self.name = name


class DummyParagraph:
    def __init__(self, style_name=None, left_indent=None, text=""):
        self.style = FakeStyle(style_name)
        self.paragraph_format = FakeFormat(left_indent=FakeIndent(left_indent) if left_indent is not None else None)
        self.text = text

    def iter_inner_content(self):
        return iter([])


def test_detect_list_level_no_style():
    para = DummyParagraph(style_name=None)
    assert md._detect_list_level(para) == (None, 0)


def test_detect_list_level_built_in_bullet():
    para = DummyParagraph(style_name="List Bullet 2")
    assert md._detect_list_level(para) == ("bullet", 2)


def test_detect_list_level_built_in_number_default():
    para = DummyParagraph(style_name="List Number")
    assert md._detect_list_level(para) == ("number", 1)


def test_detect_list_level_indentation_number():
    para = DummyParagraph(style_name="Normal", left_indent=72, text="1) item")
    assert md._detect_list_level(para) == ("number", 2)


def test_detect_list_level_indentation_bullet():
    para = DummyParagraph(style_name="Normal", left_indent=36, text="item")
    assert md._detect_list_level(para) == ("bullet", 1)


def test_detect_list_level_no_indent():
    para = DummyParagraph(style_name="Normal", left_indent=None, text="item")
    assert md._detect_list_level(para) == (None, 0)


def test_format_list_marker():
    assert md._format_list_marker("bullet") == "* "
    assert md._format_list_marker("number", 3) == "3. "


class FakeFont:
    def __init__(self, strike=False, subscript=False, superscript=False):
        self.strike = strike
        self.subscript = subscript
        self.superscript = superscript


class FakeRun:
    def __init__(
        self, text, bold=False, italic=False, underline=False, strike=False, subscript=False, superscript=False
    ):
        self.text = text
        self.bold = bold
        self.italic = italic
        self.underline = underline
        self.font = FakeFont(strike, subscript, superscript)


def test_get_run_formatting_key():
    run = FakeRun("test", bold=True, italic=False, underline=True, strike=True, subscript=True, superscript=False)
    key = md._get_run_formatting_key(run)
    assert key == (True, False, True, True, True, False)


def test_process_hyperlink(monkeypatch):
    inner = FakeRun("linktext")

    class FakeHyperlink:
        def __init__(self, url, run):
            self.url = url
            self.runs = [run]

    monkeypatch.setattr(md, "Hyperlink", FakeHyperlink)
    link = FakeHyperlink("http://example.com", inner)
    url, run = md._process_hyperlink(link)
    assert url == "http://example.com"
    assert run is inner
    url2, run2 = md._process_hyperlink(inner)
    assert url2 is None and run2 is inner


class RunParagraph:
    def __init__(self, runs):
        self._runs = runs

    def iter_inner_content(self):
        return iter(self._runs)


def test_process_paragraph_runs_simple():
    runs = [FakeRun("Hello "), FakeRun("World")]
    para = RunParagraph(runs)
    assert md._process_paragraph_runs(para) == "Hello World"


def test_process_paragraph_runs_bold():
    runs = [FakeRun("Bold", bold=True)]
    para = RunParagraph(runs)
    assert md._process_paragraph_runs(para) == "**Bold**"


def test_process_paragraph_runs_italic():
    runs = [FakeRun("Ital", italic=True)]
    para = RunParagraph(runs)
    assert md._process_paragraph_runs(para) == "*Ital*"


def test_process_paragraph_runs_underline():
    runs = [FakeRun("Under", underline=True)]
    para = RunParagraph(runs)
    assert md._process_paragraph_runs(para) == "__Under__"


def test_process_paragraph_runs_strike():
    runs = [FakeRun("Strike", strike=True)]
    para = RunParagraph(runs)
    assert md._process_paragraph_runs(para) == "~~Strike~~"


def test_process_paragraph_runs_subscript():
    runs = [FakeRun("Sub", subscript=True)]
    para = RunParagraph(runs)
    assert md._process_paragraph_runs(para) == "~Sub~"


def test_process_paragraph_runs_superscript():
    runs = [FakeRun("Sup", superscript=True)]
    para = RunParagraph(runs)
    assert md._process_paragraph_runs(para) == "^Sup^"


def test_process_paragraph_runs_multiple_formatting():
    runs = [FakeRun("Test", bold=True, italic=True)]
    para = RunParagraph(runs)
    assert md._process_paragraph_runs(para) == "***Test***"


def test_process_paragraph_runs_hyperlink(monkeypatch):
    inner = FakeRun("LinkUpdate")

    class FakeHyperlink:
        def __init__(self):
            self.url = "http://link"
            self.runs = [inner]

    monkeypatch.setattr(md, "Hyperlink", FakeHyperlink)
    runs = [FakeHyperlink()]
    para = RunParagraph(runs)
    assert md._process_paragraph_runs(para) == "[LinkUpdate](http://link)"


def test_process_paragraph_runs_whitespace_preserved():
    runs = [FakeRun("  hello  ")]
    para = RunParagraph(runs)
    assert md._process_paragraph_runs(para) == "  hello  "


class FakeCell:
    def __init__(self, texts):
        if isinstance(texts, str):
            texts_list = [texts]
        else:
            texts_list = texts
        self.paragraphs = [RunParagraph([FakeRun(text)]) for text in texts_list]


class FakeRow:
    def __init__(self, cells_texts):
        self.cells = [FakeCell(texts) for texts in cells_texts]


class FakeTable:
    def __init__(self, rows_texts):
        self.rows = [FakeRow(row) for row in rows_texts]


def test_convert_table_to_markdown():
    table = FakeTable([["H1", "H2"], ["c1", "c2"]])
    md_table = md._convert_table_to_markdown(table)
    expected = "| H1 | H2 |\n| --- | --- |\n| c1 | c2 |"
    assert md_table == expected


def test_docx_to_markdown_basic(tmp_path):
    doc = docx.Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Paragraph text")
    doc.add_paragraph("Item 1", style="List Number")
    doc.add_paragraph("Item 2", style="List Number")
    file = tmp_path / "test.docx"
    doc.save(str(file))
    md_text = md.docx_to_markdown(str(file))
    assert "# Title" in md_text
    assert "Paragraph text" in md_text
    assert "1. Item 1" in md_text
    assert "2. Item 2" in md_text


def test_docx_to_markdown_table(tmp_path):
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "H1"
    table.rows[0].cells[1].text = "H2"
    table.rows[1].cells[0].text = "c1"
    table.rows[1].cells[1].text = "c2"
    file = tmp_path / "test2.docx"
    doc.save(str(file))
    md_text = md.docx_to_markdown(str(file))
    assert "| H1 | H2 |" in md_text
    assert "| c1 | c2 |" in md_text


def test_docx_to_markdown_images(tmp_path, monkeypatch):
    data = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/w8AAn8B9FpQHLwAAAAASUVORK5CYII="
    )
    img_file = tmp_path / "img.png"
    img_file.write_bytes(data)
    doc = docx.Document()
    doc.add_picture(str(img_file))
    from all2md.options import DocxOptions

    options1 = DocxOptions(attachment_mode="alt_text")
    md_text = md.docx_to_markdown(doc, options=options1)
    assert "![image]" == md_text
    from all2md import _attachment_utils

    monkeypatch.setattr(_attachment_utils, "extract_docx_image_data", lambda parent, rid: b"fake_image_bytes")
    options2 = DocxOptions(attachment_mode="base64")
    md_text2 = md.docx_to_markdown(doc, options=options2)
    assert "![image](data:image/png;base64," in md_text2
