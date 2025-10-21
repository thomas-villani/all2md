#  Copyright (c) 2025 Tom Villani, Ph.D.
#
# tests/integration/renderers/test_docx_conversion.py
"""Integration tests for DOCX renderer.

Tests cover:
- End-to-end DOCX rendering workflows
- Custom styles and fonts
- Metadata handling
- Complete document conversion

"""

import pytest

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from all2md.ast import (
    BlockQuote,
    CodeBlock,
    Document,
    Heading,
    Image,
    Link,
    List,
    ListItem,
    Paragraph,
    Strong,
    Table,
    TableCell,
    TableRow,
    Text,
)
from all2md.options import DocxRendererOptions
from all2md.options.common import NetworkFetchOptions

if DOCX_AVAILABLE:
    from all2md.renderers.docx import DocxRenderer


def create_sample_document():
    """Create a sample AST document for testing.

    Returns
    -------
    Document
        A sample document with various elements for testing.

    """
    return Document(
        metadata={"title": "Sample Document", "author": "Test Author"},
        children=[
            Heading(level=1, content=[Text(content="Document Title")]),
            Paragraph(
                content=[
                    Text(content="This is a paragraph with "),
                    Strong(content=[Text(content="bold text")]),
                    Text(content=" and a "),
                    Link(url="https://example.com", content=[Text(content="link")]),
                    Text(content="."),
                ]
            ),
            Heading(level=2, content=[Text(content="Lists")]),
            List(
                ordered=False,
                items=[
                    ListItem(children=[Paragraph(content=[Text(content="First item")])]),
                    ListItem(children=[Paragraph(content=[Text(content="Second item")])]),
                    ListItem(children=[Paragraph(content=[Text(content="Third item")])]),
                ],
            ),
            Heading(level=2, content=[Text(content="Code Example")]),
            CodeBlock(content='def hello():\n    print("Hello, world!")', language="python"),
            Heading(level=2, content=[Text(content="Table")]),
            Table(
                header=TableRow(
                    cells=[TableCell(content=[Text(content="Name")]), TableCell(content=[Text(content="Value")])]
                ),
                rows=[
                    TableRow(
                        cells=[TableCell(content=[Text(content="Alpha")]), TableCell(content=[Text(content="1")])]
                    ),
                    TableRow(cells=[TableCell(content=[Text(content="Beta")]), TableCell(content=[Text(content="2")])]),
                ],
            ),
            Heading(level=2, content=[Text(content="Quote")]),
            BlockQuote(children=[Paragraph(content=[Text(content="This is a blockquote.")])]),
        ],
    )


@pytest.mark.integration
@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
@pytest.mark.docx
class TestDocxRendering:
    """Integration tests for DOCX rendering."""

    def test_full_document_to_docx(self, tmp_path):
        """Test rendering complete document to DOCX."""
        doc = create_sample_document()
        renderer = DocxRenderer()
        output_file = tmp_path / "full_document.docx"
        renderer.render(doc, output_file)

        assert output_file.exists()

        # Verify DOCX content
        docx_doc = DocxDocument(str(output_file))

        # Check for tables
        assert len(docx_doc.tables) >= 1

        # Check for content in paragraphs
        all_text = " ".join(p.text for p in docx_doc.paragraphs)
        assert "Document Title" in all_text
        assert "bold text" in all_text
        assert "First item" in all_text

    def test_docx_with_custom_styles(self, tmp_path):
        """Test DOCX rendering with custom styles."""
        doc = create_sample_document()
        options = DocxRendererOptions(default_font="Arial", default_font_size=12, code_font="Consolas")
        renderer = DocxRenderer(options)
        output_file = tmp_path / "custom_styles.docx"
        renderer.render(doc, output_file)

        assert output_file.exists()

    def test_docx_metadata(self, tmp_path):
        """Test DOCX metadata handling."""
        doc = create_sample_document()
        renderer = DocxRenderer()
        output_file = tmp_path / "with_metadata.docx"
        renderer.render(doc, output_file)

        docx_doc = DocxDocument(str(output_file))
        assert docx_doc.core_properties.title == "Sample Document"
        assert docx_doc.core_properties.author == "Test Author"


@pytest.mark.integration
@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
@pytest.mark.docx
class TestDocxRemoteImageFetching:
    """Integration tests for DOCX remote image fetching."""

    def test_remote_image_disabled_by_default(self, tmp_path):
        """Test that remote images are skipped by default (secure-by-default)."""
        doc = Document(
            children=[
                Paragraph(content=[Text(content="Image test:")]),
                Paragraph(content=[Image(url="https://example.com/test-image.png", alt_text="Remote image")]),
            ]
        )
        # Default options should have allow_remote_fetch=False
        renderer = DocxRenderer()
        output_file = tmp_path / "no_remote.docx"
        renderer.render(doc, output_file)

        # Should complete without error, but image is skipped
        assert output_file.exists()
        assert output_file.stat().st_size > 0

    def test_remote_image_explicit_disabled(self, tmp_path):
        """Test that remote images are skipped when explicitly disabled."""
        doc = Document(
            children=[Paragraph(content=[Image(url="https://httpbin.org/image/png", alt_text="Remote image")])]
        )
        options = DocxRendererOptions(network=NetworkFetchOptions(allow_remote_fetch=False))
        renderer = DocxRenderer(options)
        output_file = tmp_path / "remote_disabled.docx"
        renderer.render(doc, output_file)

        # Should complete without error, but image is skipped
        assert output_file.exists()
        assert output_file.stat().st_size > 0

    def test_local_image_not_affected(self, tmp_path):
        """Test that local images work regardless of network settings."""
        doc = Document(children=[Paragraph(content=[Image(url="/path/to/local/image.png", alt_text="Local image")])])
        # Even with remote fetching disabled, local images should be attempted
        options = DocxRendererOptions(
            network=NetworkFetchOptions(allow_remote_fetch=False), fail_on_resource_errors=False
        )
        renderer = DocxRenderer(options)
        output_file = tmp_path / "local_image.docx"
        renderer.render(doc, output_file)

        # Should complete (though image won't be found)
        assert output_file.exists()

    def test_base64_image_not_affected(self, tmp_path):
        """Test that base64 images work regardless of network settings."""
        # Small 1x1 red pixel PNG
        base64_image = (
            "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z"
            "8DwHwAFBQIAX8jx0gAAAABJRU5ErkJggg=="
        )
        doc = Document(children=[Paragraph(content=[Image(url=base64_image, alt_text="Base64 image")])])
        # Base64 images should work even with remote fetching disabled
        options = DocxRendererOptions(network=NetworkFetchOptions(allow_remote_fetch=False))
        renderer = DocxRenderer(options)
        output_file = tmp_path / "base64_image.docx"
        renderer.render(doc, output_file)

        assert output_file.exists()
        assert output_file.stat().st_size > 0

    def test_network_options_respected(self, tmp_path):
        """Test that network security options are properly configured."""
        options = DocxRendererOptions(
            network=NetworkFetchOptions(
                allow_remote_fetch=True, require_https=True, network_timeout=10.0, allowed_hosts=["example.com"]
            )
        )
        renderer = DocxRenderer(options)

        # Verify options are set correctly
        assert renderer.options.network.allow_remote_fetch is True
        assert renderer.options.network.require_https is True
        assert renderer.options.network.network_timeout == 10.0
        assert renderer.options.network.allowed_hosts == ["example.com"]

    def test_mixed_image_sources(self, tmp_path):
        """Test document with mixed image sources (local, remote, base64)."""
        base64_image = (
            "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z"
            "8DwHwAFBQIAX8jx0gAAAABJRU5ErkJggg=="
        )
        doc = Document(
            children=[
                Heading(level=1, content=[Text(content="Mixed Images Test")]),
                Paragraph(content=[Text(content="Base64 image:")]),
                Paragraph(content=[Image(url=base64_image, alt_text="Base64")]),
                Paragraph(content=[Text(content="Remote image (skipped):")]),
                Paragraph(content=[Image(url="https://example.com/image.png", alt_text="Remote")]),
                Paragraph(content=[Text(content="Local image (attempted):")]),
                Paragraph(content=[Image(url="local.png", alt_text="Local")]),
            ]
        )
        options = DocxRendererOptions(
            network=NetworkFetchOptions(allow_remote_fetch=False), fail_on_resource_errors=False
        )
        renderer = DocxRenderer(options)
        output_file = tmp_path / "mixed_images.docx"
        renderer.render(doc, output_file)

        # Should complete successfully
        assert output_file.exists()
        assert output_file.stat().st_size > 0
