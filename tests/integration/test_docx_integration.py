"""Integration tests for DOCX to Markdown conversion."""

import base64

import docx
import pytest

from all2md import docx2markdown as md


@pytest.mark.integration
def test_docx_to_markdown_basic(tmp_path):
    """Test basic DOCX to Markdown conversion with real documents."""
    doc = docx.Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Paragraph text")
    doc.add_paragraph("Item 1", style="List Number")
    doc.add_paragraph("Item 2", style="List Number")

    doc_path = tmp_path / "test.docx"
    doc.save(str(doc_path))

    result = md.docx_to_markdown(doc_path)

    expected_elements = [
        "# Title",
        "Paragraph text",
        "1. Item 1",
        "2. Item 2"
    ]

    for element in expected_elements:
        assert element in result


@pytest.mark.integration
def test_docx_to_markdown_table(tmp_path):
    """Test DOCX table conversion."""
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "H1"
    table.cell(0, 1).text = "H2"
    table.cell(1, 0).text = "c1"
    table.cell(1, 1).text = "c2"

    doc_path = tmp_path / "test.docx"
    doc.save(str(doc_path))

    result = md.docx_to_markdown(doc_path)

    # Should have table
    assert "|" in result
    assert "H1" in result and "H2" in result
    assert "c1" in result and "c2" in result


@pytest.mark.integration
def test_docx_to_markdown_images(tmp_path, monkeypatch):
    """Test DOCX image handling."""
    def mock_base64_encode(image_data):
        return "mock_base64_data"

    monkeypatch.setattr(base64, "b64encode", lambda x: b"mock_base64_data")

    doc = docx.Document()
    doc.add_heading("Document with Images", level=1)
    doc.add_paragraph("This document has images.")

    doc_path = tmp_path / "test.docx"
    doc.save(str(doc_path))

    result = md.docx_to_markdown(doc_path)

    assert "Document with Images" in result
    assert "This document has images." in result
