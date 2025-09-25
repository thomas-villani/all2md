"""Advanced tests for DOCX table handling edge cases."""


import docx
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

from all2md.docx2markdown import docx_to_markdown
from tests.utils import assert_markdown_valid, cleanup_test_dir, create_test_temp_dir


class TestDocxTablesAdvanced:
    """Test complex table scenarios in DOCX documents."""

    def setup_method(self):
        """Set up test environment."""
        self.temp_dir = create_test_temp_dir()

    def teardown_method(self):
        """Clean up test environment."""
        cleanup_test_dir(self.temp_dir)

    def test_merged_cells_simulation(self):
        """Test tables with merged cells (simulated due to python-docx limitations)."""
        doc = docx.Document()

        doc.add_heading("Merged Cells Test", level=1)

        # Create table with merged-like structure
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'

        # Header row
        table.rows[0].cells[0].text = "Category"
        table.rows[0].cells[1].text = "Subcategory"
        table.rows[0].cells[2].text = "Value"

        # Simulate merged cells by leaving some empty and filling others
        table.rows[1].cells[0].text = "Group A"
        table.rows[1].cells[1].text = "Item 1"
        table.rows[1].cells[2].text = "100"

        table.rows[2].cells[0].text = ""  # Would be merged with above in real scenario
        table.rows[2].cells[1].text = "Item 2"
        table.rows[2].cells[2].text = "200"

        table.rows[3].cells[0].text = "Group B"
        table.rows[3].cells[1].text = "Item 3"
        table.rows[3].cells[2].text = "300"

        temp_file = self.temp_dir / "merged_cells.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should generate valid table markdown
        assert "| Category | Subcategory | Value |" in markdown
        assert "| Group A | Item 1 | 100 |" in markdown
        assert "| Group B | Item 3 | 300 |" in markdown

        # Check for separator row
        assert "| --- | --- | --- |" in markdown or "|:---:|:---:|:---:|" in markdown

    def test_alignment_variations(self):
        """Test tables with different cell alignment settings."""
        doc = docx.Document()

        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'

        # Set up headers
        headers = ["Left", "Center", "Right"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            # Set paragraph alignment
            if i == 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif i == 1:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Set vertical alignment
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Add data rows with different alignments
        alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT]
        for row_idx in range(1, 3):
            for col_idx in range(3):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = f"Data {row_idx},{col_idx}"
                cell.paragraphs[0].alignment = alignments[col_idx]

        temp_file = self.temp_dir / "aligned_table.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should contain table structure
        assert "| Left | Center | Right |" in markdown
        assert "Data 1,0" in markdown
        assert "Data 2,2" in markdown

    def test_no_header_rows(self):
        """Test tables without header rows."""
        doc = docx.Document()

        doc.add_paragraph("Table without headers:")

        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'

        # All rows are data rows, no distinct headers
        for row_idx in range(3):
            for col_idx in range(2):
                table.rows[row_idx].cells[col_idx].text = f"Data{row_idx}{col_idx}"

        temp_file = self.temp_dir / "no_headers.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should still generate valid table
        lines = markdown.split('\n')
        table_lines = [line for line in lines if '|' in line]
        assert len(table_lines) >= 3  # At least data rows

        # First row should be treated as header by default
        assert "Data00" in table_lines[0]
        assert "Data01" in table_lines[0]

    def test_thick_borders_and_styling(self):
        """Test tables with thick borders and various styling."""
        doc = docx.Document()

        table = doc.add_table(rows=2, cols=2)
        table.style = 'Medium Grid 1 Accent 1'  # Use a more styled table

        table.rows[0].cells[0].text = "Styled Header 1"
        table.rows[0].cells[1].text = "Styled Header 2"
        table.rows[1].cells[0].text = "Styled Data 1"
        table.rows[1].cells[1].text = "Styled Data 2"

        temp_file = self.temp_dir / "styled_table.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Styling should not break basic table conversion
        assert "Styled Header 1" in markdown
        assert "Styled Data 1" in markdown
        assert "|" in markdown  # Should still have table structure

    def test_nested_tables_simulation(self):
        """Test nested tables (simulated due to DOCX limitations)."""
        doc = docx.Document()

        # Outer table
        outer_table = doc.add_table(rows=2, cols=2)
        outer_table.style = 'Table Grid'

        outer_table.rows[0].cells[0].text = "Outer Cell 1"
        outer_table.rows[0].cells[1].text = "Outer Cell 2"

        # Simulate nested content with formatted text that might represent a table
        nested_cell = outer_table.rows[1].cells[0]
        nested_cell.text = "Inner Table:\nCol1 | Col2\nVal1 | Val2"

        outer_table.rows[1].cells[1].text = "Regular Cell"

        temp_file = self.temp_dir / "nested_tables.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle the nested content appropriately
        assert "Outer Cell 1" in markdown
        assert "Inner Table" in markdown
        assert "Col1" in markdown

    def test_empty_cells_and_whitespace(self):
        """Test tables with empty cells and whitespace-only cells."""
        doc = docx.Document()

        table = doc.add_table(rows=3, cols=3)

        # Mixed content: some empty, some with text, some with whitespace
        table.rows[0].cells[0].text = "Header 1"
        table.rows[0].cells[1].text = ""  # Empty
        table.rows[0].cells[2].text = "Header 3"

        table.rows[1].cells[0].text = "Data"
        table.rows[1].cells[1].text = "   "  # Whitespace only
        table.rows[1].cells[2].text = "More Data"

        table.rows[2].cells[0].text = ""  # Empty
        table.rows[2].cells[1].text = ""  # Empty
        table.rows[2].cells[2].text = "Only Data"

        temp_file = self.temp_dir / "empty_cells.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle empty cells gracefully
        assert "Header 1" in markdown
        assert "Header 3" in markdown
        assert "Only Data" in markdown

        # Check table structure is maintained
        lines = markdown.split('\n')
        table_lines = [line for line in lines if '|' in line]
        assert len(table_lines) >= 3  # Header, separator, data rows

    def test_cells_with_complex_content(self):
        """Test table cells containing complex content (lists, formatting, etc.)."""
        doc = docx.Document()

        table = doc.add_table(rows=2, cols=2)

        # Cell with formatted text
        cell1 = table.rows[0].cells[0]
        p1 = cell1.paragraphs[0]
        run1 = p1.add_run("Bold Text")
        run1.bold = True

        # Cell with line breaks
        table.rows[0].cells[1].text = "Multi\nLine\nContent"

        # Cell with mixed formatting
        cell2 = table.rows[1].cells[0]
        p2 = cell2.paragraphs[0]
        p2.add_run("Regular ")
        bold_run = p2.add_run("Bold")
        bold_run.bold = True
        p2.add_run(" and ")
        italic_run = p2.add_run("Italic")
        italic_run.italic = True

        # Cell with simulated list content
        table.rows[1].cells[1].text = "• Item 1\n• Item 2\n• Item 3"

        temp_file = self.temp_dir / "complex_cells.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should preserve formatting within cells
        assert "**Bold Text**" in markdown
        assert "Multi" in markdown
        assert "Line" in markdown
        assert "Regular **Bold** and *Italic*" in markdown
        assert "Item 1" in markdown

    def test_table_without_grid_style(self):
        """Test tables without grid styling (no visible borders)."""
        doc = docx.Document()

        table = doc.add_table(rows=2, cols=2)
        # Don't apply grid style - default has no borders

        table.rows[0].cells[0].text = "Invisible"
        table.rows[0].cells[1].text = "Borders"
        table.rows[1].cells[0].text = "Data 1"
        table.rows[1].cells[1].text = "Data 2"

        temp_file = self.temp_dir / "no_borders.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should still be recognized as a table
        assert "Invisible" in markdown
        assert "Borders" in markdown
        assert "|" in markdown  # Should still have markdown table format

    def test_very_wide_tables(self):
        """Test tables with many columns."""
        doc = docx.Document()

        cols = 8
        table = doc.add_table(rows=2, cols=cols)

        # Create wide table
        for col in range(cols):
            table.rows[0].cells[col].text = f"H{col+1}"
            table.rows[1].cells[col].text = f"D{col+1}"

        temp_file = self.temp_dir / "wide_table.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle wide tables
        assert "H1" in markdown and "H8" in markdown
        assert "D1" in markdown and "D8" in markdown

        # Count columns in generated markdown
        lines = markdown.split('\n')
        header_line = next((line for line in lines if "H1" in line), None)
        if header_line:
            pipe_count = header_line.count('|')
            assert pipe_count >= cols + 1  # cols + 1 for proper markdown table format

    def test_table_position_and_surrounding_content(self):
        """Test tables positioned between different types of content."""
        doc = docx.Document()

        doc.add_heading("Before Table", level=1)
        doc.add_paragraph("Paragraph before table.")

        # Table in the middle
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "A"
        table.rows[0].cells[1].text = "B"
        table.rows[1].cells[0].text = "1"
        table.rows[1].cells[1].text = "2"

        doc.add_paragraph("Paragraph after table.")

        # Another paragraph
        p = doc.add_paragraph("Another paragraph with ")
        run = p.add_run("formatting")
        run.bold = True

        # Second table
        table2 = doc.add_table(rows=1, cols=2)
        table2.rows[0].cells[0].text = "Second"
        table2.rows[0].cells[1].text = "Table"

        temp_file = self.temp_dir / "positioned_tables.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should maintain proper separation and order
        assert "Before Table" in markdown
        assert "Paragraph before table" in markdown
        assert "| A | B |" in markdown
        assert "Paragraph after table" in markdown
        assert "**formatting**" in markdown
        assert "| Second | Table |" in markdown

        # Check that tables are properly separated from surrounding content
        lines = markdown.split('\n')
        table_line_indices = [i for i, line in enumerate(lines) if '| A | B |' in line]
        assert len(table_line_indices) >= 1
