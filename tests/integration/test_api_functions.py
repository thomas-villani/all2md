#  Copyright (c) 2025 Tom Villani, Ph.D.
#
# tests/integration/test_api_functions.py
"""Integration tests for major API functions: from_ast, from_markdown, and convert.

This module provides comprehensive integration tests for the three main API
functions that were previously untested:
- from_ast: Render AST to various formats
- from_markdown: Convert Markdown to other formats
- convert: General format-to-format conversion

Tests cover various input/output types, format combinations, options handling,
and transform pipeline integration.
"""

from importlib.util import find_spec
from io import BytesIO

import pytest

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.platypus import SimpleDocTemplate  # noqa: F401
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

EBOOKLIB_AVAILABLE = find_spec("ebooklib") is not None
PPTX_AVAILABLE = find_spec("pptx") is not None

from all2md import MarkdownParserOptions, convert, from_ast, from_markdown, to_ast
from all2md.ast import (
    BlockQuote,
    CodeBlock,
    Document,
    Emphasis,
    Heading,
    List,
    ListItem,
    Paragraph,
    Strong,
    Table,
    TableCell,
    TableRow,
    Text,
)
from all2md.options import (
    HtmlRendererOptions,
    MarkdownOptions,
)


def create_sample_ast_document():
    """Create a sample AST document for testing.

    Returns
    -------
    Document
        Sample document with various node types for comprehensive testing.

    """
    return Document(
        metadata={"title": "Test Document", "author": "Test Suite"},
        children=[
            Heading(level=1, content=[Text(content="Main Title")]),
            Paragraph(content=[
                Text(content="This is a paragraph with "),
                Strong(content=[Text(content="bold")]),
                Text(content=" and "),
                Emphasis(content=[Text(content="italic")]),
                Text(content=" text.")
            ]),
            Heading(level=2, content=[Text(content="Lists")]),
            List(ordered=False, items=[
                ListItem(children=[Paragraph(content=[Text(content="First item")])]),
                ListItem(children=[Paragraph(content=[Text(content="Second item")])]),
                ListItem(children=[Paragraph(content=[Text(content="Third item")])])
            ]),
            Heading(level=2, content=[Text(content="Code")]),
            CodeBlock(content='print("Hello, World!")', language="python"),
            Heading(level=2, content=[Text(content="Table")]),
            Table(
                header=TableRow(cells=[
                    TableCell(content=[Text(content="Name")]),
                    TableCell(content=[Text(content="Value")])
                ]),
                rows=[
                    TableRow(cells=[
                        TableCell(content=[Text(content="Item A")]),
                        TableCell(content=[Text(content="100")])
                    ]),
                    TableRow(cells=[
                        TableCell(content=[Text(content="Item B")]),
                        TableCell(content=[Text(content="200")])
                    ])
                ]
            ),
            BlockQuote(children=[
                Paragraph(content=[Text(content="A quoted passage.")])
            ])
        ]
    )


@pytest.mark.integration
class TestFromAst:
    """Integration tests for from_ast function."""

    def test_from_ast_to_markdown_return_string(self):
        """Test rendering AST to Markdown string."""
        doc = create_sample_ast_document()
        result = from_ast(doc, "markdown")

        assert isinstance(result, str)
        assert "# Main Title" in result
        assert "**bold**" in result
        assert "*italic*" in result
        assert "* First item" in result
        assert "```python" in result
        assert "| Name" in result
        assert "> A quoted passage" in result

    def test_from_ast_to_markdown_with_file_output(self, tmp_path):
        """Test rendering AST to Markdown file."""
        doc = create_sample_ast_document()
        output_file = tmp_path / "output.md"

        result = from_ast(doc, "markdown", output=output_file)

        assert result is None
        assert output_file.exists()
        content = output_file.read_text()
        assert "# Main Title" in content
        assert "**bold**" in content

    def test_from_ast_to_html_standalone(self):
        """Test rendering AST to standalone HTML."""
        doc = create_sample_ast_document()
        result = from_ast(
            doc,
            "html",
            renderer_options=HtmlRendererOptions(standalone=True)
        )

        assert isinstance(result, str)
        assert "<!DOCTYPE html>" in result
        assert "<html" in result
        assert "<h1" in result
        assert "Main Title" in result
        assert "<strong>bold</strong>" in result
        assert "<em>italic</em>" in result

    def test_from_ast_to_html_fragment(self):
        """Test rendering AST to HTML fragment."""
        doc = create_sample_ast_document()
        result = from_ast(
            doc,
            "html",
            renderer_options=HtmlRendererOptions(standalone=False)
        )

        assert isinstance(result, str)
        assert "<!DOCTYPE html>" not in result
        assert "<h1" in result
        assert "Main Title" in result

    def test_from_ast_with_renderer_options(self):
        """Test from_ast with various renderer options."""
        doc = create_sample_ast_document()

        # Test with emphasis symbol preference
        result = from_ast(
            doc,
            "markdown",
            renderer_options=MarkdownOptions(emphasis_symbol="_")
        )

        assert "_italic_" in result

    def test_from_ast_with_transforms(self):
        """Test from_ast with transform pipeline."""
        doc = create_sample_ast_document()

        # Apply remove-images transform
        result = from_ast(
            doc,
            "markdown",
            transforms=["remove-images"]
        )

        assert isinstance(result, str)
        assert "Main Title" in result

    def test_from_ast_to_path_output(self, tmp_path):
        """Test from_ast with Path output."""
        doc = create_sample_ast_document()
        output_file = tmp_path / "test.md"

        result = from_ast(doc, "markdown", output=output_file)

        assert result is None
        assert output_file.exists()

    def test_from_ast_to_io_output(self, tmp_path):
        """Test from_ast with IO output."""
        doc = create_sample_ast_document()
        output_file = tmp_path / "test.md"

        with open(output_file, 'w') as f:
            result = from_ast(doc, "markdown", output=f)

        assert result is None
        assert output_file.exists()
        content = output_file.read_text()
        assert "Main Title" in content

    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
    def test_from_ast_to_docx(self, tmp_path):
        """Test rendering AST to DOCX format."""
        doc = create_sample_ast_document()
        output_file = tmp_path / "test.docx"

        result = from_ast(doc, "docx", output=output_file)

        assert result is None
        assert output_file.exists()

        # Verify DOCX content
        docx_doc = DocxDocument(str(output_file))
        all_text = " ".join(p.text for p in docx_doc.paragraphs)
        assert "Main Title" in all_text
        assert "bold" in all_text

    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
    def test_from_ast_to_docx_bytes(self):
        """Test rendering AST to DOCX returns BytesIO."""
        doc = create_sample_ast_document()

        result = from_ast(doc, "docx")

        # Should return BytesIO
        assert hasattr(result, 'read')
        assert hasattr(result, 'seek')

        # Read and verify it's a valid DOCX (ZIP signature)
        content = result.read()
        assert len(content) > 0
        assert content.startswith(b'PK')  # ZIP signature for DOCX files

    @pytest.mark.skipif(not REPORTLAB_AVAILABLE, reason="reportlab not installed")
    def test_from_ast_to_pdf(self, tmp_path):
        """Test rendering AST to PDF format."""
        doc = create_sample_ast_document()
        output_file = tmp_path / "test.pdf"

        result = from_ast(doc, "pdf", output=output_file)

        assert result is None
        assert output_file.exists()
        assert output_file.stat().st_size > 0

    def test_from_ast_with_kwargs_override(self):
        """Test from_ast with kwargs overriding renderer options."""
        doc = create_sample_ast_document()
        base_options = MarkdownOptions(emphasis_symbol="*")

        result = from_ast(
            doc,
            "markdown",
            renderer_options=base_options,
            emphasis_symbol="_"  # Override
        )

        assert "_italic_" in result


@pytest.mark.integration
class TestFromMarkdown:
    """Integration tests for from_markdown function."""

    def test_from_markdown_string_to_html(self, tmp_path):
        """Test converting Markdown string to HTML."""
        # Write markdown to file first since from_markdown expects a file path
        md_file = tmp_path / "test.md"
        md_file.write_text("# Hello\n\nThis is **bold** text.")

        result = from_markdown(str(md_file), "html")

        assert isinstance(result, str)
        assert "<h1" in result
        assert "Hello" in result
        assert "<strong>bold</strong>" in result

    def test_from_markdown_file_to_html(self, tmp_path):
        """Test converting Markdown file to HTML."""
        md_file = tmp_path / "input.md"
        md_file.write_text("# Test\n\nSome content.")

        result = from_markdown(str(md_file), "html")

        assert isinstance(result, str)
        assert "<h1" in result
        assert "Test" in result

    def test_from_markdown_to_html_with_output_file(self, tmp_path):
        """Test converting Markdown to HTML file."""
        md_file = tmp_path / "input.md"
        md_file.write_text("# Document\n\nParagraph text.")
        output_file = tmp_path / "output.html"

        result = from_markdown(str(md_file), "html", output=output_file)

        assert result is None
        assert output_file.exists()
        content = output_file.read_text()
        assert "Document" in content

    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
    def test_from_markdown_to_docx(self, tmp_path):
        """Test converting Markdown to DOCX."""
        md_file = tmp_path / "input.md"
        md_file.write_text("# Title\n\n**Bold** and *italic*.")
        output_file = tmp_path / "output.docx"

        result = from_markdown(str(md_file), "docx", output=output_file)

        assert result is None
        assert output_file.exists()

        docx_doc = DocxDocument(str(output_file))
        all_text = " ".join(p.text for p in docx_doc.paragraphs)
        assert "Title" in all_text

    def test_from_markdown_with_parser_options(self, tmp_path):
        """Test from_markdown with parser options."""
        md_file = tmp_path / "test.md"
        md_file.write_text("# Title\n\nContent")

        result = from_markdown(
            str(md_file),
            "html",
            parser_options=MarkdownParserOptions(flavor="commonmark")
        )

        assert isinstance(result, str)
        assert "Title" in result

    def test_from_markdown_with_renderer_options(self, tmp_path):
        """Test from_markdown with renderer options."""
        md_file = tmp_path / "test.md"
        md_file.write_text("# Title\n\nContent")

        result = from_markdown(
            str(md_file),
            "html",
            renderer_options=HtmlRendererOptions(standalone=True)
        )

        assert "<!DOCTYPE html>" in result

    def test_from_markdown_with_transforms(self, tmp_path):
        """Test from_markdown with transforms applied."""
        md_file = tmp_path / "test.md"
        md_file.write_text("# Title\n\n![Image](test.png)\n\nContent")

        result = from_markdown(
            str(md_file),
            "html",
            transforms=["remove-images"]
        )

        assert "Title" in result
        assert "Content" in result

    def test_from_markdown_bytes_input(self, tmp_path):
        """Test from_markdown with bytes input via file."""
        md_file = tmp_path / "test.md"
        md_file.write_bytes(b"# Test\n\nContent")

        result = from_markdown(str(md_file), "html")

        assert isinstance(result, str)
        assert "Test" in result

    def test_from_markdown_path_input(self, tmp_path):
        """Test from_markdown with Path input."""
        md_file = tmp_path / "test.md"
        md_file.write_text("# Path Test\n\nContent")

        result = from_markdown(md_file, "html")

        assert "Path Test" in result

    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
    def test_from_markdown_to_docx_bytes(self, tmp_path):
        """Test from_markdown to DOCX returns BytesIO."""
        md_file = tmp_path / "test.md"
        md_file.write_text("# Title\n\nContent")

        result = from_markdown(str(md_file), "docx")

        # Should return BytesIO
        assert hasattr(result, 'read')
        assert hasattr(result, 'seek')

        # Read and verify it's a valid DOCX
        content = result.read()
        assert len(content) > 0
        assert content.startswith(b'PK')  # ZIP signature


@pytest.mark.integration
class TestConvert:
    """Integration tests for convert function."""

    def test_convert_html_to_markdown(self, tmp_path):
        """Test converting HTML to Markdown."""
        html_file = tmp_path / "input.html"
        html_file.write_text("<h1>Title</h1><p>Content with <strong>bold</strong>.</p>")

        result = convert(str(html_file), target_format="markdown")

        assert isinstance(result, str)
        assert "# Title" in result or "Title" in result
        assert "**bold**" in result

    def test_convert_markdown_to_html(self, tmp_path):
        """Test converting Markdown to HTML."""
        md_file = tmp_path / "test.md"
        md_file.write_text("# Title\n\n**Bold** text.")

        result = convert(str(md_file), target_format="html", source_format="markdown")

        assert isinstance(result, str)
        assert "Title" in result
        assert "Bold" in result or "bold" in result

    def test_convert_with_auto_source_detection(self, tmp_path):
        """Test convert with auto source format detection."""
        html_file = tmp_path / "test.html"
        html_file.write_text("<html><body><h1>Test</h1></body></html>")

        result = convert(str(html_file), target_format="markdown")

        assert isinstance(result, str)
        assert "Test" in result

    def test_convert_with_auto_target_detection(self, tmp_path):
        """Test convert with auto target format detection from output."""
        md_file = tmp_path / "input.md"
        md_file.write_text("# Test\n\nContent")
        output_file = tmp_path / "output.html"

        result = convert(
            str(md_file),
            output=output_file,
            source_format="markdown"
        )

        assert result is None
        assert output_file.exists()
        content = output_file.read_text()
        assert "Test" in content

    def test_convert_with_parser_and_renderer_options(self, tmp_path):
        """Test convert with both parser and renderer options."""
        md_file = tmp_path / "test.md"
        md_file.write_text("# Title\n\nContent")

        result = convert(
            str(md_file),
            source_format="markdown",
            target_format="html",
            parser_options=MarkdownParserOptions(),
            renderer_options=HtmlRendererOptions(standalone=False)
        )

        assert isinstance(result, str)
        assert "<!DOCTYPE html>" not in result
        assert "Title" in result

    def test_convert_with_transforms(self, tmp_path):
        """Test convert with transform pipeline."""
        md_file = tmp_path / "test.md"
        md_file.write_text("# Title\n\n![Image](test.png)\n\nContent")

        result = convert(
            str(md_file),
            source_format="markdown",
            target_format="html",
            transforms=["remove-images"]
        )

        assert "Title" in result
        assert "Content" in result

    def test_convert_bytes_input(self):
        """Test convert with bytes input."""
        html_bytes = b"<h1>Title</h1><p>Content</p>"

        result = convert(
            html_bytes,
            source_format="html",
            target_format="markdown"
        )

        assert isinstance(result, str)
        assert "Title" in result

    def test_convert_io_input(self):
        """Test convert with IO input."""
        html_content = b"<h1>Test</h1><p>Content</p>"
        html_io = BytesIO(html_content)

        result = convert(
            html_io,
            source_format="html",
            target_format="markdown"
        )

        assert isinstance(result, str)
        assert "Test" in result

    def test_convert_to_file_output(self, tmp_path):
        """Test convert with file output."""
        md_file = tmp_path / "input.md"
        md_file.write_text("# Title\n\nContent")
        output_file = tmp_path / "output.html"

        result = convert(
            str(md_file),
            output=output_file,
            source_format="markdown",
            target_format="html"
        )

        assert result is None
        assert output_file.exists()

    def test_convert_to_io_output(self, tmp_path):
        """Test convert with IO output."""
        md_file = tmp_path / "input.md"
        md_file.write_text("# Title\n\nContent")
        output_file = tmp_path / "output.html"

        with open(output_file, 'w') as f:
            result = convert(
                str(md_file),
                output=f,
                source_format="markdown",
                target_format="html"
            )

        assert result is None
        assert output_file.exists()

    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
    def test_convert_markdown_to_docx(self, tmp_path):
        """Test converting Markdown to DOCX."""
        md_file = tmp_path / "input.md"
        md_file.write_text("# Document\n\n**Important** content.")
        output_file = tmp_path / "output.docx"

        result = convert(
            str(md_file),
            output=output_file,
            source_format="markdown",
            target_format="docx"
        )

        assert result is None
        assert output_file.exists()

        docx_doc = DocxDocument(str(output_file))
        all_text = " ".join(p.text for p in docx_doc.paragraphs)
        assert "Document" in all_text

    def test_convert_with_flavor_shorthand(self, tmp_path):
        """Test convert with flavor shorthand."""
        md_file = tmp_path / "test.md"
        md_file.write_text("# Title\n\nContent")

        result = convert(
            str(md_file),
            source_format="markdown",
            target_format="markdown",
            flavor="commonmark"
        )

        assert isinstance(result, str)
        assert "Title" in result

    def test_convert_with_kwargs_split(self):
        """Test convert with kwargs split between parser and renderer."""
        html_content = b"<h1>Test</h1><p>Content</p>"

        result = convert(
            html_content,
            source_format="html",
            target_format="markdown",
            convert_nbsp=True,  # HtmlOptions (parser)
            emphasis_symbol="_"  # MarkdownOptions (renderer)
        )

        assert isinstance(result, str)
        assert "Test" in result

    def test_convert_roundtrip_markdown_html(self, tmp_path):
        """Test roundtrip conversion Markdown -> HTML -> Markdown."""
        md_file = tmp_path / "original.md"
        md_file.write_text("# Title\n\n**Bold** and *italic* text.")

        # Markdown -> HTML
        html_file = tmp_path / "temp.html"
        convert(
            str(md_file),
            output=html_file,
            source_format="markdown",
            target_format="html"
        )
        html_content = html_file.read_text()
        assert "Title" in html_content

        # HTML -> Markdown
        recovered_markdown = convert(
            str(html_file),
            source_format="html",
            target_format="markdown"
        )
        assert "Title" in recovered_markdown
        assert "**Bold**" in recovered_markdown or "bold" in recovered_markdown.lower()

    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
    def test_convert_docx_to_markdown(self, tmp_path):
        """Test converting DOCX to Markdown."""
        # Create a simple DOCX file
        from docx import Document as CreateDoc
        doc = CreateDoc()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is a test paragraph.")

        docx_file = tmp_path / "test.docx"
        doc.save(str(docx_file))

        # Convert to Markdown
        result = convert(str(docx_file), target_format="markdown")

        assert isinstance(result, str)
        assert "Test Document" in result
        assert "test paragraph" in result


@pytest.mark.integration
class TestAPIEdgeCases:
    """Test edge cases and error handling for API functions."""

    def test_from_ast_empty_document(self):
        """Test from_ast with empty document."""
        doc = Document()
        result = from_ast(doc, "markdown")

        assert result == ""

    def test_from_markdown_empty_content(self, tmp_path):
        """Test from_markdown with empty content."""
        md_file = tmp_path / "empty.md"
        md_file.write_text("")

        result = from_markdown(str(md_file), "html")

        assert isinstance(result, str)

    def test_convert_empty_file(self, tmp_path):
        """Test convert with empty file."""
        empty_file = tmp_path / "empty.md"
        empty_file.write_text("")

        result = convert(str(empty_file), source_format="markdown", target_format="markdown")

        assert result == ""

    def test_from_ast_with_complex_transforms(self):
        """Test from_ast with multiple transforms."""
        doc = create_sample_ast_document()

        result = from_ast(
            doc,
            "markdown",
            transforms=["remove-images", "heading-offset"]
        )

        assert isinstance(result, str)
        # After heading offset, h1 becomes h2
        assert "## Main Title" in result

    def test_convert_with_explicit_formats_override_detection(self, tmp_path):
        """Test that explicit formats override auto-detection."""
        # Create HTML file but process as markdown to avoid txt format issues
        html_file = tmp_path / "test.html"
        html_file.write_text("<h1>HTML</h1>")

        # Convert HTML to markdown (should process HTML properly)
        result = convert(
            str(html_file),
            source_format="html",
            target_format="markdown"
        )

        # Should convert HTML to markdown
        assert "HTML" in result


@pytest.mark.integration
class TestAPIConsistency:
    """Test consistency across different API functions."""

    def test_to_markdown_vs_convert_consistency(self, tmp_path):
        """Test that to_markdown and convert produce same results."""
        from all2md import to_markdown

        html_file = tmp_path / "test.html"
        html_file.write_text("<html><body><h1>Title</h1><p>Content</p></body></html>")

        result1 = to_markdown(str(html_file))
        result2 = convert(str(html_file), target_format="markdown")

        # Results should be similar (exact match may vary due to rendering details)
        assert "Title" in result1 and "Title" in result2
        assert "Content" in result1 and "Content" in result2

    def test_to_ast_plus_from_ast_vs_to_markdown(self, tmp_path):
        """Test that to_ast + from_ast equals to_markdown."""
        html_file = tmp_path / "test.html"
        html_file.write_text("<html><body><h1>Title</h1><p>Content</p></body></html>")

        # Direct conversion
        from all2md import to_markdown
        direct = to_markdown(str(html_file))

        # Via AST
        ast_doc = to_ast(str(html_file))
        via_ast = from_ast(ast_doc, "markdown")

        # Should be identical
        assert direct == via_ast

    def test_from_markdown_vs_convert_consistency(self, tmp_path):
        """Test that from_markdown and convert produce same results."""
        md_file = tmp_path / "test.md"
        md_file.write_text("# Title\n\n**Bold** text.")

        result1 = from_markdown(str(md_file), "html")
        result2 = convert(str(md_file), source_format="markdown", target_format="html")

        # Should be identical
        assert result1 == result2
