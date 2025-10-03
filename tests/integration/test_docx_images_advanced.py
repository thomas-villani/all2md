"""Advanced tests for DOCX image handling edge cases."""

import docx
from docx.shared import Inches

from all2md import to_markdown as docx_to_markdown
from all2md.options import DocxOptions
from tests.utils import MINIMAL_PNG_BYTES, assert_markdown_valid, cleanup_test_dir, create_test_temp_dir


class TestDocxImagesAdvanced:
    """Test complex image scenarios in DOCX documents."""

    def setup_method(self):
        """Set up test environment."""
        self.temp_dir = create_test_temp_dir()

    def teardown_method(self):
        """Clean up test environment."""
        cleanup_test_dir(self.temp_dir)

    def create_test_images(self):
        """Create test image files in various formats."""
        # PNG image
        png_path = self.temp_dir / "test.png"
        png_path.write_bytes(MINIMAL_PNG_BYTES)

        # Create a fake JPEG (just different extension)
        jpg_path = self.temp_dir / "test.jpg"
        jpg_path.write_bytes(MINIMAL_PNG_BYTES)  # Still PNG data, but different extension

        # Create unsupported format (fake)
        bmp_path = self.temp_dir / "test.bmp"
        bmp_path.write_bytes(b"BM\x00\x00\x00\x00")  # Minimal BMP header

        return png_path, jpg_path, bmp_path

    def test_alt_text_extraction(self):
        """Test extraction of alt text from images."""
        png_path, _, _ = self.create_test_images()

        doc = docx.Document()
        doc.add_heading("Image Alt Text Test", level=1)

        # Add image (python-docx doesn't directly support alt text, but we test structure)
        doc.add_paragraph("Image with alt text:")
        doc.add_picture(str(png_path), width=Inches(2))

        temp_file = self.temp_dir / "alt_text.docx"
        doc.save(str(temp_file))

        # Test alt_text mode
        options = DocxOptions(attachment_mode="alt_text")
        markdown = docx_to_markdown(str(temp_file), options=options)
        assert_markdown_valid(markdown)
        # Should contain image reference
        assert "![image]" in markdown

    def test_embedded_vs_linked_images(self):
        """Test handling of embedded vs linked images."""
        png_path, _, _ = self.create_test_images()

        doc = docx.Document()

        # Embedded image
        doc.add_paragraph("Embedded image:")
        doc.add_picture(str(png_path), width=Inches(2))

        # Linked image would require different handling in python-docx
        doc.add_paragraph("Another embedded image:")
        doc.add_picture(str(png_path), width=Inches(1))

        temp_file = self.temp_dir / "embedded_linked.docx"
        doc.save(str(temp_file))

        # Test base64 embedding
        options = DocxOptions(attachment_mode="base64")
        markdown = docx_to_markdown(str(temp_file), options=options)
        assert_markdown_valid(markdown)

        # Should contain base64 encoded images
        base64_images = [line for line in markdown.split('\n') if 'data:image/' in line]
        assert len(base64_images) >= 1

        # Test alt_text mode
        options_alt = DocxOptions(attachment_mode="alt_text")
        markdown_alt = docx_to_markdown(str(temp_file), options=options_alt)
        assert_markdown_valid(markdown_alt)

        # Should contain alt text references instead
        assert "![image]" in markdown_alt

    def test_unsupported_image_formats(self):
        """Test handling of unsupported image formats."""
        _, _, bmp_path = self.create_test_images()

        doc = docx.Document()

        try:
            # Try to add unsupported format (might fail in python-docx)
            doc.add_picture(str(bmp_path))
        except Exception:
            # If python-docx rejects it, create a document with known formats
            png_path, _, _ = self.create_test_images()
            doc.add_picture(str(png_path))

        temp_file = self.temp_dir / "unsupported_format.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle gracefully
        assert isinstance(markdown, str)

    def test_images_in_different_contexts(self):
        """Test images in various document contexts."""
        png_path, _, _ = self.create_test_images()

        doc = docx.Document()

        # Image in paragraph
        doc.add_paragraph("Text before image ")
        doc.add_picture(str(png_path), width=Inches(1))

        # Image in list item (simulated)
        doc.add_paragraph("List item with image:", style="List Bullet")
        doc.add_picture(str(png_path), width=Inches(0.5))

        # Image in table cell
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Cell with text"
        # Adding image to cell is complex in python-docx, so we simulate
        table.rows[0].cells[1].text = "[Image would be here]"
        table.rows[1].cells[0].text = "Regular cell"

        temp_file = self.temp_dir / "images_in_context.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle images in different contexts
        assert "Text before image" in markdown
        assert "List item with image" in markdown
        assert "Regular cell" in markdown

    def test_image_sizing_and_attributes(self):
        """Test images with different sizes and attributes."""
        png_path, _, _ = self.create_test_images()

        doc = docx.Document()

        # Different sized images
        doc.add_paragraph("Small image:")
        doc.add_picture(str(png_path), width=Inches(0.5))

        doc.add_paragraph("Medium image:")
        doc.add_picture(str(png_path), width=Inches(2))

        doc.add_paragraph("Large image:")
        doc.add_picture(str(png_path), width=Inches(4))

        # Image with height specification
        doc.add_paragraph("Fixed height image:")
        doc.add_picture(str(png_path), height=Inches(1))

        temp_file = self.temp_dir / "sized_images.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle all images
        image_refs = [line for line in markdown.split('\n') if '![' in line]
        assert len(image_refs) >= 4  # Should have references to all images

    def test_images_with_captions(self):
        """Test images with caption text."""
        png_path, _, _ = self.create_test_images()

        doc = docx.Document()

        # Image followed by caption-like text
        doc.add_picture(str(png_path), width=Inches(2))
        p_caption = doc.add_paragraph("Figure 1: Test image caption")
        p_caption.style = "Caption"  # Use caption style if available

        # Another image with different caption style
        doc.add_picture(str(png_path), width=Inches(1.5))
        p_caption2 = doc.add_paragraph("Image caption with emphasis:")
        run = p_caption2.add_run("Important image")
        run.italic = True

        temp_file = self.temp_dir / "images_with_captions.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should preserve captions
        assert "Figure 1: Test image caption" in markdown
        assert "Image caption" in markdown
        assert "*Important image*" in markdown

    def test_inline_vs_block_images(self):
        """Test inline images vs block-level images."""
        png_path, _, _ = self.create_test_images()

        doc = docx.Document()

        # Block-level image (standalone)
        doc.add_picture(str(png_path), width=Inches(3))

        # Simulate inline image by adding to paragraph with text
        p = doc.add_paragraph("Text with ")
        # Note: python-docx inline images are complex, so we simulate the concept
        p.add_run("[inline image]")  # Placeholder for inline image
        p.add_run(" more text after")

        temp_file = self.temp_dir / "inline_block_images.docx"
        doc.save(str(temp_file))

        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)

        # Should handle both types appropriately
        assert "Text with" in markdown
        assert "more text after" in markdown

    def test_image_removal_mode(self):
        """Test skip mode for image handling."""
        png_path, _, _ = self.create_test_images()

        doc = docx.Document()

        doc.add_paragraph("Text before image")
        doc.add_picture(str(png_path), width=Inches(2))
        doc.add_paragraph("Text after image")

        temp_file = self.temp_dir / "image_removal.docx"
        doc.save(str(temp_file))

        # Test skip mode
        options = DocxOptions(attachment_mode="skip")
        markdown = docx_to_markdown(str(temp_file), options=options)
        assert_markdown_valid(markdown)

        # Should contain text but no image references
        assert "Text before image" in markdown
        assert "Text after image" in markdown
        assert "![" not in markdown  # No image markdown

    def test_corrupted_or_missing_images(self):
        """Test handling of corrupted or missing image references."""
        # Create a document that references a non-existent image
        doc = docx.Document()

        doc.add_paragraph("Document with text content")

        # We can't easily create a corrupted image reference in python-docx,
        # but we can test the robustness of the conversion
        temp_file = self.temp_dir / "missing_images.docx"
        doc.save(str(temp_file))

        # Should handle gracefully even with issues
        markdown = docx_to_markdown(str(temp_file))
        assert_markdown_valid(markdown)
        assert "Document with text content" in markdown

    def test_download_mode_images(self):
        """Test download mode for image handling."""
        png_path, _, _ = self.create_test_images()

        doc = docx.Document()
        doc.add_paragraph("Document with downloadable image")
        doc.add_picture(str(png_path), width=Inches(2))

        temp_file = self.temp_dir / "download_images.docx"
        doc.save(str(temp_file))

        # Test download mode
        options = DocxOptions(
            attachment_mode="download",
            attachment_output_dir=str(self.temp_dir / "downloads")
        )

        markdown = docx_to_markdown(str(temp_file), options=options)
        assert_markdown_valid(markdown)

        # Should contain text and image references
        assert "Document with downloadable image" in markdown
        # Image handling depends on implementation details
