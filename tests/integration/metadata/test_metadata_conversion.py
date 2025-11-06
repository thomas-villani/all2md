"""Integration tests for metadata extraction across all parsers."""

import importlib.util
import json
from io import BytesIO

import pytest
from utils import cleanup_test_dir, create_test_temp_dir

from all2md import EmlOptions, HtmlOptions, IpynbOptions, MhtmlOptions, to_markdown
from all2md.options import MarkdownRendererOptions

# Check for optional dependencies
ODFPY_AVAILABLE = importlib.util.find_spec("odf") is not None
PYTH_AVAILABLE = importlib.util.find_spec("pyth") is not None
PYMUPDF_AVAILABLE = importlib.util.find_spec("fitz") is not None
DOCX_AVAILABLE = importlib.util.find_spec("docx") is not None
OPENPYXL_AVAILABLE = importlib.util.find_spec("openpyxl") is not None


@pytest.mark.integration
@pytest.mark.metadata
class TestMetadataIntegration:
    """Integration tests for metadata extraction functionality."""

    def setup_method(self):
        """Set up test environment."""
        self.temp_dir = create_test_temp_dir()

    def teardown_method(self):
        """Clean up test environment."""
        cleanup_test_dir(self.temp_dir)

    def test_html_metadata_extraction_integration(self):
        """Test HTML metadata extraction with real file."""
        html_content = """<!DOCTYPE html>
<html>
<head>
    <title>Integration Test Document</title>
    <meta name="author" content="Test Author">
    <meta name="description" content="Testing HTML metadata extraction in integration">
    <meta name="keywords" content="integration,test,metadata,html">
    <meta name="generator" content="Integration Test Suite">
    <meta property="og:title" content="OG Test Title">
    <meta property="og:type" content="article">
</head>
<body>
    <h1>Integration Test Document</h1>
    <p>This document tests metadata extraction in integration tests.</p>
    <p>It contains multiple paragraphs and <a href="https://example.com">links</a>.</p>
    <img src="test.png" alt="Test image">
</body>
</html>"""

        html_file = self.temp_dir / "test.html"
        html_file.write_text(html_content, encoding="utf-8")

        # Test with metadata extraction enabled
        md_options = MarkdownRendererOptions(metadata_frontmatter=True)
        parser_options = HtmlOptions(extract_metadata=True)
        result = to_markdown(str(html_file), parser_options=parser_options, renderer_options=md_options)

        # Verify YAML front matter is present
        assert result.startswith("---")
        assert "title: Integration Test Document" in result
        assert "author: Test Author" in result
        assert "description: Testing HTML metadata extraction in integration" in result
        assert "keywords: [integration, test, metadata, html]" in result or "keywords:" in result
        assert "creator: Integration Test Suite" in result or "generator: Integration Test Suite" in result
        assert "category: article" in result or "type: article" in result

        # Verify content is still present
        assert "# Integration Test Document" in result
        assert "[links](https://example.com)" in result

    def test_eml_metadata_extraction_integration(self):
        """Test EML metadata extraction with real email file."""
        eml_content = """From: sender@example.com (Test Sender)
To: recipient@example.com
Subject: Integration Test Email
Date: Fri, 26 Sep 2025 15:30:00 +0000
Message-ID: <integration-test@example.com>
X-Mailer: Integration Test Mailer
Content-Type: text/plain; charset=utf-8

This is an integration test email for metadata extraction.

It contains multiple lines and should extract proper metadata.

Best regards,
Test Sender
"""

        eml_file = self.temp_dir / "test.eml"
        eml_file.write_bytes(eml_content.encode("utf-8"))

        # Test with metadata extraction enabled
        md_options = MarkdownRendererOptions(metadata_frontmatter=True)
        parser_options = EmlOptions(extract_metadata=True)
        result = to_markdown(str(eml_file), parser_options=parser_options, renderer_options=md_options)

        # Verify YAML front matter is present
        assert result.startswith("---")
        assert "title: Integration Test Email" in result
        assert "sender@example.com" in result  # Author field (may be quoted)
        # Date may be quoted or unquoted in YAML
        assert (
            "creation_date: Fri, 26 Sep 2025 15:30:00 +0000" in result
            or "creation_date: 2025-09-26" in result
            or "creation_date: '2025-09-26'" in result
        )
        assert "creator: Integration Test Mailer" in result
        assert "recipient@example.com" in result  # To field (format may vary)
        assert "integration-test@example.com" in result  # Message ID (format may vary)

        # Verify content is still present
        assert "This is an integration test email" in result

    def test_ipynb_metadata_extraction_integration(self):
        """Test Jupyter notebook metadata extraction with real notebook."""
        notebook_content = {
            "cells": [
                {
                    "cell_type": "markdown",
                    "metadata": {},
                    "source": ["# Integration Test Notebook\n", "\n", "This notebook tests metadata extraction."],
                },
                {
                    "cell_type": "code",
                    "execution_count": 1,
                    "metadata": {},
                    "outputs": [
                        {"name": "stdout", "output_type": "stream", "text": ["Hello from integration test!\n"]}
                    ],
                    "source": ["print('Hello from integration test!')"],
                },
                {"cell_type": "markdown", "metadata": {}, "source": ["## Section 2\n", "\n", "More content here."]},
            ],
            "metadata": {
                "title": "Integration Test Notebook",
                "authors": [{"name": "Integration Test Author"}],
                "kernelspec": {"display_name": "Python 3", "language": "python", "name": "python3"},
                "language_info": {"name": "python", "version": "3.9.0", "mimetype": "text/x-python"},
            },
            "nbformat": 4,
            "nbformat_minor": 4,
        }

        notebook_file = self.temp_dir / "test.ipynb"
        notebook_file.write_text(json.dumps(notebook_content), encoding="utf-8")

        # Test with metadata extraction enabled
        md_options = MarkdownRendererOptions(metadata_frontmatter=True)
        parser_options = IpynbOptions(extract_metadata=True)
        result = to_markdown(str(notebook_file), parser_options=parser_options, renderer_options=md_options)

        # Verify YAML front matter is present
        assert result.startswith("---")
        assert "title: Integration Test Notebook" in result
        assert "author: Integration Test Author" in result
        assert "language: python" in result
        assert "kernel: Python 3" in result
        assert "language_version: 3.9.0" in result or 'language_version: "3.9.0"' in result
        assert "cell_count: 3" in result
        assert "code_cells: 1" in result
        assert "markdown_cells: 2" in result

        # Verify content is still present
        assert "# Integration Test Notebook" in result
        assert "print('Hello from integration test!')" in result

    def test_mhtml_metadata_extraction_integration(self):
        """Test MHTML metadata extraction with real MHTML file."""
        mhtml_content = b"""MIME-Version: 1.0
Content-Type: multipart/related; boundary="integration-test-boundary"
From: test@example.com
Subject: Integration Test MHTML
Date: Fri, 26 Sep 2025 16:00:00 +0000

--integration-test-boundary
Content-Type: text/html; charset=utf-8

<!DOCTYPE html>
<html>
<head>
    <title>MHTML Integration Test</title>
    <meta name="author" content="MHTML Test Author">
    <meta name="description" content="Integration test for MHTML metadata">
    <meta name="keywords" content="mhtml,integration,test">
</head>
<body>
    <h1>MHTML Integration Test</h1>
    <p>This is an MHTML document for integration testing.</p>
    <p>Word count: This document contains multiple words for testing.</p>
    <a href="https://example.com">External link</a>
    <a href="https://test.com">Another link</a>
</body>
</html>

--integration-test-boundary--
"""

        mhtml_file = self.temp_dir / "test.mht"
        mhtml_file.write_bytes(mhtml_content)

        # Test with metadata extraction enabled
        md_options = MarkdownRendererOptions(metadata_frontmatter=True)
        parser_options = MhtmlOptions(extract_metadata=True)
        result = to_markdown(str(mhtml_file), parser_options=parser_options, renderer_options=md_options)

        # Verify YAML front matter is present
        assert result.startswith("---")
        assert "title: Integration Test MHTML" in result or "title: MHTML Integration Test" in result
        assert "test@example.com" in result  # Author field (may be quoted)
        # Keywords format may vary (list or inline)
        assert "mhtml" in result and "integration" in result and "test" in result

        # Verify content is still present
        assert "# MHTML Integration Test" in result
        assert "[External link](https://example.com)" in result

    @pytest.mark.skipif(not ODFPY_AVAILABLE, reason="odfpy not installed")
    def test_odf_metadata_extraction_integration(self):
        """Test ODF metadata extraction."""
        from odf.dc import Title
        from odf.meta import CreationDate, InitialCreator, Keyword
        from odf.opendocument import OpenDocumentText
        from odf.text import H, P

        from all2md import OdtOptions

        # Create an ODT with metadata using odfpy
        odt_file = self.temp_dir / "test.odt"
        doc = OpenDocumentText()

        # Add content
        title = H(outlinelevel=1, text="ODT Integration Test Document")
        doc.text.addElement(title)

        p1 = P(text="This is a test ODT document for metadata extraction.")
        doc.text.addElement(p1)

        p2 = P(text="It contains multiple paragraphs and rich content.")
        doc.text.addElement(p2)

        # Set metadata
        doc.meta.addElement(Title(text="ODT Integration Test"))
        doc.meta.addElement(InitialCreator(text="ODT Test Author"))
        doc.meta.addElement(Keyword(text="odt"))
        doc.meta.addElement(Keyword(text="integration"))
        doc.meta.addElement(Keyword(text="test"))
        doc.meta.addElement(Keyword(text="metadata"))
        doc.meta.addElement(CreationDate(text="2025-09-26T15:00:00"))

        doc.save(str(odt_file))

        # Test with metadata extraction enabled
        md_options = MarkdownRendererOptions(metadata_frontmatter=True)
        parser_options = OdtOptions(extract_metadata=True)
        result = to_markdown(str(odt_file), parser_options=parser_options, renderer_options=md_options)

        # Verify YAML front matter is present
        assert result.startswith("---")
        assert "title: ODT Integration Test" in result
        assert "author: ODT Test Author" in result
        # Keywords format may vary (list or inline)
        assert "odt" in result and "integration" in result and "test" in result and "metadata" in result
        # Check for date
        assert "2025-09-26" in result
        # Check for custom metadata
        assert "document_type: text" in result

        # Verify content is still present
        assert "# ODT Integration Test Document" in result
        assert "This is a test ODT document for metadata extraction" in result

    @pytest.mark.skipif(not PYTH_AVAILABLE, reason="pyth not installed")
    def test_rtf_metadata_extraction_integration(self):
        """Test RTF metadata extraction."""
        from all2md import RtfOptions

        # Create an RTF document with content for metadata extraction
        # Note: RTF metadata extraction via pyth is limited to content analysis
        # (word count, paragraph count, etc.) rather than document properties
        rtf_content = (
            "{\\rtf1\\ansi\\deff0"
            "{\\fonttbl{\\f0 Arial;}}"
            "\\pard\\fs28 RTF Integration Test\\par"
            "\\pard\\fs24\\par"
            "\\pard This is a test RTF document for metadata extraction.\\par"
            "\\pard It contains multiple paragraphs and various content.\\par"
            "\\pard RTF metadata extraction analyzes document structure.\\par"
            "\\pard This helps verify content-based metadata capabilities.\\par"
            "}"
        )

        rtf_file = self.temp_dir / "test.rtf"
        rtf_file.write_bytes(rtf_content.encode("utf-8"))

        # Test with metadata extraction enabled
        md_options = MarkdownRendererOptions(metadata_frontmatter=True)
        parser_options = RtfOptions(extract_metadata=True)
        result = to_markdown(str(rtf_file), parser_options=parser_options, renderer_options=md_options)

        # Verify YAML front matter is present
        assert result.startswith("---")

        # RTF metadata extraction focuses on content analysis
        # Note: Title extraction from RTF is based on first line heuristics
        # and may not always detect a title. We check for other metadata instead.

        # Check for custom metadata (content-based)
        assert "document_type: rtf" in result
        assert "format: Rich Text Format" in result
        assert "paragraph_count:" in result
        assert "word_count:" in result
        assert "character_count:" in result

        # Verify content is still present
        assert "RTF Integration Test" in result
        assert "This is a test RTF document for metadata extraction" in result

    @pytest.mark.skipif(not PYMUPDF_AVAILABLE, reason="PyMuPDF not installed")
    def test_pdf_metadata_extraction_integration(self):
        """Test PDF metadata extraction."""
        import fitz

        from all2md import PdfOptions

        # Create a PDF with metadata using PyMuPDF
        pdf_file = self.temp_dir / "test.pdf"
        doc = fitz.open()
        page = doc.new_page(width=595, height=842)

        # Add content
        page.insert_text((50, 50), "Integration Test PDF Document", fontsize=16)
        page.insert_text((50, 100), "This is a test PDF for metadata extraction.", fontsize=12)
        page.insert_text((50, 120), "It contains multiple lines of content.", fontsize=12)

        # Set metadata
        doc.set_metadata(
            {
                "title": "PDF Integration Test",
                "author": "PDF Test Author",
                "subject": "Testing PDF metadata extraction",
                "keywords": "pdf, integration, test, metadata",
                "creator": "PyMuPDF Integration Suite",
                "producer": "all2md Test Framework",
            }
        )

        doc.save(str(pdf_file))
        doc.close()

        # Test with metadata extraction enabled
        md_options = MarkdownRendererOptions(metadata_frontmatter=True)
        parser_options = PdfOptions(extract_metadata=True)
        result = to_markdown(str(pdf_file), parser_options=parser_options, renderer_options=md_options)

        # Verify YAML front matter is present
        assert result.startswith("---")
        assert "title: PDF Integration Test" in result
        assert "author: PDF Test Author" in result
        # Note: subject field is converted to description in YAML output
        assert "description: Testing PDF metadata extraction" in result
        # Keywords format may vary (list or inline)
        assert "pdf" in result and "integration" in result and "test" in result and "metadata" in result
        assert "creator: PyMuPDF Integration Suite" in result

        # Verify content is still present
        assert "Integration Test PDF Document" in result
        assert "This is a test PDF for metadata extraction" in result

    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
    def test_docx_metadata_extraction_integration(self):
        """Test DOCX metadata extraction."""
        import datetime

        import docx

        from all2md import DocxOptions

        # Create a DOCX with metadata using python-docx
        docx_file = self.temp_dir / "test.docx"
        doc = docx.Document()

        # Add content
        doc.add_heading("DOCX Integration Test Document", level=1)
        doc.add_paragraph("This is a test Word document for metadata extraction.")
        doc.add_paragraph("It contains multiple paragraphs and rich content.")

        # Set core properties (metadata)
        core_props = doc.core_properties
        core_props.title = "DOCX Integration Test"
        core_props.author = "DOCX Test Author"
        core_props.subject = "Testing DOCX metadata extraction"
        core_props.keywords = "docx, integration, test, metadata"
        core_props.created = datetime.datetime(2025, 9, 26, 15, 0, 0)
        core_props.modified = datetime.datetime(2025, 9, 26, 16, 0, 0)
        core_props.category = "Test Documents"

        doc.save(str(docx_file))

        # Test with metadata extraction enabled
        md_options = MarkdownRendererOptions(metadata_frontmatter=True)
        parser_options = DocxOptions(extract_metadata=True)
        result = to_markdown(str(docx_file), parser_options=parser_options, renderer_options=md_options)

        # Verify YAML front matter is present
        assert result.startswith("---")
        assert "title: DOCX Integration Test" in result
        assert "author: DOCX Test Author" in result
        # Note: subject field is converted to description in YAML output
        assert "description: Testing DOCX metadata extraction" in result
        # Keywords format may vary (list or inline)
        assert "docx" in result and "integration" in result and "test" in result and "metadata" in result
        # Check for date (format may vary)
        assert "2025-09-26" in result or "2025" in result
        assert "category: Test Documents" in result

        # Verify content is still present
        assert "# DOCX Integration Test Document" in result
        assert "This is a test Word document for metadata extraction" in result

    @pytest.mark.skipif(not OPENPYXL_AVAILABLE, reason="openpyxl not installed")
    def test_xlsx_metadata_extraction_integration(self):
        """Test XLSX metadata extraction."""
        import datetime

        import openpyxl

        from all2md import XlsxOptions

        # Create an XLSX with metadata using openpyxl
        xlsx_file = self.temp_dir / "test.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Integration Test"

        # Add content
        ws["A1"] = "Product"
        ws["B1"] = "Price"
        ws["A2"] = "Widget"
        ws["B2"] = 10.99
        ws["A3"] = "Gadget"
        ws["B3"] = 25.50

        # Set workbook properties (metadata)
        wb.properties.title = "XLSX Integration Test"
        wb.properties.creator = "XLSX Test Author"
        wb.properties.subject = "Testing XLSX metadata extraction"
        wb.properties.keywords = "xlsx, integration, test, metadata"
        wb.properties.created = datetime.datetime(2025, 9, 26, 15, 0, 0)

        wb.save(str(xlsx_file))
        wb.close()

        # Test with metadata extraction enabled
        md_options = MarkdownRendererOptions(metadata_frontmatter=True)
        parser_options = XlsxOptions(extract_metadata=True)
        result = to_markdown(str(xlsx_file), parser_options=parser_options, renderer_options=md_options)

        # Verify YAML front matter is present
        assert result.startswith("---")
        assert "title: XLSX Integration Test" in result
        assert "author: XLSX Test Author" in result
        # Note: subject field is converted to description in YAML output
        assert "description: Testing XLSX metadata extraction" in result
        # Keywords format may vary (list or inline)
        assert "xlsx" in result and "integration" in result and "test" in result and "metadata" in result
        # Check for custom metadata
        assert "sheet_count:" in result
        assert "sheet_names:" in result

        # Verify content is still present (table content)
        assert "Product" in result and "Price" in result
        assert "Widget" in result and "Gadget" in result

    def test_metadata_extraction_disabled_by_default(self):
        """Test that metadata extraction is disabled by default."""
        html_content = """<!DOCTYPE html>
<html>
<head>
    <title>Default Test</title>
    <meta name="author" content="Default Author">
</head>
<body>
    <h1>Default Test</h1>
    <p>Testing default behavior.</p>
</body>
</html>"""

        html_file = self.temp_dir / "default.html"
        html_file.write_text(html_content, encoding="utf-8")

        # Test without explicit metadata options (should default to False)
        result = to_markdown(str(html_file))

        # Should NOT have YAML front matter
        assert not result.startswith("---")
        assert "title:" not in result
        assert "author:" not in result

        # Should still have content
        assert "# Default Test" in result
        assert "Testing default behavior." in result

    def test_metadata_extraction_with_file_like_objects(self):
        """Test metadata extraction works with file-like objects."""
        html_content = """<!DOCTYPE html>
<html>
<head>
    <title>BytesIO Test</title>
    <meta name="author" content="BytesIO Author">
</head>
<body>
    <h1>BytesIO Test</h1>
</body>
</html>"""

        # Test with BytesIO
        html_bytes = BytesIO(html_content.encode("utf-8"))
        md_options = MarkdownRendererOptions(metadata_frontmatter=True)
        parser_options = HtmlOptions(extract_metadata=True)
        result = to_markdown(
            html_bytes, source_format="html", parser_options=parser_options, renderer_options=md_options
        )

        # Verify metadata extraction works with file-like objects
        assert result.startswith("---")
        assert "title: BytesIO Test" in result
        assert "author: BytesIO Author" in result
        assert "# BytesIO Test" in result
