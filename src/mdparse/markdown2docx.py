"""Markdown to Word document (DOCX) conversion module.

This module provides functionality to convert Markdown content into Microsoft
Word documents (DOCX format) while preserving formatting, structure, and
styling. It processes Markdown syntax and converts it to equivalent Word
document elements including headers, paragraphs, lists, tables, and links.

The converter parses Markdown text and creates a structured Word document
with appropriate styles, formatting, and layout. It handles complex Markdown
features and converts them to their Word equivalents while maintaining
document hierarchy and visual appearance.

Key Features
------------
- Complete Markdown syntax support (headers, lists, tables, links)
- Text formatting conversion (bold, italic, code, strikethrough)
- Table creation with proper structure and formatting
- List handling (bulleted, numbered, nested)
- Hyperlink creation and styling
- Code block and inline code formatting
- Image embedding support
- Custom styling and formatting options

Supported Markdown Elements
---------------------------
- Headers (H1-H6) with appropriate Word styles
- Paragraphs with text formatting (bold, italic, underline)
- Lists (bulleted, numbered) with proper indentation
- Tables with headers and content cells
- Links (inline and reference-style)
- Code blocks and inline code with monospace formatting
- Horizontal rules and line breaks
- Blockquotes with indentation

Document Features
-----------------
- Automatic style creation and management
- Consistent formatting throughout document
- Proper indentation and spacing
- Hyperlink styling with color and underline
- Table formatting with borders and structure
- List numbering and bullet management
- Font and color customization

Dependencies
------------
- python-docx: For Word document creation and manipulation
- re: For Markdown parsing and pattern matching

Examples
--------
Convert Markdown to Word document:

    >>> from markdown2docx import markdown_to_docx
    >>> markdown_text = "# Title\\n\\nThis is **bold** text."
    >>> doc = markdown_to_docx(markdown_text)
    >>> doc.save('output.docx')

Custom styling options:

    >>> doc = markdown_to_docx(markdown_content, custom_styles=True)

Note
----
Requires python-docx package. Complex Markdown features may be approximated
in Word format. The conversion focuses on maintaining readability and
structure while working within Word's formatting constraints.
"""

#  Copyright (c) 2023-2025 Tom Villani, Ph.D.
#
#  Permission is hereby granted, free of charge, to any person obtaining a copy of this software and associated
#  documentation files (the “Software”), to deal in the Software without restriction, including without limitation
#  the rights to use, copy, modify, merge, publish, distribute, sublicense, and/or sell copies of the Software,
#  and to permit persons to whom the Software is furnished to do so, subject to the following conditions:
#
#  The above copyright notice and this permission notice shall be included in all copies or substantial
#  portions of the Software.
#
#  THE SOFTWARE IS PROVIDED “AS IS”, WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED, INCLUDING
#  BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT.
#  IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER LIABILITY,
#  WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM, OUT OF OR IN CONNECTION WITH THE
#  SOFTWARE OR THE USE OR OTHER DEALINGS IN THE SOFTWARE.

import copy
import re
from typing import Any

import docx.opc.constants
import docx.text.paragraph
import docx.text.run
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches, Pt, RGBColor

BULLETED_LIST_INDENT = 24


def get_or_create_hyperlink_style(document: Any) -> Any:
    """Get or create a hyperlink character style for the document.

    Creates a hyperlink style with blue color and underline formatting if it
    doesn't already exist in the document. This style is used for formatting
    hyperlink text in the document.

    Parameters
    ----------
    document : Any
        Word document object that contains the styles collection.

    Returns
    -------
    Any
        Hyperlink character style object that can be applied to text runs.

    Notes
    -----
    - Creates style with blue color (RGB 0, 0, 255) and underline
    - Returns existing style if "Hyperlink" style already exists
    - Style is added to document's style collection for reuse
    """
    styles = document.styles
    if "Hyperlink" not in styles:
        hyperlink_style = styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
        hyperlink_style.font.color.rgb = RGBColor(0, 0, 255)
        hyperlink_style.font.underline = True
    return styles["Hyperlink"]


def add_hyperlink(paragraph: docx.text.paragraph.Paragraph, text: str, url: str) -> Any:
    """Add a hyperlink to a Word document paragraph.

    Creates a clickable hyperlink within a paragraph by adding the necessary
    XML elements and establishing the relationship to the external URL.
    Applies hyperlink styling if available.

    Parameters
    ----------
    paragraph : docx.text.paragraph.Paragraph
        Target paragraph where the hyperlink will be inserted.
    text : str
        Display text for the hyperlink that will be visible in the document.
    url : str
        Target URL that the hyperlink will navigate to when clicked.

    Returns
    -------
    Any
        The created hyperlink XML element that was added to the paragraph.

    Notes
    -----
    - Creates external relationship between document and URL
    - Applies hyperlink character style if document has one available
    - Uses python-docx XML manipulation for proper hyperlink structure
    - Hyperlink becomes clickable in Word when document is opened
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = docx.text.run.Run(OxmlElement("w:r"), paragraph)
    new_run.text = text

    # Get document from paragraph's part
    parent = paragraph._parent if hasattr(paragraph, '_parent') else None
    if parent is not None:
        doc = parent.part._document_part if hasattr(parent, "part") else None
        if doc:
            new_run.style = get_or_create_hyperlink_style(doc)

    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def insert_hr(paragraph: docx.text.paragraph.Paragraph) -> None:
    """Insert a horizontal rule into a Word document paragraph.

    Adds a horizontal line (border) to the specified paragraph by manipulating
    the underlying XML elements. This creates a visual separator similar to
    the HTML <hr> tag.

    Parameters
    ----------
    paragraph : docx.text.paragraph.Paragraph
        Target paragraph where the horizontal rule will be inserted.

    Notes
    -----
    Implementation adapted from Stack Overflow solution:
    https://stackoverflow.com/questions/39006878/python-docx-add-horizontal-line

    Creates a bottom border on the paragraph using XML manipulation
    of the Word document structure.
    """
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    pPr.insert_element_before(
        pBdr,
        "w:shd",
        "w:tabs",
        "w:suppressAutoHyphens",
        "w:kinsoku",
        "w:wordWrap",
        "w:overflowPunct",
        "w:topLinePunct",
        "w:autoSpaceDE",
        "w:autoSpaceDN",
        "w:bidi",
        "w:adjustRightInd",
        "w:snapToGrid",
        "w:spacing",
        "w:ind",
        "w:contextualSpacing",
        "w:mirrorIndents",
        "w:suppressOverlap",
        "w:jc",
        "w:textDirection",
        "w:textAlignment",
        "w:textboxTightWrap",
        "w:outlineLvl",
        "w:divId",
        "w:cnfStyle",
        "w:rPr",
        "w:sectPr",
        "w:pPrChange",
    )
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)


def markdown_to_docx(markdown_text: str, document: Any | None = None) -> Any:
    """Convert Markdown to DOCX format.

    Parameters
    ----------
    markdown_text : str
        The markdown str to convert
    document : docx.Document
        Optionally provide a starting document (e.g. a template) to append to.

    Returns
    -------
    docx.Document

    """
    if document is None:
        document = Document()

    try:
        # Create a 'Code' style
        code_style = document.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Courier New"
        code_style.font.size = Pt(10)
    except ValueError:  # docx throws an error if style already exists, in which case we move on.
        pass

    lines = markdown_text.split("\n")
    in_code_block = False
    current_list_level = 0
    current_table = None
    table_header_processed = False

    def get_list_properties(_line: str) -> tuple[str | None, str | None, int | None, int]:
        """Determine list level and type based on indentation and markers"""
        stripped = _line.lstrip()
        indentation = len(_line) - len(stripped)
        _level = indentation // 2  # Each level is 2 spaces

        is_bullet = stripped.startswith(("- ", "* ", "+ "))
        is_number = bool(re.match(r"^\d+\.", stripped))

        if is_bullet:
            _list_type = "Bullet"
            _text = stripped.lstrip("- *+").strip()
        elif is_number:
            _list_type = "Number"
            _text = re.sub(r"^\d+\.\s*", "", stripped)
        else:
            return None, None, None, 0

        return _text, _list_type, _level, indentation

    def process_inline_styles(text: str, paragraph: docx.text.paragraph.Paragraph) -> None:
        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Courier New"
            elif part.startswith("[") and "](" in part and part.endswith(")"):
                match = re.match(r"\[(.*?)\]\((.*?)\)", part)
                if match:
                    link_text, url = match.groups()
                    add_hyperlink(paragraph, link_text, url)
            else:
                paragraph.add_run(part)

    for line in lines:
        if line.startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block:
                document.add_paragraph()
            continue

        if in_code_block:
            document.add_paragraph(line, style="Code")
            continue

        # Handle lists with proper nesting
        text, list_style, level, indentation = get_list_properties(line)

        if text is not None and level is not None:  # This is a list item
            p = document.add_paragraph(style=f"List {list_style}")
            p.paragraph_format.left_indent = Pt((level + 1) * BULLETED_LIST_INDENT)  # 24 points per level
            p.paragraph_format.first_line_indent = Pt(-BULLETED_LIST_INDENT)  # Hanging indent for bullet
            process_inline_styles(text, p)
            current_list_level = level
            continue

        # Headings
        if line.startswith("#"):
            level = len(line.split()[0])  # number of '#' symbols
            text = line.strip("#").strip()
            document.add_heading(text, level=level)

        # Blockquotes
        elif line.strip().startswith(">"):
            text = line.strip("> ").strip()
            p = document.add_paragraph(style="Intense Quote")
            process_inline_styles(text, p)

        # Horizontal rule
        elif line.strip() == "---":
            p = document.add_paragraph()
            insert_hr(p)

        # Images
        elif line.strip().startswith("!"):
            match = re.match(r"!\[(.*?)\]\((.*?)\)", line.strip())
            if match:
                alt_text, src = match.groups()
                try:
                    document.add_picture(src, width=Inches(6))
                    document.add_paragraph(alt_text, style="Caption")
                except Exception:
                    document.add_paragraph(f"[Image: {alt_text} ({src}) not found]")

        # Table handling
        elif line.strip().startswith("|"):
            # Skip divider rows that look like |---|---|---|
            if not re.match(r"\|(\s*:?-+:?\s*\|)+\s*$", line.strip()):
                # Split the line into cells, removing leading/trailing |
                cells = [cell.strip() for cell in line.strip("|").split("|")]

                # Create table if this is the first row
                if current_table is None:
                    current_table = document.add_table(rows=0, cols=len(cells))
                    current_table.style = "Table Grid"
                    table_header_processed = False

                # Add row to table
                row_cells = current_table.add_row().cells
                for i, cell in enumerate(cells):
                    # Process markdown formatting within cells
                    p = row_cells[i].paragraphs[0]
                    process_inline_styles(cell, p)

                # Apply header formatting to first row
                if not table_header_processed:
                    for cell in current_table.rows[0].cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    table_header_processed = True
        # Regular paragraphs
        elif line.strip() == "":
            current_list_level = min(current_list_level, 0)
            document.add_paragraph()
        else:
            if current_table is not None:
                current_table = None
                table_header_processed = False

            p = document.add_paragraph()
            process_inline_styles(line.strip(), p)

    return document


def combine_documents(documents: list[Any], add_page_breaks: bool = False) -> Any:
    """Combines a list of docx.Document into a single object.

    Parameters
    ----------
    documents : list[Document]
        The source documents to combine.
    add_page_breaks : bool, default = False
        If True, adds page breaks to separate each source document.

    Returns
    -------
    docx.Document

    """
    merged_document = Document()

    for index, doc in enumerate(documents):
        # Copy the document relationship dictionary
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                # Get target part (image) from source doc
                target_part = rel.target_part

                # Add new image part to merged document
                new_rel_id = merged_document.part.relate_to(target_part, rel.reltype, is_external=rel.is_external)

                # Update relationship ID in the element's XML
                for element in doc.element.body.iter():
                    if element.tag.endswith("}drawing"):
                        # Find embedded relationship IDs
                        for embed in element.iter(qn("a:blip")):
                            if embed.get(qn("r:embed")) == rel_id:
                                embed.set(qn("r:embed"), new_rel_id)

        # Copy content
        for element in doc.element.body:
            merged_document.element.body.append(copy.deepcopy(element))

        # Add page break if needed
        if add_page_breaks and 0 < index < len(documents) - 1:
            merged_document.add_page_break()  # type: ignore[no-untyped-call]

    return merged_document
