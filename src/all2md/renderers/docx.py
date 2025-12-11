#  Copyright (c) 2025 Tom Villani, Ph.D.
#
# src/all2md/renderers/docx.py
"""DOCX rendering from AST.

This module provides the DocxRenderer class which converts AST nodes
to Microsoft Word (.docx) format. The renderer uses the python-docx library
to generate properly formatted Word documents.

The rendering process uses the visitor pattern to traverse the AST and
generate DOCX content with appropriate styles and formatting.

"""

from __future__ import annotations

import logging
import tempfile
from io import BytesIO
from pathlib import Path
from typing import IO, TYPE_CHECKING, Any, Union
from urllib.parse import urlparse

if TYPE_CHECKING:
    from docx.table import _Cell
    from docx.text.paragraph import Paragraph

from all2md.ast.nodes import (
    BlockQuote,
    Code,
    CodeBlock,
    Comment,
    CommentInline,
    DefinitionDescription,
    DefinitionList,
    DefinitionTerm,
    Emphasis,
    FootnoteDefinition,
    FootnoteReference,
    Heading,
    HTMLBlock,
    HTMLInline,
    Image,
    LineBreak,
    Link,
    List,
    ListItem,
    MathBlock,
    MathInline,
    Node,
    Strikethrough,
    Strong,
    Subscript,
    Superscript,
    Table,
    TableCell,
    TableRow,
    Text,
    ThematicBreak,
    Underline,
)
from all2md.ast.nodes import (
    Document as ASTDocument,
)
from all2md.ast.nodes import (
    Paragraph as ASTParagraph,
)
from all2md.ast.visitors import NodeVisitor
from all2md.constants import DEPS_DOCX_RENDER
from all2md.exceptions import RenderingError
from all2md.options.docx import DocxRendererOptions
from all2md.renderers.base import BaseRenderer
from all2md.utils.decorators import requires_dependencies
from all2md.utils.images import decode_base64_image_to_file
from all2md.utils.network_security import fetch_image_securely, is_network_disabled

logger = logging.getLogger(__name__)


class DocxRenderer(NodeVisitor, BaseRenderer):
    """Render AST nodes to DOCX format.

    This class implements the visitor pattern to traverse an AST and
    generate a Microsoft Word document. It uses python-docx for document
    generation and supports most common formatting features.

    Parameters
    ----------
    options : DocxRendererOptions or None, default = None
        DOCX formatting options

    Examples
    --------
    Basic usage:

        >>> from all2md.ast import Document, Heading, Text
        >>> from all2md.options import DocxRendererOptions
        >>> from all2md.renderers.docx import DocxRenderer
        >>> doc = Document(children=[
        ...     Heading(level=1, content=[Text(content="Title")])
        ... ])
        >>> options = DocxRendererOptions()
        >>> renderer = DocxRenderer(options)
        >>> renderer.render(doc, "output.docx")

    """

    def __init__(self, options: DocxRendererOptions | None = None):
        """Initialize the DOCX renderer with options."""
        BaseRenderer._validate_options_type(options, DocxRendererOptions, "docx")
        options = options or DocxRendererOptions()
        BaseRenderer.__init__(self, options)
        self.options: DocxRendererOptions = options
        self.document: Any = None  # Word document (python-docx Document object)
        self._current_paragraph: Paragraph | None = None
        self._list_level: int = 0
        self._in_table: bool = False
        self._temp_files: list[str] = []
        self._list_ordered_stack: list[bool] = []  # Track ordered/unordered at each level
        self._blockquote_depth: int = 0  # Track blockquote nesting depth

    @requires_dependencies("docx_render", DEPS_DOCX_RENDER)
    def render(self, doc: ASTDocument, output: Union[str, Path, IO[bytes]]) -> None:
        """Render the AST to a DOCX file.

        Parameters
        ----------
        doc : Document
            AST Document node to render
        output : str, Path, or IO[bytes]
            Output destination (file path or file-like object)

        Raises
        ------
        RenderingError
            If DOCX generation fails

        """
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor

        # Store imports as instance variables for use in other methods
        self._Document = Document
        self._WD_STYLE_TYPE = WD_STYLE_TYPE
        self._WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
        self._OxmlElement = OxmlElement
        self._qn = qn
        self._Inches = Inches
        self._Pt = Pt
        self._RGBColor = RGBColor

        try:
            # Create new Word document (with template if specified)
            if self.options.template_path:
                self.document = self._Document(self.options.template_path)
            else:
                self.document = self._Document()

            # Set default font
            self._set_document_defaults()

            # Render document
            doc.accept(self)

            # Save document
            if isinstance(output, (str, Path)):
                self.document.save(str(output))
            else:
                self.document.save(output)
        except Exception as e:
            raise RenderingError(f"Failed to render DOCX: {e!r}", rendering_stage="rendering", original_error=e) from e
        finally:
            # Clean up temp files
            self._cleanup_temp_files()

    @requires_dependencies("docx_render", DEPS_DOCX_RENDER)
    def render_to_bytes(self, doc: ASTDocument) -> bytes:
        """Render the AST to DOCX bytes.

        Parameters
        ----------
        doc : Document
            AST Document node to render

        Returns
        -------
        bytes
            DOCX file content as bytes

        Raises
        ------
        RenderingError
            If DOCX generation fails

        """
        # Create a BytesIO buffer and render to it
        buffer = BytesIO()
        self.render(doc, buffer)

        # Return the bytes content
        return buffer.getvalue()

    def _set_document_defaults(self) -> None:
        """Set default document styles and formatting."""
        if not self.document:
            return

        # Set default font for Normal style
        style = self.document.styles["Normal"]
        font = style.font
        font.name = self.options.default_font
        font.size = self._Pt(self.options.default_font_size)

        # Set creator metadata if configured
        if self.options.creator:
            self.document.core_properties.last_modified_by = self.options.creator

    def _cleanup_temp_files(self) -> None:
        """Remove temporary files created during rendering."""
        for temp_file in self._temp_files:
            try:
                Path(temp_file).unlink(missing_ok=True)
            except Exception as e:
                logger.debug(f"Failed to cleanup temp file {temp_file}: {e}")
        self._temp_files.clear()

    def visit_document(self, node: ASTDocument) -> None:
        """Render a Document node.

        Parameters
        ----------
        node : Document
            Document to render

        """
        # Render metadata as document properties if present
        if node.metadata:
            self._set_document_properties(node.metadata)

        # Render children
        for child in node.children:
            child.accept(self)

    def _set_document_properties(self, metadata: dict) -> None:
        """Set document properties from metadata.

        Parameters
        ----------
        metadata : dict
            Document metadata

        """
        if not self.document:
            return

        core_props = self.document.core_properties
        if "title" in metadata:
            core_props.title = str(metadata["title"])
        if "author" in metadata:
            core_props.author = str(metadata["author"])
        if "subject" in metadata:
            core_props.subject = str(metadata["subject"])
        if "keywords" in metadata:
            keywords = metadata["keywords"]
            if isinstance(keywords, list):
                core_props.keywords = ", ".join(str(k) for k in keywords)
            else:
                core_props.keywords = str(keywords)

    def visit_heading(self, node: Heading) -> None:
        """Render a Heading node.

        Parameters
        ----------
        node : Heading
            Heading to render

        """
        if not self.document:
            return

        # Add heading with appropriate level
        level = min(9, max(1, node.level))  # Word supports levels 1-9

        if self.options.use_styles:
            # Use built-in heading styles
            heading = self.document.add_heading(level=level)
            # Clear the heading text (add_heading adds empty text)
            heading.text = ""
        else:
            # Use direct formatting
            heading = self.document.add_paragraph()

        # Apply blockquote indentation if inside a blockquote
        if self._blockquote_depth > 0:
            heading.paragraph_format.left_indent = self._Inches(0.5 * self._blockquote_depth)

        self._current_paragraph = heading

        # Render content
        for child in node.content:
            child.accept(self)

        # Apply direct formatting AFTER content is rendered (when use_styles=False)
        if not self.options.use_styles:
            if self.options.heading_font_sizes and level in self.options.heading_font_sizes:
                for run in heading.runs:
                    run.font.size = self._Pt(self.options.heading_font_sizes[level])
                    run.font.bold = True
            else:
                # No custom sizes specified - just make it bold
                for run in heading.runs:
                    run.font.bold = True

        self._current_paragraph = None

    def visit_paragraph(self, node: ASTParagraph) -> None:
        """Render a Paragraph node.

        Parameters
        ----------
        node : Paragraph
            Paragraph to render

        """
        if not self.document:
            return

        # Don't create new paragraph if we're already in one (e.g., heading)
        if self._current_paragraph is None:
            self._current_paragraph = self.document.add_paragraph()

            # Apply blockquote indentation if inside a blockquote
            if self._blockquote_depth > 0:
                self._current_paragraph.paragraph_format.left_indent = self._Inches(0.5 * self._blockquote_depth)

        # Render content
        for child in node.content:
            child.accept(self)

        # Only clear if we created the paragraph
        if not self._in_table:
            self._current_paragraph = None

    def visit_code_block(self, node: CodeBlock) -> None:
        """Render a CodeBlock node.

        Parameters
        ----------
        node : CodeBlock
            Code block to render

        """
        if not self.document:
            return

        # Add paragraph with code formatting
        para = self.document.add_paragraph()
        run = para.add_run(node.content)
        run.font.name = self.options.code_font
        run.font.size = self._Pt(self.options.code_font_size)

        # Apply blockquote indentation if inside a blockquote
        if self._blockquote_depth > 0:
            para.paragraph_format.left_indent = self._Inches(0.5 * self._blockquote_depth)

        # Set paragraph background (light gray)
        self._set_paragraph_shading(para, "F0F0F0")

    def _set_paragraph_shading(self, paragraph: Paragraph, color: str) -> None:
        """Set paragraph background color.

        Parameters
        ----------
        paragraph : Paragraph
            Paragraph to shade
        color : str
            Hex color code (e.g., "F0F0F0")

        """
        shading_elm = self._OxmlElement("w:shd")
        shading_elm.set(self._qn("w:fill"), color)
        paragraph._element.get_or_add_pPr().append(shading_elm)

    def visit_block_quote(self, node: BlockQuote) -> None:
        """Render a BlockQuote node.

        Parameters
        ----------
        node : BlockQuote
            Block quote to render

        """
        # Increase blockquote depth for indentation tracking
        self._blockquote_depth += 1

        # Render children (they will check _blockquote_depth and indent themselves)
        for child in node.children:
            child.accept(self)

        # Decrease blockquote depth
        self._blockquote_depth -= 1

    def visit_list(self, node: List) -> None:
        """Render a List node.

        Parameters
        ----------
        node : List
            List to render

        """
        self._list_level += 1
        self._list_ordered_stack.append(node.ordered)

        for _i, item in enumerate(node.items):
            item.accept(self)

        self._list_level -= 1
        self._list_ordered_stack.pop()

    def visit_list_item(self, node: ListItem) -> None:
        """Render a ListItem node.

        Parameters
        ----------
        node : ListItem
            List item to render

        """
        if not self.document:
            return

        # Create paragraph for list item
        para = self.document.add_paragraph()

        # Determine list style based on ordered/unordered
        is_ordered = self._list_ordered_stack[-1] if self._list_ordered_stack else False
        if is_ordered:
            para.style = "List Number"
        else:
            para.style = "List Bullet"

        self._current_paragraph = para

        # Render children
        # Special handling: first Paragraph child should render inline with the list marker,
        # while subsequent Paragraphs create new paragraphs (for multi-paragraph list items)
        is_first_paragraph = True
        for child in node.children:
            if isinstance(child, ASTParagraph):
                if is_first_paragraph:
                    # First paragraph: render inline content directly into list item's paragraph
                    is_first_paragraph = False
                    # Render the paragraph's inline content directly
                    for inline_child in child.content:
                        inline_child.accept(self)
                else:
                    # Subsequent paragraphs: create new paragraphs (multi-paragraph list items)
                    saved_para = self._current_paragraph
                    self._current_paragraph = None
                    child.accept(self)
                    self._current_paragraph = saved_para
            elif isinstance(child, List):
                # Nested lists: handle normally
                saved_para = self._current_paragraph
                self._current_paragraph = None
                child.accept(self)
                self._current_paragraph = saved_para
            else:
                # Other inline content (should be rare, but handle it)
                child.accept(self)

        self._current_paragraph = None

    def visit_table(self, node: Table) -> None:
        """Render a Table node.

        Parameters
        ----------
        node : Table
            Table to render

        """
        if not self.document:
            return

        # Collect all rows
        all_rows = []
        if node.header:
            all_rows.append(node.header)
        all_rows.extend(node.rows)

        if not all_rows:
            return

        # Compute grid dimensions accounting for colspan/rowspan
        num_rows = len(all_rows)
        num_cols = self._compute_table_columns(all_rows)

        if num_cols == 0:
            return

        # Create table with proper dimensions
        table = self.document.add_table(rows=num_rows, cols=num_cols)

        # Apply table style if requested
        if self.options.table_style:
            table.style = self.options.table_style

        self._in_table = True

        # Track which grid cells are occupied by spanning cells
        occupied = [[False] * num_cols for _ in range(num_rows)]

        # Render all rows
        for row_idx, ast_row in enumerate(all_rows):
            col_idx = 0
            for ast_cell in ast_row.cells:
                # Skip occupied cells
                while col_idx < num_cols and occupied[row_idx][col_idx]:
                    col_idx += 1

                if col_idx >= num_cols:
                    break

                # Render cell content
                docx_cell = table.rows[row_idx].cells[col_idx]
                is_header = row_idx == 0 and node.header is not None
                self._render_table_cell(docx_cell, ast_cell, is_header=is_header)

                # Handle cell spanning
                colspan = ast_cell.colspan
                rowspan = ast_cell.rowspan

                # Mark occupied cells
                for r in range(row_idx, min(row_idx + rowspan, num_rows)):
                    for c in range(col_idx, min(col_idx + colspan, num_cols)):
                        occupied[r][c] = True

                # Merge cells if needed
                if colspan > 1 or rowspan > 1:
                    end_row = min(row_idx + rowspan - 1, num_rows - 1)
                    end_col = min(col_idx + colspan - 1, num_cols - 1)
                    if end_row > row_idx or end_col > col_idx:
                        docx_cell.merge(table.rows[end_row].cells[end_col])

                col_idx += colspan

        self._in_table = False

    def _render_table_cell(self, docx_cell: _Cell, ast_cell: TableCell, is_header: bool = False) -> None:
        """Render a single table cell.

        Parameters
        ----------
        docx_cell : _Cell
            python-docx table cell
        ast_cell : TableCell
            AST table cell node
        is_header : bool, default = False
            Whether this is a header cell

        """
        # Clear default paragraph
        if len(docx_cell.paragraphs) > 0:
            docx_cell.paragraphs[0].text = ""
            self._current_paragraph = docx_cell.paragraphs[0]
        else:
            self._current_paragraph = docx_cell.add_paragraph()

        # Render cell content
        for child in ast_cell.content:
            child.accept(self)

        # Make header cells bold
        if is_header:
            for para in docx_cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True

        self._current_paragraph = None

    def visit_table_row(self, node: TableRow) -> None:
        """Render a TableRow node.

        Parameters
        ----------
        node : TableRow
            Table row to render

        """
        # Handled by visit_table
        pass

    def visit_table_cell(self, node: TableCell) -> None:
        """Render a TableCell node.

        Parameters
        ----------
        node : TableCell
            Table cell to render

        """
        # Handled by visit_table
        pass

    def visit_thematic_break(self, node: ThematicBreak) -> None:
        """Render a ThematicBreak node.

        Parameters
        ----------
        node : ThematicBreak
            Thematic break to render

        """
        if not self.document:
            return

        # Add horizontal line using a simple text separator
        para = self.document.add_paragraph()
        run = para.add_run("─" * 78)  # Box drawing horizontal character
        run.font.color.rgb = self._RGBColor(192, 192, 192)  # Light gray

    def visit_html_block(self, node: HTMLBlock) -> None:
        """Render an HTMLBlock node.

        Parameters
        ----------
        node : HTMLBlock
            HTML block to render

        """
        # Skip HTML content in DOCX
        pass

    def _extract_text_from_inlines(self, nodes: list[Node]) -> str:
        """Extract plain text from inline nodes without creating temporary elements.

        This method recursively extracts text content from inline nodes,
        providing an efficient alternative to rendering to a temporary paragraph.

        Parameters
        ----------
        nodes : list of Node
            Inline nodes to extract text from

        Returns
        -------
        str
            Plain text content

        """
        text_parts: list[str] = []

        def collect_text(node_list: list[Node]) -> None:
            """Recursively collect text from nodes."""
            for node in node_list:
                if isinstance(node, Text):
                    text_parts.append(node.content)
                elif isinstance(node, Code):
                    text_parts.append(node.content)
                elif hasattr(node, "content"):
                    if isinstance(node.content, list):
                        collect_text(node.content)
                    elif isinstance(node.content, str):
                        text_parts.append(node.content)

        collect_text(nodes)
        return "".join(text_parts)

    def _render_inlines(
        self,
        paragraph: Paragraph,
        nodes: list[Node],
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        strike: bool = False,
        superscript: bool = False,
        subscript: bool = False,
        code_font: bool = False,
    ) -> None:
        """Render inline nodes directly into a paragraph with formatting.

        This method efficiently renders inline content by applying formatting
        flags recursively, avoiding the overhead of creating temporary documents
        and paragraphs.

        Parameters
        ----------
        paragraph : Paragraph
            Target paragraph to render into
        nodes : list of Node
            Inline nodes to render
        bold : bool, default = False
            Apply bold formatting
        italic : bool, default = False
            Apply italic formatting
        underline : bool, default = False
            Apply underline formatting
        strike : bool, default = False
            Apply strikethrough formatting
        superscript : bool, default = False
            Apply superscript formatting
        subscript : bool, default = False
            Apply subscript formatting
        code_font : bool, default = False
            Apply code font formatting

        """
        for node in nodes:
            if isinstance(node, Text):
                # Create run with text and apply all formatting
                run = paragraph.add_run(node.content)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
                if underline:
                    run.underline = True
                if strike:
                    run.font.strike = True
                if superscript:
                    run.font.superscript = True
                if subscript:
                    run.font.subscript = True
                if code_font:
                    run.font.name = self.options.code_font
                    run.font.size = self._Pt(self.options.code_font_size)

            elif isinstance(node, Strong):
                # Recursively render with bold flag
                self._render_inlines(
                    paragraph,
                    node.content,
                    bold=True,
                    italic=italic,
                    underline=underline,
                    strike=strike,
                    superscript=superscript,
                    subscript=subscript,
                    code_font=code_font,
                )

            elif isinstance(node, Emphasis):
                # Recursively render with italic flag
                self._render_inlines(
                    paragraph,
                    node.content,
                    bold=bold,
                    italic=True,
                    underline=underline,
                    strike=strike,
                    superscript=superscript,
                    subscript=subscript,
                    code_font=code_font,
                )

            elif isinstance(node, Underline):
                # Recursively render with underline flag
                self._render_inlines(
                    paragraph,
                    node.content,
                    bold=bold,
                    italic=italic,
                    underline=True,
                    strike=strike,
                    superscript=superscript,
                    subscript=subscript,
                    code_font=code_font,
                )

            elif isinstance(node, Strikethrough):
                # Recursively render with strike flag
                self._render_inlines(
                    paragraph,
                    node.content,
                    bold=bold,
                    italic=italic,
                    underline=underline,
                    strike=True,
                    superscript=superscript,
                    subscript=subscript,
                    code_font=code_font,
                )

            elif isinstance(node, Superscript):
                # Recursively render with superscript flag
                self._render_inlines(
                    paragraph,
                    node.content,
                    bold=bold,
                    italic=italic,
                    underline=underline,
                    strike=strike,
                    superscript=True,
                    subscript=subscript,
                    code_font=code_font,
                )

            elif isinstance(node, Subscript):
                # Recursively render with subscript flag
                self._render_inlines(
                    paragraph,
                    node.content,
                    bold=bold,
                    italic=italic,
                    underline=underline,
                    strike=strike,
                    superscript=superscript,
                    subscript=True,
                    code_font=code_font,
                )

            elif isinstance(node, Code):
                # Render as code with code font
                self._render_inlines(
                    paragraph,
                    [Text(content=node.content)],
                    bold=bold,
                    italic=italic,
                    underline=underline,
                    strike=strike,
                    superscript=superscript,
                    subscript=subscript,
                    code_font=True,
                )

            elif isinstance(node, Link):
                # Extract link text using efficient in-memory extraction
                link_text = self._extract_text_from_inlines(node.content)
                # Add hyperlink to target paragraph
                self._add_hyperlink(paragraph, node.url, link_text)

            elif isinstance(node, LineBreak):
                paragraph.add_run().add_break()

            else:
                # For other inline nodes, try to process if they have content
                if hasattr(node, "content") and isinstance(node.content, list):
                    self._render_inlines(
                        paragraph,
                        node.content,
                        bold=bold,
                        italic=italic,
                        underline=underline,
                        strike=strike,
                        superscript=superscript,
                        subscript=subscript,
                        code_font=code_font,
                    )

    def visit_text(self, node: Text) -> None:
        """Render a Text node.

        Parameters
        ----------
        node : Text
            Text to render

        """
        if self._current_paragraph is None:
            if self.document:
                self._current_paragraph = self.document.add_paragraph()

        if self._current_paragraph:
            self._current_paragraph.add_run(node.content)

    def visit_emphasis(self, node: Emphasis) -> None:
        """Render an Emphasis node.

        Parameters
        ----------
        node : Emphasis
            Emphasis to render

        """
        if self._current_paragraph is None:
            if self.document:
                self._current_paragraph = self.document.add_paragraph()

        # Use efficient inline rendering
        if self._current_paragraph:
            self._render_inlines(self._current_paragraph, node.content, italic=True)

    def visit_strong(self, node: Strong) -> None:
        """Render a Strong node.

        Parameters
        ----------
        node : Strong
            Strong to render

        """
        if self._current_paragraph is None:
            if self.document:
                self._current_paragraph = self.document.add_paragraph()

        # Use efficient inline rendering
        if self._current_paragraph:
            self._render_inlines(self._current_paragraph, node.content, bold=True)

    def visit_code(self, node: Code) -> None:
        """Render a Code node.

        Parameters
        ----------
        node : Code
            Code to render

        """
        if self._current_paragraph is None:
            if self.document:
                self._current_paragraph = self.document.add_paragraph()

        # Use efficient inline rendering with code font
        if self._current_paragraph:
            self._render_inlines(self._current_paragraph, [Text(content=node.content)], code_font=True)

    def visit_link(self, node: Link) -> None:
        """Render a Link node.

        Parameters
        ----------
        node : Link
            Link to render

        """
        if self._current_paragraph is None:
            if self.document:
                self._current_paragraph = self.document.add_paragraph()

        # Use efficient inline rendering (handles links internally)
        if self._current_paragraph:
            self._render_inlines(self._current_paragraph, [node])

    def _add_hyperlink(self, paragraph: Paragraph, url: str, text: str) -> None:
        """Add a hyperlink to a paragraph.

        This method uses OOXML private APIs to construct hyperlinks, as python-docx
        does not provide a high-level API for this functionality. The implementation
        may be brittle across python-docx versions.

        Parameters
        ----------
        paragraph : Paragraph
            Paragraph to add link to
        url : str
            URL to link to
        text : str
            Link text

        """
        # This is a complex operation in python-docx requiring XML manipulation
        part = paragraph.part
        r_id = part.relate_to(
            url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
        )

        # Create hyperlink element
        hyperlink = self._OxmlElement("w:hyperlink")
        hyperlink.set(self._qn("r:id"), r_id)

        # Create run element
        new_run = self._OxmlElement("w:r")
        rPr = self._OxmlElement("w:rPr")

        # Add hyperlink style
        r_style = self._OxmlElement("w:rStyle")
        r_style.set(self._qn("w:val"), "Hyperlink")
        rPr.append(r_style)
        new_run.append(rPr)

        # Create text element with xml:space="preserve" to prevent whitespace collapse
        t = self._OxmlElement("w:t")
        t.set(self._qn("xml:space"), "preserve")
        t.text = text
        new_run.append(t)

        hyperlink.append(new_run)
        paragraph._element.append(hyperlink)

    def visit_image(self, node: Image) -> None:
        """Render an Image node.

        Parameters
        ----------
        node : Image
            Image to render

        """
        if not self.document or not node.url:
            return

        try:
            # Handle different image sources
            if node.url.startswith("data:"):
                # Base64 encoded image
                image_file = self._decode_base64_image(node.url)
            elif urlparse(node.url).scheme in ("http", "https"):
                # Remote URL - use secure fetching if enabled
                image_file = self._fetch_remote_image(node.url)
            else:
                # Local file
                image_file = node.url

            # Add image to document
            if image_file:
                para = self.document.add_paragraph()
                run = para.add_run()
                run.add_picture(image_file, width=self._Inches(4))

                # Add caption if alt text exists
                if node.alt_text:
                    caption_para = self.document.add_paragraph(node.alt_text)
                    caption_para.alignment = self._WD_ALIGN_PARAGRAPH.CENTER
                    caption_para.runs[0].italic = True
        except Exception as e:
            # If image loading fails, log and optionally raise
            logger.warning(f"Failed to add image to DOCX: {e}")
            if self.options.fail_on_resource_errors:
                raise RenderingError(
                    f"Failed to add image to DOCX: {e!r}", rendering_stage="image_processing", original_error=e
                ) from e

    def _decode_base64_image(self, data_uri: str) -> str | None:
        """Decode base64 image to temporary file.

        Parameters
        ----------
        data_uri : str
            Data URI with base64 encoded image

        Returns
        -------
        str or None
            Path to temporary file, or None if decoding failed

        """
        # Use centralized image utility
        temp_path = decode_base64_image_to_file(data_uri, delete_on_exit=False)
        if temp_path:
            self._temp_files.append(temp_path)
        return temp_path

    def _fetch_remote_image(self, url: str) -> str | None:
        """Fetch remote image securely.

        Parameters
        ----------
        url : str
            Remote image URL

        Returns
        -------
        str or None
            Path to temporary file containing the image, or None if fetching failed

        """
        if is_network_disabled():
            logger.debug(f"Network disabled, skipping remote image: {url}")
            return None

        if not self.options.network.allow_remote_fetch:
            logger.debug(f"Remote fetching disabled, skipping image: {url}")
            return None

        try:
            image_data = fetch_image_securely(
                url=url,
                allowed_hosts=self.options.network.allowed_hosts,
                require_https=self.options.network.require_https,
                max_size_bytes=self.options.max_asset_size_bytes,
                timeout=self.options.network.network_timeout,
                require_head_success=self.options.network.require_head_success,
            )

            # Determine extension from URL
            parsed = urlparse(url)
            path_lower = parsed.path.lower()
            if path_lower.endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".svg")):
                ext = path_lower.split(".")[-1]
            else:
                ext = "png"

            # Write to temp file
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as f:
                f.write(image_data)
                temp_path = f.name

            self._temp_files.append(temp_path)
            return temp_path

        except Exception as e:
            logger.warning(f"Failed to fetch remote image {url}: {e}")
            if self.options.fail_on_resource_errors:
                raise RenderingError(
                    f"Failed to fetch remote image {url}: {e!r}", rendering_stage="image_processing", original_error=e
                ) from e
            return None

    def visit_line_break(self, node: LineBreak) -> None:
        """Render a LineBreak node.

        Parameters
        ----------
        node : LineBreak
            Line break to render

        """
        if self._current_paragraph:
            self._current_paragraph.add_run().add_break()

    def visit_strikethrough(self, node: Strikethrough) -> None:
        """Render a Strikethrough node.

        Parameters
        ----------
        node : Strikethrough
            Strikethrough to render

        """
        if self._current_paragraph is None:
            if self.document:
                self._current_paragraph = self.document.add_paragraph()

        # Use efficient inline rendering
        if self._current_paragraph:
            self._render_inlines(self._current_paragraph, node.content, strike=True)

    def visit_underline(self, node: Underline) -> None:
        """Render an Underline node.

        Parameters
        ----------
        node : Underline
            Underline to render

        """
        if self._current_paragraph is None:
            if self.document:
                self._current_paragraph = self.document.add_paragraph()

        # Use efficient inline rendering
        if self._current_paragraph:
            self._render_inlines(self._current_paragraph, node.content, underline=True)

    def visit_superscript(self, node: Superscript) -> None:
        """Render a Superscript node.

        Parameters
        ----------
        node : Superscript
            Superscript to render

        """
        if self._current_paragraph is None:
            if self.document:
                self._current_paragraph = self.document.add_paragraph()

        # Use efficient inline rendering
        if self._current_paragraph:
            self._render_inlines(self._current_paragraph, node.content, superscript=True)

    def visit_subscript(self, node: Subscript) -> None:
        """Render a Subscript node.

        Parameters
        ----------
        node : Subscript
            Subscript to render

        """
        if self._current_paragraph is None:
            if self.document:
                self._current_paragraph = self.document.add_paragraph()

        # Use efficient inline rendering
        if self._current_paragraph:
            self._render_inlines(self._current_paragraph, node.content, subscript=True)

    def visit_html_inline(self, node: HTMLInline) -> None:
        """Render an HTMLInline node.

        Parameters
        ----------
        node : HTMLInline
            Inline HTML to render

        """
        # Skip inline HTML in DOCX
        pass

    def visit_footnote_reference(self, node: FootnoteReference) -> None:
        """Render a FootnoteReference node.

        Parameters
        ----------
        node : FootnoteReference
            Footnote reference to render

        """
        # Footnote references could be rendered as endnotes in Word
        # For now, render as superscript text
        if self._current_paragraph is None:
            if self.document:
                self._current_paragraph = self.document.add_paragraph()

        if self._current_paragraph:
            run = self._current_paragraph.add_run(f"[{node.identifier}]")
            run.font.superscript = True

    def visit_math_inline(self, node: MathInline) -> None:
        """Render a MathInline node.

        Parameters
        ----------
        node : MathInline
            Inline math to render

        """
        # Render math with best available representation (fallback to plain text)
        if self._current_paragraph is None:
            if self.document:
                self._current_paragraph = self.document.add_paragraph()

        if self._current_paragraph:
            content, notation = node.get_preferred_representation("latex")
            if notation == "latex":
                text = f"${content}$"
            else:
                text = content

            self._current_paragraph.add_run(text)

    def visit_footnote_definition(self, node: FootnoteDefinition) -> None:
        """Render a FootnoteDefinition node.

        Parameters
        ----------
        node : FootnoteDefinition
            Footnote definition to render

        """
        # Render as a separate paragraph
        if not self.document:
            return

        para = self.document.add_paragraph()
        run = para.add_run(f"[{node.identifier}]: ")
        run.font.superscript = True

        self._current_paragraph = para
        for child in node.content:
            child.accept(self)
        self._current_paragraph = None

    def visit_definition_list(self, node: DefinitionList) -> None:
        """Render a DefinitionList node.

        Parameters
        ----------
        node : DefinitionList
            Definition list to render

        """
        for term, descriptions in node.items:
            # Render term as bold paragraph
            if self.document:
                term_para = self.document.add_paragraph()
                self._current_paragraph = term_para

                for child in term.content:
                    child.accept(self)

                # Make term bold
                for run in term_para.runs:
                    run.bold = True

                # Render descriptions as indented paragraphs
                for desc in descriptions:
                    desc_para = self.document.add_paragraph()
                    desc_para.paragraph_format.left_indent = self._Inches(0.5)
                    self._current_paragraph = desc_para

                    for child in desc.content:
                        child.accept(self)

                self._current_paragraph = None

    def visit_definition_term(self, node: DefinitionTerm) -> None:
        """Render a DefinitionTerm node.

        Parameters
        ----------
        node : DefinitionTerm
            Definition term to render

        """
        # Handled by visit_definition_list
        pass

    def visit_definition_description(self, node: DefinitionDescription) -> None:
        """Render a DefinitionDescription node.

        Parameters
        ----------
        node : DefinitionDescription
            Definition description to render

        """
        # Handled by visit_definition_list
        pass

    def visit_math_block(self, node: MathBlock) -> None:
        """Render a MathBlock node.

        Parameters
        ----------
        node : MathBlock
            Math block to render

        """
        # Render math as code block (proper equation rendering is complex)
        if not self.document:
            return

        para = self.document.add_paragraph()
        content, notation = node.get_preferred_representation("latex")
        if notation == "latex":
            text = f"$$\n{content}\n$$"
        else:
            text = content

        run = para.add_run(text)
        run.font.name = self.options.code_font
        run.font.size = self._Pt(self.options.code_font_size)

    def visit_comment(self, node: Comment) -> None:
        """Render a Comment node (block-level).

        Parameters
        ----------
        node : Comment
            Comment block to render

        """
        if not self.document:
            return

        # Check comment_mode option
        comment_mode = self.options.comment_mode

        if comment_mode == "ignore":
            # Skip rendering comment entirely
            return

        # Extract comment metadata
        author = node.metadata.get("author", "")
        date = node.metadata.get("date", "")
        label = node.metadata.get("label", "")
        comment_type = node.metadata.get("comment_type", "")

        if comment_mode == "visible":
            # Render as visible text paragraph with attribution
            para = self.document.add_paragraph()

            # Build attribution prefix
            prefix_parts = []
            if comment_type:
                prefix_parts.append(comment_type.upper())
            if label:
                prefix_parts.append(f"#{label}")

            prefix = " ".join(prefix_parts) if prefix_parts else "Comment"

            # Add attribution
            if author:
                if date:
                    attribution = f"{prefix} by {author} ({date}):"
                else:
                    attribution = f"{prefix} by {author}:"

                # Add attribution as bold
                run = para.add_run(attribution + " ")
                run.font.bold = True

            # Add comment content
            run = para.add_run(node.content)
            run.font.italic = True

            return

        # Mode is "native" - try to use native DOCX comments if supported
        try:
            # Create a paragraph to attach the comment to
            para = self.document.add_paragraph(node.content)

            # Extract comment metadata
            author = node.metadata.get("author", "")
            initials = node.metadata.get("initials", "")
            if not initials and author:
                # Generate initials from author name
                parts = author.split()
                initials = "".join(p[0].upper() for p in parts if p)[:3]

            # Add native DOCX comment using python-docx 1.2.0+ API
            if hasattr(self.document, "add_comment"):
                self.document.add_comment(
                    runs=para.runs,
                    text=node.content,
                    author=author,
                    initials=initials,
                )
                # Clear the paragraph content since comment is attached
                para.text = ""
            else:
                # Fallback: render as highlighted text with comment styling
                self._render_comment_fallback(para, node, is_inline=False)

        except Exception as e:
            logger.warning(f"Failed to add block comment to DOCX: {e}")
            # Fallback: render as styled paragraph
            if self.document:
                para = self.document.add_paragraph()
                run = para.add_run(f"[Comment: {node.content}]")
                run.font.italic = True
                run.font.color.rgb = self._RGBColor(128, 128, 128)

    def visit_comment_inline(self, node: CommentInline) -> None:
        """Render a CommentInline node.

        Parameters
        ----------
        node : CommentInline
            Inline comment to render

        """
        if not self.document or self._current_paragraph is None:
            return

        # Check comment_mode option
        comment_mode = self.options.comment_mode

        if comment_mode == "ignore":
            # Skip rendering comment entirely
            return

        # Extract comment metadata
        author = node.metadata.get("author", "")
        date = node.metadata.get("date", "")
        label = node.metadata.get("label", "")
        comment_type = node.metadata.get("comment_type", "")

        if comment_mode == "visible":
            # Render as visible inline text with attribution
            # Build attribution prefix
            prefix_parts = []
            if comment_type:
                prefix_parts.append(comment_type.upper())
            if label:
                prefix_parts.append(f"#{label}")

            prefix = " ".join(prefix_parts) if prefix_parts else "Comment"

            # Build full text
            if author:
                if date:
                    full_text = f"[{prefix} by {author} ({date}): {node.content}]"
                else:
                    full_text = f"[{prefix} by {author}: {node.content}]"
            else:
                full_text = f"[{node.content}]"

            # Add as italic run
            run = self._current_paragraph.add_run(full_text)
            run.font.italic = True

            return

        # Mode is "native" - try to use native DOCX comments if supported
        try:
            # Extract comment metadata
            author = node.metadata.get("author", "")
            initials = node.metadata.get("initials", "")
            if not initials and author:
                # Generate initials from author name
                parts = author.split()
                initials = "".join(p[0].upper() for p in parts if p)[:3]

            # For inline comments, we need to create a temporary run to attach the comment to
            # Add a placeholder character that will be commented
            placeholder_run = self._current_paragraph.add_run("\u200b")  # Zero-width space

            # Add native DOCX comment using python-docx 1.2.0+ API
            if hasattr(self.document, "add_comment"):
                self.document.add_comment(
                    runs=[placeholder_run],
                    text=node.content,
                    author=author,
                    initials=initials,
                )
            else:
                # Fallback: render as highlighted text with comment styling
                self._render_comment_fallback(self._current_paragraph, node, is_inline=True)

        except Exception as e:
            logger.warning(f"Failed to add inline comment to DOCX: {e}")
            # Fallback: render as styled text
            if self._current_paragraph:
                run = self._current_paragraph.add_run(f"[{node.content}]")
                run.font.italic = True
                run.font.color.rgb = self._RGBColor(128, 128, 128)

    def _render_comment_fallback(self, paragraph: Paragraph, node: Comment | CommentInline, is_inline: bool) -> None:
        """Render comment as styled text when native comments are not supported.

        This fallback method renders comments as highlighted text with metadata,
        providing a visual alternative when python-docx doesn't support comments.

        Parameters
        ----------
        paragraph : Paragraph
            Target paragraph to render into
        node : Comment or CommentInline
            Comment node to render
        is_inline : bool
            Whether this is an inline comment

        """
        # Build comment text with metadata
        comment_text = node.content
        author = node.metadata.get("author", "")
        date = node.metadata.get("date", "")
        label = node.metadata.get("label", "")

        # Format comment with metadata
        if author or date or label:
            prefix_parts = []
            if label:
                prefix_parts.append(f"Comment {label}")
            else:
                prefix_parts.append("Comment")

            if author:
                prefix_parts.append(f"by {author}")

            if date:
                prefix_parts.append(f"({date})")

            prefix = " ".join(prefix_parts)
            comment_text = f"[{prefix}: {comment_text}]"
        else:
            comment_text = f"[Comment: {comment_text}]"

        # Add styled run
        run = paragraph.add_run(comment_text)
        run.font.italic = True
        run.font.color.rgb = self._RGBColor(255, 165, 0)  # Orange color for visibility

        # Add highlight/background if not inline
        if not is_inline:
            # Set paragraph shading for block comments
            self._set_paragraph_shading(paragraph, "FFF8DC")  # Cornsilk color
