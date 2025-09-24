"""Word document to Markdown conversion module.

This module provides functionality to convert Microsoft Word documents (DOCX format)
to Markdown while preserving formatting, structure, and embedded content. It handles
complex document elements including tables, lists, hyperlinks, images, and various
text formatting options.

The converter processes Word documents by analyzing their internal structure,
extracting styled content, and converting it to equivalent Markdown syntax.
Special attention is paid to preserving document hierarchy, list structures,
and table layouts.

Key Features
------------
- Comprehensive text formatting preservation (bold, italic, underline, etc.)
- Intelligent list detection and conversion (bulleted and numbered)
- Table structure preservation with Markdown table syntax
- Hyperlink extraction and conversion
- Image embedding as base64 data URIs
- Document structure preservation (headers, paragraphs)
- Style-based formatting detection

Supported Elements
------------------
- Text formatting (bold, italic, underline, strikethrough)
- Headers (converted to Markdown headers)
- Lists (bulleted, numbered, with proper nesting)
- Tables (with cell content and basic formatting)
- Hyperlinks (internal and external)
- Images (embedded as base64)
- Line breaks and paragraph separation

Dependencies
------------
- python-docx: For parsing Word document structure
- logging: For debug and error reporting

Examples
--------
Basic conversion:

    >>> from all2md.docx2markdown import docx_to_markdown
    >>> with open('document.docx', 'rb') as f:
    ...     markdown = docx_to_markdown(f)
    >>> print(markdown)

Note
----
Requires python-docx package. Some advanced Word features may not have
direct Markdown equivalents and will be approximated or omitted.
"""

#  Copyright (c) 2023-2025 Tom Villani, Ph.D.
#
#  Permission is hereby granted, free of charge, to any person obtaining a copy of this software and associated
#  documentation files (the “Software”), to deal in the Software without restriction, including without limitation
#  the rights to use, copy, modify, merge, publish, distribute, sublicense, and/or sell copies of the Software,
#  and to permit persons to whom the Software is furnished to do so, subject to the following conditions:
#
#  The above copyright notice and this permission notice shall be included in all copies or substantial
#  portions of the Software.
#
#  THE SOFTWARE IS PROVIDED “AS IS”, WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED, INCLUDING
#  BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT.
#  IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER LIABILITY,
#  WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM, OUT OF OR IN CONNECTION WITH THE
#  SOFTWARE OR THE USE OR OTHER DEALINGS IN THE SOFTWARE.

import logging
import re
from pathlib import Path
from typing import IO, Any, Union

import docx
import docx.document
from docx.table import Table
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph

from ._attachment_utils import extract_docx_image_data, process_attachment
from .constants import DEFAULT_INDENTATION_PT_PER_LEVEL
from .exceptions import MdparseConversionError
from .options import DocxOptions, MarkdownOptions

logger = logging.getLogger(__name__)


def _detect_list_level(paragraph: Paragraph) -> tuple[str | None, int]:
    """Detect the list level of a paragraph based on its style, numbering, and indentation.

    Returns tuple of (list_type, level) where list_type is 'bullet' or 'number' and level is integer depth
    """
    if not paragraph.style or not paragraph.style.name:
        # Check for Word native numbering properties
        if hasattr(paragraph, "_p") and paragraph._p is not None:
            try:
                # Check for numPr (numbering properties) element
                num_pr = paragraph._p.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
                if num_pr is not None:
                    # Get numbering level
                    ilvl_elem = num_pr.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl")
                    level = (
                        int(ilvl_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "0")) + 1
                        if ilvl_elem is not None
                        else 1
                    )

                    # Get numbering ID to determine list type
                    num_id_elem = num_pr.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId")
                    if num_id_elem is not None:
                        # For now, detect type from paragraph text pattern
                        text = paragraph.text.strip()
                        if re.match(r"^\d+[.)]", text) or re.match(r"^[a-zA-Z][.)]", text):
                            return "number", level
                        else:
                            return "bullet", level
            except Exception:
                pass
        return None, 0

    # Check for built-in list styles
    style_name = paragraph.style.name
    if match := re.match(r"List\s*Bullet\s?(?P<level>\d+)?", style_name, re.I):
        return "bullet", int(match.group("level") or 1)
    elif match := re.match(r"List\s*Number\s?(?P<level>\d+)?", style_name, re.I):
        return "number", int(match.group("level") or 1)

    # Check indentation level
    try:
        indent = paragraph.paragraph_format.left_indent
        if indent:
            # Convert Pt to level (assume DEFAULT_INDENTATION_PT_PER_LEVEL per level)
            level = int(indent.pt / DEFAULT_INDENTATION_PT_PER_LEVEL)
            if level > 0:
                # Try to detect if numbered based on paragraph text
                if re.match(r"^\d+[.)]", paragraph.text.strip()):
                    return "number", level
                return "bullet", level
    except AttributeError:
        pass

    return None, 0


def _process_hyperlink(run: Any) -> tuple[str | None, Any]:
    """Extract hyperlink URL from a run."""
    if isinstance(run, Hyperlink):
        return run.url, run.runs[0]
    return None, run


def _get_run_formatting_key(run: Any) -> tuple[bool, bool, bool, bool, bool, bool]:
    """Get a tuple of formatting attributes to use as a key for grouping similarly formatted runs."""
    return (
        bool(run.bold),
        bool(run.italic),
        bool(run.underline),
        bool(run.font.strike),
        bool(run.font.subscript),
        bool(run.font.superscript),
        # bool(_process_hyperlink(run))  # Include hyperlink in formatting key
    )


def _format_list_marker(list_type: str, number: int = 1) -> str:
    """Generate a properly formatted list marker."""
    if list_type == "bullet":
        return "* "
    else:
        return f"{number}. "


def _process_paragraph_runs(paragraph: Paragraph, md_options: MarkdownOptions | None = None) -> str:
    """Process all runs in a paragraph, combining similarly formatted runs."""
    grouped_runs: list[tuple[str, tuple[bool, bool, bool, bool, bool, bool, bool] | None, str | None]] = []
    current_text: list[str] = []
    current_format: tuple[bool, bool, bool, bool, bool, bool, bool] | None = None
    current_url: str | None = None

    for run in paragraph.iter_inner_content():
        url, run_to_parse = _process_hyperlink(run)
        format_key: tuple[bool, bool, bool, bool, bool, bool, bool] = (
            *_get_run_formatting_key(run_to_parse),
            url is not None,
        )

        # Start new group if format changes or hyperlink changes
        if format_key != current_format or url != current_url:
            if current_text:
                grouped_runs.append(("".join(current_text), current_format, current_url))
                current_text = []
            current_format = format_key
            current_url = url

        current_text.append(run_to_parse.text)

    # Add final group
    if current_text:
        grouped_runs.append(("".join(current_text), current_format, current_url))

    # Process each group
    text_parts = []
    for text, format_key, url in grouped_runs:  # type: ignore[assignment]
        # Skip empty groups
        if not text.strip():
            text_parts.append(text)
            continue

        markers = []
        if format_key:
            if format_key[0]:  # bold
                markers.append(("**", "**"))
            if format_key[1]:  # italic
                markers.append(("*", "*"))
            if format_key[2]:  # underline
                markers.append(("__", "__"))
            if format_key[3]:  # strike
                markers.append(("~~", "~~"))
            if format_key[4]:  # subscript
                markers.append(("~", "~"))
            if format_key[5]:  # superscript
                markers.append(("^", "^"))

        # Preserve whitespace
        content = text.strip()
        prefix = text[: len(text) - len(text.lstrip())]
        suffix = text[len(text.rstrip()) :]

        if markers:
            for start, end in markers:
                content = f"{start}{content}{end}"

        # Add hyperlink if present
        if url:
            content = f"[{content}]({url})"

        text_parts.append(prefix + content + suffix)

    return "".join(text_parts)


def _convert_table_to_markdown(table: Table, md_options: MarkdownOptions | None = None) -> str:
    """Convert a docx table to markdown format."""
    markdown_rows = []

    # Process header row
    if len(table.rows) > 0:
        header_cell_md = []
        for cell in table.rows[0].cells:
            cell_md = "".join(_process_paragraph_runs(p) for p in cell.paragraphs)
            header_cell_md.append(cell_md)

        # header = [cell.text.strip() for cell in table.rows[0].cells]
        markdown_rows.append("| " + " | ".join(header_cell_md) + " |")
        # Add separator row
        markdown_rows.append("| " + " | ".join(["---"] * len(header_cell_md)) + " |")

        # Process data rows
        for row in table.rows[1:]:
            row_md = []
            for cell in row.cells:
                cell_md = "".join(_process_paragraph_runs(p) for p in cell.paragraphs)
                row_md.append(cell_md)

            # cells = [cell.text.strip() for cell in row.cells]
            markdown_rows.append("| " + " | ".join(row_md) + " |")

    return "\n".join(markdown_rows)


def _iter_block_items(parent: Any, options: DocxOptions) -> Any:
    """
    Generate a sequence of Paragraph and Table elements in order, handling images.
    """
    parent_elm = parent.element.body if isinstance(parent, docx.document.Document) else parent._element

    for child in parent_elm.iterchildren():
        if child.tag.endswith("tbl"):
            yield Table(child, parent)
        elif child.tag.endswith("p"):
            paragraph = Paragraph(child, parent)

            # Check if paragraph contains an image
            has_image = False
            img_data = []

            for run in paragraph.runs:
                for pic in run._element.findall(
                    ".//pic:pic",
                    {"pic": "http://schemas.openxmlformats.org/drawingml/2006/picture"},
                ):
                    has_image = True

                    # Get image info
                    title = None

                    if (t := run._element.xpath(".//wp:docPr/@descr")) or (
                        t := run._element.xpath(".//wp:docPr/@title")
                    ):
                        title = t[0]

                    # Get image data
                    blip = pic.xpath(".//a:blip")[0]
                    blip_rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    raw_image_data = extract_docx_image_data(parent, blip_rId)

                    # Handle pre-formatted data URIs (for backward compatibility with tests)
                    if isinstance(raw_image_data, str):
                        # This is already a formatted URI, use it directly
                        processed_image = f"![{title or 'image'}]({raw_image_data})"
                        img_data.append((title or "image", processed_image))
                        continue
                    elif not isinstance(raw_image_data, (bytes, type(None))):
                        raw_image_data = None

                    # Process image using unified attachment handling
                    image_filename = f"image_{len(img_data) + 1}.png"
                    processed_image = process_attachment(
                        attachment_data=raw_image_data,
                        attachment_name=image_filename,
                        alt_text=title or "image",
                        attachment_mode=options.attachment_mode,
                        attachment_output_dir=options.attachment_output_dir,
                        attachment_base_url=options.attachment_base_url,
                        is_image=True,
                    )

                    img_data.append((title, processed_image))

            if has_image and img_data:
                # Create a new paragraph with default style for each image
                from docx.oxml.shared import OxmlElement

                for _title, processed_image in img_data:
                    p = OxmlElement("w:p")
                    r = OxmlElement("w:r")
                    t = OxmlElement("w:t")

                    # The processed_image already contains the proper markdown
                    t.text = processed_image

                    r.append(t)
                    p.append(r)

                    clean_paragraph = Paragraph(p, parent)
                    clean_paragraph.style = parent.styles["Normal"]
                    yield clean_paragraph
            elif not has_image:
                yield paragraph


def docx_to_markdown(
    input_data: Union[str, Path, docx.document.Document, IO[bytes]], options: DocxOptions | None = None
) -> str:
    """Convert Word document (DOCX) to Markdown format.

    Processes Microsoft Word documents and converts them to well-formatted
    Markdown while preserving text formatting, document structure, tables,
    lists, and embedded images. Handles complex document elements and
    maintains hierarchical organization.

    Parameters
    ----------
    input_data : str, file-like object, or docx.Document
        Word document to convert. Can be:
        - String path to DOCX file
        - File-like object containing DOCX data
        - Already opened python-docx Document object
    options : DocxOptions or None, default None
        Configuration options for DOCX conversion. If None, uses default settings.

    Returns
    -------
    str
        Markdown representation of the Word document with preserved
        formatting, structure, and content.

    Raises
    ------
    MdparseInputError
        If input type is not supported or document cannot be opened
    MdparseConversionError
        If document processing fails

    Examples
    --------
    Convert a Word file to Markdown:

        >>> with open('document.docx', 'rb') as f:
        ...     markdown = docx_to_markdown(f)
        >>> print(markdown)

    Convert with embedded images:

        >>> markdown = docx_to_markdown('document.docx', convert_images_to_base64=True)

    Notes
    -----
    - Preserves text formatting (bold, italic, underline, strikethrough)
    - Converts lists with proper nesting and numbering
    - Handles tables with Markdown table syntax
    - Processes hyperlinks and converts to Markdown link format
    - Maintains document structure with appropriate heading levels
    """
    # Handle backward compatibility and merge options
    if options is None:
        options = DocxOptions()

    # Validate and convert input - for now use simplified approach
    try:
        if isinstance(input_data, docx.document.Document):
            doc = input_data
        elif isinstance(input_data, Path):
            doc = docx.Document(str(input_data))
        else:
            doc = docx.Document(input_data)
    except Exception as e:
        raise MdparseConversionError(
            f"Failed to open DOCX document: {str(e)}", conversion_stage="document_opening", original_error=e
        ) from e

    # Get Markdown options (create default if not provided)
    md_options = options.markdown_options or MarkdownOptions()

    markdown_lines = []
    list_stack: list[tuple[str, int, int]] = []  # Track nested lists: (type, level, current_number)

    for block in _iter_block_items(doc, options=options):
        if isinstance(block, Paragraph):
            text = ""

            # Skip empty paragraphs
            if not block.text.strip():
                markdown_lines.append("")
                continue

            # Handle heading styles
            style_name = block.style.name if block.style else ""
            heading_match = re.match(r"Heading (\d+)", style_name)
            if heading_match:
                level = int(heading_match.group(1))
                text = "#" * level + " " + block.text.strip()
                markdown_lines.append(text)
                continue

            # Handle lists
            list_type, level = _detect_list_level(block)
            if list_type:
                # Close any deeper nested lists
                while list_stack and list_stack[-1][1] > level:
                    list_stack.pop()

                # Check if we're continuing the same list type at this level
                if list_stack and list_stack[-1][1] == level:
                    if list_stack[-1][0] == list_type:
                        # Continue numbering for numbered lists
                        if list_type == "number":
                            list_stack[-1] = (list_type, level, list_stack[-1][2] + 1)
                    else:
                        # Different list type at same level - start new list
                        list_stack[-1] = (list_type, level, 1)
                else:
                    # Start new list
                    while len(list_stack) > 0 and list_stack[-1][1] >= level:
                        list_stack.pop()
                    list_stack.append((list_type, level, 1))

                # Calculate proper indentation
                indent = ""
                for _i, (lst_type, _lst_level, _) in enumerate(list_stack[:-1]):
                    # Add 3 spaces for numbered lists, 2 for bullet lists
                    indent += "   " if lst_type == "number" else "  "

                # Add marker with proper numbering or bullet
                if list_type == "number":
                    marker = _format_list_marker(list_type, list_stack[-1][2])
                else:
                    marker = _format_list_marker(list_type)

                text = indent + marker + _process_paragraph_runs(block, md_options).strip()

            else:
                # Not a list - clear list stack
                if list_stack:
                    list_stack = []
                    markdown_lines.append("")

                text = _process_paragraph_runs(block, md_options)

            if text:
                markdown_lines.append(text)

        elif isinstance(block, Table):
            # Add spacing before table
            if markdown_lines and markdown_lines[-1]:
                markdown_lines.append("")

            # Convert and add table
            markdown_lines.append(_convert_table_to_markdown(block, md_options))

            # Add spacing after table
            markdown_lines.append("")

    # Clean up multiple blank lines
    markdown = "\n\n".join(markdown_lines)
    markdown = re.sub(r"\n{3,}", "\n\n", markdown)

    return markdown.strip()
