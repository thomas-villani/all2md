#  Copyright (c) 2025 Tom Villani, Ph.D.
#
# src/all2md/converters/docx2ast.py
"""DOCX to AST converter.

This module provides conversion from Microsoft Word DOCX documents to AST representation.
It replaces direct markdown string generation with structured AST building,
enabling multiple rendering strategies and improved testability.

"""

from __future__ import annotations

from dataclasses import dataclass
import logging
import re
from typing import TYPE_CHECKING, Any


from all2md.constants import DEFAULT_INDENTATION_PT_PER_LEVEL
from all2md.utils.attachments import extract_docx_image_data, process_attachment


@dataclass
class ImageData:
    """Structured representation of image data from DOCX.

    Parameters
    ----------
    url : str
        Image URL or data URI
    alt_text : str
        Alternative text for the image
    title : str or None
        Optional title for the image
    """

    url: str
    alt_text: str
    title: str | None = None


if TYPE_CHECKING:
    import docx.document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.text.hyperlink import Hyperlink


from all2md.ast import (
    Document,
    Emphasis,
    Heading,
    Image,
    Link,
    List,
    ListItem,
    Node,
    Paragraph as AstParagraph,
    Strikethrough,
    Strong,
    Subscript,
    Superscript,
    Table as AstTable,
    TableCell,
    TableRow,
    Text,
    Underline,
)
from all2md.options import DocxOptions

logger = logging.getLogger(__name__)


class DocxToAstConverter:
    """Convert DOCX to AST representation.

    This converter parses Word documents using python-docx and builds an AST
    that can be rendered to various markdown flavors.

    Parameters
    ----------
    options : DocxOptions or None, default = None
        Conversion options
    doc : docx.document.Document or None
        Document to convert (stored for list detection)
    base_filename : str
        Base filename for image attachments
    attachment_sequencer : callable or None
        Sequencer for generating attachment filenames

    """

    def __init__(
        self,
        options: DocxOptions | None = None,
        doc: "docx.document.Document | None" = None,
        base_filename: str = "document",
        attachment_sequencer: Any = None,
    ):
        self.options = options or DocxOptions()
        self.doc = doc
        self.base_filename = base_filename
        self.attachment_sequencer = attachment_sequencer
        self._list_stack: list[tuple[str, int, list[ListItem]]] = []  # (type, level, items)
        self._numbering_defs: dict[str, dict[str, str]] | None = None

    def convert_to_ast(self, doc: "docx.document.Document") -> Document:
        """Convert DOCX document to AST Document.

        Parameters
        ----------
        doc : docx.document.Document
            DOCX document to convert

        Returns
        -------
        Document
            AST document node

        """
        self.doc = doc
        self._numbering_defs = None  # Will be loaded lazily if needed
        self._list_stack = []

        children: list[Node] = []

        # Import here to avoid circular dependencies
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        for block in _iter_block_items(
            doc,
            options=self.options,
            base_filename=self.base_filename,
            attachment_sequencer=self.attachment_sequencer,
        ):
            if isinstance(block, ImageData):
                # Handle ImageData objects directly
                image_node = Image(url=block.url, alt_text=block.alt_text, title=block.title)
                children.append(AstParagraph(content=[image_node]))
            elif isinstance(block, Paragraph):
                nodes = self._process_paragraph_to_ast(block)
                if nodes:
                    if isinstance(nodes, list):
                        children.extend(nodes)
                    else:
                        children.append(nodes)
            elif isinstance(block, Table):
                # Check if tables should be preserved
                if self.options.preserve_tables:
                    table_node = self._process_table_to_ast(block)
                    if table_node:
                        children.append(table_node)
                else:
                    # Flatten table to paragraphs
                    for row in block.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                inline_nodes = self._process_paragraph_runs_to_inline(paragraph)
                                if inline_nodes:
                                    children.append(AstParagraph(content=inline_nodes))

        # Finalize any remaining list at the end of document
        if self._list_stack:
            final_list = self._finalize_current_list()
            if final_list:
                children.append(final_list)
            self._list_stack = []

        # Add footnotes, endnotes, and comments at the end if requested
        if self.options.include_footnotes:
            footnotes_nodes = self._process_footnotes(doc)
            if footnotes_nodes:
                children.extend(footnotes_nodes)

        if self.options.include_endnotes:
            endnotes_nodes = self._process_endnotes(doc)
            if endnotes_nodes:
                children.extend(endnotes_nodes)

        if self.options.include_comments:
            comments_nodes = self._process_comments(doc)
            if comments_nodes:
                children.extend(comments_nodes)

        return Document(children=children)

    def _process_paragraph_to_ast(self, paragraph: "Paragraph") -> Node | list[Node] | None:
        """Process a DOCX paragraph to AST nodes.

        Parameters
        ----------
        paragraph : Paragraph
            DOCX paragraph to process

        Returns
        -------
        Node, list of Node, or None
            Resulting AST node(s)

        """
        # Skip empty paragraphs
        text = paragraph.text.strip()
        if not text:
            return None

        # Handle heading styles
        style_name = paragraph.style.name if paragraph.style else ""
        heading_match = re.match(r"Heading (\d+)", style_name)
        if heading_match:
            level = int(heading_match.group(1))
            content = self._process_paragraph_runs_to_inline(paragraph)
            return Heading(level=level, content=content)

        # Handle lists

        list_type, level = _detect_list_level(paragraph, self.doc)
        if list_type:
            # This paragraph is part of a list
            # We'll accumulate list items and return them when list ends
            return self._process_list_item_paragraph(paragraph, list_type, level)

        # Not a list - clear list stack and return any accumulated list
        if self._list_stack:
            accumulated_list = self._finalize_current_list()
            self._list_stack = []
            # Return the accumulated list and the current paragraph
            content = self._process_paragraph_runs_to_inline(paragraph)
            return [accumulated_list, AstParagraph(content=content)]

        # Regular paragraph
        content = self._process_paragraph_runs_to_inline(paragraph)
        if content:
            return AstParagraph(content=content)
        return None

    def _process_list_item_paragraph(
        self, paragraph: "Paragraph", list_type: str, level: int
    ) -> Node | None:
        """Process a paragraph that is part of a list.

        This method accumulates list items and handles nesting properly.
        Returns a completed list when we transition to a non-list item.

        Parameters
        ----------
        paragraph : Paragraph
            List item paragraph
        list_type : str
            'bullet' or 'number'
        level : int
            Nesting level (1-based)

        Returns
        -------
        Node or None
            Completed list node if transitioning out, None if still accumulating

        """
        # Process paragraph content
        content = self._process_paragraph_runs_to_inline(paragraph)
        item_node = ListItem(children=[AstParagraph(content=content)])

        # Handle level changes
        if not self._list_stack:
            # Start new list at this level
            self._list_stack = [(list_type, level, [item_node])]
            return None

        current_type, current_level, current_items = self._list_stack[-1]

        if level > current_level:
            # Nested list - deeper level
            # Start new list at the deeper level
            self._list_stack.append((list_type, level, [item_node]))
            return None

        elif level < current_level:
            # Going back to shallower level - need to finalize deeper lists
            # Pop and nest all lists deeper than the target level
            while self._list_stack and self._list_stack[-1][1] > level:
                # Pop the deeper list and nest it
                popped_type, popped_level, popped_items = self._list_stack.pop()
                nested_list = List(ordered=(popped_type == "number"), items=popped_items, tight=True)

                # Add nested list to the last item of the parent level
                if self._list_stack:
                    parent_items = self._list_stack[-1][2]
                    if parent_items:
                        parent_items[-1].children.append(nested_list)

            # Now at the correct level - check if type matches
            if self._list_stack and self._list_stack[-1][0] == list_type and self._list_stack[-1][1] == level:
                # Same level and type - add item
                self._list_stack[-1][2].append(item_node)
                return None  # Still accumulating
            else:
                # Different type at this level - finalize old, start new
                result_node = None
                if self._list_stack and self._list_stack[-1][1] == level:
                    old_type, old_level, old_items = self._list_stack.pop()
                    result_node = List(ordered=(old_type == "number"), items=old_items, tight=True)
                self._list_stack.append((list_type, level, [item_node]))
                return result_node

        else:
            # Same level
            if current_type == list_type:
                # Same type - add to current list
                current_items.append(item_node)
                return None
            else:
                # Different type at same level - finalize old, start new
                old_type, old_level, old_items = self._list_stack.pop()
                self._list_stack.append((list_type, level, [item_node]))
                return List(ordered=(old_type == "number"), items=old_items, tight=True)

    def _finalize_current_list(self) -> List | None:
        """Finalize all lists in the stack, nesting them properly.

        Returns
        -------
        List or None
            Completed top-level list node with all nesting

        """
        if not self._list_stack:
            return None

        # Pop and nest all lists from deepest to shallowest
        while len(self._list_stack) > 1:
            # Pop deeper list
            deeper_type, deeper_level, deeper_items = self._list_stack.pop()
            nested_list = List(ordered=(deeper_type == "number"), items=deeper_items, tight=True)

            # Add to parent's last item
            parent_items = self._list_stack[-1][2]
            if parent_items:
                parent_items[-1].children.append(nested_list)

        # Return the top-level list
        if self._list_stack:
            list_type, level, items = self._list_stack.pop()
            return List(ordered=(list_type == "number"), items=items, tight=True)

        return None

    def _process_paragraph_runs_to_inline(self, paragraph: "Paragraph") -> list[Node]:
        """Process paragraph runs to inline AST nodes.

        Parameters
        ----------
        paragraph : Paragraph
            Paragraph containing runs

        Returns
        -------
        list of Node
            List of inline AST nodes

        """
        result: list[Node] = []

        # Group runs by formatting to avoid excessive nesting
        grouped_runs: list[tuple[str, tuple[bool, bool, bool, bool, bool, bool, bool] | None, str | None]] = []
        current_text: list[str] = []
        current_format: tuple[bool, bool, bool, bool, bool, bool, bool] | None = None
        current_url: str | None = None

        for run in paragraph.iter_inner_content():
            # Check for hyperlink
            url, run_to_parse = self._process_hyperlink(run)

            # Get formatting
            format_key = self._get_run_formatting_key(run_to_parse, url is not None)

            # Start new group if format changes
            if format_key != current_format or url != current_url:
                if current_text:
                    grouped_runs.append(("".join(current_text), current_format, current_url))
                    current_text = []
                current_format = format_key
                current_url = url

            from docx.text.hyperlink import Hyperlink
            # Extract text
            if isinstance(run_to_parse, Hyperlink):
                hyperlink_text = "".join(r.text for r in run_to_parse.runs)
                current_text.append(hyperlink_text)
            else:
                current_text.append(run_to_parse.text)

        # Add final group
        if current_text:
            grouped_runs.append(("".join(current_text), current_format, current_url))

        # Convert groups to AST nodes
        for text, format_key, url in grouped_runs:
            if not text:
                continue

            # Build inline nodes with formatting
            inline_node: Node = Text(content=text)

            if format_key:
                # Apply formatting layers from innermost to outermost
                # Order: text -> underline/sub/super -> bold -> italic -> strike -> link

                if format_key[2]:  # underline
                    inline_node = Underline(content=[inline_node])
                if format_key[4]:  # subscript
                    inline_node = Subscript(content=[inline_node])
                if format_key[5]:  # superscript
                    inline_node = Superscript(content=[inline_node])
                if format_key[0]:  # bold
                    inline_node = Strong(content=[inline_node])
                if format_key[1]:  # italic
                    inline_node = Emphasis(content=[inline_node])
                if format_key[3]:  # strike
                    inline_node = Strikethrough(content=[inline_node])

            # Wrap in link if URL present
            if url:
                inline_node = Link(url=url, content=[inline_node])

            result.append(inline_node)

        return result

    def _process_hyperlink(self, run: Any) -> tuple[str | None, Any]:
        """Process a run to extract hyperlink information.

        Parameters
        ----------
        run : Any
            Run object to process

        Returns
        -------
        tuple of (url, run_to_parse)
            URL if hyperlink, and the run to parse for text

        """
        from docx.text.hyperlink import Hyperlink
        if isinstance(run, Hyperlink):
            return run.address, run
        return None, run

    def _get_run_formatting_key(self, run: Any, is_hyperlink: bool) -> tuple[bool, bool, bool, bool, bool, bool, bool]:
        """Get formatting key for a run.

        Parameters
        ----------
        run : Any
            Run to analyze
        is_hyperlink : bool
            Whether this run is part of a hyperlink

        Returns
        -------
        tuple of bool
            (bold, italic, underline, strike, subscript, superscript, is_hyperlink)

        """
        from docx.text.hyperlink import Hyperlink
        # Handle Hyperlink object
        if isinstance(run, Hyperlink):
            if run.runs:
                first_run = run.runs[0]
                return (
                    first_run.bold or False,
                    first_run.italic or False,
                    first_run.underline or False,
                    first_run.font.strike or False,
                    first_run.font.subscript or False,
                    first_run.font.superscript or False,
                    is_hyperlink,
                )
            return (False, False, False, False, False, False, is_hyperlink)

        # Regular run
        return (
            run.bold or False,
            run.italic or False,
            run.underline or False,
            run.font.strike or False,
            run.font.subscript or False,
            run.font.superscript or False,
            is_hyperlink,
        )


    def _process_table_to_ast(self, table: "Table") -> AstTable | None:
        """Process a DOCX table to AST Table node.

        Parameters
        ----------
        table : Table
            DOCX table to convert

        Returns
        -------
        AstTable or None
            Table node if table has content

        """
        if len(table.rows) == 0:
            return None

        # First row is header
        header_cells: list[TableCell] = []
        for cell in table.rows[0].cells:
            cell_content = []
            for p in cell.paragraphs:
                inline_nodes = self._process_paragraph_runs_to_inline(p)
                if inline_nodes:
                    cell_content.extend(inline_nodes)
            header_cells.append(TableCell(content=cell_content))

        header_row = TableRow(cells=header_cells, is_header=True)

        # Data rows
        data_rows: list[TableRow] = []
        for row in table.rows[1:]:
            row_cells: list[TableCell] = []
            for cell in row.cells:
                cell_content = []
                for p in cell.paragraphs:
                    inline_nodes = self._process_paragraph_runs_to_inline(p)
                    if inline_nodes:
                        cell_content.extend(inline_nodes)
                row_cells.append(TableCell(content=cell_content))
            data_rows.append(TableRow(cells=row_cells))

        return AstTable(header=header_row, rows=data_rows)

    def _process_footnotes(self, doc: "docx.document.Document") -> list[Node]:
        """Process footnotes from document.

        Parameters
        ----------
        doc : docx.document.Document
            Document to extract footnotes from

        Returns
        -------
        list[Node]
            List of AST nodes for footnotes

        """
        nodes: list[Node] = []

        try:
            if hasattr(doc.part, 'footnotes_part'):
                footnotes_part = doc.part.footnotes_part
                if footnotes_part and hasattr(footnotes_part, 'element'):
                    # Add section heading
                    nodes.append(Heading(level=2, content=[Text(content="Footnotes")]))

                    # Iterate over footnotes
                    for footnote in footnotes_part.element.findall(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footnote'
                    ):
                        footnote_id = footnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                        # Skip special footnotes (separator, continuation separator)
                        if footnote_id in ('-1', '0'):
                            continue

                        # Extract text from footnote paragraphs
                        footnote_text_parts = []
                        for p_elem in footnote.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                            for t_elem in p_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                                if t_elem.text:
                                    footnote_text_parts.append(t_elem.text)

                        if footnote_text_parts:
                            footnote_text = ' '.join(footnote_text_parts)
                            nodes.append(AstParagraph(content=[
                                Strong(content=[Text(content=f"[{footnote_id}]")]),
                                Text(content=f" {footnote_text}")
                            ]))
        except Exception as e:
            logger.debug(f"Could not process footnotes: {e}")

        return nodes

    def _process_endnotes(self, doc: "docx.document.Document") -> list[Node]:
        """Process endnotes from document.

        Parameters
        ----------
        doc : docx.document.Document
            Document to extract endnotes from

        Returns
        -------
        list[Node]
            List of AST nodes for endnotes

        """
        nodes: list[Node] = []

        try:
            if hasattr(doc.part, 'endnotes_part'):
                endnotes_part = doc.part.endnotes_part
                if endnotes_part and hasattr(endnotes_part, 'element'):
                    # Add section heading
                    nodes.append(Heading(level=2, content=[Text(content="Endnotes")]))

                    # Iterate over endnotes
                    for endnote in endnotes_part.element.findall(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}endnote'
                    ):
                        endnote_id = endnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                        # Skip special endnotes
                        if endnote_id in ('-1', '0'):
                            continue

                        # Extract text from endnote paragraphs
                        endnote_text_parts = []
                        for p_elem in endnote.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                            for t_elem in p_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                                if t_elem.text:
                                    endnote_text_parts.append(t_elem.text)

                        if endnote_text_parts:
                            endnote_text = ' '.join(endnote_text_parts)
                            nodes.append(AstParagraph(content=[
                                Strong(content=[Text(content=f"[{endnote_id}]")]),
                                Text(content=f" {endnote_text}")
                            ]))
        except Exception as e:
            logger.debug(f"Could not process endnotes: {e}")

        return nodes

    def _process_comments(self, doc: "docx.document.Document") -> list[Node]:
        """Process comments from document.

        Parameters
        ----------
        doc : docx.document.Document
            Document to extract comments from

        Returns
        -------
        list[Node]
            List of AST nodes for comments

        """
        nodes: list[Node] = []

        try:
            if hasattr(doc.part, 'comments_part'):
                comments_part = doc.part.comments_part
                if comments_part and hasattr(comments_part, 'element'):
                    # Add section heading
                    nodes.append(Heading(level=2, content=[Text(content="Comments")]))

                    # Iterate over comments
                    for comment in comments_part.element.findall(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment'
                    ):
                        comment_id = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                        author = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown')
                        date = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')

                        # Extract text from comment paragraphs
                        comment_text_parts = []
                        for p_elem in comment.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                            for t_elem in p_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                                if t_elem.text:
                                    comment_text_parts.append(t_elem.text)

                        if comment_text_parts:
                            comment_text = ' '.join(comment_text_parts)
                            # Format: **[Comment ID]** Author (Date): Comment text
                            header = f"[{comment_id}] {author}"
                            if date:
                                header += f" ({date})"
                            nodes.append(AstParagraph(content=[
                                Strong(content=[Text(content=header)]),
                                Text(content=f": {comment_text}")
                            ]))
        except Exception as e:
            logger.debug(f"Could not process comments: {e}")

        return nodes


def _get_numbering_definitions(doc: "docx.document.Document") -> dict[str, dict[str, str]]:
    """Extract and cache numbering definitions from document.

    Returns a mapping of numId -> {level -> format_type} where format_type is 'bullet' or 'decimal'.
    """
    numbering_defs: dict[str, dict[str, str]] = {}

    if not hasattr(doc, '_part') or not hasattr(doc._part, 'numbering_part'):
        return numbering_defs

    numbering_part = doc._part.numbering_part
    if not numbering_part:
        return numbering_defs

    try:
        numbering_xml = numbering_part._element

        # First, collect abstract numbering definitions
        abstract_nums = {}
        for elem in numbering_xml.iter():
            if elem.tag.endswith('abstractNum'):
                abstract_num_id = elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId')
                if abstract_num_id:
                    levels = {}
                    for level_elem in elem.iter():
                        if level_elem.tag.endswith('lvl'):
                            level_id = level_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
                            if level_id is not None:
                                for child in level_elem.iter():
                                    if child.tag.endswith('numFmt'):
                                        fmt_val = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                        if fmt_val:
                                            # Map Word numbering formats to our types
                                            if fmt_val in ('bullet', 'none'):
                                                levels[level_id] = 'bullet'
                                            elif fmt_val in (
                                                'decimal', 'lowerLetter', 'upperLetter', 'lowerRoman', 'upperRoman'
                                            ):
                                                levels[level_id] = 'number'
                                            break
                    if levels:
                        abstract_nums[abstract_num_id] = levels

        # Then, map number IDs to abstract numbers
        for elem in numbering_xml.iter():
            if elem.tag.endswith('num'):
                num_id = elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
                if num_id:
                    for child in elem.iter():
                        if child.tag.endswith('abstractNumId'):
                            abs_id = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            if abs_id in abstract_nums:
                                numbering_defs[num_id] = abstract_nums[abs_id]
                            break

    except Exception as e:
        logger.debug(f"Error parsing numbering definitions: {e}")

    return numbering_defs


def _detect_list_level(paragraph: "Paragraph", doc: "docx.document.Document" = None) -> tuple[str | None, int]:
    """Detect the list level of a paragraph based on its style, numbering, and indentation.

    Returns tuple of (list_type, level) where list_type is 'bullet' or 'number' and level is integer depth
    """
    # Check for Word native numbering properties first (works for all list styles including "List Paragraph")
    if hasattr(paragraph, "_p") and paragraph._p is not None:
        try:
            # Check for numPr (numbering properties) element
            num_pr = paragraph._p.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
            if num_pr is not None:
                # Get numbering level (Word uses 0-based indexing, we use 1-based)
                ilvl_elem = num_pr.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl")
                level = (
                    int(ilvl_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "0")) + 1
                    if ilvl_elem is not None
                    else 1
                )

                # Get numbering ID to determine list type
                num_id_elem = num_pr.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId")
                if num_id_elem is not None:
                    num_id = num_id_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")

                    # Look up the numbering definition if document is available
                    if doc and num_id:
                        numbering_defs = _get_numbering_definitions(doc)
                        if num_id in numbering_defs:
                            # Get the format for this level (use level-1 since Word is 0-based)
                            level_key = str(level - 1)
                            if level_key in numbering_defs[num_id]:
                                return numbering_defs[num_id][level_key], level
                            # If specific level not found, use level 0 as fallback
                            elif '0' in numbering_defs[num_id]:
                                return numbering_defs[num_id]['0'], level

                    # Fallback: detect type from paragraph text pattern
                    text = paragraph.text.strip()
                    if re.match(r"^\d+[.)]", text) or re.match(r"^[a-zA-Z][.)]", text):
                        return "number", level
                    else:
                        return "bullet", level
        except Exception:
            pass

    # Check for built-in list styles
    style_name = paragraph.style.name if paragraph.style else None
    if not style_name:
        return None, 0

    base_type = None
    style_level = 1

    # Handle "List Paragraph" style - check for numbering properties above
    if style_name == "List Paragraph":
        # If we got here, numbering properties weren't found or processed
        # This might be a list paragraph without proper numbering - treat as bullet by default
        return "bullet", 1
    elif match := re.match(r"List\s*Bullet\s?(?P<level>\d+)?", style_name, re.I):
        base_type = "bullet"
        style_level = int(match.group("level") or 1)
    elif match := re.match(r"List\s*Number\s?(?P<level>\d+)?", style_name, re.I):
        base_type = "number"
        style_level = int(match.group("level") or 1)

    # Check indentation level for additional nesting
    indent_level = 0
    try:
        indent = paragraph.paragraph_format.left_indent
        if indent:
            # Convert Pt to level (assume DEFAULT_INDENTATION_PT_PER_LEVEL per level)
            indent_level = int(indent.pt / DEFAULT_INDENTATION_PT_PER_LEVEL)
    except AttributeError:
        pass

    # If we have a list style, combine with indentation
    if base_type:
        final_level = max(style_level, style_level + indent_level)
        return base_type, final_level

    # Check indentation level for paragraphs without list styles
    if indent_level > 0:
        # Try to detect if numbered based on paragraph text
        if re.match(r"^\d+[.)]", paragraph.text.strip()):
            return "number", indent_level
        return "bullet", indent_level

    return None, 0


def _iter_block_items(
    parent: Any, options: DocxOptions, base_filename: str = "document", attachment_sequencer=None
) -> Any:
    """
    Generate a sequence of Paragraph and Table elements in order, handling images.
    """
    import docx.document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent_elm = parent.element.body if isinstance(parent, docx.document.Document) else parent._element

    for child in parent_elm.iterchildren():
        if child.tag.endswith("tbl"):
            yield Table(child, parent)
        elif child.tag.endswith("p"):
            paragraph = Paragraph(child, parent)

            # Check if paragraph contains an image
            has_image = False
            img_data = []

            for run in paragraph.runs:
                for pic in run._element.findall(
                        ".//pic:pic",
                        {"pic": "http://schemas.openxmlformats.org/drawingml/2006/picture"},
                ):
                    has_image = True

                    # Get image info
                    title = None

                    if options.include_image_captions:
                        if (t := run._element.xpath(".//wp:docPr/@descr")) or (
                                t := run._element.xpath(".//wp:docPr/@title")
                        ):
                            title = t[0]

                    # Get image data and detected format
                    blip = pic.xpath(".//a:blip")[0]
                    blip_rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    raw_image_data, detected_extension = extract_docx_image_data(parent, blip_rId)

                    # Handle pre-formatted data URIs (for backward compatibility with tests)
                    if isinstance(raw_image_data, str):
                        # This is already a formatted URI, use it directly
                        img_data.append(ImageData(url=raw_image_data, alt_text=title or "image", title=title))
                        continue
                    elif not isinstance(raw_image_data, (bytes, type(None))):
                        logger.warning(f"Invalid image data type for image '{title or 'unnamed'}', skipping")
                        raw_image_data = None

                    # Use detected extension or fallback to png
                    extension = detected_extension or "png"

                    # Log format detection result
                    if detected_extension:
                        logger.debug(f"Detected image format: {detected_extension}")
                    else:
                        logger.debug("No image format detected, using PNG as fallback")

                    # Process image using unified attachment handling
                    # Use sequencer if available, otherwise fall back to manual counting
                    if attachment_sequencer:
                        image_filename, _ = attachment_sequencer(
                            base_stem=base_filename,
                            format_type="general",
                            extension=extension
                        )
                    else:
                        from all2md.utils.attachments import generate_attachment_filename
                        image_filename = generate_attachment_filename(
                            base_stem=base_filename,
                            format_type="general",
                            sequence_num=len(img_data) + 1,
                            extension=extension
                        )
                    processed_image_url = process_attachment(
                        attachment_data=raw_image_data,
                        attachment_name=image_filename,
                        alt_text=title or "image",
                        attachment_mode=options.attachment_mode,
                        attachment_output_dir=options.attachment_output_dir,
                        attachment_base_url=options.attachment_base_url,
                        is_image=True,
                        alt_text_mode=options.alt_text_mode,
                    )

                    if processed_image_url:
                        img_data.append(ImageData(url=processed_image_url, alt_text=title or "image", title=title))

            if has_image and img_data:
                # Yield ImageData objects directly
                for image_data in img_data:
                    yield image_data
            elif not has_image:
                yield paragraph
