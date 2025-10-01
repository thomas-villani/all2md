#  Copyright (c) 2025 Tom Villani, Ph.D.
#
# src/all2md/converters/docx2markdown.py
"""Word document to Markdown conversion module.

This module provides functionality to convert Microsoft Word documents (DOCX format)
to Markdown while preserving formatting, structure, and embedded content. It handles
complex document elements including tables, lists, hyperlinks, images, and various
text formatting options.

The converter processes Word documents by analyzing their internal structure,
extracting styled content, and converting it to equivalent Markdown syntax.
Special attention is paid to preserving document hierarchy, list structures,
and table layouts.

Key Features
------------
- Comprehensive text formatting preservation (bold, italic, underline, etc.)
- Intelligent list detection and conversion (bulleted and numbered)
- Table structure preservation with Markdown table syntax
- Hyperlink extraction and conversion
- Image embedding as base64 data URIs
- Document structure preservation (headers, paragraphs)
- Style-based formatting detection

Supported Elements
------------------
- Text formatting (bold, italic, underline, strikethrough)
- Headers (converted to Markdown headers)
- Lists (bulleted, numbered, with proper nesting)
- Tables (with cell content and basic formatting)
- Hyperlinks (internal and external)
- Images (embedded as base64)
- Line breaks and paragraph separation

Dependencies
------------
- python-docx: For parsing Word document structure
- logging: For debug and error reporting

Examples
--------
Basic conversion:

    >>> from all2md.converters.docx2markdown import docx_to_markdown
    >>> with open('document.docx', 'rb') as f:
    ...     markdown = docx_to_markdown(f)
    >>> print(markdown)

Note
----
Requires python-docx package. Some advanced Word features may not have
direct Markdown equivalents and will be approximated or omitted.
"""

import logging
import re
from pathlib import Path
from typing import IO, TYPE_CHECKING, Any, Union

if TYPE_CHECKING:
    import docx
    import docx.document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

# Make Hyperlink available for testing while maintaining lazy loading
try:
    from docx.text.hyperlink import Hyperlink
except ImportError:
    Hyperlink = None  # type: ignore

from all2md.constants import DEFAULT_INDENTATION_PT_PER_LEVEL
from all2md.converter_metadata import ConverterMetadata
from all2md.exceptions import MarkdownConversionError
from all2md.options import DocxOptions, MarkdownOptions
from all2md.utils.attachments import create_attachment_sequencer, extract_docx_image_data, process_attachment
from all2md.utils.inputs import escape_markdown_special, format_special_text
from all2md.utils.metadata import (
    OFFICE_FIELD_MAPPING,
    DocumentMetadata,
    map_properties_to_metadata,
    prepend_metadata_if_enabled,
)
from all2md.utils.security import validate_zip_archive

logger = logging.getLogger(__name__)

# Converter metadata for registration
CONVERTER_METADATA = ConverterMetadata(
    format_name="docx",
    extensions=[".docx"],
    mime_types=["application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
    magic_bytes=[
        (b"PK\x03\x04", 0),  # ZIP signature (docx is ZIP-based)
    ],
    converter_module="all2md.converters.docx2markdown",
    converter_function="docx_to_markdown",
    required_packages=[("python-docx", "")],
    optional_packages=[],
    import_error_message=(
        "DOCX conversion requires 'python-docx'. "
        "Install with: pip install python-docx"
    ),
    options_class="DocxOptions",
    description="Convert Microsoft Word DOCX documents to Markdown",
    priority=8
)


def _get_numbering_definitions(doc: "docx.document.Document") -> dict[str, dict[str, str]]:
    """Extract and cache numbering definitions from document.

    Returns a mapping of numId -> {level -> format_type} where format_type is 'bullet' or 'decimal'.
    """
    numbering_defs: dict[str, dict[str, str]] = {}

    if not hasattr(doc, '_part') or not hasattr(doc._part, 'numbering_part'):
        return numbering_defs

    numbering_part = doc._part.numbering_part
    if not numbering_part:
        return numbering_defs

    try:
        numbering_xml = numbering_part._element

        # First, collect abstract numbering definitions
        abstract_nums = {}
        for elem in numbering_xml.iter():
            if elem.tag.endswith('abstractNum'):
                abstract_num_id = elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId')
                if abstract_num_id:
                    levels = {}
                    for level_elem in elem.iter():
                        if level_elem.tag.endswith('lvl'):
                            level_id = level_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
                            if level_id is not None:
                                for child in level_elem.iter():
                                    if child.tag.endswith('numFmt'):
                                        fmt_val = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                        if fmt_val:
                                            # Map Word numbering formats to our types
                                            if fmt_val in ('bullet', 'none'):
                                                levels[level_id] = 'bullet'
                                            elif fmt_val in (
                                                'decimal', 'lowerLetter', 'upperLetter', 'lowerRoman', 'upperRoman'
                                            ):
                                                levels[level_id] = 'number'
                                            break
                    if levels:
                        abstract_nums[abstract_num_id] = levels

        # Then, map number IDs to abstract numbers
        for elem in numbering_xml.iter():
            if elem.tag.endswith('num'):
                num_id = elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
                if num_id:
                    for child in elem.iter():
                        if child.tag.endswith('abstractNumId'):
                            abs_id = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            if abs_id in abstract_nums:
                                numbering_defs[num_id] = abstract_nums[abs_id]
                            break

    except Exception as e:
        logger.debug(f"Error parsing numbering definitions: {e}")

    return numbering_defs


def _detect_list_level(paragraph: "Paragraph", doc: "docx.document.Document" = None) -> tuple[str | None, int]:
    """Detect the list level of a paragraph based on its style, numbering, and indentation.

    Returns tuple of (list_type, level) where list_type is 'bullet' or 'number' and level is integer depth
    """
    # Check for Word native numbering properties first (works for all list styles including "List Paragraph")
    if hasattr(paragraph, "_p") and paragraph._p is not None:
        try:
            # Check for numPr (numbering properties) element
            num_pr = paragraph._p.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
            if num_pr is not None:
                # Get numbering level (Word uses 0-based indexing, we use 1-based)
                ilvl_elem = num_pr.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl")
                level = (
                    int(ilvl_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "0")) + 1
                    if ilvl_elem is not None
                    else 1
                )

                # Get numbering ID to determine list type
                num_id_elem = num_pr.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId")
                if num_id_elem is not None:
                    num_id = num_id_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")

                    # Look up the numbering definition if document is available
                    if doc and num_id:
                        numbering_defs = _get_numbering_definitions(doc)
                        if num_id in numbering_defs:
                            # Get the format for this level (use level-1 since Word is 0-based)
                            level_key = str(level - 1)
                            if level_key in numbering_defs[num_id]:
                                return numbering_defs[num_id][level_key], level
                            # If specific level not found, use level 0 as fallback
                            elif '0' in numbering_defs[num_id]:
                                return numbering_defs[num_id]['0'], level

                    # Fallback: detect type from paragraph text pattern
                    text = paragraph.text.strip()
                    if re.match(r"^\d+[.)]", text) or re.match(r"^[a-zA-Z][.)]", text):
                        return "number", level
                    else:
                        return "bullet", level
        except Exception:
            pass

    # Check for built-in list styles
    style_name = paragraph.style.name if paragraph.style else None
    if not style_name:
        return None, 0

    base_type = None
    style_level = 1

    # Handle "List Paragraph" style - check for numbering properties above
    if style_name == "List Paragraph":
        # If we got here, numbering properties weren't found or processed
        # This might be a list paragraph without proper numbering - treat as bullet by default
        return "bullet", 1
    elif match := re.match(r"List\s*Bullet\s?(?P<level>\d+)?", style_name, re.I):
        base_type = "bullet"
        style_level = int(match.group("level") or 1)
    elif match := re.match(r"List\s*Number\s?(?P<level>\d+)?", style_name, re.I):
        base_type = "number"
        style_level = int(match.group("level") or 1)

    # Check indentation level for additional nesting
    indent_level = 0
    try:
        indent = paragraph.paragraph_format.left_indent
        if indent:
            # Convert Pt to level (assume DEFAULT_INDENTATION_PT_PER_LEVEL per level)
            indent_level = int(indent.pt / DEFAULT_INDENTATION_PT_PER_LEVEL)
    except AttributeError:
        pass

    # If we have a list style, combine with indentation
    if base_type:
        final_level = max(style_level, style_level + indent_level)
        return base_type, final_level

    # Check indentation level for paragraphs without list styles
    if indent_level > 0:
        # Try to detect if numbered based on paragraph text
        if re.match(r"^\d+[.)]", paragraph.text.strip()):
            return "number", indent_level
        return "bullet", indent_level

    return None, 0


def _process_hyperlink(run: Any) -> tuple[str | None, Any]:
    """Extract hyperlink URL from a run."""
    if Hyperlink is not None and isinstance(run, Hyperlink):
        return run.url, run
    return None, run


def _get_run_formatting_key(run: Any) -> tuple[bool, bool, bool, bool, bool, bool]:
    """Get a tuple of formatting attributes to use as a key for grouping similarly formatted runs."""
    return (
        bool(run.bold if hasattr(run, "bold") else False),
        bool(run.italic if hasattr(run, "italic") else False),
        bool(run.underline if hasattr(run, "underline") else False),
        bool(run.font.strike if hasattr(run, "font") else False),
        bool(run.font.subscript if hasattr(run, "font") else False),
        bool(run.font.superscript if hasattr(run, "font") else False),
        # bool(_process_hyperlink(run))  # Include hyperlink in formatting key
    )


def _format_list_marker(list_type: str, number: int = 1, level: int = 1, bullet_symbols: str = "*-+") -> str:
    """Generate a properly formatted list marker."""
    if list_type == "bullet":
        # Use different bullet symbols for different levels
        symbol_index = (level - 1) % len(bullet_symbols)
        return bullet_symbols[symbol_index] + " "
    else:
        return f"{number}. "


def _process_paragraph_runs(paragraph: "Paragraph", md_options: MarkdownOptions | None = None) -> str:
    """Process all runs in a paragraph, combining similarly formatted runs."""

    grouped_runs: list[tuple[str, tuple[bool, bool, bool, bool, bool, bool, bool] | None, str | None]] = []
    current_text: list[str] = []
    current_format: tuple[bool, bool, bool, bool, bool, bool, bool] | None = None
    current_url: str | None = None

    for run in paragraph.iter_inner_content():
        url, run_to_parse = _process_hyperlink(run)
        format_key: tuple[bool, bool, bool, bool, bool, bool, bool] = (
            *_get_run_formatting_key(run_to_parse),
            url is not None,
        )

        # Start new group if format changes or hyperlink changes
        if format_key != current_format or url != current_url:
            if current_text:
                grouped_runs.append(("".join(current_text), current_format, current_url))
                current_text = []
            current_format = format_key
            current_url = url

        # Handle hyperlink text extraction - concatenate all runs in hyperlink
        if Hyperlink is not None and isinstance(run_to_parse, Hyperlink):
            # Extract text from all runs in the hyperlink
            hyperlink_text = "".join(run.text for run in run_to_parse.runs)
            current_text.append(hyperlink_text)
        else:
            current_text.append(run_to_parse.text)

    # Add final group
    if current_text:
        grouped_runs.append(("".join(current_text), current_format, current_url))

    # Process each group
    text_parts = []
    for text, format_key, url in grouped_runs:  # type: ignore[assignment]
        # Skip empty groups
        if not text.strip():
            text_parts.append(text)
            continue

        # Preserve whitespace
        content = text.strip()
        prefix = text[: len(text) - len(text.lstrip())]
        suffix = text[len(text.rstrip()):]

        # Apply escaping if enabled, but skip for already-formatted markdown images
        # Check if content is a markdown image: ![alt](url) or ![alt]
        is_markdown_image = re.match(r'^!\[.*?\](?:\(.*?\))?$', content)
        if md_options and md_options.escape_special and not is_markdown_image:
            content = escape_markdown_special(content)

        # Apply formatting using format_special_text for special formatting and markers for others
        if format_key:
            # Handle special formatting types that require format_special_text
            if format_key[2]:  # underline
                underline_mode = md_options.underline_mode if md_options else "html"
                content = format_special_text(content, "underline", underline_mode)
            if format_key[4]:  # subscript
                subscript_mode = md_options.subscript_mode if md_options else "html"
                content = format_special_text(content, "subscript", subscript_mode)
            if format_key[5]:  # superscript
                superscript_mode = md_options.superscript_mode if md_options else "html"
                content = format_special_text(content, "superscript", superscript_mode)

            # Handle standard markdown formatting
            if format_key[0]:  # bold
                content = f"**{content}**"
            if format_key[1]:  # italic
                content = f"*{content}*"
            if format_key[3]:  # strike
                content = f"~~{content}~~"

        # Add hyperlink if present
        if url:
            content = f"[{content}]({url})"

        text_parts.append(prefix + content + suffix)

    return "".join(text_parts)


def _convert_table_to_markdown(table: "Table", md_options: MarkdownOptions | None = None) -> str:
    """Convert a docx table to markdown format."""
    markdown_rows = []

    # Process header row
    if len(table.rows) > 0:
        header_cell_md = []
        for cell in table.rows[0].cells:
            cell_md = "".join(_process_paragraph_runs(p) for p in cell.paragraphs)
            header_cell_md.append(cell_md)

        # header = [cell.text.strip() for cell in table.rows[0].cells]
        markdown_rows.append("| " + " | ".join(header_cell_md) + " |")
        # Add separator row
        markdown_rows.append("| " + " | ".join(["---"] * len(header_cell_md)) + " |")

        # Process data rows
        for row in table.rows[1:]:
            row_md = []
            for cell in row.cells:
                cell_md = "".join(_process_paragraph_runs(p) for p in cell.paragraphs)
                row_md.append(cell_md)

            # cells = [cell.text.strip() for cell in row.cells]
            markdown_rows.append("| " + " | ".join(row_md) + " |")

    return "\n".join(markdown_rows)


def _iter_block_items(
    parent: Any, options: DocxOptions, base_filename: str = "document", attachment_sequencer=None
) -> Any:
    """
    Generate a sequence of Paragraph and Table elements in order, handling images.
    """
    import docx.document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent_elm = parent.element.body if isinstance(parent, docx.document.Document) else parent._element

    for child in parent_elm.iterchildren():
        if child.tag.endswith("tbl"):
            yield Table(child, parent)
        elif child.tag.endswith("p"):
            paragraph = Paragraph(child, parent)

            # Check if paragraph contains an image
            has_image = False
            img_data = []

            for run in paragraph.runs:
                for pic in run._element.findall(
                        ".//pic:pic",
                        {"pic": "http://schemas.openxmlformats.org/drawingml/2006/picture"},
                ):
                    has_image = True

                    # Get image info
                    title = None

                    if (t := run._element.xpath(".//wp:docPr/@descr")) or (
                            t := run._element.xpath(".//wp:docPr/@title")
                    ):
                        title = t[0]

                    # Get image data and detected format
                    blip = pic.xpath(".//a:blip")[0]
                    blip_rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    raw_image_data, detected_extension = extract_docx_image_data(parent, blip_rId)

                    # Handle pre-formatted data URIs (for backward compatibility with tests)
                    if isinstance(raw_image_data, str):
                        # This is already a formatted URI, use it directly
                        processed_image = f"![{title or 'image'}]({raw_image_data})"
                        img_data.append((title or "image", processed_image))
                        continue
                    elif not isinstance(raw_image_data, (bytes, type(None))):
                        logger.warning(f"Invalid image data type for image '{title or 'unnamed'}', skipping")
                        raw_image_data = None

                    # Use detected extension or fallback to png
                    extension = detected_extension or "png"

                    # Log format detection result
                    if detected_extension:
                        logger.debug(f"Detected image format: {detected_extension}")
                    else:
                        logger.debug("No image format detected, using PNG as fallback")

                    # Process image using unified attachment handling
                    # Use sequencer if available, otherwise fall back to manual counting
                    if attachment_sequencer:
                        image_filename, _ = attachment_sequencer(
                            base_stem=base_filename,
                            format_type="general",
                            extension=extension
                        )
                    else:
                        from all2md.utils.attachments import generate_attachment_filename
                        image_filename = generate_attachment_filename(
                            base_stem=base_filename,
                            format_type="general",
                            sequence_num=len(img_data) + 1,
                            extension=extension
                        )
                    processed_image = process_attachment(
                        attachment_data=raw_image_data,
                        attachment_name=image_filename,
                        alt_text=title or "image",
                        attachment_mode=options.attachment_mode,
                        attachment_output_dir=options.attachment_output_dir,
                        attachment_base_url=options.attachment_base_url,
                        is_image=True,
                        alt_text_mode=options.alt_text_mode,
                    )

                    img_data.append((title, processed_image))

            if has_image and img_data:
                # Create a new paragraph with default style for each image
                from docx.oxml.shared import OxmlElement

                for _title, processed_image in img_data:
                    p = OxmlElement("w:p")
                    r = OxmlElement("w:r")
                    t = OxmlElement("w:t")

                    # The processed_image already contains the proper markdown
                    t.text = processed_image

                    r.append(t)
                    p.append(r)

                    clean_paragraph = Paragraph(p, parent)
                    clean_paragraph.style = parent.styles["Normal"]
                    yield clean_paragraph
            elif not has_image:
                yield paragraph


def extract_docx_metadata(doc: "docx.document.Document") -> DocumentMetadata:
    """Extract metadata from DOCX document.

    Parameters
    ----------
    doc : docx.document.Document
        python-docx Document object

    Returns
    -------
    DocumentMetadata
        Extracted metadata
    """
    if not hasattr(doc, 'core_properties'):
        return DocumentMetadata()

    props = doc.core_properties

    # Use the utility function for standard metadata extraction
    metadata = map_properties_to_metadata(props, OFFICE_FIELD_MAPPING)

    # Add DOCX-specific custom metadata
    custom_properties = ['last_modified_by', 'revision', 'version', 'comments']
    for prop_name in custom_properties:
        if hasattr(props, prop_name):
            value = getattr(props, prop_name)
            if value:
                metadata.custom[prop_name] = value

    return metadata


def docx_to_markdown(
        input_data: Union[str, Path, "docx.document.Document", IO[bytes]], options: DocxOptions | None = None
) -> str:
    """Convert Word document (DOCX) to Markdown format.

    Processes Microsoft Word documents and converts them to well-formatted
    Markdown while preserving text formatting, document structure, tables,
    lists, and embedded images. Handles complex document elements and
    maintains hierarchical organization.

    Parameters
    ----------
    input_data : str, file-like object, or docx.Document
        Word document to convert. Can be:
        - String path to DOCX file
        - File-like object containing DOCX data
        - Already opened python-docx Document object
    options : DocxOptions or None, default None
        Configuration options for DOCX conversion. If None, uses default settings.

    Returns
    -------
    str
        Markdown representation of the Word document with preserved
        formatting, structure, and content.

    Raises
    ------
    InputError
        If input type is not supported or document cannot be opened
    MarkdownConversionError
        If document processing fails

    Examples
    --------
    Convert a Word file to Markdown:

        >>> with open('document.docx', 'rb') as f:
        ...     markdown = docx_to_markdown(f)
        >>> print(markdown)

    Convert with embedded images:

        >>> markdown = docx_to_markdown('document.docx', convert_images_to_base64=True)

    Notes
    -----
    - Preserves text formatting (bold, italic, underline, strikethrough)
    - Converts lists with proper nesting and numbering
    - Handles tables with Markdown table syntax
    - Processes hyperlinks and converts to Markdown link format
    - Maintains document structure with appropriate heading levels
    """
    # Handle backward compatibility and merge options
    try:
        import docx
        import docx.document
        from docx.table import Table
        from docx.text.hyperlink import Hyperlink  # noqa: F401
        from docx.text.paragraph import Paragraph
    except ImportError as e:
        from all2md.exceptions import DependencyError
        raise DependencyError(
            converter_name="docx",
            missing_packages=[("python-docx", "")],
        ) from e

    if options is None:
        options = DocxOptions()

    # Extract base filename for standardized attachment naming
    if isinstance(input_data, (str, Path)):
        base_filename = Path(input_data).stem
    else:
        # For non-file inputs, use a default name
        base_filename = "document"

    # Validate ZIP archive security for file-based inputs
    if isinstance(input_data, (str, Path)):
        try:
            validate_zip_archive(input_data)
        except Exception as e:
            raise MarkdownConversionError(
                f"DOCX archive failed security validation: {str(e)}",
                conversion_stage="archive_validation",
                original_error=e
            ) from e

    # Validate and convert input - for now use simplified approach
    try:
        if isinstance(input_data, docx.document.Document):
            doc = input_data
        elif isinstance(input_data, Path):
            doc = docx.Document(str(input_data))
        else:
            doc = docx.Document(input_data)
    except Exception as e:
        raise MarkdownConversionError(
            f"Failed to open DOCX document: {str(e)}", conversion_stage="document_opening", original_error=e
        ) from e

    # Extract metadata if requested
    metadata = None
    if options.extract_metadata:
        metadata = extract_docx_metadata(doc)

    # Use new AST-based conversion path
    from all2md.converters.docx2ast import DocxToAstConverter
    from all2md.ast import MarkdownRenderer

    # Create attachment sequencer for consistent filename generation
    attachment_sequencer = create_attachment_sequencer()

    # Convert DOCX to AST
    ast_converter = DocxToAstConverter(
        options=options,
        doc=doc,
        base_filename=base_filename,
        attachment_sequencer=attachment_sequencer,
    )
    ast_document = ast_converter.convert_to_ast(doc)

    # Get MarkdownOptions (use provided or create default)
    md_opts = options.markdown_options if options.markdown_options else MarkdownOptions()

    # Render AST to markdown using MarkdownOptions directly
    renderer = MarkdownRenderer(md_opts)
    markdown = renderer.render(ast_document)

    # Prepend metadata if enabled
    markdown = prepend_metadata_if_enabled(markdown.strip(), metadata, options.extract_metadata)

    return markdown
