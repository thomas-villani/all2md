#  Copyright (c) 2025 Tom Villani, Ph.D.
#
# src/all2md/parsers/docx.py
"""DOCX to AST converter.

This module provides conversion from Microsoft Word DOCX documents to AST representation.
It replaces direct markdown string generation with structured AST building,
enabling multiple rendering strategies and improved testability.

"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import IO, TYPE_CHECKING, Any, Callable, cast

from all2md.constants import DEFAULT_INDENTATION_PT_PER_LEVEL, DEPS_DOCX
from all2md.exceptions import MalformedFileError
from all2md.utils.attachments import create_attachment_sequencer, extract_docx_image_data, process_attachment
from all2md.utils.metadata import (
    OFFICE_FIELD_MAPPING,
    DocumentMetadata,
    map_properties_to_metadata,
)
from all2md.utils.parser_helpers import append_attachment_footnotes

if TYPE_CHECKING:
    import docx.document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

from all2md.ast import (
    CodeBlock,
    Comment,
    CommentInline,
    Document,
    Emphasis,
    FootnoteReference,
    Heading,
    Image,
    LineBreak,
    Link,
    List,
    ListItem,
    MathBlock,
    MathInline,
    Node,
    Strikethrough,
    Strong,
    Subscript,
    Superscript,
    TableCell,
    TableRow,
    Text,
    ThematicBreak,
    Underline,
)
from all2md.ast import (
    Paragraph as AstParagraph,
)
from all2md.ast import (
    Table as AstTable,
)
from all2md.converter_metadata import ConverterMetadata
from all2md.options.docx import DocxOptions
from all2md.parsers.base import BaseParser
from all2md.progress import ProgressCallback
from all2md.utils.decorators import requires_dependencies
from all2md.utils.footnotes import FootnoteCollector

logger = logging.getLogger(__name__)

WORDPROCESSING_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WORD_TAG_PREFIX = f"{{{WORDPROCESSING_NS}}}"
WORD_ID_ATTR = f"{WORD_TAG_PREFIX}id"
WORD_PARAGRAPH_TAG = f"{WORD_TAG_PREFIX}p"
WORD_FOOTNOTE_REFERENCE_TAG = f"{WORD_TAG_PREFIX}footnoteReference"
WORD_ENDNOTE_REFERENCE_TAG = f"{WORD_TAG_PREFIX}endnoteReference"
WORD_COMMENT_REFERENCE_TAG = f"{WORD_TAG_PREFIX}commentReference"
WORD_COMMENT_RANGE_START_TAG = f"{WORD_TAG_PREFIX}commentRangeStart"
WORD_COMMENT_RANGE_END_TAG = f"{WORD_TAG_PREFIX}commentRangeEnd"
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
MATH_TAG_PREFIX = f"{{{MATH_NS}}}"


@dataclass
class ImageData:
    """Structured representation of image data from DOCX.

    Parameters
    ----------
    url : str
        Image URL or data URI
    alt_text : str
        Alternative text for the image
    title : str or None
        Optional title for the image
    footnote_label : str or None
        Footnote label if using footnote mode
    footnote_content : str or None
        Footnote content text
    source_data : str or None
        Source of the image data (e.g., "base64", "downloaded")

    """

    url: str
    alt_text: str
    title: str | None = None
    footnote_label: str | None = None
    footnote_content: str | None = None
    source_data: str | None = None


@dataclass
class CommentData:
    """Comment data from DOCX document.

    Parameters
    ----------
    identifier : str
        Comment ID
    label : str
        Comment label/number
    author : str
        Comment author name
    date : str
        Comment date
    text : str
        Comment text content

    """

    identifier: str
    label: str
    author: str
    date: str
    text: str


class DocxToAstConverter(BaseParser):
    """Convert DOCX to AST representation.

    This converter parses Word documents using python-docx and builds an AST
    that can be rendered to various markdown flavors.

    Parameters
    ----------
    options : DocxOptions or None, default = None
        Conversion options
    doc : docx.document.Document or None
        Document to convert (stored for list detection)
    base_filename : str
        Base filename for image attachments
    attachment_sequencer : callable or None
        Sequencer for generating attachment filenames

    """

    def __init__(self, options: DocxOptions | None = None, progress_callback: ProgressCallback | None = None):
        """Initialize the DOCX parser with options and progress callback."""
        BaseParser._validate_options_type(options, DocxOptions, "docx")
        options = options or DocxOptions()
        super().__init__(options, progress_callback)
        self.options: DocxOptions = options

        # Internally used to stash info between functions
        self._list_stack: list[tuple[str, int, list[ListItem]]] = []  # (type, level, items)
        self._numbering_defs: dict[str, dict[str, str]] | None = None
        self._footnote_collector: FootnoteCollector | None = None
        self._comments_map: dict[str, CommentData] = {}
        self._attachment_footnotes: dict[str, str] = {}  # label -> content for footnote definitions

    @requires_dependencies("docx", DEPS_DOCX)
    def parse(self, input_data: str | Path | IO[bytes] | bytes) -> Document:
        """Parse DOCX document into AST.

        This method handles loading the DOCX file and converting it to AST.
        Performs security validation, dependency checking, and error handling.

        Parameters
        ----------
        input_data : str, Path, IO[bytes], or bytes
            DOCX file to parse

        Returns
        -------
        Document
            AST document node

        Raises
        ------
        DependencyError
            If python-docx is not installed
        MalformedFileError
            If document loading fails

        """
        import docx.document

        base_filename = "document"

        # Extract base filename from input_data if it's a Path/str
        if isinstance(input_data, (str, Path)):
            base_filename = Path(input_data).stem
        elif not hasattr(input_data, "read"):
            # For non-file inputs that aren't file-like, keep existing base_filename
            pass

        # Validate ZIP archive security for all input types
        self._validate_zip_security(input_data, suffix=".docx")

        # Load the document with error handling
        try:
            if isinstance(input_data, docx.document.Document):
                doc = input_data
            elif isinstance(input_data, Path):
                doc = docx.Document(str(input_data))  # type: ignore[assignment]
            else:
                doc = docx.Document(input_data)  # type: ignore[assignment,arg-type]
        except Exception as e:
            raise MalformedFileError(
                f"Failed to open DOCX document: {str(e)}",
                file_path=str(input_data) if isinstance(input_data, (str, Path)) else None,
                original_error=e,
            ) from e

        return self.convert_to_ast(doc, base_filename)

    def extract_metadata(self, document: "docx.document.Document") -> DocumentMetadata:
        """Extract metadata from DOCX document.

        Parameters
        ----------
        document : docx.document.Document
            python-docx Document object

        Returns
        -------
        DocumentMetadata
            Extracted metadata including title, author, dates, keywords, etc.
            Returns empty DocumentMetadata if no metadata is available.

        Notes
        -----
        Uses the OFFICE_FIELD_MAPPING to map core document properties to
        standard metadata fields. Also extracts DOCX-specific custom metadata
        properties such as last_modified_by, revision, version, and comments.

        """
        if not hasattr(document, "core_properties"):
            return DocumentMetadata()

        props = document.core_properties

        # Use the utility function for standard metadata extraction
        metadata = map_properties_to_metadata(props, OFFICE_FIELD_MAPPING)

        # Add DOCX-specific custom metadata
        custom_properties = ["last_modified_by", "revision", "version", "comments"]
        for prop_name in custom_properties:
            if hasattr(props, prop_name):
                value = getattr(props, prop_name)
                if value:
                    metadata.custom[prop_name] = value

        return metadata

    def _process_document_blocks(self, doc: "docx.document.Document", base_filename: str, children: list[Node]) -> None:
        """Process document blocks and append to children list.

        Parameters
        ----------
        doc : docx.document.Document
            DOCX document to process
        base_filename : str
            Base filename for attachments
        children : list[Node]
            List to append processed nodes to

        """
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        for block in _iter_block_items(
            doc,
            options=self.options,
            base_filename=base_filename,
            attachment_sequencer=create_attachment_sequencer(),
        ):
            if isinstance(block, ImageData):
                self._process_image_block(block, children)
            elif isinstance(block, Paragraph):
                self._process_paragraph_block(block, doc, children)
            elif isinstance(block, Table):
                self._process_table_block(block, children)

    def _process_image_block(self, block: ImageData, children: list[Node]) -> None:
        """Process an image block.

        Parameters
        ----------
        block : ImageData
            Image data to process
        children : list[Node]
            List to append processed nodes to

        """
        # Collect footnote info if present
        if block.footnote_label and block.footnote_content:
            self._attachment_footnotes[block.footnote_label] = block.footnote_content
        # Handle ImageData objects directly
        image_node = Image(url=block.url, alt_text=block.alt_text, title=block.title)
        if block.source_data:
            image_node.metadata["source_data"] = block.source_data
        children.append(AstParagraph(content=[image_node]))

    def _process_paragraph_block(self, block: "Paragraph", doc: "docx.document.Document", children: list[Node]) -> None:
        """Process a paragraph block.

        Parameters
        ----------
        block : Paragraph
            Paragraph to process
        doc : docx.document.Document
            Parent document
        children : list[Node]
            List to append processed nodes to

        """
        nodes = self._process_paragraph_to_ast(block, doc)
        if nodes:
            if isinstance(nodes, list):
                children.extend(nodes)
            else:
                children.append(nodes)

    def _process_table_block(self, block: "Table", children: list[Node]) -> None:
        """Process a table block.

        Parameters
        ----------
        block : Table
            Table to process
        children : list[Node]
            List to append processed nodes to

        """
        if self.options.preserve_tables:
            table_node = self._process_table_to_ast(block)
            if table_node:
                children.append(table_node)
        else:
            # Flatten table to paragraphs
            for row in block.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        inline_nodes = self._process_paragraph_runs_to_inline(paragraph)
                        if inline_nodes:
                            children.append(AstParagraph(content=inline_nodes))

    def _finalize_lists_and_notes(self, doc: "docx.document.Document", children: list[Node]) -> None:
        """Finalize lists and add footnotes/endnotes/comments to document.

        Parameters
        ----------
        doc : docx.document.Document
            DOCX document
        children : list[Node]
            List to append processed nodes to

        """
        # Finalize any remaining list at the end of document
        if self._list_stack:
            final_list = self._finalize_current_list()
            if final_list:
                children.append(final_list)
            self._list_stack = []

        # Add footnotes and endnotes if requested
        if self.options.include_footnotes:
            self._process_footnotes(doc)

        if self.options.include_endnotes:
            self._process_endnotes(doc)

        # Add footnote/endnote definitions
        if self._footnote_collector:
            priority: list[str] = []
            if self.options.include_footnotes:
                priority.append("footnote")
            if self.options.include_endnotes:
                priority.append("endnote")

            for definition in self._footnote_collector.iter_definitions(note_type_priority=priority):
                children.append(definition)

        # Add comments if configured as footnotes
        if self.options.include_comments and self.options.comments_position == "footnotes":
            comments_nodes = self._process_comments()
            if comments_nodes:
                children.extend(comments_nodes)

        # Append attachment footnote definitions if any were collected
        if self.options.attachments_footnotes_section:
            append_attachment_footnotes(
                children, self._attachment_footnotes, self.options.attachments_footnotes_section
            )

    def convert_to_ast(self, doc: "docx.document.Document", base_filename: str = "document") -> Document:
        """Convert DOCX document to AST Document.

        Parameters
        ----------
        doc : docx.document.Document
            DOCX document to convert
        base_filename : str, default="document"
            Base filename for attachments

        Returns
        -------
        Document
            AST document node

        """
        self._numbering_defs = None
        self._list_stack = []
        self._footnote_collector = FootnoteCollector()
        self._comments_map = {}
        self._attachment_footnotes = {}

        # Emit started event
        self._emit_progress("started", "Converting DOCX document", current=0, total=0)

        children: list[Node] = []

        # Load comments if requested
        if self.options.include_comments:
            self._comments_map = self._load_comments(doc)

        # Process all document blocks
        self._process_document_blocks(doc, base_filename, children)

        # Finalize lists and add notes/comments
        self._finalize_lists_and_notes(doc, children)

        # Emit finished event
        self._emit_progress("finished", "DOCX conversion completed", current=1, total=1)

        # Extract and attach metadata if requested
        metadata_dict = {}
        if self.options.extract_metadata:
            metadata = self.extract_metadata(doc)
            metadata_dict = metadata.to_dict()

        document = Document(children=children, metadata=metadata_dict)
        self._footnote_collector = None
        self._comments_map = {}
        return document

    def _process_paragraph_to_ast(
        self, paragraph: "Paragraph", doc: "docx.document.Document"
    ) -> Node | list[Node] | None:
        """Process a DOCX paragraph to AST nodes.

        Parameters
        ----------
        paragraph : Paragraph
            DOCX paragraph to process
        doc : docx.document.Document
            The document being processed (needed for heading match)

        Returns
        -------
        Node, list of Node, or None
            Resulting AST node(s)

        """
        # Handle heading styles
        style_name = paragraph.style.name if paragraph.style else ""
        heading_match = re.match(r"Heading (\d+)", style_name)
        if heading_match:
            level = int(heading_match.group(1))
            content = self._process_paragraph_runs_to_inline(paragraph)
            return Heading(level=level, content=content)

        # Handle code block styles
        if style_name and self.options.code_style_names:
            for code_style in self.options.code_style_names:
                # Support both exact match and partial match (case-insensitive)
                if code_style.lower() in style_name.lower():
                    code_text = paragraph.text
                    return CodeBlock(content=code_text)

        # Handle horizontal rules (thematic breaks)
        # Detect paragraphs that look like horizontal rules:
        # 1. Empty or contain only special characters like --- or ___
        # 2. Have a bottom border that spans the paragraph
        text = paragraph.text.strip()
        if text in ("---", "___", "***", "—" * 3, "–" * 3) or (not text and self._has_bottom_border(paragraph)):
            return ThematicBreak()

        # Handle lists

        list_type, level = _detect_list_level(paragraph, doc)
        if list_type:
            # This paragraph is part of a list
            # We'll accumulate list items and return them when list ends
            return self._process_list_item_paragraph(paragraph, list_type, level)

        math_blocks = self._extract_math_blocks_from_paragraph(paragraph)

        # Not a list - clear list stack and return any accumulated list
        if self._list_stack:
            accumulated_list = self._finalize_current_list()
            self._list_stack = []
            content = self._process_paragraph_runs_to_inline(paragraph)
            nodes: list[Node] = []
            if accumulated_list:
                nodes.append(accumulated_list)
            if content:
                nodes.append(AstParagraph(content=content))
            if math_blocks:
                nodes.extend(math_blocks)
            if not nodes:
                return None
            if len(nodes) == 1:
                return nodes[0]
            return nodes

        # Regular paragraph
        content = self._process_paragraph_runs_to_inline(paragraph)
        if content:
            # Skip paragraphs that only contain whitespace
            has_non_whitespace = any(not isinstance(node, Text) or node.content.strip() for node in content)
            if has_non_whitespace:
                if math_blocks:
                    return [AstParagraph(content=content), *math_blocks]
                return AstParagraph(content=content)

        if math_blocks:
            if len(math_blocks) == 1:
                return math_blocks[0]
            return cast(list[Node], math_blocks)

        return None

    def _process_list_item_paragraph(self, paragraph: "Paragraph", list_type: str, level: int) -> Node | None:
        """Process a paragraph that is part of a list.

        This method accumulates list items and handles nesting properly.
        Returns a completed list when we transition to a non-list item.

        Parameters
        ----------
        paragraph : Paragraph
            List item paragraph
        list_type : str
            'bullet' or 'number'
        level : int
            Nesting level (1-based)

        Returns
        -------
        Node or None
            Completed list node if transitioning out, None if still accumulating

        """
        # Process paragraph content
        content = self._process_paragraph_runs_to_inline(paragraph)
        item_node = ListItem(children=[AstParagraph(content=content)])

        # Handle level changes
        if not self._list_stack:
            # Start new list at this level
            self._list_stack = [(list_type, level, [item_node])]
            return None

        current_type, current_level, current_items = self._list_stack[-1]

        if level > current_level:
            # Nested list - deeper level
            # Start new list at the deeper level
            self._list_stack.append((list_type, level, [item_node]))
            return None

        elif level < current_level:
            # Going back to shallower level - need to finalize deeper lists
            # Pop and nest all lists deeper than the target level
            while self._list_stack and self._list_stack[-1][1] > level:
                # Pop the deeper list and nest it
                popped_type, popped_level, popped_items = self._list_stack.pop()
                nested_list = List(ordered=(popped_type == "number"), items=popped_items, tight=True)

                # Add nested list to the last item of the parent level
                if self._list_stack:
                    parent_items = self._list_stack[-1][2]
                    if parent_items:
                        parent_items[-1].children.append(nested_list)

            # Now at the correct level - check if type matches
            if self._list_stack and self._list_stack[-1][0] == list_type and self._list_stack[-1][1] == level:
                # Same level and type - add item
                self._list_stack[-1][2].append(item_node)
                return None  # Still accumulating
            else:
                # Different type at this level - finalize old, start new
                result_node = None
                if self._list_stack and self._list_stack[-1][1] == level:
                    old_type, old_level, old_items = self._list_stack.pop()
                    result_node = List(ordered=(old_type == "number"), items=old_items, tight=True)
                self._list_stack.append((list_type, level, [item_node]))
                return result_node

        else:
            # Same level
            if current_type == list_type:
                # Same type - add to current list
                current_items.append(item_node)
                return None
            else:
                # Different type at same level - finalize old, start new
                old_type, old_level, old_items = self._list_stack.pop()
                self._list_stack.append((list_type, level, [item_node]))
                return List(ordered=(old_type == "number"), items=old_items, tight=True)

    def _finalize_current_list(self) -> List | None:
        """Finalize all lists in the stack, nesting them properly.

        Returns
        -------
        List or None
            Completed top-level list node with all nesting

        """
        if not self._list_stack:
            return None

        # Pop and nest all lists from deepest to shallowest
        while len(self._list_stack) > 1:
            # Pop deeper list
            deeper_type, deeper_level, deeper_items = self._list_stack.pop()
            nested_list = List(ordered=(deeper_type == "number"), items=deeper_items, tight=True)

            # Add to parent's last item
            parent_items = self._list_stack[-1][2]
            if parent_items:
                parent_items[-1].children.append(nested_list)

        # Return the top-level list
        if self._list_stack:
            list_type, level, items = self._list_stack.pop()
            return List(ordered=(list_type == "number"), items=items, tight=True)

        return None

    def _append_text_with_line_breaks(
        self,
        text: str,
        current_text: list[str],
        result: list[Node],
        flush_callback: Callable[[], None],
    ) -> None:
        """Append text to current buffer, handling line breaks.

        Parameters
        ----------
        text : str
            Text to append (may contain newlines)
        current_text : list[str]
            Current text buffer to append to
        result : list[Node]
            Result list to add line break nodes to
        flush_callback : Callable
            Function to flush current text buffer

        """
        if "\n" in text:
            parts = text.split("\n")
            for i, part in enumerate(parts):
                if part:
                    current_text.append(part)
                if i < len(parts) - 1:
                    flush_callback()
                    result.append(LineBreak(soft=False))
        else:
            current_text.append(text)

    def _process_paragraph_runs_to_inline(self, paragraph: "Paragraph") -> list[Node]:
        """Process paragraph runs to inline AST nodes.

        Parameters
        ----------
        paragraph : Paragraph
            Paragraph containing runs

        Returns
        -------
        list of Node
            List of inline AST nodes

        Notes
        -----
        This method cannot use the `group_and_format_runs` helper because it handles
        DOCX-specific complexities including:
        - Hyperlink grouping (different URLs = different groups)
        - Inline math extraction (OMML to LaTeX)
        - Footnote and endnote references
        - Comment references
        - Special text extraction for Hyperlink objects

        """
        result: list[Node] = []
        current_text: list[str] = []
        current_format: tuple[bool, bool, bool, bool, bool, bool, bool] | None = None
        current_url: str | None = None

        def flush_group() -> None:
            if not current_text:
                return
            text_value = "".join(current_text)
            result.append(self._build_formatted_inline_node(text_value, current_format, current_url))
            current_text.clear()

        from docx.text.hyperlink import Hyperlink

        for run in paragraph.iter_inner_content():
            url, run_to_parse = self._process_hyperlink(run)
            format_key = self._get_run_formatting_key(run_to_parse, url is not None)

            if format_key != current_format or url != current_url:
                flush_group()
                current_format = format_key
                current_url = url

            math_nodes = self._extract_math_from_run(run_to_parse)
            if math_nodes:
                flush_group()
                result.extend(math_nodes)
                continue

            note_nodes = self._extract_note_reference_nodes(run_to_parse)
            comment_nodes = self._extract_comment_reference_nodes(run_to_parse)
            if note_nodes or comment_nodes:
                flush_group()
                if note_nodes:
                    result.extend(note_nodes)
                if comment_nodes:
                    result.extend(comment_nodes)

            if isinstance(run_to_parse, Hyperlink):
                hyperlink_text = "".join(r.text for r in run_to_parse.runs)
                if hyperlink_text:
                    self._append_text_with_line_breaks(hyperlink_text, current_text, result, flush_group)
            else:
                run_text = run_to_parse.text
                if run_text:
                    self._append_text_with_line_breaks(run_text, current_text, result, flush_group)

        flush_group()
        return result

    def _extract_math_blocks_from_paragraph(self, paragraph: "Paragraph") -> list[MathBlock]:
        if not hasattr(paragraph, "_element"):
            return []

        blocks: list[MathBlock] = []
        for child in paragraph._element:
            tag = getattr(child, "tag", None)
            if not isinstance(tag, str):
                continue
            if _omml_local_name(tag) != "oMathPara":
                continue
            latex = _omml_to_latex(child)
            if latex:
                blocks.append(MathBlock(content=latex))
        return blocks

    def _extract_math_from_run(self, run: Any) -> list[MathInline]:
        element = getattr(run, "_element", None)
        if element is None:
            return []

        nodes: list[MathInline] = []
        for math_elem in element.findall(f".//{MATH_TAG_PREFIX}oMath"):
            parent = getattr(math_elem, "getparent", lambda: None)()
            if parent is not None and _omml_local_name(getattr(parent, "tag", "")) == "oMathPara":
                continue
            latex = _omml_to_latex(math_elem)
            if latex:
                nodes.append(MathInline(content=latex))
        return nodes

    def _build_formatted_inline_node(
        self,
        text: str,
        format_key: tuple[bool, bool, bool, bool, bool, bool, bool] | None,
        url: str | None,
    ) -> Node:
        inline_node: Node = Text(content=text)

        if format_key:
            if format_key[2]:  # underline
                inline_node = Underline(content=[inline_node])
            if format_key[4]:  # subscript
                inline_node = Subscript(content=[inline_node])
            if format_key[5]:  # superscript
                inline_node = Superscript(content=[inline_node])
            if format_key[0]:  # bold
                inline_node = Strong(content=[inline_node])
            if format_key[1]:  # italic
                inline_node = Emphasis(content=[inline_node])
            if format_key[3]:  # strike
                inline_node = Strikethrough(content=[inline_node])

        if url:
            inline_node = Link(url=url, content=[inline_node])

        return inline_node

    def _extract_note_reference_nodes(self, run: Any) -> list[Node]:
        from docx.text.hyperlink import Hyperlink

        if isinstance(run, Hyperlink):
            return []

        collector = self._footnote_collector
        if collector is None or not hasattr(run, "_element"):
            return []

        nodes: list[Node] = []
        for element in run._element.iter():
            if element.tag == WORD_FOOTNOTE_REFERENCE_TAG:
                identifier = element.get(WORD_ID_ATTR)
                canonical_id = self._register_note_reference(identifier, note_type="footnote")
                if canonical_id:
                    nodes.append(FootnoteReference(identifier=canonical_id))
            elif element.tag == WORD_ENDNOTE_REFERENCE_TAG:
                identifier = element.get(WORD_ID_ATTR)
                canonical_id = self._register_note_reference(identifier, note_type="endnote")
                if canonical_id:
                    nodes.append(FootnoteReference(identifier=canonical_id))

        return nodes

    def _extract_comment_reference_nodes(self, run: Any) -> list[Node]:
        from docx.text.hyperlink import Hyperlink

        if not self.options.include_comments or self.options.comments_position != "inline":
            return []

        if isinstance(run, Hyperlink) or not hasattr(run, "_element"):
            return []

        nodes: list[Node] = []
        for element in run._element.iter():
            if element.tag == WORD_COMMENT_REFERENCE_TAG:
                comment_id = element.get(WORD_ID_ATTR)
                nodes.extend(self._render_comment_inline(comment_id))

        return nodes

    def _register_note_reference(self, identifier: str | None, *, note_type: str) -> str | None:
        if note_type == "footnote" and not self.options.include_footnotes:
            return None
        if note_type == "endnote" and not self.options.include_endnotes:
            return None

        collector = self._footnote_collector
        if collector is None:
            return None

        raw_identifier = identifier
        if note_type == "endnote" and identifier is not None:
            raw_identifier = f"end{identifier}"

        return collector.register_reference(raw_identifier, note_type=note_type)

    def _render_comment_inline(self, identifier: str | None) -> list[Node]:
        if identifier is None:
            return []

        comment = self._comments_map.get(str(identifier))
        if comment is None:
            return []

        if not comment.text:
            return []

        # Create CommentInline node with rich metadata
        # Renderer will decide how to present it based on its own options
        return [
            CommentInline(
                content=comment.text,
                metadata={
                    "comment_type": "docx_review",
                    "identifier": comment.identifier,
                    "label": comment.label,
                    "author": comment.author,
                    "date": comment.date,
                },
            )
        ]

    def _format_comment_header(
        self,
        comment: CommentData,
        *,
        include_id: bool,
        include_prefix: bool,
    ) -> str:
        segments: list[str] = []

        if include_prefix:
            segments.append("Comment")

        if include_id:
            segments.append(comment.label)

        if comment.author:
            segments.append(comment.author)

        header = " ".join(segments).strip()

        if comment.date:
            header = f"{header} ({comment.date})" if header else comment.date

        if not header:
            return comment.label

        return header

    def _process_hyperlink(self, run: Any) -> tuple[str | None, Any]:
        """Process a run to extract hyperlink information.

        Parameters
        ----------
        run : Any
            Run object to process

        Returns
        -------
        tuple of (url, run_to_parse)
            URL if hyperlink, and the run to parse for text

        """
        from docx.text.hyperlink import Hyperlink

        if isinstance(run, Hyperlink):
            return run.address, run
        return None, run

    def _get_run_formatting_key(self, run: Any, is_hyperlink: bool) -> tuple[bool, bool, bool, bool, bool, bool, bool]:
        """Get formatting key for a run.

        Parameters
        ----------
        run : Any
            Run to analyze
        is_hyperlink : bool
            Whether this run is part of a hyperlink

        Returns
        -------
        tuple of bool
            (bold, italic, underline, strike, subscript, superscript, is_hyperlink)

        """
        from docx.text.hyperlink import Hyperlink

        # Handle Hyperlink object
        if isinstance(run, Hyperlink):
            if run.runs:
                first_run = run.runs[0]
                return (
                    first_run.bold or False,
                    first_run.italic or False,
                    first_run.underline or False,
                    first_run.font.strike or False,
                    first_run.font.subscript or False,
                    first_run.font.superscript or False,
                    is_hyperlink,
                )
            return (False, False, False, False, False, False, is_hyperlink)

        # Regular run
        return (
            run.bold or False,
            run.italic or False,
            run.underline or False,
            run.font.strike or False,
            run.font.subscript or False,
            run.font.superscript or False,
            is_hyperlink,
        )

    def _has_bottom_border(self, paragraph: "Paragraph") -> bool:
        """Check if paragraph has a bottom border that could represent a horizontal rule.

        Parameters
        ----------
        paragraph : Paragraph
            Paragraph to check for bottom border

        Returns
        -------
        bool
            True if paragraph has a significant bottom border

        """
        try:
            # Access paragraph formatting element
            if not hasattr(paragraph, "_element"):
                return False

            # Check for paragraph border bottom
            pPr = paragraph._element.pPr
            if pPr is None:
                return False

            # Check pBdr (paragraph borders)
            pBdr = pPr.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr")
            if pBdr is None:
                return False

            # Check for bottom border
            bottom = pBdr.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom")
            if bottom is not None:
                # Check if border is substantial (not just a thin line)
                val = bottom.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                if val and val != "none":
                    return True

            return False
        except AttributeError as exc:
            # Missing expected XML attributes or elements
            logger.debug(f"Could not determine paragraph border (missing attribute): {exc}")
            return False

    def _process_table_to_ast(self, table: "Table") -> AstTable | None:
        """Process a DOCX table to AST Table node.

        Parameters
        ----------
        table : Table
            DOCX table to convert

        Returns
        -------
        AstTable or None
            Table node if table has content

        """
        if len(table.rows) == 0:
            return None

        # First row is header
        header_cells: list[TableCell] = []
        for cell in table.rows[0].cells:
            cell_content = []
            for p in cell.paragraphs:
                inline_nodes = self._process_paragraph_runs_to_inline(p)
                if inline_nodes:
                    cell_content.extend(inline_nodes)
            header_cells.append(TableCell(content=cell_content))

        header_row = TableRow(cells=header_cells, is_header=True)

        # Data rows
        data_rows: list[TableRow] = []
        for row in table.rows[1:]:
            row_cells: list[TableCell] = []
            for cell in row.cells:
                cell_content = []
                for p in cell.paragraphs:
                    inline_nodes = self._process_paragraph_runs_to_inline(p)
                    if inline_nodes:
                        cell_content.extend(inline_nodes)
                row_cells.append(TableCell(content=cell_content))
            data_rows.append(TableRow(cells=row_cells))

        return AstTable(header=header_row, rows=data_rows)

    def _process_notes(
        self,
        doc: "docx.document.Document",
        relationship_type: str,
        part_attr_name: str,
        tag_name: str,
        note_type: str,
        id_prefix: str = "",
    ) -> None:
        """Populate the collector with note definitions (footnotes or endnotes).

        Parameters
        ----------
        doc : docx.document.Document
            The Word document
        relationship_type : str
            The relationship type constant (RT.FOOTNOTES or RT.ENDNOTES)
        part_attr_name : str
            The attribute name for the part ("footnotes_part" or "endnotes_part")
        tag_name : str
            The XML tag name to search for ("footnote" or "endnote")
        note_type : str
            The note type identifier ("footnote" or "endnote")
        id_prefix : str
            Optional prefix to add to note IDs (e.g., "end" for endnotes)

        """
        collector = self._footnote_collector
        if collector is None:
            return

        try:
            note_part = self._get_note_part(doc, relationship_type, part_attr_name)
            if note_part is None:
                return

            element = self._get_note_part_element(note_part)
            if element is None:
                return

            for note in element.findall(f".//{WORD_TAG_PREFIX}{tag_name}"):
                note_id = note.get(WORD_ID_ATTR)
                if note_id in {"-1", "0"}:
                    continue

                content_nodes = self._build_note_definition_content(note, note_part)
                if content_nodes:
                    collector.register_definition(f"{id_prefix}{note_id}", content_nodes, note_type=note_type)
        except (AttributeError, KeyError) as exc:
            logger.debug(f"Could not access {note_type}s (missing element): {exc}")

    def _process_footnotes(self, doc: "docx.document.Document") -> None:
        """Populate the collector with footnote definitions from the document."""
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
        except ImportError:
            return

        self._process_notes(doc, RT.FOOTNOTES, "footnotes_part", "footnote", "footnote")

    def _process_endnotes(self, doc: "docx.document.Document") -> None:
        """Populate the collector with endnote definitions from the document."""
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
        except ImportError:
            return

        self._process_notes(doc, RT.ENDNOTES, "endnotes_part", "endnote", "endnote", id_prefix="end")

    def _get_note_part(
        self,
        doc: "docx.document.Document",
        relationship_type: str,
        attr_name: str | None,
    ) -> Any | None:
        """Return a related note part if available."""
        if attr_name:
            part = getattr(doc.part, attr_name, None)
            if part is not None:
                return part

        try:
            return doc.part.part_related_by(relationship_type)
        except (KeyError, AttributeError):
            package = getattr(doc.part, "package", None)
            if package is None:
                return None
            try:
                return package.part_related_by(relationship_type)
            except KeyError:
                return None

    def _get_note_part_element(self, note_part: Any) -> Any | None:
        """Return the XML element root for the supplied note part."""
        element = getattr(note_part, "element", None)
        if element is not None:
            return element

        blob = getattr(note_part, "blob", None)
        if not blob:
            return None

        try:
            from lxml import etree
        except ImportError as exc:
            logger.debug(f"lxml not available for parsing note XML: {exc}")
            return None

        try:
            return etree.fromstring(blob)
        except (ValueError, UnicodeDecodeError) as exc:
            logger.debug(f"Could not parse note part XML (invalid data): {exc}")
            return None

    def _build_note_definition_content(self, note_element: Any, note_part: Any) -> list[Node]:
        """Create block content for a single footnote or endnote definition."""
        content_nodes: list[Node] = []
        for paragraph_element in note_element.findall(f".//{WORD_PARAGRAPH_TAG}"):
            inline_nodes = self._note_paragraph_to_inline(paragraph_element, note_part)
            if inline_nodes:
                content_nodes.append(AstParagraph(content=inline_nodes))

        return content_nodes

    def _note_paragraph_to_inline(self, paragraph_element: Any, note_part: Any) -> list[Node]:
        """Convert a note paragraph XML element into inline nodes."""
        try:
            from docx.text.paragraph import Paragraph

            paragraph = Paragraph(paragraph_element, note_part)  # type: ignore[call-arg]
            return self._process_paragraph_runs_to_inline(paragraph)
        except (TypeError, AttributeError):
            # Expected fallback when Paragraph construction fails
            return self._extract_inline_nodes_from_xml(paragraph_element)

    def _extract_inline_nodes_from_xml(self, paragraph_element: Any) -> list[Node]:
        """Fallback text extraction when python-docx objects are unavailable."""
        text_fragments: list[str] = []
        for text_element in paragraph_element.findall(f".//{WORD_TAG_PREFIX}t"):
            if text_element.text:
                text_fragments.append(text_element.text)

        if not text_fragments:
            return []

        combined_text = "".join(text_fragments)
        return [Text(content=combined_text)]

    def _load_comments(self, doc: "docx.document.Document") -> dict[str, CommentData]:
        comments: dict[str, CommentData] = {}

        if not self.options.include_comments:
            return comments

        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
        except ImportError:
            RT = None  # type: ignore[assignment,misc]

        comments_part = getattr(doc.part, "comments_part", None)
        if comments_part is None and RT is not None:
            comments_part = self._get_note_part(doc, RT.COMMENTS, "comments_part")  # type: ignore[attr-defined]

        if comments_part is None:
            return comments

        element = self._get_note_part_element(comments_part)
        if element is None:
            return comments

        for index, comment_element in enumerate(element.findall(f".//{WORD_TAG_PREFIX}comment"), start=1):
            comment_id = comment_element.get(WORD_ID_ATTR)
            if comment_id is None:
                continue

            author = comment_element.get(f"{WORD_TAG_PREFIX}author", "Unknown")
            date = comment_element.get(f"{WORD_TAG_PREFIX}date", "")

            text_parts: list[str] = []
            for paragraph_node in comment_element.findall(f".//{WORD_PARAGRAPH_TAG}"):
                for text_element in paragraph_node.findall(f".//{WORD_TAG_PREFIX}t"):
                    if text_element.text:
                        text_parts.append(text_element.text)

            text = " ".join(text_parts).strip()
            comments[comment_id] = CommentData(
                identifier=comment_id,
                label=f"comment{index}",
                author=author or "",
                date=date or "",
                text=text,
            )

        return comments

    def _process_comments(self) -> list[Node]:
        """Process collected comments as block-level Comment nodes.

        Creates Comment nodes with metadata. The renderer will decide how to
        present them (as HTML comments, blockquotes, etc.) based on its own options.
        """
        nodes: list[Node] = []

        if not self._comments_map:
            return nodes

        for comment in self._comments_map.values():
            if not comment.text:
                continue

            # Create Comment node with rich metadata
            # Renderer will decide presentation based on its comment_mode option
            nodes.append(
                Comment(
                    content=comment.text,
                    metadata={
                        "comment_type": "docx_review",
                        "identifier": comment.identifier,
                        "label": comment.label,
                        "author": comment.author,
                        "date": comment.date,
                    },
                )
            )

        return nodes


def _get_numbering_definitions(doc: "docx.document.Document") -> dict[str, dict[str, str]]:
    """Extract and cache numbering definitions from document.

    Returns a mapping of numId -> {level -> format_type} where format_type is 'bullet' or 'decimal'.
    """
    numbering_defs: dict[str, dict[str, str]] = {}

    if not hasattr(doc, "_part") or not hasattr(doc._part, "numbering_part"):
        return numbering_defs

    numbering_part = doc._part.numbering_part
    if not numbering_part:
        return numbering_defs

    try:
        numbering_xml = numbering_part._element

        # First, collect abstract numbering definitions
        abstract_nums = {}
        for elem in numbering_xml.iter():
            if elem.tag.endswith("abstractNum"):
                abstract_num_id = elem.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId"
                )
                if abstract_num_id:
                    levels = {}
                    for level_elem in elem.iter():
                        if level_elem.tag.endswith("lvl"):
                            level_id = level_elem.get(
                                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl"
                            )
                            if level_id is not None:
                                for child in level_elem.iter():
                                    if child.tag.endswith("numFmt"):
                                        fmt_val = child.get(
                                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                                        )
                                        if fmt_val:
                                            # Map Word numbering formats to our types
                                            if fmt_val in ("bullet", "none"):
                                                levels[level_id] = "bullet"
                                            elif fmt_val in (
                                                "decimal",
                                                "lowerLetter",
                                                "upperLetter",
                                                "lowerRoman",
                                                "upperRoman",
                                            ):
                                                levels[level_id] = "number"
                                            break
                    if levels:
                        abstract_nums[abstract_num_id] = levels

        # Then, map number IDs to abstract numbers
        for elem in numbering_xml.iter():
            if elem.tag.endswith("num"):
                num_id = elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId")
                if num_id:
                    for child in elem.iter():
                        if child.tag.endswith("abstractNumId"):
                            abs_id = child.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                            if abs_id in abstract_nums:
                                numbering_defs[num_id] = abstract_nums[abs_id]
                            break

    except Exception as e:
        logger.debug(f"Error parsing numbering definitions: {e}")

    return numbering_defs


def _detect_list_level(paragraph: "Paragraph", doc: "docx.document.Document" | None = None) -> tuple[str | None, int]:
    """Detect the list level of a paragraph based on its style, numbering, and indentation.

    Returns tuple of (list_type, level) where list_type is 'bullet' or 'number' and level is integer depth
    """
    # Check for Word native numbering properties first (works for all list styles including "List Paragraph")
    if hasattr(paragraph, "_p") and paragraph._p is not None:
        try:
            # Check for numPr (numbering properties) element
            num_pr = paragraph._p.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
            if num_pr is not None:
                # Get numbering level (Word uses 0-based indexing, we use 1-based)
                ilvl_elem = num_pr.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl")
                level = (
                    int(ilvl_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "0")) + 1
                    if ilvl_elem is not None
                    else 1
                )

                # Get numbering ID to determine list type
                num_id_elem = num_pr.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId")
                if num_id_elem is not None:
                    num_id = num_id_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")

                    # Look up the numbering definition if document is available
                    if doc and num_id:
                        numbering_defs = _get_numbering_definitions(doc)
                        if num_id in numbering_defs:
                            # Get the format for this level (use level-1 since Word is 0-based)
                            level_key = str(level - 1)
                            if level_key in numbering_defs[num_id]:
                                return numbering_defs[num_id][level_key], level
                            # If specific level not found, use level 0 as fallback
                            elif "0" in numbering_defs[num_id]:
                                return numbering_defs[num_id]["0"], level

                    # Fallback: detect type from paragraph text pattern
                    text = paragraph.text.strip()
                    if re.match(r"^\d+[.)]", text) or re.match(r"^[a-zA-Z][.)]", text):
                        return "number", level
                    else:
                        return "bullet", level
        except Exception:
            pass

    # Check for built-in list styles
    style_name = paragraph.style.name if paragraph.style else None
    if not style_name:
        return None, 0

    base_type = None
    style_level = 1

    # Handle "List Paragraph" style - check for numbering properties above
    if style_name == "List Paragraph":
        # If we got here, numbering properties weren't found or processed
        # This might be a list paragraph without proper numbering - treat as bullet by default
        return "bullet", 1
    elif match := re.match(r"List\s*Bullet\s?(?P<level>\d+)?", style_name, re.I):
        base_type = "bullet"
        style_level = int(match.group("level") or 1)
    elif match := re.match(r"List\s*Number\s?(?P<level>\d+)?", style_name, re.I):
        base_type = "number"
        style_level = int(match.group("level") or 1)

    # Check indentation level for additional nesting
    indent_level = 0
    try:
        indent = paragraph.paragraph_format.left_indent
        if indent:
            # Convert Pt to level (assume DEFAULT_INDENTATION_PT_PER_LEVEL per level)
            indent_level = int(indent.pt / DEFAULT_INDENTATION_PT_PER_LEVEL)
    except AttributeError:
        pass

    # If we have a list style, combine with indentation
    if base_type:
        final_level = max(style_level, style_level + indent_level)
        return base_type, final_level

    # Check indentation level for paragraphs without list styles
    if indent_level > 0:
        # Try to detect if numbered based on paragraph text
        if re.match(r"^\d+[.)]", paragraph.text.strip()):
            return "number", indent_level
        return "bullet", indent_level

    return None, 0


def _omml_local_name(tag: str | None) -> str:
    if not tag:
        return ""
    if "}" in tag:
        return tag.split("}", 1)[1]
    return tag


def _iter_omml_children(element: Any) -> list[Any]:
    children: list[Any] = []
    iterator = getattr(element, "__iter__", None)
    if iterator is None:
        return children
    for child in iterator():
        tag = getattr(child, "tag", None)
        if not isinstance(tag, str):
            continue
        if _omml_local_name(tag) == "rPr":
            continue
        children.append(child)
    return children


def _omml_find_child(element: Any, name: str) -> Any | None:
    for child in _iter_omml_children(element):
        if _omml_local_name(getattr(child, "tag", "")) == name:
            return child
    return None


def _omml_handle_container(element: Any) -> str:
    """Handle container elements that just concatenate children."""
    return "".join(_omml_to_latex(child) for child in _iter_omml_children(element))


def _omml_handle_text(element: Any) -> str:
    """Handle text elements."""
    return (getattr(element, "text", "") or "").strip()


def _omml_handle_fraction(element: Any) -> str:
    """Handle fraction elements (f)."""
    numerator = _omml_to_latex(_omml_find_child(element, "num"))
    denominator = _omml_to_latex(_omml_find_child(element, "den"))
    if numerator and denominator:
        return f"\\frac{{{numerator}}}{{{denominator}}}"
    return numerator or denominator


def _omml_handle_superscript(element: Any) -> str:
    """Handle superscript elements (sSup)."""
    base_expr = _omml_to_latex(_omml_find_child(element, "e"))
    sup_expr = _omml_to_latex(_omml_find_child(element, "sup"))
    if base_expr and sup_expr:
        return f"{base_expr}^{{{sup_expr}}}"
    return base_expr or sup_expr


def _omml_handle_subscript(element: Any) -> str:
    """Handle subscript elements (sSub)."""
    base_expr = _omml_to_latex(_omml_find_child(element, "e"))
    sub_expr = _omml_to_latex(_omml_find_child(element, "sub"))
    if base_expr and sub_expr:
        return f"{base_expr}_{{{sub_expr}}}"
    return base_expr or sub_expr


def _omml_handle_subsuperscript(element: Any) -> str:
    """Handle combined subscript and superscript elements (sSubSup)."""
    base_expr = _omml_to_latex(_omml_find_child(element, "e"))
    sub_expr = _omml_to_latex(_omml_find_child(element, "sub"))
    sup_expr = _omml_to_latex(_omml_find_child(element, "sup"))
    if base_expr:
        if sub_expr:
            base_expr = f"{base_expr}_{{{sub_expr}}}"
        if sup_expr:
            base_expr = f"{base_expr}^{{{sup_expr}}}"
        return base_expr
    return sub_expr or sup_expr


def _omml_handle_radical(element: Any) -> str:
    """Handle radical elements (rad)."""
    base_expr = _omml_to_latex(_omml_find_child(element, "base"))
    degree_expr = _omml_to_latex(_omml_find_child(element, "deg"))
    if base_expr and degree_expr:
        return f"\\sqrt[{degree_expr}]{{{base_expr}}}"
    if base_expr:
        return f"\\sqrt{{{base_expr}}}"
    return base_expr


def _omml_handle_nary(element: Any) -> str:
    """Handle n-ary operator elements (nary)."""
    char_node = _omml_find_child(element, "chr")
    symbol = "\\sum"
    if char_node is not None:
        symbol = char_node.get(f"{MATH_TAG_PREFIX}val", symbol) or symbol
    sub_expr = _omml_to_latex(_omml_find_child(element, "sub"))
    sup_expr = _omml_to_latex(_omml_find_child(element, "sup"))
    base_expr = _omml_to_latex(_omml_find_child(element, "e"))
    result = symbol
    if sub_expr:
        result += f"_{{{sub_expr}}}"
    if sup_expr:
        result += f"^{{{sup_expr}}}"
    if base_expr:
        result += base_expr
    return result


def _omml_handle_default(element: Any) -> str:
    """Handle default elements by concatenating text, children, and tail."""
    text = (getattr(element, "text", "") or "").strip()
    children_text = "".join(_omml_to_latex(child) for child in _iter_omml_children(element))
    tail = (getattr(element, "tail", "") or "").strip()
    return (text + children_text + tail).strip()


_OMML_HANDLERS: dict[str, Any] = {
    "oMathPara": lambda el: _omml_handle_container(el).strip(),
    "oMath": lambda el: _omml_handle_container(el).strip(),
    "r": _omml_handle_container,
    "t": _omml_handle_text,
    "f": _omml_handle_fraction,
    "num": _omml_handle_container,
    "den": _omml_handle_container,
    "e": _omml_handle_container,
    "base": _omml_handle_container,
    "sup": _omml_handle_container,
    "sub": _omml_handle_container,
    "sSup": _omml_handle_superscript,
    "sSub": _omml_handle_subscript,
    "sSubSup": _omml_handle_subsuperscript,
    "rad": _omml_handle_radical,
    "nary": _omml_handle_nary,
}


def _omml_to_latex(element: Any) -> str:
    """Convert Office Math Markup Language (OMML) element to LaTeX.

    Parameters
    ----------
    element : Any
        OMML element to convert

    Returns
    -------
    str
        LaTeX representation of the element

    """
    if element is None:
        return ""

    name = _omml_local_name(getattr(element, "tag", ""))
    if not name:
        return ""

    # Dispatch to handler function
    handler = _OMML_HANDLERS.get(name, _omml_handle_default)
    return handler(element)


def _iter_block_items(
    parent: Any, options: DocxOptions, base_filename: str = "document", attachment_sequencer: Any = None
) -> Any:
    """Generate a sequence of Paragraph and Table elements in order, handling images."""
    import docx.document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent_elm = parent.element.body if isinstance(parent, docx.document.Document) else parent._element

    for child in parent_elm.iterchildren():
        if child.tag.endswith("tbl"):
            yield Table(child, parent)  # type: ignore[call-arg]
        elif child.tag.endswith("p"):
            paragraph = Paragraph(child, parent)  # type: ignore[call-arg]

            # Check if paragraph contains an image
            has_image = False
            img_data: list[ImageData] = []

            for run in paragraph.runs:
                for pic in run._element.findall(
                    ".//pic:pic",
                    {"pic": "http://schemas.openxmlformats.org/drawingml/2006/picture"},
                ):
                    has_image = True

                    # Get image info
                    title = None

                    if options.include_image_captions:
                        if (t := run._element.xpath(".//wp:docPr/@descr")) or (
                            t := run._element.xpath(".//wp:docPr/@title")
                        ):
                            title = t[0]

                    # Get image data and detected format
                    blip = pic.xpath(".//a:blip")[0]
                    blip_rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    raw_image_data, detected_extension = extract_docx_image_data(parent, blip_rId)

                    # Use detected extension or fallback to png
                    extension = detected_extension or "png"

                    # Log format detection result
                    if detected_extension:
                        logger.debug(f"Detected image format: {detected_extension}")
                    else:
                        logger.debug("No image format detected, using PNG as fallback")

                    # Process image using unified attachment handling
                    # Use sequencer if available, otherwise fall back to manual counting
                    if attachment_sequencer:
                        image_filename, _ = attachment_sequencer(
                            base_stem=base_filename, format_type="general", extension=extension
                        )
                    else:
                        from all2md.utils.attachments import generate_attachment_filename

                        image_filename = generate_attachment_filename(
                            base_stem=base_filename,
                            format_type="general",
                            sequence_num=len(img_data) + 1,
                            extension=extension,
                        )
                    result = process_attachment(
                        attachment_data=raw_image_data,
                        attachment_name=image_filename,
                        alt_text=title or "image",
                        attachment_mode=options.attachment_mode,
                        attachment_output_dir=options.attachment_output_dir,
                        attachment_base_url=options.attachment_base_url,
                        is_image=True,
                        alt_text_mode=options.alt_text_mode,
                    )

                    if result.get("markdown"):
                        img_data.append(
                            ImageData(
                                url=result.get("url", ""),
                                alt_text=title or "image",
                                title=title,
                                footnote_label=result.get("footnote_label"),
                                footnote_content=result.get("footnote_content"),
                                source_data=result.get("source_data"),
                            )
                        )

            # Always yield the paragraph to preserve text content
            yield paragraph
            # Also yield any images found separately
            if has_image and img_data:
                for image_data in img_data:
                    yield image_data


# Converter metadata for registration
CONVERTER_METADATA = ConverterMetadata(
    format_name="docx",
    extensions=[".docx"],
    mime_types=["application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
    magic_bytes=[
        (b"PK\x03\x04", 0),  # ZIP signature (docx is ZIP-based)
    ],
    parser_class=DocxToAstConverter,
    renderer_class="all2md.renderers.docx.DocxRenderer",
    renders_as_string=False,
    parser_required_packages=[("python-docx", "docx", "")],
    renderer_required_packages=[("python-docx", "docx", ">=1.2.0")],
    optional_packages=[],
    import_error_message=("DOCX conversion requires 'python-docx'. " "Install with: pip install python-docx"),
    parser_options_class=DocxOptions,
    renderer_options_class="all2md.options.docx.DocxRendererOptions",
    description="Convert Microsoft Word DOCX documents to/from AST",
    priority=8,
)
